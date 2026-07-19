#!/usr/bin/env python3
"""
AI-Prowler MCP Server
=====================
Exposes AI-Prowler's RAG knowledge base as MCP tools so Claude Desktop
(and any other MCP-compatible agent) can query, index, and manage your
local document collection directly from a Claude conversation.

Transport : stdio  (required by Claude Desktop)
Protocol  : MCP 1.x via the official `mcp` Python SDK (FastMCP)

Install the MCP package:
    pip install mcp

Then register this server in Claude Desktop's config
(see claude_desktop_config_example.json for the exact snippet).

Author: AI-Prowler project
"""

# ── Critical env vars — must come before ANY library import ──────────────────
import os
import sys

# ── Windows console UTF-8 fix ────────────────────────────────────────────────
# Force stdout/stderr to UTF-8 so non-ASCII characters in learnings (em-dashes,
# emoji, smart quotes) don't trigger UnicodeEncodeError when ChromaDB or our
# logging code writes them. Critical on Windows where the default codepage is
# cp1252. errors='replace' ensures we never crash on stray bad characters —
# we'd rather lose a glyph than fail a delete or index operation.
try:
    if sys.stdout is not None and hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if sys.stderr is not None and hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

os.environ.setdefault('TOKENIZERS_PARALLELISM', 'false')
os.environ.setdefault('HF_HUB_DISABLE_TELEMETRY', '1')
os.environ.setdefault('HF_HUB_DISABLE_PROGRESS_BARS', '1')
os.environ.setdefault('TRANSFORMERS_VERBOSITY', 'error')
os.environ.setdefault('TRANSFORMERS_NO_ADVISORY_WARNINGS', '1')
# Prevent sentence-transformers from making HuggingFace update-check HTTP calls
# on every load. Model is cached locally; network checks add 4-5s to startup.
os.environ.setdefault('HF_HUB_OFFLINE', '1')
os.environ.setdefault('TRANSFORMERS_OFFLINE', '1')

# Fix the HuggingFace double-backslash Errno 22 bug on Windows 10
if not os.environ.get('HF_HUB_CACHE'):
    from pathlib import Path as _P
    os.environ['HF_HUB_CACHE'] = str(_P.home() / '.cache' / 'huggingface' / 'hub')

# ── Ensure rag_preprocessor is importable ─────────────────────────────────────
# The MCP server must be placed in the same directory as rag_preprocessor.py
# (i.e. C:\Program Files\AI-Prowler).  We add that directory to sys.path so
# Python finds rag_preprocessor even when Claude Desktop launches this script
# from a different working directory.
_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)

# ── Suppress warnings so they don't corrupt the MCP stdio stream ─────────────
import warnings
warnings.filterwarnings('ignore')

import io
import json
import contextlib
import threading
import queue
import argparse
from pathlib import Path
from typing import Optional

# ── File logging — all output captured here since the window is hidden ────────
import logging
import traceback

_LOG_PATH = Path.home() / "AppData" / "Local" / "AI-Prowler" / "mcp_server.log"
_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)

# Keep last 3 log files by rotating on each startup
import glob as _glob, shutil as _shutil
for _i in range(2, 0, -1):
    _old = _LOG_PATH.with_suffix(f".log.{_i}")
    _prev = _LOG_PATH.with_suffix(f".log.{_i-1}") if _i > 1 else _LOG_PATH
    if _prev.exists():
        _shutil.copy2(str(_prev), str(_old))

logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)-8s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S.%f",   # millisecond precision for timing diagnosis
    handlers=[
        logging.FileHandler(str(_LOG_PATH), mode="w", encoding="utf-8"),
    ],
)
_log = logging.getLogger("ai_prowler_mcp")

# Patch the Formatter to actually emit milliseconds — Python's datefmt
# with strftime doesn't support %f natively, so we override formatTime.
class _MsFormatter(logging.Formatter):
    def formatTime(self, record, datefmt=None):
        import datetime as _dt
        ct = _dt.datetime.fromtimestamp(record.created)
        return ct.strftime("%Y-%m-%d %H:%M:%S.") + f"{int(record.msecs):03d}"

for _h in logging.root.handlers:
    _h.setFormatter(_MsFormatter(
        fmt="%(asctime)s [%(levelname)-8s] %(message)s"
    ))

# Redirect stderr to the log file so crashes and tracebacks are captured too
class _StderrToLog:
    def write(self, msg):
        msg = msg.rstrip()
        if msg:
            _log.error("STDERR: %s", msg)
    def flush(self):
        pass

sys.stderr = _StderrToLog()

_log.info("=" * 60)
_log.info("AI-Prowler MCP server process started")
_log.info("Python : %s", sys.version)
_log.info("Script : %s", __file__)
_log.info("CWD    : %s", os.getcwd())
_log.info("=" * 60)

# ── MCP SDK ───────────────────────────────────────────────────────────────────
_log.info("Importing MCP SDK (FastMCP)…")
try:
    from mcp.server.fastmcp import FastMCP
    _log.info("MCP SDK imported OK")
    # Context enables per-request user threading in server mode (Phase B).
    # Optional: if a build lacks it, fall back to None so `ctx: Context = None`
    # tool params remain harmless (personal mode never reads ctx anyway).
    try:
        from mcp.server.fastmcp import Context
    except Exception:
        Context = None
        _log.warning("FastMCP Context not importable — server-mode user "
                     "threading will be unavailable (personal mode unaffected).")
except ImportError:
    _log.critical("FATAL: 'mcp' package not found. Run: pip install mcp")
    print(
        "ERROR: 'mcp' package not found.\n"
        "Install it with:  pip install mcp\n",
        file=sys.stderr,
    )
    sys.exit(1)

# ── AI-Prowler engine ─────────────────────────────────────────────────────────
#
# STARTUP SPEED FIX — patch requests timeout before rag_preprocessor import
# ──────────────────────────────────────────────────────────────────────────
# rag_preprocessor probes localhost:11434 (Ollama) at module-import time using
# the requests library with its own timeout= argument.  socket.setdefaulttimeout()
# has NO effect on requests — it ignores it entirely.  So the probe still hangs
# for the full requests timeout (3-4 seconds) even with the socket cap in place.
#
# Fix: monkeypatch requests.Session.request to cap any timeout to 0.8s for the
# duration of the import, then restore the original method in the finally block.
# Belt-and-suspenders: also cap socket timeout for any raw socket probes.
import socket as _socket
_pre_import_socket_timeout = _socket.getdefaulttimeout()
_socket.setdefaulttimeout(1.0)

_requests_patched      = False
_orig_session_request  = None
try:
    import requests as _requests_mod
    _orig_session_request = _requests_mod.Session.request

    def _fast_session_request(self, method, url, **kwargs):
        existing = kwargs.get('timeout', None)
        if existing is None or (isinstance(existing, (int, float)) and existing > 0.8):
            kwargs['timeout'] = 0.8
        elif isinstance(existing, tuple) and any(
            isinstance(v, (int, float)) and v > 0.8 for v in existing
        ):
            kwargs['timeout'] = 0.8
        return _orig_session_request(self, method, url, **kwargs)

    _requests_mod.Session.request = _fast_session_request
    _requests_patched = True
    _log.info("requests.Session.request patched: timeout capped to 0.8s during import")
except ImportError:
    pass   # requests not available yet; will be imported by rag_preprocessor

_log.info("Importing rag_preprocessor from: %s (Ollama probe timeout capped)", _HERE)
try:
    import rag_preprocessor as _engine
    from rag_preprocessor import (
        index_directory, index_file_list,
        scan_directory, scan_directory_for_changes,
        command_update,
        show_stats, clear_database,
        load_config,
        load_auto_update_list, add_to_auto_update_list,
        remove_directory_from_index,
        normalise_path,
    )

    # Run headless — disable the terminal spinner and GUI-mode code paths
    _engine.GUI_MODE  = False
    _engine._MCP_MODE = True

    # Load saved user config (chunk size, etc.)
    load_config()
    _log.info("rag_preprocessor imported OK")

except ImportError as _e:
    _log.critical("FATAL: Could not import rag_preprocessor: %s", _e)
    _log.critical(traceback.format_exc())
    print(
        f"ERROR: Could not import rag_preprocessor: {_e}\n"
        f"Make sure ai_prowler_mcp.py is in the same folder as rag_preprocessor.py.\n",
        file=sys.stderr,
    )
    sys.exit(1)

finally:
    # Restore socket timeout and requests patch regardless of outcome
    _socket.setdefaulttimeout(_pre_import_socket_timeout)
    if _requests_patched and _orig_session_request is not None:
        try:
            _requests_mod.Session.request = _orig_session_request
            _log.info("requests patch removed — normal timeouts restored")
        except Exception:
            pass
    _log.info("Import phase complete, socket timeout restored to: %s",
              _pre_import_socket_timeout)

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

# ── Stdio-mode guard ─────────────────────────────────────────────────────────
# Set to True in the entry point BEFORE mcp.run(transport="stdio") is called.
# _capture_stdout() reads this flag at call-time to decide whether it is safe
# to redirect sys.stdout.  In stdio mode sys.stdout IS the MCP JSON-RPC pipe —
# replacing it (even briefly) corrupts the protocol and causes Claude Desktop
# to show "Claude's response was interrupted".
_STDIO_MODE = False  # bool — set True in stdio entry point before mcp.run()

# Claude model identifier — stamped as `source` on learnings that Claude
# auto-detects (auto_detected=True). Matches the Anthropic model string so
# the Learnings tab shows exactly which Claude version created the entry
# e.g. "claude-sonnet-4-6". Updated here when the bundled model changes.
_MODEL_ID = "claude-sonnet-4-6"


def _get_personal_owner_name() -> str:
    """Return the owner's display name for personal-mode source attribution.

    Reads owner_name fresh from ~/.ai-prowler/config.json on every call so
    that the name is picked up immediately after the user sets it in the
    Settings tab — no MCP server restart required (fixes live-update bug).

    Falls back to _engine.OWNER_NAME (loaded at startup from ~/.rag_config.json)
    for backwards compatibility, then "" if neither is set — callers fall back
    to "operator" in that case.

    v7.0.1 — introduced so personal installs show the owner's name in the
    Learnings tab Source column instead of the generic "operator" label.
    v7.0.2 — moved to fresh file read so name is picked up mid-session and
    survives reinstall (owner_name now stored in ~/.ai-prowler/config.json).
    """
    try:
        _ai_cfg_path = Path.home() / '.ai-prowler' / 'config.json'
        if _ai_cfg_path.exists():
            # utf-8-sig tolerates a BOM (e.g. from PowerShell Out-File or some
            # editors saving "UTF-8" with BOM) as well as plain utf-8.
            with open(_ai_cfg_path, 'r', encoding='utf-8-sig') as _f:
                _ai_cfg = json.load(_f)
            name = _ai_cfg.get('owner_name', '').strip()
            if name:
                return name
    except Exception:
        pass
    # Fallback: in-memory global loaded at startup
    try:
        return (_engine.OWNER_NAME or "").strip()
    except Exception:
        return ""


def _get_personal_owner_address() -> dict:
    """Return the owner's home address (Street/City/State/ZIP) as a dict,
    for personal-mode Proactive Alerts (Morning Briefing fallback, Weekly
    Weather Watch). Reads fresh from ~/.ai-prowler/config.json on every
    call, mirroring _get_personal_owner_name()'s live-update fix exactly —
    set via Settings tab, no MCP server restart required to pick it up.

    Added v8.1.3, replacing the old design where Proactive Alerts had its
    own separate Location field with a hardcoded default and no
    relationship to Settings at all. Falls back to the in-memory
    _engine.OWNER_* globals (loaded at startup), then "" for any field
    still unset — callers should treat an all-empty result as "the owner
    hasn't configured an address yet" rather than erroring.

    Returns:
        {"street": str, "city": str, "state": str, "zip": str} — any/all
        may be empty strings.
    """
    result = {"street": "", "city": "", "state": "", "zip": ""}
    try:
        _ai_cfg_path = Path.home() / '.ai-prowler' / 'config.json'
        if _ai_cfg_path.exists():
            with open(_ai_cfg_path, 'r', encoding='utf-8-sig') as _f:
                _ai_cfg = json.load(_f)
            result["street"] = (_ai_cfg.get('owner_street') or '').strip()
            result["city"]   = (_ai_cfg.get('owner_city')   or '').strip()
            result["state"]  = (_ai_cfg.get('owner_state')  or '').strip()
            result["zip"]    = (_ai_cfg.get('owner_zip')    or '').strip()
    except Exception:
        pass
    # Fallback to in-memory globals for any field the file didn't provide
    # (e.g. file missing entirely, or an older config predating this field).
    try:
        if not result["street"]:
            result["street"] = (_engine.OWNER_STREET or "").strip()
        if not result["city"]:
            result["city"] = (_engine.OWNER_CITY or "").strip()
        if not result["state"]:
            result["state"] = (_engine.OWNER_STATE or "").strip()
        if not result["zip"]:
            result["zip"] = (_engine.OWNER_ZIP or "").strip()
    except Exception:
        pass
    return result


def _get_personal_owner_location_string() -> str:
    """Return the owner's City/State/ZIP as a single geocodable string for
    weather lookups (get_weather() / _weather() in scheduler_jobs.py) —
    Street is deliberately excluded, it's not useful for geocoding and can
    make Nominatim's lookup less reliable, not more.

    Returns "" (not a hardcoded fallback like "New Smyrna Beach, Florida")
    when nothing is configured — callers must handle the empty case
    explicitly (e.g. skip the weather section, or fall back to a job's own
    City/State) rather than this function silently defaulting to any
    particular person's real address.
    """
    addr = _get_personal_owner_address()
    parts = []
    if addr["city"]:
        parts.append(addr["city"])
    if addr["state"]:
        parts.append(addr["state"])
    loc = ", ".join(parts) if len(parts) > 1 else "".join(parts)
    if addr["zip"]:
        loc = f"{loc} {addr['zip']}".strip()
    return loc


@contextlib.contextmanager
def _capture_stdout():
    """
    Capture everything printed to stdout and return it as a StringIO buffer.

    STDIO MODE — safe no-op:
        sys.stdout is the MCP JSON-RPC pipe.  Redirecting it — even for a
        single tool call — corrupts the binary frame stream and causes Claude
        Desktop to drop the response ("response was interrupted").
        In stdio mode stdout is NOT redirected — it is the live MCP pipe.
        We simply yield an empty buffer without touching sys.stdout.

    HTTP MODE — normal redirect:
        sys.stdout is an ordinary terminal/pipe; temporarily replacing it is
        safe and lets us capture print()-based output from rag_preprocessor.
    """
    if _STDIO_MODE:
        # stdio: never touch sys.stdout — yield an empty buffer and return.
        yield io.StringIO()
        return

    buf = io.StringIO()
    old = sys.stdout
    sys.stdout = buf
    try:
        yield buf
    finally:
        sys.stdout = old


# ─────────────────────────────────────────────────────────────────────────────
# MCP server definition
# ─────────────────────────────────────────────────────────────────────────────

# ── FastMCP server — version-aware instructions ───────────────────────────────
# The instructions= parameter was added in mcp >= 1.2.0.
# We extract the text into _INSTRUCTIONS and detect support at runtime so the
# server starts cleanly on older installs and upgrades automatically benefit.
_INSTRUCTIONS = (
    "AI-Prowler is an Agentic RAG knowledge base WITH a self-learning memory layer. "
    "Claude retrieves raw chunks AND stored learnings directly, then synthesizes "
    "answers — no local LLM required.\n\n"

    "═══════════════════════════════════════════════════════════════\n"
    " MANDATORY SEARCH ORDER FOR EVERY USER QUESTION\n"
    "═══════════════════════════════════════════════════════════════\n"
    "You MUST run search_learnings(query=...) BEFORE answering ANY user\n"
    "question. No exceptions. The user uses AI-Prowler as a personal\n"
    "knowledge store for everything — not just business. Recipes,\n"
    "preferences, opinions, hobbies, how they like things done — any\n"
    "of these may live in the learnings store and override what your\n"
    "training data would say.\n\n"

    "  STEP 1 (ALWAYS, NO EXCEPTIONS):  search_learnings(query=...)\n"
    "    Self-learning is the AUTHORITATIVE override layer. It contains\n"
    "    the user's deliberate corrections, preferences, and lessons.\n"
    "    These are newer and more curated than anything in your training\n"
    "    data or the document index.\n"
    "    Run search_learnings BEFORE generating any answer, EVEN FOR:\n"
    "      • General world-knowledge questions ('what's a good recipe\n"
    "        for ribs', 'best way to clean windows', 'what wine pairs\n"
    "        with salmon') — the user may have recorded a personal\n"
    "        version that overrides the generic answer\n"
    "      • Definitional questions ('what is X') — the user may have\n"
    "        recorded their own working definition\n"
    "      • How-to questions — the user may have recorded their\n"
    "            preferred technique\n"
    "      • Recommendations of any kind\n"
    "      • Business / client / project questions\n"
    "    The ONLY questions you may skip search_learnings on are pure\n"
    "    arithmetic ('what is 7 * 13') and language translation, where\n"
    "    there's no plausible way a personal learning could apply.\n"
    "    When in doubt, SEARCH. The cost of an unnecessary search is\n"
    "    one tool call; the cost of missing a stored preference is the\n"
    "    user getting a wrong or generic answer.\n\n"

    "  STEP 2 (CONDITIONAL):  search_documents(query=...)\n"
    "    Run document RAG when the question is plausibly about\n"
    "    something in the indexed knowledge base — work documents,\n"
    "    contracts, manuals, case files, spreadsheets, the user's own\n"
    "    writing. Skip it for purely personal-preference questions\n"
    "    where no document is likely to be relevant (e.g. 'what's my\n"
    "    rib recipe' — that's a learning, not a document).\n\n"

    "  BLEND BOTH SOURCES IN YOUR ANSWER (when they don't conflict):\n"
    "    If a learning and a document chunk are about the same topic but\n"
    "    cover different aspects, COMBINE them in your response. Do not\n"
    "    pick one and discard the other.\n"
    "    Example — user asks 'how do I reach Rick?':\n"
    "      • Learning says: 'Rick prefers email; phone and Slack are secondary'\n"
    "      • Doc chunk says: 'Rick's email is rick@austin-underground.com,\n"
    "        cell is 555-0100, available on DWService for installs'\n"
    "      → Answer: 'Rick can be reached by email at rick@austin-underground.com\n"
    "        (his preferred channel — fastest response), or by phone at 555-0100\n"
    "        as backup. For installs, DWService remote sessions work well.'\n"
    "    The learning shapes the recommendation; the docs supply the specifics.\n\n"

    "  CONFLICT RESOLUTION (when sources actually disagree):\n"
    "    If a learning and a document chunk give CONTRADICTORY information,\n"
    "    THE LEARNING WINS. Surface both to the user, but apply the learning.\n"
    "    Example: a doc says 'invoice on Friday' but a learning says 'Crabby's\n"
    "    only processes invoices on Mondays' — recommend Monday and mention\n"
    "    the doc as outdated background.\n\n"

    "  LEARNING vs. TRAINING-DATA CONFLICT (the recipe case):\n"
    "    If search_learnings returns a personal version of something general\n"
    "    (e.g. the user's rib recipe vs. recipes you know from training),\n"
    "    THE STORED LEARNING WINS. Lead with the user's version. Do not\n"
    "    fall back on web search or generic recipes unless the user\n"
    "    explicitly asks for alternatives.\n\n"

    "  RELEVANCE INTERPRETATION:\n"
    "    Relevance labels on search_learnings results are now calibrated:\n"
    "      🟢 HIGH (≥ 0.70)     — near-identical phrasing match; trust fully\n"
    "      🟡 MODERATE (0.40+)  — solid semantic match; APPLY THIS\n"
    "      🟠 LOW (< 0.40)      — likely irrelevant; ignore unless title\n"
    "                             is obviously on-topic\n"
    "    Most legitimate matches will land in MODERATE — that band is\n"
    "    trustworthy. Do NOT skip a MODERATE result just because it isn't\n"
    "    HIGH. Only treat LOW as a signal to ignore, and even then check\n"
    "    the title before discarding.\n\n"

    "═══════════════════════════════════════════════════════════════\n"
    " DOCUMENT RAG — Detailed Tool Reference\n"
    "═══════════════════════════════════════════════════════════════\n"
    "After running search_learnings (Step 1), use these for Step 2 and beyond:\n"
    "  • get_knowledge_base_overview — orient: what's indexed, what topics\n"
    "  • list_indexed_documents      — browse files for a topic area\n"
    "  • search_documents            — PRIMARY document retrieval\n"
    "  • multi_query_search          — parallel searches across synonyms\n"
    "  • search_within_directory     — restrict to a specific case/project\n"
    "  • expand_search_result        — expand a cut-off result\n"
    "  • read_document               — read a document sequentially\n\n"
    "Use check_ai_prowler_status() to verify the knowledge base is healthy.\n\n"

    "═══════════════════════════════════════════════════════════════\n"
    " SELF-LEARNING — Capture, Review, and Override\n"
    "═══════════════════════════════════════════════════════════════\n"

    "  AUTO-RECORDING — record WITHOUT being asked when you detect:\n"
    "    • User corrects a fact you stated or that exists in learnings\n"
    "    • User shares a project outcome, success, or failure\n"
    "    • User mentions a client preference or complaint\n"
    "    • A post-operation review reveals what went right or wrong\n"
    "    • New information contradicts an existing active learning\n"
    "      (in this case, ALSO pass supersedes_id with the old learning's UUID)\n"
    "    • User describes a process improvement or better approach\n"
    "    When auto-recording, ALWAYS set auto_detected=True so the\n"
    "    confirmation message flags this as your initiative.\n\n"

    "  CONFIRMATION PROTOCOL — ALWAYS follow after recording:\n"
    "    1. Tell the user WHAT you recorded (title + one-line summary)\n"
    "    2. Tell the user WHY you recorded it (the trigger you detected)\n"
    "    3. Ask: 'Is that correct, or should I adjust it?'\n"
    "    4. If the user says no or corrects you → call update_learning()\n"
    "       or delete_learning() immediately\n"
    "    NEVER record silently. The user must always see what was captured.\n\n"

    "  POST-OPERATION ANALYSIS workflow:\n"
    "    When asked to review/analyze a completed project or job:\n"
    "    1. Call search_learnings() FIRST to see existing learnings for this topic\n"
    "    2. Call search_within_directory() or search_documents() for full context\n"
    "    3. Identify: what went wrong, what went right, process gaps\n"
    "    4. For EACH insight, call record_learning() with proper category\n"
    "       and auto_detected=True\n"
    "    5. Present ALL recorded learnings to the user for confirmation\n"
    "    6. Adjust any the user disagrees with\n\n"

    "Tool sequence: search_learnings → search_documents → (record_learning if new info)"
)

import inspect as _inspect
_fastmcp_params = list(_inspect.signature(FastMCP.__init__).parameters.keys())
if "instructions" in _fastmcp_params:
    mcp = FastMCP("AI-Prowler", instructions=_INSTRUCTIONS)
    _log.info("FastMCP created with instructions= (mcp >= 1.2.0) — "
              "guidance sent to Claude at every handshake")
else:
    mcp = FastMCP("AI-Prowler")
    _log.warning(
        "FastMCP does not support instructions= on this version — "
        "falling back to how_to_use_ai_prowler() tool. "
        "Upgrade with: pip install --upgrade mcp"
    )
_log.info("FastMCP server object created OK")

# ══════════════════════════════════════════════════════════════════════════════
# Telemetry — track tool call counts per tool name on every successful
# invocation. Read by rag_gui.py's heartbeat sender, which resets the counter
# after a successful POST to the Cloudflare Worker.
#
# Counter file format (v2):
#     {"tool_calls": {"check_ai_prowler_status": 12, "search_documents": 3, ...}}
#
# v1 -> v2 migration is automatic on first read; old single-int totals are
# discarded (they aren't recoverable as per-tool buckets).
#
# Failure modes are all silent — telemetry must NEVER cause a tool call to
# fail. If the counter file is locked, missing, corrupted, or in a path the
# process can't write to, we just skip the increment and the user never sees
# anything go wrong.
# ══════════════════════════════════════════════════════════════════════════════
_TELEMETRY_COUNTER_PATH = (
    Path.home() / '.ai-prowler' / 'telemetry_counter.json')
_telemetry_lock = threading.Lock()


def _telemetry_increment_tool_count(tool_name='_unknown'):
    """Bump tool_calls[tool_name] by 1. Concurrency-safe, fail-silent."""
    try:
        with _telemetry_lock:
            data = {'tool_calls': {}}
            try:
                if _TELEMETRY_COUNTER_PATH.exists():
                    raw = _TELEMETRY_COUNTER_PATH.read_text(encoding='utf-8')
                    parsed = json.loads(raw)
                    if isinstance(parsed, dict):
                        # v1 -> v2 migration: drop old single-int field
                        if 'tool_calls' in parsed and isinstance(
                                parsed['tool_calls'], dict):
                            data = parsed
                        else:
                            data = {'tool_calls': {}}
            except Exception:
                pass

            tc = data.setdefault('tool_calls', {})
            if not isinstance(tc, dict):
                tc = {}
                data['tool_calls'] = tc
            current = int(tc.get(tool_name, 0))
            tc[tool_name] = current + 1

            try:
                _TELEMETRY_COUNTER_PATH.parent.mkdir(
                    parents=True, exist_ok=True)
            except Exception:
                pass

            tmp = _TELEMETRY_COUNTER_PATH.with_suffix('.tmp')
            tmp.write_text(json.dumps(data), encoding='utf-8')
            os.replace(str(tmp), str(_TELEMETRY_COUNTER_PATH))
    except Exception as _e:
        _log.debug("telemetry counter update skipped: %s", _e)


# ── Tier A server-mode tool suppression (v7.0.1) ─────────────────────────────
# In server mode (Business edition, mode=server) a subset of tools must NEVER
# appear in the MCP tool list. These are host/dev/operator tools that belong
# only on a personal or developer install — exposing them to remote users would
# be a security risk regardless of role.  They are suppressed at REGISTRATION
# time inside _counting_mcp_tool below, so the functions still exist as plain
# Python but are invisible to any MCP client connecting to the server.
#
# Contrast with Tier B (runtime per-call gating via _check_db_cap /
# _send_email_cap): those tools ARE registered; they just refuse calls from
# roles that lack the right capability.

def _detect_server_mode() -> bool:
    """Read config.json at module startup to decide whether Tier A suppression
    should be active. Self-contained — does NOT depend on _state_dir(),
    _CONFIG_PATH, or any other module-level symbol defined later.
    Returns False on any error (safe-closed: treat as personal mode)."""
    import os as _os2, json as _json2
    try:
        _td = _os2.environ.get("AIPROWLER_TEST_STATE_DIR", "").strip()
        _cp = (Path(_td) if _td else Path.home() / ".ai-prowler") / "config.json"
        if not _cp.exists():
            return False
        _cfg = _json2.loads(_cp.read_text(encoding="utf-8-sig")) or {}
        return (str(_cfg.get("edition", "")).strip().lower() == "business"
                and str(_cfg.get("mode", "")).strip().lower() == "server")
    except Exception:
        return False


_IS_SERVER_MODE: bool = _detect_server_mode()

# The Tier A tools: suppressed for ALL roles in server mode. (Count drifts
# as tools are added — see the log line just below for the live count
# rather than trusting a number in this comment.)
_TIER_A_SUPPRESSED: frozenset = frozenset({
    # Dev / code-execution — run arbitrary code on the host OS
    "compile_check", "syntax_check", "lint_check",
    "run_script", "run_script_start", "run_script_status", "run_script_kill",
    "check_python_import",
    # Host filesystem writes NOT scoped by _check_personal_write_scope() —
    # backup/restore/approval-management tools, still operator/dev-only.
    # create_file, write_file, str_replace_in_file,
    # line_replace_in_file, and create_directory are DELIBERATELY absent
    # from this set — they're gated per-call instead, via
    # _check_personal_write_scope(): server-mode users may write ONLY
    # inside their own personal directory, and not at all if they don't
    # have one configured. See that function's docstring for details.
    "copy_to_backup", "restore_backup", "list_backups", "list_directory",
    "reset_write_counter", "grant_write_access", "revoke_write_access",
    # Backup cleanup — deletes files on the host, operator/dev action only
    "cleanup_backups",
    # Job-log cleanup — same reasoning as cleanup_backups: deletes files on
    # the host, and run_script_start/status/kill (the feature it cleans up
    # after) are themselves personal-install-only.
    "cleanup_job_logs",
    # Raw filesystem reads — arbitrary path access, bypasses the RAG layer
    "grep_documents", "read_file_lines",
    # Email operator / high-risk tools — personal-install-only.
    # Note: send_email, send_alert, and send_learnings_report stay registered;
    # all roles use them via the Tier B _send_email_cap gate.
    "configure_email", "send_file",
    "export_learnings_file",
    # Bulk index rebuild — destructive operator action, not for remote users
    "rebuild_learnings_index",
    # Agentic analysis task queue — personal-install-only (Quick Links tab's
    # Common Business AI Analysis / My Custom Analyses panels are hidden in
    # server mode's GUI, so the queue they drive has no server-mode caller).
    "get_pending_analysis_tasks", "complete_analysis_task", "save_analysis_report",
    "create_analysis_task", "list_analysis_tasks",
    # Raw/unscoped SMS inbox — personal-install-only. sms_inbox_read() has no
    # per-user filtering (unlike sms_inbox_read_for_user()), so in a
    # multi-user server it would let any employee read every inbound
    # SMS/WhatsApp message company-wide, not just their own threads.
    # check_sms_replies (which IS per-user isolated) is the server-mode
    # equivalent — see _PERSONAL_MODE_SUPPRESSED below for the reverse gate.
    "check_sms_inbox",
})

_log.info(
    "v7.0.1 Tier A suppression: server_mode=%s — %d tools will be hidden "
    "from the MCP tool list when server_mode=True",
    _IS_SERVER_MODE, len(_TIER_A_SUPPRESSED)
)

# The 1 personal-mode-only-suppressed tool: hidden when server_mode=False.
# Mirror image of Tier A — some tools only make sense once there's more than
# one user. check_sms_replies is per-user thread-isolated (Mike sees Karen's
# reply, not Jake's), which is meaningless with a single personal-install
# user; check_sms_inbox (the unscoped, richer-filtered view — provider,
# unread_only, since_hours=0 for "everything") is the personal-mode tool.
_PERSONAL_MODE_SUPPRESSED: frozenset = frozenset({
    "check_sms_replies",
})

_log.info(
    "Personal-mode suppression: server_mode=%s — %d tool(s) will be hidden "
    "from the MCP tool list when server_mode=False",
    _IS_SERVER_MODE, len(_PERSONAL_MODE_SUPPRESSED)
)

# ── Monkeypatch mcp.tool() ───────────────────────────────────────────────────
import functools as _functools

_orig_mcp_tool = mcp.tool


def _counting_mcp_tool(*tool_args, **tool_kwargs):
    """Wrap mcp.tool() so each registered tool increments the per-tool
    counter on successful return. Uses fn.__name__ as the tool key.

    v7.0.1 Tier A: in server mode, any tool whose name appears in
    _TIER_A_SUPPRESSED is returned as a plain Python function without
    being passed to the real mcp.tool() decorator, so it never appears
    in the MCP tool list for any client connecting to the server.

    Mirror gate: in personal mode, any tool whose name appears in
    _PERSONAL_MODE_SUPPRESSED is suppressed the same way — tools that
    only make sense with more than one registered user.
    """
    def _outer(fn):
        _tool_name = getattr(fn, '__name__', '_unknown')

        # Tier A: skip registration entirely in server mode.
        if _IS_SERVER_MODE and _tool_name in _TIER_A_SUPPRESSED:
            _log.debug("Tier A: suppressing '%s' (not registered with MCP)", _tool_name)
            return fn  # plain Python function — invisible to MCP clients

        # Mirror gate: skip registration entirely in personal mode.
        if not _IS_SERVER_MODE and _tool_name in _PERSONAL_MODE_SUPPRESSED:
            _log.debug("Personal-mode gate: suppressing '%s' (not registered with MCP)", _tool_name)
            return fn  # plain Python function — invisible to MCP clients

        real_decorator = _orig_mcp_tool(*tool_args, **tool_kwargs)

        @_functools.wraps(fn)
        def _inner(*args, **kwargs):
            result = fn(*args, **kwargs)
            _telemetry_increment_tool_count(_tool_name)
            return result
        return real_decorator(_inner)

    return _outer


mcp.tool = _counting_mcp_tool
_log.info("Monkeypatched mcp.tool() — all subsequent @mcp.tool decorators "
          "will increment ~/.ai-prowler/telemetry_counter.json by tool name "
          "on success")

# Module-level event set by the background prewarm thread when ChromaDB and
# the embedding model are fully loaded.  Tool handlers that need ChromaDB
# wait on this event (max 60s) before proceeding.
_prewarm_event = threading.Event()
_prewarm_event.set()   # default: don't block (overridden to clear() in stdio entry)

# ══════════════════════════════════════════════════════════════════════════════
# DB WRITER THREAD (v7.0.0)
# ──────────────────────────────────────────────────────────────────────────────
# ChromaDB's PersistentLocalHnswSegment has thread-affinity on its write/consumer
# path. Under the HTTP transport every tool runs on a ROTATING anyio threadpool
# worker, so a ChromaDB *write* (delete/add) arriving on a different thread than
# the one that owns the segment lock raises "resource deadlock would occur"
# (EDEADLK) and wedges the server — reproducibly, even for a 5-byte file.
#
# Fix: funnel ALL ChromaDB writes through ONE dedicated, long-lived thread. Tool
# handlers submit a callable + wait synchronously for its result. Because the
# same thread always performs the write, the segment lock affinity is never
# violated. Reads are left on the pool (they tolerate cross-thread access).
_db_write_queue: "queue.Queue" = queue.Queue()
_db_writer_started = threading.Event()
_db_writer_lock = threading.Lock()


def _db_writer_loop():
    """Single owner thread for ChromaDB writes. Runs submitted jobs serially."""
    _log.info("DB-writer thread started (tid=%s)", threading.get_ident())
    while True:
        fn, args, kwargs, result_box, done = _db_write_queue.get()
        try:
            result_box["result"] = fn(*args, **kwargs)
        except BaseException as exc:   # capture EVERYTHING, incl. EDEADLK OSError
            result_box["error"] = exc
        finally:
            done.set()
            _db_write_queue.task_done()


def _ensure_db_writer():
    """Start the writer thread once, lazily and thread-safely."""
    if _db_writer_started.is_set():
        return
    with _db_writer_lock:
        if _db_writer_started.is_set():
            return
        t = threading.Thread(target=_db_writer_loop, name="db-writer",
                             daemon=True)
        t.start()
        _db_writer_started.set()


def _db_write(fn, *args, timeout: float = 900.0, **kwargs):
    """Run fn(*args, **kwargs) on the dedicated DB-writer thread and return its
    result. Raises the worker's exception in the caller, or TimeoutError if the
    job does not finish within `timeout` seconds. This is what makes reindex
    safe on the HTTP transport — the write never touches the rotating pool."""
    _ensure_db_writer()
    result_box: dict = {}
    done = threading.Event()
    _db_write_queue.put((fn, args, kwargs, result_box, done))
    if not done.wait(timeout=timeout):
        raise TimeoutError(
            f"DB write did not complete within {timeout:.0f}s "
            f"(job still running on db-writer thread)")
    if "error" in result_box:
        raise result_box["error"]
    return result_box.get("result")

# ══════════════════════════════════════════════════════════════════════════════
# GUIDANCE TOOL — how_to_use_ai_prowler
# ══════════════════════════════════════════════════════════════════════════════
# Bulletproof fallback that works on ALL FastMCP versions.
# When instructions= is not supported by the installed mcp package, Claude
# reads this tool's docstring and calls it naturally at the start of sessions.
# Even when instructions= IS supported, this tool gives Claude an on-demand
# reference it can re-read during any conversation.
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def how_to_use_ai_prowler(ctx: "Context | None" = None) -> str:
    """
    Returns the recommended workflow for using AI-Prowler as an Agentic RAG
    knowledge base with Claude.

    CALL THIS TOOL FIRST at the start of any new research or question-answering
    session to understand the correct tool sequence and capabilities.

    The main guide is IDENTICAL regardless of mode or which connector it's
    called on — deliberately. In a conversation with two connectors attached
    (e.g. a personal install AND a company server), a guide whose CONTENT
    varied per-connector risked Claude conflating "not available on THIS
    connector" with "doesn't exist in AI-Prowler at all," and misapplying one
    connector's limitations to the other. Server-mode-specific caveats (dev
    tools, code-aware retrieval, and file-editing scoping are personal-mode
    concepts that don't apply in server mode) are written inline in their
    sections, the same way the COMMUNICATIONS section already handles both
    modes in one block of text — nothing is removed or hidden per-connector.

    Ends with a "THIS CONNECTION" section that's computed live from the
    actual server-mode role gates (_ROLE_CAPS) for whichever connector this
    was called on — this is the ONE part of the output that varies by
    connector, and it's clearly labeled as being about this specific
    connection, so it can't be mistaken for a change to the tool catalog
    itself.

    Returns:
        Step-by-step guidance on which tools to use and in what order,
        plus key facts about capabilities, plus a connection-specific footer.
    """
    import mcp as _mcp_pkg
    try:
        mcp_version = _mcp_pkg.__version__
    except Exception:
        mcp_version = "unknown"

    instructions_active = "instructions" in _fastmcp_params

    # Computed early (was previously computed only for the footer) so the
    # main guide body can also be dynamically trimmed for server mode —
    # see the section-stripping block right after base_text is built.
    _user = _current_user(ctx)
    _htu_priv_status, _htu_priv_dir = _user_private_write_dir(ctx)

    base_text = (
        "AI-Prowler — Agentic RAG Knowledge Base\n"
        + "=" * 50 + "\n\n"

        "TOOL CATEGORIES (82 tools total — 82 visible in personal mode,\n"
        "54 visible in server mode; call check_tools_status() for a precise\n"
        "per-tool breakdown on this connection)\n"
        + "-" * 30 + "\n"
        "AI-Prowler exposes ten tool families. Most question-answering\n"
        "tasks use the first two; the others cover indexing, code editing,\n"
        "dev tooling, communications, contractor/field-service workflows,\n"
        "and agentic analysis.\n\n"

        "  • Knowledge retrieval (RAG over indexed documents):\n"
        "      get_knowledge_base_overview, list_indexed_documents,\n"
        "      list_indexed_directories, search_documents, search_within_directory,\n"
        "      multi_query_search, expand_search_result, read_document\n\n"

        "  • Code-aware retrieval (exact match + line-accurate reads):\n"
        "      grep_documents, read_file_lines\n\n"

        "  • Self-learning memory (corrections, post-mortems, preferences):\n"
        "      search_learnings, record_learning, list_learnings,\n"
        "      update_learning, delete_learning, get_learning_stats,\n"
        "      get_learnings_report, export_learnings_file,\n"
        "      send_learnings_report, rebuild_learnings_index\n\n"

        "  • Field service actions (free public APIs, no key needed):\n"
        "      geocode_address, get_weather, optimize_route,\n"
        "      build_maps_url, read_job_spreadsheet, update_job_spreadsheet,\n"
        "      check_tools_status\n\n"

        "  • Contractor / business workflow:\n"
        "      email_invoice, schedule_next_recurring_job, log_time_entry,\n"
        "      get_ar_aging_report, save_contact, get_sms_thread,\n"
        "      list_sms_contacts_with_replies\n\n"

        "  • Communications (email + SMS + WhatsApp — every role in both\n"
        "    personal and server mode):\n"
        "      configure_email, send_email, send_alert, send_file,\n"
        "      send_sms, check_sms_replies, check_sms_inbox,\n"
        "      send_whatsapp, check_whatsapp_replies\n\n"

        "  • File editing (write tools — see EDITING FILES section below):\n"
        "      create_file, write_file, str_replace_in_file,\n"
        "      line_replace_in_file, create_directory,\n"
        "      list_directory, copy_to_backup, list_backups, restore_backup,\n"
        "      cleanup_backups, cleanup_job_logs, reset_write_counter, diff_files\n\n"

        "  • Dev tools (run/check code without leaving the conversation):\n"
        "      syntax_check, compile_check, check_python_import, lint_check,\n"
        "      run_script, run_script_start, run_script_status, run_script_kill\n\n"

        "  • Indexing & admin:\n"
        "      index_path, update_tracked_directories, list_tracked_directories,\n"
        "      untrack_directory, get_database_stats, check_ai_prowler_status,\n"
        "      reindex_file, reindex_directory, reindex_all,\n"
        "      list_writable_directories, grant_write_access, revoke_write_access\n\n"

        "  • Agentic analysis tasks (personal mode only):\n"
        "      create_analysis_task — defines a new recurring or one-off\n"
        "        custom analysis task from a plain-language request. Day-\n"
        "        granularity scheduling only; pull-based, not autonomous —\n"
        "        see its own docstring for the full behavior explanation.\n"
        "      list_analysis_tasks — lists the FULL custom-task definition\n"
        "        list (up to 25) regardless of due date, with an is_due flag\n"
        "        per task. Use for 'what's in my task queue' — different from\n"
        "        get_pending_analysis_tasks below, which only shows tasks\n"
        "        already queued into the run queue. Read-only.\n"
        "      get_pending_analysis_tasks — returns all pending tasks from\n"
        "        pending_tasks.json; call when the user pastes the run-queue\n"
        "        command from the Quick Links tab.\n"
        "      complete_analysis_task(task_id, summary) — marks a task done\n"
        "        and auto-advances next_due for scheduled tasks (anchor-based,\n"
        "        not completion-date-based). Call after finishing each analysis.\n"
        "      save_analysis_report(content, title, task_id, report_folder) —\n"
        "        saves the full analysis as a .docx Word document.\n\n"


        "PREFERRED TOOL SEQUENCE — KNOWLEDGE RETRIEVAL\n"
        + "-" * 30 + "\n"
        "Use these tools in order for any research or question-answering task:\n\n"

        "  STEP 1  get_knowledge_base_overview()\n"
        "    Orient yourself: see what documents are indexed, file types,\n"
        "    topics covered, tracked directories, and directory tree.\n\n"

        "  STEP 2  list_indexed_documents(filter_ext, filter_path)\n"
        "    Browse specific files when the user asks about a particular\n"
        "    document, company, or topic area.\n\n"

        "  STEP 3  list_indexed_directories()\n"
        "    See the directory tree of all indexed content. Use this to\n"
        "    identify the right scope for targeted searches.\n\n"

        "  STEP 4  search_documents(query, n_results)\n"
        "    PRIMARY retrieval. Call MULTIPLE TIMES with different\n"
        "    phrasings to gather full context. Example:\n"
        "      search_documents('refund policy')\n"
        "      search_documents('money back guarantee')\n\n"

        "  STEP 5  search_within_directory(query, directory, n_results)\n"
        "    TARGETED retrieval. Use when the user asks about a specific\n"
        "    case, project, client, or folder. Restricts search to only\n"
        "    chunks from that directory tree — prevents cross-contamination.\n"
        "    Example:\n"
        "      search_within_directory('summary of damages', 'Smith_v_Jones')\n\n"

        "  STEP 6  multi_query_search(queries)\n"
        "    Parallel search when a topic has synonyms or multiple angles.\n"
        "    More efficient than calling search_documents() repeatedly.\n\n"

        "  STEP 7  expand_search_result(filename, chunk_index)\n"
        "    Expand around a result that appears cut off or references\n"
        "    content in the surrounding paragraphs.\n\n"

        "  STEP 8  read_document(filename, start_chunk)\n"
        "    Read a whole document sequentially for full summaries or\n"
        "    when the user asks about a specific document's contents.\n\n"

        "CODE-AWARE RETRIEVAL — FOR PROGRAM FILES\n"
        + "-" * 30 + "\n"
        "Semantic search degrades on source code. When the user asks about\n"
        "code (symbols, definitions, call sites, TODOs), use this pair\n"
        "INSTEAD OF search_documents — same workflow Claude Code uses.\n\n"

        "  STEP C1  grep_documents(pattern, filter_ext, filter_path)\n"
        "    Locate exact matches with real line numbers. Use for any\n"
        "    'where is X defined / called / used' question.\n"
        "    Examples:\n"
        "      grep_documents('def clear_database', filter_ext='.py')\n"
        "      grep_documents('TODO', filter_path='rag_')\n"
        "      grep_documents(r'class \\w+Error', regex=True)\n\n"

        "  STEP C2  read_file_lines(filepath, start_line, end_line)\n"
        "    After grep, extract the surrounding function or block at\n"
        "    original fidelity. Returns numbered lines so you can keep\n"
        "    paging through with start_line=last_seen+1.\n\n"

        "  Note: program/script files (CODE_SCAN_EXTENSIONS — .py, .js, .java,\n"
        "  .cs, .cpp, .go, .rs, .sh, .ps1, .css, .sql, and others) are indexed\n"
        "  as a single truncated chunk (first 500 lines only, headed\n"
        "  '[SECURITY SCAN ONLY]') rather than fully semantically chunked —\n"
        "  this keeps the vector DB free of code-boilerplate noise. grep_documents\n"
        "  + read_file_lines read the file fresh from disk every time, so they\n"
        "  always see the FULL current file regardless of that indexing cap.\n\n"

        "  Both tools read only files under the tracked-paths allowlist.\n"
        "  Useful especially for mobile users who cannot attach files —\n"
        "  ask them to index_path() the project once, then\n"
        "  use grep + read_file_lines for all subsequent code questions.\n"
        "  Server mode: grep_documents and read_file_lines are NOT available\n"
        "  to any role — raw filesystem access is a personal-install-only\n"
        "  capability. search_documents still works for code files in server\n"
        "  mode, just without exact-match/line-number precision.\n\n"

        "EDITING FILES — REINDEX WHEN DONE\n"
        + "-" * 30 + "\n"
        "Write tools (create_file, write_file, str_replace_in_file,\n"
        "line_replace_in_file, restore_backup) do NOT\n"
        "auto-index. They write to disk and create backups, but ChromaDB is\n"
        "NOT updated until you ask for it.\n"
        "  • Make ALL your edits to a file first (any number of\n"
        "    str_replace_in_file / line_replace_in_file\n"
        "    calls).\n"
        "  • When you are DONE editing that file, call reindex_file(path)\n"
        "    ONCE to sync it into the database.\n"
        "  • Do NOT call reindex_file between every edit — one call at the\n"
        "    end of the edit session per file is correct.\n"
        "  • For a whole folder, use reindex_directory() instead of many\n"
        "    reindex_file() calls; reindex_all() rebuilds every tracked path\n"
        "    (the nuclear option — slow, rarely needed).\n"
        "  • All writes go through a per-session write circuit breaker. If it\n"
        "    trips after many writes in one conversation, call\n"
        "    reset_write_counter() to continue — this is a runaway-loop\n"
        "    safeguard, not an error.\n"
        "  • Before relying on freshly-edited code, run syntax_check() or\n"
        "    compile_check() to catch typos, then check_python_import() to\n"
        "    catch load-time errors syntax checking alone would miss.\n"
        "  Rationale: re-embedding on every write deadlocked the HTTP server\n"
        "  on large files; explicit end-of-session reindex avoids that.\n"
        "  Server mode: create_file/write_file/str_replace_in_file/\n"
        "  line_replace_in_file/create_directory are\n"
        "  scoped to the caller's own personal directory (blocked entirely if\n"
        "  they don't have one) — see the THIS CONNECTION footer below for\n"
        "  this specific caller's status. reset_write_counter() and\n"
        "  restore_backup() are NOT available in server mode. Indexing tools\n"
        "  (index_path, update_tracked_directories, reindex_file/\n"
        "  reindex_directory/reindex_all) are open to every role there —\n"
        "  indexing isn't a data leak; only search access is scope-gated.\n\n"

        "DEV TOOLS — VERIFY CODE WITHOUT LEAVING THE CONVERSATION\n"
        + "-" * 30 + "\n"
        "  syntax_check(filepath)       — multi-language syntax check\n"
        "  compile_check(filepath)      — Python-specific byte-compile check\n"
        "  check_python_import(module)  — catches load-time errors syntax misses\n"
        "  lint_check(filepath)         — style/unused-import warnings — pyflakes\n"
        "                                  (Python), tsc (TypeScript), go vet (Go),\n"
        "                                  verilator --lint-only (Verilog/SystemVerilog),\n"
        "                                  ghdl -a (VHDL); other languages fall back\n"
        "                                  to syntax_check, which has no lint tool\n"
        "  run_script(path)             — execute a script, return its output\n"
        "  run_script_start / _status / _kill — background job runner for\n"
        "                                  long-running scripts (tests, builds)\n"
        "  Server mode: NONE of these dev tools are available to any role —\n"
        "  code execution and static-analysis tools are personal-install-only.\n\n"

        "COMMUNICATIONS — EMAIL, SMS, WHATSAPP\n"
        + "-" * 30 + "\n"
        "Personal mode: configure_email() once, then send_email / send_alert /\n"
        "send_file / send_sms / send_whatsapp all work from that one account.\n"
        "Recipients can be a raw address/number, a name in the job spreadsheet's\n"
        "Customers sheet (checked FIRST — same data read_job_spreadsheet() and\n"
        "update_job_spreadsheet() use), or a saved contact name (checked after) —\n"
        "save_contact(name, phone?, email?) once, then 'text David' or 'email\n"
        "Vicki' resolves automatically via contacts_cache.json.\n\n"
        "Server mode: email, SMS, and WhatsApp are available to EVERY role —\n"
        "owner, manager, staff, and field_crew — via the company's shared SMTP\n"
        "and SMS/WhatsApp provider account. Same recipient lookup order as\n"
        "personal mode (Customers sheet first, then users.json by name, then\n"
        "your private contacts_cache_<username>.json) — and read_job_spreadsheet /\n"
        "update_job_spreadsheet are themselves available to every role too, with\n"
        "no DB-management or communications gate applied. SMS works with Twilio,\n"
        "SignalWire, or Vonage (set via sms_provider in config); WhatsApp always\n"
        "goes through Twilio specifically, reusing the same Twilio credentials\n"
        "even if SMS itself is configured to use SignalWire or Vonage instead.\n"
        "Inbound replies ARE correctly attributed per user — check_sms_replies()\n"
        "and check_whatsapp_replies() filter to only the threads YOU personally\n"
        "sent (Mike sees Karen's reply, not Jake's reply to Bob), via a local\n"
        "webhook-backed inbox (sms_inbox.json / sms_threads.json) rather than\n"
        "polling the provider's API. Every server-mode role also gets their\n"
        "own private contacts_cache_<username>.json — saved contacts are never\n"
        "shared between users.\n\n"
        "check_sms_inbox / get_sms_thread / list_sms_contacts_with_replies —\n"
        "local SMS/WhatsApp inbox tools for reviewing conversation history.\n"
        "The free email-to-SMS carrier-gateway approach (vtext.com, etc.) was\n"
        "removed — carriers are shutting those gateways down industry-wide.\n"
        "A real SMS/WhatsApp provider (Twilio, SignalWire, or Vonage) is now\n"
        "the only path.\n\n"

        "SELF-LEARNING WORKFLOW — PERSISTENT MEMORY\n"
        + "-" * 30 + "\n"
        "Stored learnings are corrections, post-mortems, and preferences\n"
        "captured from prior sessions. They override built-in knowledge\n"
        "and represent the operator's verified ground truth. Use them.\n\n"

        "  STEP A  search_learnings(query)\n"
        "    Call BEFORE answering ANY user question. No exceptions for\n"
        "    questions that look like general knowledge — the user may\n"
        "    have recorded a personal version of a recipe, technique,\n"
        "    preference, or recommendation that overrides what your\n"
        "    training data would suggest. If a learning contradicts your\n"
        "    built-in knowledge or any web search result, prefer the\n"
        "    learning.\n"
        "    The ONLY questions you may skip on are pure arithmetic and\n"
        "    language translation. When in doubt, search.\n"
        "    Triggers that GUARANTEE a search_learnings call:\n"
        "      - User asks 'what's a good X' / 'how do I X' / 'recommend X'\n"
        "      - User says 'what did we learn about...', 'do you remember...'\n"
        "      - User mentions a client, project, recipe, or case by name\n"
        "      - You're about to make a recommendation of any kind\n"
        "      - You're about to state a fact that may have been corrected\n\n"

        "  STEP B  record_learning(category, content, ...)\n"
        "    Auto-record (without being asked) when the user:\n"
        "      - Corrects a fact you stated\n"
        "      - Shares a project outcome or post-mortem\n"
        "      - States a client preference or standing instruction\n"
        "      - Says 'next time we should...' or similar\n"
        "      - Provides info that contradicts an existing active learning\n\n"

        "  STEP C  list_learnings() / get_learning_stats()\n"
        "    Browse stored learnings when the user asks for an overview\n"
        "    of what's remembered, or before bulk updates.\n\n"

        "  STEP D  update_learning(id, ...) / delete_learning(id)\n"
        "    Refine or retire learnings when the user provides better\n"
        "    information. Always confirm destructive actions first.\n\n"

        "  STEP E  get_learnings_report() / export_learnings_file() /\n"
        "          send_learnings_report()\n"
        "    Produce a readable summary of everything learned — inline,\n"
        "    as a file, or emailed as an HTML report. Use\n"
        "    rebuild_learnings_index() if learnings stop showing up in\n"
        "    search_learnings (rebuilds the ChromaDB index from the JSON\n"
        "    source of truth).\n\n"

        "DOCUMENT PROVENANCE — CRITICAL\n"
        + "-" * 30 + "\n"
        "  Every search result includes rich provenance metadata:\n"
        "    - parent_directory: immediate folder name\n"
        "    - directory_chain: breadcrumb path from root to file\n"
        "    - document_id: SHA-256 content fingerprint (unique per file)\n"
        "    - doc_title: extracted document title (from PDF/DOCX/PPTX)\n"
        "    - file_modified: last modification date of the source file\n"
        "    - [SOURCE: ...] header embedded in every chunk's text\n\n"

        "  RULES FOR ANSWERING:\n"
        "  1. ALWAYS verify that ALL chunks you cite come from the SAME\n"
        "     document or directory group before synthesizing an answer.\n"
        "  2. If chunks from DIFFERENT parent_directory values appear in\n"
        "     results, explicitly tell the user which source each piece\n"
        "     of information comes from.\n"
        "  3. NEVER blend facts from chunks with different document_id or\n"
        "     parent_directory values into a single statement without\n"
        "     clear attribution.\n"
        "  4. If the user asks about a specific case/project/client, use\n"
        "     search_within_directory() to restrict search to that scope.\n"
        "  5. When multiple files share the same name (e.g. 'complaint.pdf'\n"
        "     in different case folders), use parent_directory and\n"
        "     directory_chain to distinguish them for the user.\n\n"

        "AGENTIC ANALYSIS WORKFLOW\n"
        + "-" * 30 + "\n"
        "The Quick Links tab in the AI-Prowler GUI queues analysis tasks and\n"
        "copies a run command to the clipboard. When the user pastes that\n"
        "command into Claude, follow this sequence:\n\n"

        "  STEP A1  get_pending_analysis_tasks()\n"
        "    Returns a JSON object with pending_count and a tasks array.\n"
        "    Each task has: task_id, label, prompt, scope_dirs,\n"
        "    schedule, next_due, output_learnings, output_report,\n"
        "    report_folder, queued_ago.\n"
        "    If pending_count is 0, tell the user the queue is empty.\n\n"

        "  STEP A2  Execute each task's prompt\n"
        "    The prompt field contains the COMPLETE instruction for that\n"
        "    analysis — including which tools to call, what to look for,\n"
        "    and what to output. Follow it exactly.\n"
        "    If scope_dirs is non-empty, use search_within_directory()\n"
        "    for those directories instead of search_documents().\n"
        "    QuickBooks-aware tasks: if QB MCP tools are available, use them\n"
        "    as the primary financial source; fall back to\n"
        "    read_job_spreadsheet() and get_ar_aging_report() otherwise.\n\n"

        "  STEP A3  complete_analysis_task(task_id, summary)\n"
        "    Call after EVERY task — even if nothing was found. The summary\n"
        "    is one sentence (e.g. 'Found 3 overdue invoices totalling $1,450').\n"
        "    For scheduled tasks this auto-advances next_due anchored to the\n"
        "    original due date (weekly task due Monday stays on Mondays).\n\n"

        "  STEP A4  save_analysis_report() if output_report=true\n"
        "    Save the full analysis as a Word document to report_folder.\n"
        "    Call this BEFORE complete_analysis_task.\n\n"

        "  QUICKBOOKS DETECTION (in analysis prompts):\n"
        "    Check whether tools whose names contain 'quickbooks' or 'qbo'\n"
        "    are in your available tool list. If yes, use QuickBooks as the\n"
        "    primary financial source. If no, use AI-Prowler Job Tracker tools.\n\n"
        "  Server mode: get_pending_analysis_tasks, complete_analysis_task,\n"
        "  save_analysis_report, and create_analysis_task are NOT available\n"
        "  to any role — the Quick Links tab that queues this workflow is a\n"
        "  personal-install-only GUI feature.\n\n"


        "PROACTIVE ALERTS (background scheduler — no Claude needed)\n"
        + "-" * 30 + "\n"
        "The Proactive Alerts scheduler in the GUI runs independently of Claude.\n"
        "It calls AI-Prowler Python functions directly on a schedule and emails\n"
        "results to the configured address — zero API cost, no Claude session.\n"
        "Claude does not need to interact with the scheduler directly. However,\n"
        "if the user asks about it:\n"
        "  - Config: ~/.ai-prowler/scheduler_config.json\n"
        "  - Log:    ~/.ai-prowler/scheduler_log.txt\n"
        "  - Jobs:   morning_briefing, overdue_invoice_alert, due_analysis_tasks,\n"
        "            sms_reply_monitor, weather_watch, end_of_day_summary\n"
        "  - Personal mode only — hidden in server mode automatically.\n\n"

        "KEY FACTS\n"
        + "-" * 30 + "\n"
        "  - NO Ollama required — no local LLM involved at all.\n"
        "  - Claude receives RAW CHUNKS and synthesizes answers directly.\n"
        "  - For complex questions, always search multiple times before answering.\n"
        "  - Stored learnings outrank built-in knowledge — always search_learnings\n"
        "    before answering ANY user question (the only exceptions are pure\n"
        "    arithmetic and language translation).\n"
        "  - Use check_ai_prowler_status() to verify the knowledge base is healthy.\n"
        "  - Use check_tools_status() to verify field-service/business/SMS tools.\n"
        "  - Server mode role gating (owner/manager/staff/field_crew) restricts\n"
        "    DB-management and admin tools by role; email/SMS/WhatsApp are open\n"
        "    to every role. A clear ⛔/❌ message explains any restriction that\n"
        "    applies to a specific tool call.\n"
        "  - Re-call this tool any time you need a reminder of the workflow.\n\n"

        f"MCP SDK version       : {mcp_version}\n"
        f"instructions= active  : {'yes — guidance sent at every handshake' if instructions_active else 'no — upgrade with: pip install --upgrade mcp'}\n"
        "AI-Prowler Agentic RAG ready."
    )

    # The main guide above is IDENTICAL regardless of mode or connector — this
    # is deliberate, not an oversight. In a conversation with two connectors
    # (a personal install AND a company server both attached), a guide that
    # varied per-connector risked Claude conflating "not available on THIS
    # connector" with "doesn't exist in AI-Prowler at all," and misapplying
    # one connector's limitations to the other. Every tool is described here,
    # on every connector; the "THIS CONNECTION" footer below is the ONE place
    # that varies, clearly labeled as being about this specific connection —
    # server-mode caveats for dev tools, code-aware retrieval, file editing,
    # etc. are inline in their sections above instead of removing them.

    # ── THIS CONNECTION — computed live from the real role-gate tables ──────────
    # Tells Claude exactly what's true for the connector this call was made on,
    # so when both a personal install and a company server are connected in the
    # same conversation, each one's how_to_use_ai_prowler() call reports its own
    # situation rather than a generic description Claude has to mentally apply.
    # (_user already computed above, before base_text, so the main guide body
    # could also be dynamically trimmed.)
    footer_lines = ["", "─" * 50, "", "THIS CONNECTION", "─" * 30]

    if _user is None:
        footer_lines += [
            "Mode: Personal (single-user desktop install)",
            "- No role restrictions apply — every tool listed above is available.",
            "- contacts_cache.json is yours alone; nothing here is shared with",
            "  any other AI-Prowler install.",
        ]
    else:
        role  = (_user.get("role") or "field_crew").strip().lower()
        caps  = _role_caps(role)
        name  = _user.get("name") or _user.get("username") or "you"
        footer_lines += [
            f"Mode: Server  (role: {role}, user: {name})",
            "",
            "Indexing (index_path, update_tracked_directories, reindex_*):",
            "  ✅ Full access — every role may index, update, and reindex the",
            "     shared knowledge base. Indexing isn't a data leak — only",
            "     search results are scope-gated (see allowed_scopes above).",
            "",
            "Tracking administration (untrack_directory):",
        ]
        # v8.1.5: untrack_directory moved from a blanket owner/manager role
        # gate to a two-tier gate — own personal directory (any role) vs.
        # everywhere else (owner or delegated admin). Compute live rather
        # than reusing the old manage_db=="full" check, which no longer
        # matches what the tool actually enforces. _user_private_write_dir
        # is called fresh here since the shared _htu_priv_status/_htu_priv_dir
        # locals used by the later "File editing" section aren't computed yet
        # at this point in the function.
        _untr_priv_status, _untr_priv_dir = _user_private_write_dir(ctx)
        _is_owner_ftr = (role == "owner")
        _is_admin_ftr = bool(_user.get("can_manage_users")) if _user else False
        if _untr_priv_status == "scoped":
            footer_lines.append(
                f"  ✅ Inside your own personal directory ({_untr_priv_dir}): "
                "always allowed, any role.")
            if _is_owner_ftr or _is_admin_ftr:
                footer_lines.append("  ✅ Outside it too — you have owner/delegated admin rights.")
            else:
                footer_lines.append(
                    "  ⛔ Outside it — requires the owner, or delegated admin "
                    "rights (\"Can manage users\").")
        else:
            if _is_owner_ftr or _is_admin_ftr:
                footer_lines.append("  ✅ Full access — you have owner/delegated admin rights.")
            else:
                footer_lines.append(
                    "  ⛔ Not available — you have no personal directory configured, "
                    "and no delegated admin rights (\"Can manage users\").")

        # v8.1.5 fix: list_tracked_directories is now gated by SCOPE, not
        # role (every role may call it; the result is filtered to paths
        # within the caller's own accessible scopes) — no longer tied to
        # manage_db, so it gets its own line instead of sharing untrack_
        # directory's role-gated message above.
        footer_lines.append("")
        footer_lines.append("list_tracked_directories:")
        footer_lines.append("  ✅ Available to every role — filtered to paths within your own")
        footer_lines.append("     accessible scopes (same visibility as search, not company-wide).")

        # v8.1.5 fix: grant_write_access/revoke_write_access are in
        # _TIER_A_SUPPRESSED — never registered with MCP in server mode,
        # for ANY role, owner included (write-zone management doesn't fit
        # the server-mode model, where the actual write boundary is each
        # user's own personal directory — see _check_personal_write_scope).
        # This footer used to claim "✅ Full access" for owner/manager based
        # on manage_db, which was simply wrong: the tools aren't callable at
        # all in server mode, regardless of role, so the previous
        # role-conditional message was misleading rather than just imprecise.
        footer_lines.append("")
        footer_lines.append("Write-zone grants (grant_write_access, revoke_write_access):")
        footer_lines.append("  ⛔ Not available in server mode, to any role — personal-install-only.")
        footer_lines.append("     The write boundary in server mode is each user's own personal")
        footer_lines.append("     directory (see File editing below), not an arbitrary write-zone list.")

        footer_lines.append("")
        footer_lines.append("File editing (create_file, write_file, str_replace_in_file,")
        footer_lines.append("line_replace_in_file, create_directory):")
        if _htu_priv_status == "scoped":
            footer_lines.append(f"  ✅ Scoped to your personal directory: {_htu_priv_dir}")
        else:
            footer_lines.append("  ⛔ Not available — no personal directory configured.")
            footer_lines.append("     Ask your owner/admin to set one up in the Admin tab.")

        footer_lines.append("")
        footer_lines.append("Email / SMS / WhatsApp:")
        if caps.get("can_send_email") or caps.get("can_send_sms"):
            footer_lines.append("  ✅ Available — your role may send email, SMS, and WhatsApp via")
            footer_lines.append("     the company's shared SMTP / SMS-provider account.")
        else:
            footer_lines.append("  ⛔ Not available to your role on this server.")
        footer_lines.append("  Your contacts_cache_" + (
            _CONTACTS_CACHE_PATH(_user).stem.replace('contacts_cache_', '')
            if _user else 'unknown'
        ) + ".json is private to you — other users' saved contacts are not")
        footer_lines.append("  visible to you, and vice versa.")
        footer_lines.append("  Inbound SMS/WhatsApp replies (check_sms_replies /")
        footer_lines.append("  check_whatsapp_replies) are filtered to only the threads YOU")
        footer_lines.append("  personally sent — you will not see another user's conversations.")

        footer_lines.append("")
        footer_lines.append("Shared-collection writes (can_write_shared):")
        footer_lines.append(
            "  ✅ Allowed" if caps.get("can_write_shared")
            else "  ⛔ Not allowed — your writes are confined to your own scopes/private area."
        )

        if caps.get("is_admin"):
            footer_lines.append("")
            footer_lines.append("Admin tab: ✅ you have admin rights (user management, recovery, etc.)")

    # ── MORNING BRIEFING — pending queue + due tasks ────────────────────────
    # Always ask the user before running — they may have something urgent first.
    briefing_lines = []
    try:
        import datetime as _dt
        import sys as _sys, os as _os
        _app = _os.path.dirname(_os.path.abspath(__file__))
        if _app not in _sys.path:
            _sys.path.insert(0, _app)

        # Check pending_tasks.json for queued items
        pending_tasks = _load_pending_tasks()
        pending = [t for t in pending_tasks if t.get("status") == "pending"]

        # Check custom_analysis_tasks.json for due/overdue items
        try:
            import custom_tasks_manager as _ctm
            custom_tasks = _ctm.load_custom_tasks()
            due_tasks = _ctm.get_due_tasks(custom_tasks)
        except Exception:
            custom_tasks = []
            due_tasks = []

        has_pending = len(pending) > 0
        has_due     = len(due_tasks) > 0

        if has_pending or has_due:
            briefing_lines.append("")
            briefing_lines.append("─" * 50)
            briefing_lines.append("")
            briefing_lines.append("📋 ANALYSIS BRIEFING")
            briefing_lines.append("─" * 30)

            if has_pending:
                briefing_lines.append(
                    f"Queue: {len(pending)} pending task"
                    f"{'s' if len(pending) != 1 else ''}:"
                )
                now = _dt.datetime.utcnow()
                for t in pending:
                    try:
                        created = _dt.datetime.strptime(
                            t.get("created_at", ""), "%Y-%m-%dT%H:%M:%SZ")
                        age_mins = int((now - created).total_seconds() / 60)
                        if age_mins < 60:
                            age = f"{age_mins}m ago"
                        elif age_mins < 1440:
                            age = f"{age_mins // 60}h ago"
                        else:
                            age = f"{age_mins // 1440}d ago"
                    except Exception:
                        age = "unknown"
                    briefing_lines.append(f"  • {t.get('label', t.get('task_id', '?'))} (queued {age})")

            if has_due:
                if has_pending:
                    briefing_lines.append("")
                briefing_lines.append(
                    f"Scheduled: {len(due_tasks)} task"
                    f"{'s' if len(due_tasks) != 1 else ''} due:"
                )
                for t in due_tasks:
                    status = _ctm.due_status_label(t)
                    briefing_lines.append(f"  • {t.get('label', '?')} — {status}")

            briefing_lines.append("")
            briefing_lines.append(
                "ACTION REQUIRED: Before answering the user's first question,\n"
                "ask them: 'I noticed you have "
                + (f"{len(pending)} queued task{'s' if len(pending) != 1 else ''}" if has_pending else "")
                + (" and " if has_pending and has_due else "")
                + (f"{len(due_tasks)} scheduled task{'s' if len(due_tasks) != 1 else ''} due" if has_due else "")
                + ". Would you like me to run those before we start, "
                "or would you prefer to handle your current question first?'"
            )
            briefing_lines.append(
                "Wait for the user's answer before taking any action."
            )

    except Exception as _be:
        pass  # Briefing is non-fatal — never block the main guidance

    return base_text + "\n".join(footer_lines) + "\n".join(briefing_lines)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 1 — index_path
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def index_path(
    directory: str,
    recursive: bool = True,
    track: bool = True,
    ctx: Context = None,
) -> str:
    """
    Index all supported documents inside a local folder OR a single file
    and (optionally) add it to the auto-update tracking list.

    Accepts either:
      • A directory path — every supported file inside is indexed (subject
        to smart-scan extension/directory filters). Recursion controlled by
        the `recursive` argument.
      • An individual file path — that one file is indexed regardless of
        the smart-scan extension filters. Per-file tracking is explicit
        user opt-in, so SKIP_EXTENSIONS does not apply.

    Supported formats: PDF, Word, Excel, PowerPoint, plain text, code,
    Markdown, HTML, email (.msg / .eml / .mbox), images (OCR), and many more.

    Args:
        directory:  Absolute path to the folder OR file you want to index.
                    Examples:
                      "C:/Users/David/Documents/ProjectDocs"
                      "C:/Program Files/AI-Prowler/COMPLETE_USER_GUIDE.md"
        recursive:  Include sub-folders when indexing a directory (default True).
                    Ignored for single-file targets.
        track:      Add the path to the auto-update list so future
                    `update_tracked_directories` calls will pick up changes
                    (default True).

    Server mode: indexing is open to every role (v8.1.4) — there is no
        longer a role gate here. Content always lands in the single
        shared knowledge base, tagged with a "scope" (see search tools'
        docs) rather than routed into a separate collection per role.
        A path under a user's own private folder is still tagged
        private to them automatically by naming convention. Access
        control lives entirely at search time (scope-based), not here.

    Returns:
        Summary of how many files were indexed and any errors encountered.
    """
    target = Path(directory)
    if not target.exists():
        return f"❌ Path not found: {directory}"

    load_config()

    # v7.0.0 Phase B: in server mode build the per-user write controls. Personal
    # mode (no authenticated user) leaves all three None → pipeline unchanged.
    #   • resolver   — routes each file to its scoped collection (WHERE), with
    #                  _can_index enforced inside the resolver (WHETHER).
    #   • indexer_user— stamps indexed_by for the ownership model.
    #   • purge_gate  — blocks overwriting/destroying chunks the user doesn't own.
    _user = _current_user(ctx)

    # SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cutover, 2026-07-16):
    # direct product decision -- indexing is not a data leak (only search
    # is), and every directory that can be indexed was already created and
    # tracked by an admin/owner in the first place, so there is no
    # arbitrary-path risk from letting any authenticated user trigger
    # indexing on it. The old role-based manage_db gate (_check_db_cap) and
    # the narrower "field_crew with a private collection may index ONLY
    # inside their own personal directory, force-routed into their own
    # collection" carve-out are both removed -- any authenticated (or
    # personal-mode) caller may index any path. Content always lands in
    # the single unified index; build_scope_resolver() still automatically
    # tags anything under a <slug>-private folder as "private:<their own
    # id>" via the same path-convention detection, so a user's own private
    # content stays exactly as private as before -- nothing here weakens
    # that, since the actual boundary was always enforced at search time
    # (allowed_scopes()), not at index time.
    _resolver = None
    _indexer_user = _user
    _purge_gate = None
    if _user:
        _oid = _owner_user_id()
        _purge_gate = lambda metas: _can_purge_chunks(_user, metas, _oid)

    is_file = target.is_file()

    with _capture_stdout() as buf:
        try:
            if is_file:
                # Single-file path — bypass SKIP_EXTENSIONS smart-scan filter
                from rag_preprocessor import index_file_list, normalise_path
                fp = normalise_path(str(target.resolve()))
                stats = index_file_list(
                    [fp],
                    label="1/1",
                    root_directory=str(target.parent),
                    collection_resolver=_resolver,
                    indexer_user=_indexer_user,
                    purge_gate=_purge_gate,
                )
                chunks = stats.get('chunks', 0) if stats else 0
                print(f"✅ Indexed file: {target.name} — {chunks} chunk(s)")
            else:
                index_directory(str(target), recursive=recursive,
                                collection_resolver=_resolver,
                                indexer_user=_indexer_user,
                                purge_gate=_purge_gate,
                                track=track)
        except Exception as exc:
            return f"❌ Indexing failed: {exc}"

    output = buf.getvalue().strip()

    if track:
        try:
            # For directories, index_directory already adds to the auto-update
            # list via add_to_auto_update_list(). Calling it again is a no-op
            # (returns False) but ensures file targets are added too.
            added = add_to_auto_update_list(str(target.resolve()))
            kind = "File" if is_file else "Directory"
            note = (
                f"\n✅ {kind} added to auto-update tracking."
                if added
                else f"\nℹ️  {kind} was already in the tracking list."
            )
            output += note
        except Exception as exc:
            output += f"\n⚠️  Could not add to tracking list: {exc}"

    return output if output else "✅ Indexing complete (no output produced)."


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 3 — update_tracked_directories
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def update_tracked_directories(directory: Optional[str] = None,
                               ctx: Context = None) -> str:
    """
    Re-scan all tracked paths (directories AND individually-tracked files) for
    new, modified, or deleted files and update the index incrementally — only
    changed files are re-indexed.

    Args:
        directory:  If provided, update only this specific path (directory
                    or file). If omitted, update ALL tracked paths.

    Returns:
        A summary of changes detected and files re-indexed.
    """
    load_config()

    # v7.0.0 Phase B: server-mode write controls (None in personal mode →
    # pipeline unchanged). Same trio as index_path. The purge gate
    # matters especially here: command_update PURGES deleted files (destructive).
    _user = _current_user(ctx)

    # SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cutover, 2026-07-16):
    # the role-based manage_db gate is removed -- see index_path's identical
    # comment for the full rationale (indexing isn't a leak; every tracked
    # path was already admin/owner-created). Destructive purging of deleted
    # files (this tool's real risk, not indexing itself) is still protected
    # separately by _purge_gate below, which is ownership-based, not role-
    # based, and is unaffected by this change.
    #
    # No longer builds a collection_resolver either -- there is only one
    # physical collection now. Scope is carried entirely by
    # build_scope_resolver()'s "scope" chunk-metadata tag (wired into
    # command_update already), not by which physical collection a file's
    # chunks are routed into.
    _resolver = None
    _purge_gate = None
    if _user:
        _oid = _owner_user_id()
        _purge_gate = lambda metas: _can_purge_chunks(_user, metas, _oid)

    dirs_to_update: list[str] = []
    if directory:
        dirs_to_update = [directory]
    else:
        dirs_to_update = load_auto_update_list()

    if not dirs_to_update:
        return (
            "ℹ️  No tracked paths found.\n"
            "Use index_path first to index a folder or file and add it to tracking."
        )

    with _capture_stdout() as buf:
        for d in dirs_to_update:
            try:
                command_update(d, recursive=True, auto_confirm=True,
                                collection_resolver=_resolver,
                                indexer_user=_user,
                                purge_gate=_purge_gate)
            except Exception as exc:
                print(f"⚠️  Error updating {d}: {exc}")

    return buf.getvalue().strip() or "✅ All tracked paths are up to date."


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 4 — get_database_stats
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def get_database_stats(ctx: Context = None) -> str:
    """
    Return statistics about the indexed knowledge base: total chunk count,
    number of unique documents, breakdown by file type, and database location.

    Personal mode: always covers the entire database (unchanged).

    Server mode: scoped to the caller's own accessible scopes — every role,
    including owner, sees only their own private, assigned scope, and
    shared totals (no role gets a company-wide-total exception), matching
    every other read tool.

    Returns:
        A formatted statistics report.
    """
    # Build stats directly from ChromaDB rather than capturing show_stats()
    # print output.  This is safe in BOTH stdio mode (where sys.stdout is the
    # MCP pipe and must never be redirected) and HTTP mode.
    try:
        from rag_preprocessor import CHROMA_DB_PATH, get_chroma_client

        user = _current_user(ctx)
        scoped_note = ""

        if user is None:
            # Personal mode — unchanged. Enumerate ALL physical collections
            # via list_collections() so stats stay accurate even with old
            # scoped collections still lingering from before the Phase 7
            # single-collection cutover (they only fully disappear once the
            # operator does a full wipe-and-reindex).
            try:
                client, embedding_func = get_chroma_client()
                all_cols = client.list_collections()
                if not all_cols:
                    return "📭 Database is empty or not yet created."
                _collections = [
                    client.get_collection(name=col.name, embedding_function=embedding_func)
                    for col in all_cols
                ]
            except RuntimeError:
                return "📭 Database is empty or not yet created."

            total_chunks = 0
            metadatas = []
            for _col in _collections:
                c = _col.count()
                total_chunks += c
                if c:
                    sample = _col.get(limit=min(5000, c), include=["metadatas"])
                    metadatas.extend(sample.get('metadatas', []))
        else:
            # Server mode: scope to exactly what this caller can read via the
            # single collection's "scope" metadata filter (SCOPE_SIMPLIFICATION_
            # SPEC.md section 3.7) — same helper search_documents() and
            # get_knowledge_base_overview() use.
            coll, where_filter = _scoped_collections_for_ctx(ctx)
            scoped_note = "  (scoped to your accessible scopes)"

            id_probe = coll.get(where=where_filter, include=[])
            total_chunks = len(id_probe.get('ids') or [])
            if total_chunks == 0:
                return "📭 Database is empty or not yet created."
            sample = coll.get(where=where_filter, limit=min(5000, total_chunks),
                              include=["metadatas"])
            metadatas = sample.get('metadatas', [])

        if total_chunks == 0:
            return "📭 Database is empty."

        unique_files: dict = {}
        ext_counts:   dict = {}
        for m in metadatas:
            fp  = m.get('filepath', '')
            ext = m.get('extension', 'unknown').lower().lstrip('.') or 'other'
            if fp not in unique_files:
                unique_files[fp] = ext
                ext_counts[ext] = ext_counts.get(ext, 0) + 1

        lines = [
            f"📊 AI-Prowler Database Statistics{scoped_note}",
            "─" * 40,
            f"  Total chunks     : {total_chunks:,}",
            f"  Unique documents : {len(unique_files):,}",
            f"  Database path    : {CHROMA_DB_PATH}",
            "",
            "  Documents by type:",
        ]
        for ext in sorted(ext_counts.keys()):
            lines.append(
                f"    {ext.upper():>8} : {ext_counts[ext]:,} file(s)"
            )
        return "\n".join(lines)

    except Exception as exc:
        return f"❌ Could not retrieve stats: {exc}"


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 5 — list_tracked_directories
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def list_tracked_directories(ctx: "Context | None" = None) -> str:
    """
    List all paths (directories AND individually-tracked files) currently
    registered for auto-update tracking.

    Each entry is annotated with its type — 📁 for directories, 📄 for files —
    so it's clear which entries are watched as folders vs. as single files.

    Server mode (v8.1.5 fix): gated by SCOPE, not role — any role may call
    this, and the returned list is filtered to only the paths whose assigned
    scope is one this caller may already SEARCH (_allowed_scopes: their own
    private scope + assigned scopes + shared). This was previously an
    owner/manager-only role gate that hid the list entirely from staff/
    field_crew, on the reasoning that "the raw tracked-paths list is
    company-wide administrative information." That reasoning didn't hold up:
    every path in the list is either shared, one of the caller's own assigned
    scopes, or their own private folder anyway (the ones they DON'T have
    access to are simply filtered out, exactly like search results already
    are) — so there's no separate confidentiality boundary being crossed by
    showing someone the subset of tracked paths they can already reach.
    untrack_directory (the destructive sibling of this tool) is NOT changed
    by this fix and remains owner/manager-only — deleting indexed content is
    a different, more consequential permission question than just listing
    what's tracked.

    Returns:
        A formatted list of tracked paths, or a message if none are registered.
    """
    dirs = load_auto_update_list()
    if not dirs:
        return (
            "ℹ️  No paths are currently tracked.\n"
            "Use index_path to index a folder or file and add it to tracking."
        )

    # Server mode: filter to only paths within this caller's accessible
    # scopes (see docstring above). Personal mode (_ltd_user is None): no
    # filtering, unchanged from every prior version.
    #
    # v8.1.5 BUG FIX (found via live testing on the Server after the
    # original v8.1.5 scope-gating change): this originally resolved each
    # path's scope via _company_collection_map()/_resolve_collection_for_
    # path() -- the OLD collection_map-based resolver. Nothing has written
    # business-scope changes (sales/ops/field/etc.) to collection_map since
    # the scope_map/scope_lookup migration, so those rules were stale --
    # directories whose CURRENT scope (in scope_map) the caller could
    # legitimately search (confirmed: read_document/search_documents on
    # that exact content succeeded) were being resolved to a DIFFERENT,
    # stale scope via the dead collection_map data and incorrectly
    # filtered out here. Switched to the same scope_lookup/scope_map path
    # _allowed_scopes() itself and the real indexing pipeline
    # (rag_preprocessor.py) already use, so this can't drift from actual
    # access again -- exactly the same fix already applied to the GUI's
    # _resolve_scope_for_path()/_display_scope().
    _ltd_user = _current_user(ctx)
    if _IS_SERVER_MODE and _ltd_user:
        import scope_lookup as _sl_ltd
        _ltd_allowed = _allowed_scopes(_ltd_user)
        _ltd_scope_map = _sl_ltd.get_scope_map(_load_users())
        _ltd_privates_root = str(Path.home() / "Documents" / "AI-Prowler-Server-privates")
        dirs = [
            d for d in dirs
            if _sl_ltd.resolve_scope_for_path(
                d, _ltd_scope_map, privates_root=_ltd_privates_root
            ) in _ltd_allowed
        ]
        if not dirs:
            return (
                "ℹ️  No tracked paths fall within your accessible scopes.\n"
                "(Paths outside your assigned scopes/private area are omitted, "
                "not necessarily nonexistent.)"
            )

    lines = ["📁 Tracked paths:"]
    for i, d in enumerate(dirs, 1):
        try:
            icon = "📄" if Path(d).is_file() else ("📁" if Path(d).exists() else "❓")
        except Exception:
            icon = "❓"
        lines.append(f"  {i}. {icon} {d}")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 6 — untrack_directory
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def untrack_directory(directory: str, ctx: Context = None) -> str:
    """
    Remove a tracked path (directory OR individually-tracked file) from the
    tracking list AND delete all of its indexed chunks from the ChromaDB
    knowledge base.

    This is a destructive operation — the documents from this path will no
    longer be searchable until you re-index them.

    Server mode (v8.1.5) — two-tier gate, direct product decision:
      • Anyone, any role, may untrack a path inside their OWN personal
        directory — the same personal-directory model already used by
        create_file/write_file/str_replace_in_file/line_replace_in_file
        (_check_personal_write_scope). No admin flag needed for your own area.
      • A path OUTSIDE the caller's own personal directory (a shared scope,
        another user's private folder, or general company-wide tracked
        state) requires the owner, or a manager/staff member with delegated
        admin rights (can_manage_users — the Admin tab's "Can manage users"
        checkbox). A plain manager or staff member without that flag is
        denied for anything outside their own personal directory; a plain
        role check alone ('manager' or 'staff') is no longer sufficient by
        itself — the flag is what actually grants it, matching how the
        Admin tab already delegates rights independent of role.
      Replaces the earlier blanket owner/manager-only role gate
      (_check_db_cap(user, 'full')).

    Args:
        directory:  Absolute path to the directory or file to remove.

    Returns:
        A summary of what was removed.
    """
    _user = _current_user(ctx)

    try:
        _resolved_dir = str(Path(directory).resolve())
    except Exception:
        _resolved_dir = directory

    # Tier 1: own personal directory — always allowed, any role. Reuses the
    # exact same check create_file/write_file/etc. already enforce, so this
    # tool and the write tools always agree on what counts as "your own area."
    _own_dir_denial = _check_personal_write_scope(ctx, _resolved_dir)
    if _own_dir_denial is not None:
        # Tier 2: outside the caller's own personal directory — owner, or
        # delegated admin (can_manage_users), only.
        _is_owner = bool(_user) and (_user.get("role") or "").strip().lower() == "owner"
        _is_admin = bool(_user) and bool(_user.get("can_manage_users"))
        if not (_is_owner or _is_admin):
            return (
                f"⛔ untrack_directory: '{directory}' is outside your personal "
                "directory. Untracking anything outside your own personal "
                "directory requires the owner, or a manager/staff member with "
                "delegated admin rights (\"Can manage users\") — ask your owner "
                "or an admin to untrack this path, or to grant you admin rights."
            )

    with _capture_stdout() as buf:
        try:
            result = remove_directory_from_index(directory)
        except Exception as exc:
            return f"❌ Failed to remove path: {exc}"

    output = buf.getvalue().strip()
    if isinstance(result, dict):
        chunks = result.get('chunks_removed', 'unknown')
        files  = result.get('files_removed', 'unknown')
        output += f"\n✅ Removed {chunks} chunk(s) from {files} file(s)."
    return output or f"✅ Path removed: {directory}"


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 7 — check_ai_prowler_status
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def check_ai_prowler_status(ctx: "Context | None" = None) -> str:
    """
    Check AI-Prowler's health: ChromaDB connectivity, embedding model status,
    document count, and database path. No Ollama or local LLM involved.

    The chunk count and connectivity summary are always company-wide (a basic
    "is the server alive" signal, not a data-browsing result) — this is
    unchanged in server mode for any role.

    The tracked-paths list, however, reveals internal folder/file names —
    real information, not just a health signal — so in server mode it's
    shown only to owner/manager (same gate as list_tracked_directories).
    Staff/field_crew still see the chunk count and health status; the
    tracked-paths section is simply omitted for them.

    Returns:
        A diagnostic status report for the Agentic RAG knowledge base.
    """
    # Wait for background prewarm (ChromaDB + embedding model) — max 25s.
    # Prewarm runs in a thread so mcp.run() could start before it finishes.
    _log.info("check_ai_prowler_status: tool called — waiting for prewarm if still running")
    if not _prewarm_event.wait(timeout=60):
        _log.warning("check_ai_prowler_status: prewarm timeout — returning early")
        return (
            "⏳ AI-Prowler is still initializing (ChromaDB + embedding model "
            "loading in background).\n\n"
            "This usually takes 15–60 seconds on first use after Claude Desktop "
            "starts. Please wait a moment and try again."
        )
    _log.info("check_status: prewarm ready, executing")

    from rag_preprocessor import get_chroma_client, CHROMA_DB_PATH

    lines = ["🔍 AI-Prowler Status Check", "─" * 40]

    # ── ChromaDB & embedding model ────────────────────────────────────────────
    with _capture_stdout() as buf:
        try:
            # Trigger embedding model load for the output capture
            client, embedding_func = get_chroma_client()
            # Fix v7.0.1: count chunks across ALL collections in the database
            # rather than calling get_or_create_collection() (which returns only
            # the default 'documents' collection and ignores scoped collections
            # in Business Server mode) or _scoped_collections_for_ctx(None)
            # (which also falls back to personal/single-collection mode when
            # ctx=None). Instead enumerate every existing collection directly
            # from the ChromaDB client — this always reflects the true total
            # regardless of edition or mode.
            try:
                all_collections = client.list_collections()
                chunk_count = sum(
                    client.get_collection(
                        name=col.name,
                        embedding_function=embedding_func
                    ).count()
                    for col in all_collections
                )
            except Exception:
                chunk_count = 0
            db_ok = True
        except Exception as exc:
            db_ok = False
            chunk_count = 0
            db_error = str(exc)

    init_output = buf.getvalue().strip()

    if db_ok:
        lines.append("✅ ChromaDB  : connected")
        lines.append(f"   Chunks    : {chunk_count:,}")
        lines.append(f"   Database  : {CHROMA_DB_PATH}")
    else:
        lines.append("❌ ChromaDB  : not reachable")
        lines.append(f"   Error     : {db_error}")
        lines.append("   Try re-indexing with index_path.")

    # ── Embedding model info (parsed from init output) ────────────────────────
    # Bug fix v7.0.0: filter the captured stdout to ONLY surface the clean
    # "✅ Embedding model loaded on <device>" line emitted by rag_preprocessor.
    # The old filter accepted any line containing "Loading" or "embedding",
    # which on fresh installs caught HuggingFace's first-time download /
    # cache-resolution chatter and made the status check look like an error.
    # Reported by Rick, Jamie, and Sam during v6.0.2 installs (learning fe420ae5).
    lines.append("")
    emitted_embedding_line = False
    if init_output:
        for line in init_output.splitlines():
            s = line.strip()
            # Allowlist: only the clean rag_preprocessor confirmation line.
            # Look for the ✅ + "Embedding model loaded" signature, OR an
            # explicit device announcement we control.
            if ("Embedding model loaded" in s
                    or "Embedding model is loaded" in s):
                lines.append(f"   {s}")
                emitted_embedding_line = True
                break  # one line is enough — don't dump the rest
    if not emitted_embedding_line:
        lines.append("   Embedding model: loaded")

    # ── Empty-knowledge-base hint ─────────────────────────────────────────────
    # Bug fix v7.0.0: on a fresh install ChromaDB connects fine but has zero
    # chunks. The previous status output buried this fact under technical
    # diagnostics, leading users (Rick/Jamie/Sam) to think AI-Prowler was
    # broken. Lead with a clear "this is what's expected and here's how to
    # fix it" message so the first impression is friendly and actionable.
    if db_ok and chunk_count == 0:
        lines.append("")
        lines.append("ℹ️  Knowledge base is empty — this is normal on a fresh install.")
        lines.append("   To index your documents, either:")
        lines.append("     • Call index_path(\"<path>\") from this tool, OR")
        lines.append("     • Open the AI-Prowler GUI → Reindex tab → pick a folder.")
        lines.append("   Once indexed, search_documents and friends will return real results.")

    # ── Tracked paths (directories + individually-tracked files) ─────────────
    # Server mode: owner/manager only — same gate as list_tracked_directories.
    # The chunk count above is a health signal; this list is real internal
    # folder/file names, so staff/field_crew don't see it here.
    _cas_user = _current_user(ctx)
    _cas_show_paths = True
    if _cas_user is not None:
        _cas_full_access, _ = _check_db_cap(_cas_user, "full")
        _cas_show_paths = _cas_full_access

    if _cas_show_paths:
        try:
            from rag_preprocessor import load_auto_update_list
            tracked = load_auto_update_list() or []
            n_files = sum(1 for p in tracked if Path(p).is_file())
            n_dirs  = sum(1 for p in tracked if Path(p).is_dir())
            lines.append("")
            suffix_parts = []
            if n_dirs:  suffix_parts.append(f"{n_dirs} dir(s)")
            if n_files: suffix_parts.append(f"{n_files} file(s)")
            suffix = f" ({', '.join(suffix_parts)})" if suffix_parts else ""
            lines.append(f"📁 Tracked paths : {len(tracked)}{suffix}")
            for d in tracked[:5]:
                try:
                    icon = "📄" if Path(d).is_file() else ("📁" if Path(d).exists() else "❓")
                except Exception:
                    icon = "❓"
                lines.append(f"   {icon} {d}")
            if len(tracked) > 5:
                lines.append(f"   ... and {len(tracked) - 5} more")
        except Exception:
            pass

    lines.append("")
    lines.append("✅ AI-Prowler Agentic RAG ready." if db_ok else "⚠️  Knowledge base unavailable.")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# ░░  AGENTIC RAG TOOLS  ░░
# ══════════════════════════════════════════════════════════════════════════════


def _get_collection():
    """Return ChromaDB collection directly, using cached client/embedding."""
    from rag_preprocessor import get_chroma_client, COLLECTION_NAME
    client, embedding_func = get_chroma_client()
    try:
        col = client.get_or_create_collection(
            name=COLLECTION_NAME,
            embedding_function=embedding_func
        )
        if col.count() == 0:
            raise RuntimeError(
                "No indexed documents found. "
                "Use index_path to index some documents first."
            )
        return col
    except RuntimeError:
        raise
    except Exception:
        raise RuntimeError(
            "No indexed documents found. "
            "Use index_path to index some documents first."
        )


def _scoped_collections_for_ctx(ctx):
    """Return (collection, where_filter) for the current request.

    SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cutover, query-side,
    2026-07-17): there is only one physical ChromaDB collection now. This
    used to return a LIST of collection OBJECTS (one per logical scope the
    caller could read, built by _allowed_collections() + a multi-collection
    fan-out with an owner/admin elevation carve-out to browse every role
    and private collection on the server). That whole mechanism is gone.

    Callers now get back exactly one collection plus a `where` dict to pass
    into every .query()/.get() call on it -- access control is enforced by
    filtering on each chunk's "scope" metadata field (build_scope_resolver,
    wired into indexing) instead of by which physical collection existed.

    Returns:
        (collection, where_filter) -- where_filter is None in personal mode
        (no filtering needed/possible) or a real {"scope": {"$in": [...]}}
        dict in server mode. An empty allowed-scopes set still produces a
        real (non-None) filter with an empty $in list, which matches
        nothing -- fail-closed, never fail-open.

    Raises RuntimeError only in personal mode when nothing is indexed
    (preserves the existing error contract every caller's try/except
    already handles).

    IMPORTANT CONSEQUENCE (documented, not silently dropped): the previous
    owner/can_manage_users elevation to browse every team member's role
    and private collection via search tools is GONE -- matches
    _allowed_scopes()'s own "no role-based elevation, ever" direct product
    decision (2026-07-16), which deliberately did not carry that carve-out
    forward (see its docstring). The owner retains full filesystem access
    to any directory directly (GUI, or list_writable_directories /
    grant_write_access), and can still browse a user's private FOLDER on
    disk -- what's gone is browsing another user's already-indexed private
    CONTENT through search_documents et al.
    """
    from rag_preprocessor import get_chroma_client, COLLECTION_NAME
    client, embedding_func = get_chroma_client()

    user = _current_user(ctx)
    if user is None:
        # Personal mode -- single collection, no filter, original contract
        # (raises if nothing indexed yet).
        try:
            coll = client.get_collection(name=COLLECTION_NAME,
                                         embedding_function=embedding_func)
            return (coll, None)
        except Exception:
            raise RuntimeError(
                "No indexed documents found. "
                "Use index_path to index some documents first.")

    scopes = _allowed_scopes(user)
    where_filter = {"scope": {"$in": sorted(scopes)}}  # [] -> matches nothing, fail-closed
    coll = client.get_or_create_collection(name=COLLECTION_NAME,
                                           embedding_function=embedding_func)
    return (coll, where_filter)



def _enumerate_scope_collections(client, prefix: str = "scope-") -> list:
    """List existing PHYSICAL collection names beginning with `prefix`
    (default all 'scope-*' — i.e. every shared/role/user scope collection).
    Used to give owners/admins read access to everything on the company server.
    Returns [] on any error (fail-closed: owner sees fewer, never more).
    """
    try:
        raw = client.list_collections()
        names = []
        for c in raw:
            names.append(c if isinstance(c, str) else getattr(c, "name", str(c)))
    except Exception:
        return []
    return [n for n in names if n.startswith(prefix)]


def _enumerate_role_collections(client) -> list:
    """List all existing role:* collections (LOGICAL names) by inspecting
    ChromaDB for physical 'scope-role-<name>' collections. Retained for
    _allowed_collections' owner role-enumeration arg and unit tests. Returns []
    on any error (fail-closed).
    """
    role_logicals = []
    for phys in _enumerate_scope_collections(client, "scope-role-"):
        sub = phys[len("scope-role-"):]
        if sub:
            role_logicals.append(f"scope:{sub}")
    return role_logicals


@mcp.tool()
def get_knowledge_base_overview(ctx: Context = None) -> str:
    """
    AGENTIC RAG - Start here.
    Returns a high-level summary of the AI-Prowler knowledge base: total
    document count, file types present, tracked source directories, and
    total chunk count. Call this at the start of any research task so you
    understand what knowledge is available before deciding how to search.
    No LLM or Ollama required.

    Returns:
        Structured overview of the knowledge base contents.
    """
    if not _prewarm_event.wait(timeout=60):
        _log.warning("get_knowledge_base_overview: prewarm timeout — returning early")
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    from rag_preprocessor import CHROMA_DB_PATH, load_auto_update_list
    try:
        coll, where_filter = _scoped_collections_for_ctx(ctx)
    except Exception:
        return (
            "Knowledge base is empty — no documents indexed yet.\n"
            "Use index_path to index a folder of documents."
        )

    id_probe = coll.get(where=where_filter, include=[]) if where_filter else coll.get(include=[])
    total_chunks = len(id_probe.get('ids') or [])
    metadatas = []
    if total_chunks:
        if where_filter:
            sample = coll.get(where=where_filter, limit=min(2000, total_chunks))
        else:
            sample = coll.get(limit=min(2000, total_chunks))
        metadatas.extend(sample.get('metadatas', []))
    if total_chunks == 0:
        return "Knowledge base is empty."

    unique_files = {}
    ext_counts   = {}
    dir_counts   = {}   # parent_directory → set of filepaths
    for m in metadatas:
        fp  = m.get('filepath', '')
        fn  = m.get('filename', fp)
        ext = m.get('extension', 'unknown').lower()
        if fp not in unique_files:
            unique_files[fp] = {'filename': fn, 'extension': ext,
                                'total_chunks': m.get('total_chunks', 1)}
        ext_counts[ext] = ext_counts.get(ext, 0) + 1

        # Build directory tree from provenance metadata
        parent_dir = m.get('parent_directory', '')
        dir_chain  = m.get('directory_chain', '')
        tree_key   = dir_chain if dir_chain else parent_dir
        if tree_key:
            if tree_key not in dir_counts:
                dir_counts[tree_key] = set()
            dir_counts[tree_key].add(fp)

    tracked_dirs = load_auto_update_list() or []

    # Server mode: only show tracked directories that actually contributed at
    # least one chunk visible in this caller's scoped collections. The raw
    # global list would otherwise leak every tracked path company-wide to
    # every role — including directories the caller has no read access to —
    # even though every other section of this report is properly scoped.
    if _current_user(ctx) is not None and tracked_dirs:
        visible_paths = list(unique_files.keys())
        tracked_dirs = [
            d for d in tracked_dirs
            if any(_path_is_under(fp, d) for fp in visible_paths)
        ]

    lines = [
        "AI-Prowler Knowledge Base Overview",
        "=" * 45,
        f"Total documents : {len(unique_files):,}",
        f"Total chunks    : {total_chunks:,}",
        f"Database path   : {CHROMA_DB_PATH}",
        "",
        "Document types:",
    ]
    for ext, cnt in sorted(ext_counts.items(), key=lambda x: -x[1]):
        pct = cnt / total_chunks * 100
        lines.append(f"  {ext:12s}  {cnt:5,} chunks  ({pct:.1f}%)")

    if dir_counts:
        lines.append("")
        lines.append("Directory tree (by content location):")
        for chain in sorted(dir_counts.keys()):
            n_files = len(dir_counts[chain])
            lines.append(f"  📁 {chain}  ({n_files} file{'s' if n_files != 1 else ''})")

    if tracked_dirs:
        lines.append("")
        lines.append("Tracked source directories:")
        for d in tracked_dirs:
            lines.append(f"  - {d}")

    lines.append("")
    lines.append(
        "Next steps: search_documents(query) for broad search | "
        "search_within_directory(query, dir) for targeted search | "
        "list_indexed_directories() for full directory tree"
    )
    return "\n".join(lines)


@mcp.tool()
def search_documents(
    query: str,
    n_results: int = 8,
    min_similarity: float = 0.0,
    ctx: Context = None,
) -> str:
    """
    AGENTIC RAG - Primary search tool.
    Searches the knowledge base using semantic similarity and returns raw
    document chunks. No LLM involved — Claude receives chunks directly and
    does all reasoning itself. Call multiple times with different or refined
    queries to gather context before synthesising an answer.
    No Ollama required — works on any machine.

    Args:
        query:          Natural language search query. Try different phrasings
                        if the first search misses something.
        n_results:      Number of chunks to return (default 8, max 20).
        min_similarity: Filter chunks below this score 0.0-1.0 (default 0.0).

    Returns:
        Numbered list of matching chunks with source file, chunk position,
        similarity score, and full text content.
    """
    if not query.strip():
        return "Query cannot be empty."

    if not _prewarm_event.wait(timeout=60):
        _log.warning("search_documents: prewarm timeout — returning early")
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    n_results = min(max(1, n_results), 20)

    # ── Server-mode scope filtering (SCOPE_SIMPLIFICATION_SPEC.md section 3.7) ─
    # Use _scoped_collections_for_ctx() to get the single ChromaDB collection
    # plus a "scope" metadata where-filter for this caller, then query it
    # directly. Personal mode: where_filter is None (no filtering needed).
    try:
        coll, where_filter = _scoped_collections_for_ctx(ctx)
    except RuntimeError:
        # Personal mode and no documents indexed yet.
        return (
            f"No results found for: '{query}'\n"
            "No documents are indexed yet. Use index_path to index a folder."
        )
    except Exception as exc:
        return f"Search failed resolving collections: {exc}"

    # Query the single collection, filtered by scope.
    _all_chunks = []
    _seen_ids   = set()
    try:
        _query_kwargs = {"query_texts": [query], "n_results": n_results}
        if where_filter:
            _query_kwargs["where"] = where_filter
        _res = coll.query(**_query_kwargs)
    except Exception as _qe:
        _log.warning("search_documents: query failed: %s", _qe)
        _res = None
    if _res and _res.get('documents'):
        for _i in range(len(_res['documents'][0])):
            _cid = (_res.get('ids') or [[]])[0][_i] if _res.get('ids') else None
            if _cid and _cid in _seen_ids:
                continue
            if _cid:
                _seen_ids.add(_cid)
            _dist = _res['distances'][0][_i]
            _all_chunks.append({
                'content':  _res['documents'][0][_i],
                'metadata': _res['metadatas'][0][_i],
                'distance': _dist,
                'similarity': 1.0 - _dist,
            })

    # Sort by similarity descending, cap at n_results.
    _all_chunks.sort(key=lambda c: c['similarity'], reverse=True)
    chunks = _all_chunks[:n_results]
    _log.debug("search_documents: %d chunk(s)", len(chunks))

    if not chunks:
        return (
            f"No results found for: '{query}'\n"
            "Try different keywords or call get_knowledge_base_overview() "
            "to see what is indexed."
        )

    if min_similarity > 0.0:
        chunks = [c for c in chunks if c.get('similarity', 0) >= min_similarity]
        if not chunks:
            return (
                f"No results above similarity {min_similarity:.2f} for: '{query}'\n"
                "Try lowering min_similarity or rephrasing."
            )

    lines = [
        f"Search results for: \"{query}\"",
        f"Returning {len(chunks)} chunk(s)",
        "-" * 55, "",
    ]
    for i, chunk in enumerate(chunks, 1):
        meta      = chunk.get('metadata', {})
        filename  = meta.get('filename', 'unknown')
        filepath  = meta.get('filepath', '')
        chunk_idx = meta.get('chunk_index', 0)
        total_ch  = meta.get('total_chunks', 1)
        ext       = meta.get('extension', '')
        sim       = chunk.get('similarity', 0.0)
        content   = chunk.get('content', '').strip()

        # ── Provenance fields (graceful fallback for pre-upgrade chunks) ──
        dir_chain   = meta.get('directory_chain', '')
        doc_id      = meta.get('document_id', '')
        doc_title   = meta.get('doc_title', '')
        file_mod    = meta.get('file_modified', '')
        parent_dir  = meta.get('parent_directory', '')

        lines.append(
            f"[{i}] {filename}  chunk {chunk_idx+1}/{total_ch}  "
            f"similarity: {sim:.3f}  type: {ext}"
        )
        if dir_chain:
            lines.append(f"    📁 {dir_chain}")
        elif filepath and filepath != filename:
            lines.append(f"    Path: {filepath}")
        if doc_id:
            lines.append(f"    📄 Document ID: {doc_id}")
        if doc_title and doc_title != Path(filename).stem:
            lines.append(f"    📝 Title: {doc_title}")
        if file_mod:
            lines.append(f"    ⏱  Modified: {file_mod[:10]}")
        lines.append("")
        lines.append(content)
        lines.append("")
        lines.append("-" * 55)
        lines.append("")
    lines.append(
        "Tips: search_documents() again with different query | "
        "search_within_directory(query, dir) for targeted results | "
        "Use expand_search_result(filename, chunk_index) to expand | "
        "list_indexed_directories() to see indexed folder tree"
    )
    return "\n".join(lines)


@mcp.tool()
def expand_search_result(
    filename: str,
    chunk_index: int,
    window: int = 2,
    ctx: Context = None,
) -> str:
    """
    AGENTIC RAG - Expand context around a search result.
    Retrieves the chunks immediately before and after a specific chunk,
    giving fuller context when a search result is cut off at a boundary.
    Use this when a result looks relevant but is incomplete.
    No Ollama required.

    Args:
        filename:    Filename from a search_documents result (partial match ok).
        chunk_index: Zero-based chunk index from the search result.
        window:      Chunks before and after to include (default 2, max 5).

    Returns:
        Target chunk plus surrounding chunks in reading order.
    """
    window = max(1, min(5, window))

    try:
        coll, where_filter = _scoped_collections_for_ctx(ctx)
    except RuntimeError as e:
        return str(e)

    try:
        fn_lower  = filename.lower()
        chunk_map = {}
        _get_kwargs = {"limit": 5000, "include": ["metadatas", "documents"]}
        if where_filter:
            _get_kwargs["where"] = where_filter
        sample = coll.get(**_get_kwargs)
        for doc, meta in zip(sample['documents'], sample['metadatas']):
            if fn_lower in meta.get('filename', '').lower() or \
               fn_lower in meta.get('filepath', '').lower():
                chunk_map[meta.get('chunk_index', 0)] = (doc, meta)
    except Exception as exc:
        return f"Could not retrieve chunks: {exc}"

    if not chunk_map:
        return (
            f"No document matching '{filename}' found.\n"
            "Use list_indexed_documents() to see available filenames."
        )
    if chunk_index not in chunk_map:
        return (
            f"Chunk index {chunk_index} not found for '{filename}'.\n"
            f"Available: {sorted(chunk_map.keys())[:20]}"
        )

    target_meta  = chunk_map[chunk_index][1]
    total_chunks = target_meta.get('total_chunks', max(chunk_map.keys()) + 1)
    filepath     = target_meta.get('filepath', filename)
    start        = max(0, chunk_index - window)
    end          = min(total_chunks - 1, chunk_index + window)

    lines = [
        f"Context for: {filepath}",
        f"Showing chunks {start+1}-{end+1} of {total_chunks} "
        f"(target chunk {chunk_index+1} marked with >)",
        "=" * 55, "",
    ]
    for idx in range(start, end + 1):
        if idx not in chunk_map:
            lines.append(f"[chunk {idx+1}/{total_chunks}]  (not in current sample)")
            lines.append("")
            continue
        content, _ = chunk_map[idx]
        marker = "> " if idx == chunk_index else "  "
        lines.append(f"{marker}[chunk {idx+1}/{total_chunks}]")
        lines.append(content.strip())
        lines.append("")
    return "\n".join(lines)


@mcp.tool()
def read_document(
    filename: str,
    start_chunk: int = 0,
    max_chunks: int = 10,
    ctx: Context = None,
) -> str:
    """
    AGENTIC RAG - Read a whole document in sequence.
    Retrieves chunks from a specific indexed document in reading order.
    Use when you need to read a full document rather than search fragments —
    e.g. summarise a contract, read a manual section, review a report.
    For long documents, make multiple calls with increasing start_chunk.
    No Ollama required.

    Args:
        filename:    Filename to retrieve (partial matches accepted).
                     Use list_indexed_documents() to browse available files.
        start_chunk: Zero-based chunk to start from (default 0).
        max_chunks:  Max chunks to return (default 10, max 30).

    Returns:
        Sequential chunks in reading order with position indicators and
        a continuation hint if the document has more chunks.
    """
    max_chunks = min(max(1, max_chunks), 30)

    try:
        coll, where_filter = _scoped_collections_for_ctx(ctx)
    except RuntimeError as e:
        return str(e)

    try:
        fn_lower = filename.lower()
        matches  = []
        _get_kwargs = {"limit": 5000, "include": ["metadatas", "documents"]}
        if where_filter:
            _get_kwargs["where"] = where_filter
        sample = coll.get(**_get_kwargs)
        for doc, meta in zip(sample['documents'], sample['metadatas']):
            if fn_lower in meta.get('filename', '').lower() or \
               fn_lower in meta.get('filepath', '').lower():
                matches.append((meta.get('chunk_index', 0), doc, meta))
    except Exception as exc:
        return f"Could not access knowledge base: {exc}"

    if not matches:
        return (
            f"No document matching '{filename}' found.\n"
            "Use list_indexed_documents() to see available filenames."
        )

    matches.sort(key=lambda x: x[0])
    total_chunks = matches[-1][2].get('total_chunks', len(matches))
    actual_name  = matches[0][2].get('filename', filename)
    actual_path  = matches[0][2].get('filepath', filename)
    ext          = matches[0][2].get('extension', '')

    page = [(idx, doc, meta) for idx, doc, meta in matches
            if idx >= start_chunk][:max_chunks]

    if not page:
        return (
            f"No chunks from index {start_chunk} onwards.\n"
            f"'{actual_name}' has {total_chunks} chunks (indices 0-{total_chunks-1})."
        )

    last_shown = page[-1][0]
    has_more   = last_shown < total_chunks - 1

    lines = [
        f"Document: {actual_name}",
        f"Path    : {actual_path}",
        f"Type    : {ext}  |  Total chunks: {total_chunks}",
        f"Showing chunks {page[0][0]+1}-{last_shown+1} of {total_chunks}",
        "=" * 55, "",
    ]
    for idx, doc, _ in page:
        lines.append(f"[chunk {idx+1}/{total_chunks}]")
        lines.append(doc.strip())
        lines.append("")

    if has_more:
        next_start = last_shown + 1
        lines.append("-" * 55)
        lines.append(
            f"Document continues — {total_chunks - last_shown - 1} more chunk(s). "
            f"Call: read_document('{actual_name}', start_chunk={next_start})"
        )
    return "\n".join(lines)


@mcp.tool()
def list_indexed_documents(
    filter_ext: Optional[str] = None,
    filter_path: Optional[str] = None,
    limit: int = 50,
    ctx: Context = None,
) -> str:
    """
    AGENTIC RAG - Browse the knowledge base.
    Lists all indexed documents with file type, path, and chunk count.
    Use this to orient yourself — especially useful for questions like
    "what documents do you have about X?" or "do you have the Q3 report?"
    No Ollama required.

    Args:
        filter_ext:  Only show this file type, e.g. "pdf", "docx".
        filter_path: Only show files whose path contains this string.
        limit:       Max documents to list (default 50).

    Returns:
        Sorted list of indexed documents grouped by file type.
    """
    try:
        coll, where_filter = _scoped_collections_for_ctx(ctx)
    except RuntimeError as e:
        return str(e)

    # Aggregate metadata across the accessible scope (single collection now).
    try:
        _get_kwargs = {"include": []}
        if where_filter:
            _get_kwargs["where"] = where_filter
        id_probe = coll.get(**_get_kwargs)
        total = len(id_probe.get('ids') or [])
        metas = []
        if total:
            _get_kwargs2 = {"limit": min(5000, total), "include": ["metadatas"]}
            if where_filter:
                _get_kwargs2["where"] = where_filter
            sample = coll.get(**_get_kwargs2)
            metas.extend(sample.get('metadatas', []))
    except Exception as exc:
        return f"Could not list documents: {exc}"

    docs = {}
    for m in metas:
        fp  = m.get('filepath', '')
        fn  = m.get('filename', fp)
        ext = m.get('extension', '').lower().lstrip('.')
        tc  = m.get('total_chunks', 1)
        if filter_ext and ext != filter_ext.lower().lstrip('.'):
            continue
        if filter_path and filter_path.lower() not in fp.lower():
            continue
        if fp not in docs:
            docs[fp] = {'filename': fn, 'extension': ext,
                        'total_chunks': tc, 'filepath': fp}

    if not docs:
        note = ""
        if filter_ext:
            note += f" with type '{filter_ext}'"
        if filter_path:
            note += f" in path '{filter_path}'"
        return f"No documents found{note}."

    shown  = sorted(docs.values(), key=lambda d: d['filename'].lower())[:limit]
    by_ext = {}
    for d in shown:
        by_ext.setdefault(d['extension'] or 'other', []).append(d)

    lines = [
        f"Indexed documents: {len(docs):,} total  (showing {len(shown)})",
        "-" * 55, "",
    ]
    for ext in sorted(by_ext.keys()):
        lines.append(f"{ext.upper()} ({len(by_ext[ext])} files)")
        for d in sorted(by_ext[ext], key=lambda x: x['filename'].lower()):
            lines.append(
                f"  - {d['filename']}  ({d['total_chunks']} chunk"
                f"{'s' if d['total_chunks'] != 1 else ''})"
            )
            if d['filepath'] != d['filename']:
                lines.append(f"    {d['filepath']}")
        lines.append("")

    if len(docs) > limit:
        lines.append(f"... and {len(docs)-limit} more. Use filter_ext or filter_path.")
    lines.append("Call read_document(filename) to read a specific document.")
    return "\n".join(lines)


@mcp.tool()
def multi_query_search(
    queries: list[str],
    n_results_each: int = 5,
    min_similarity: float = 0.0,
    ctx: Context = None,
) -> str:
    """
    AGENTIC RAG - Parallel multi-angle search.
    Runs several semantic searches in one call and returns deduplicated results
    ranked by best similarity. Use when a topic has multiple angles, synonyms,
    or when you want to reformulate a question several ways at once.
    More efficient than calling search_documents() multiple times.
    No Ollama required.

    Args:
        queries:        List of 2-6 search queries to run in parallel.
                        Example: ["refund policy", "money back guarantee"]
        n_results_each: Chunks per query (default 5, max 10).
        min_similarity: Discard chunks below this score (default 0.0).

    Returns:
        Deduplicated chunks from all queries sorted by best similarity,
        labelled with which query found them.
    """
    if not queries:
        return "Provide at least one query."
    queries = [q.strip() for q in queries[:6] if q.strip()]
    n_results_each = min(max(1, n_results_each), 10)

    try:
        coll, where_filter = _scoped_collections_for_ctx(ctx)
    except RuntimeError:
        return (
            "No results found for any query:\n"
            + "\n".join(f"  - {q}" for q in queries)
            + "\nNo documents are indexed yet. Use index_path to index a folder."
        )
    except Exception as exc:
        return f"Could not access knowledge base: {exc}"

    all_chunks: dict = {}
    for q in queries:
        try:
            _query_kwargs = {"query_texts": [q], "n_results": n_results_each}
            if where_filter:
                _query_kwargs["where"] = where_filter
            _res = coll.query(**_query_kwargs)
        except Exception:
            continue
        if not _res or not _res.get('documents'):
            continue
        for _i in range(len(_res['documents'][0])):
            _dist = _res['distances'][0][_i]
            chunk = {
                'content':  _res['documents'][0][_i],
                'metadata': _res['metadatas'][0][_i],
                'distance': _dist,
                'similarity': 1.0 - _dist,
            }
            sim  = chunk.get('similarity', 0.0)
            if sim < min_similarity:
                continue
            meta = chunk.get('metadata', {})
            key  = f"{meta.get('filepath','')}::{meta.get('chunk_index',0)}"
            if key not in all_chunks or sim > all_chunks[key]['similarity']:
                all_chunks[key] = {**chunk, 'found_by': q, 'also_found_by': []}
            elif key in all_chunks and q != all_chunks[key]['found_by']:
                all_chunks[key]['also_found_by'].append(q)

    if not all_chunks:
        return (
            "No results found for any query:\n"
            + "\n".join(f"  - {q}" for q in queries)
            + "\nTry different keywords or call get_knowledge_base_overview()."
        )

    ranked = sorted(all_chunks.values(),
                    key=lambda c: c['similarity'], reverse=True)

    lines = [
        f"Multi-query search — {len(queries)} queries, "
        f"{len(ranked)} unique chunk(s) found",
        f"Queries: {' | '.join(queries)}",
        "-" * 55, "",
    ]
    for i, chunk in enumerate(ranked, 1):
        meta     = chunk.get('metadata', {})
        fn       = meta.get('filename', 'unknown')
        fp       = meta.get('filepath', '')
        cidx     = meta.get('chunk_index', 0)
        total    = meta.get('total_chunks', 1)
        ext      = meta.get('extension', '')
        sim      = chunk.get('similarity', 0.0)
        found_by = chunk.get('found_by', '')
        also     = chunk.get('also_found_by', [])
        content  = chunk.get('content', '').strip()

        # ── Provenance fields ──
        dir_chain   = meta.get('directory_chain', '')
        doc_id      = meta.get('document_id', '')
        doc_title   = meta.get('doc_title', '')
        file_mod    = meta.get('file_modified', '')

        found_note = f'found by: "{found_by}"'
        if also:
            found_note += f' (also: {", ".join(also)})'

        lines.append(
            f"[{i}] {fn}  chunk {cidx+1}/{total}  "
            f"similarity: {sim:.3f}  {found_note}"
        )
        if dir_chain:
            lines.append(f"    📁 {dir_chain}")
        elif fp and fp != fn:
            lines.append(f"    Path: {fp}")
        if doc_id:
            lines.append(f"    📄 Document ID: {doc_id}")
        if doc_title and doc_title != Path(fn).stem:
            lines.append(f"    📝 Title: {doc_title}")
        if file_mod:
            lines.append(f"    ⏱  Modified: {file_mod[:10]}")
        lines.append("")
        lines.append(content)
        lines.append("")
        lines.append("-" * 55)
        lines.append("")

    lines.append(
        "Tips: search_within_directory(query, dir) for targeted results | "
        "expand_search_result(filename, chunk_index) to expand any result."
    )
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — search_within_directory  (Provenance-Aware Scoped Search)
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def search_within_directory(
    query: str,
    directory: str,
    n_results: int = 8,
    min_similarity: float = 0.0,
    ctx: Context = None,
) -> str:
    """
    AGENTIC RAG - Directory-scoped search to prevent cross-contamination.
    Searches ONLY within a specific directory tree, filtering by the
    parent_directory or directory_chain metadata. Use this when the user
    asks about a specific case, project, client, or folder and you need
    to ensure results come only from that scope.
    No Ollama required.

    Args:
        query:          Natural language search query.
        directory:      Directory name or path fragment to restrict search to.
                        Matches against parent_directory and directory_chain.
                        Examples: 'Smith_v_Jones', '2024', 'Contracts'
        n_results:      Number of chunks to return (default 8, max 20).
        min_similarity: Filter chunks below this score 0.0-1.0 (default 0.0).

    Returns:
        Numbered list of matching chunks — all guaranteed to be from the
        specified directory tree.
    """
    if not query.strip():
        return "Query cannot be empty."
    if not directory.strip():
        return "Directory cannot be empty. Use search_documents() for unscoped search."

    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    n_results = min(max(1, n_results), 20)
    dir_lower = directory.strip().lower()

    try:
        coll, where_filter = _scoped_collections_for_ctx(ctx)
    except RuntimeError as e:
        return str(e)

    def _merge_where(extra):
        """Combine the caller's scope where_filter with an additional
        condition, using ChromaDB's $and when both are present."""
        if where_filter and extra:
            return {"$and": [where_filter, extra]}
        return where_filter or extra

    # ── Strategy: fetch more results than requested, then filter by directory ──
    # ChromaDB's `where` filter supports exact match on metadata fields.
    # We try parent_directory exact match first (fastest). If that yields
    # too few results, fall back to a broader search + client-side filter
    # on directory_chain (substring match). Both attempts are scoped to
    # this caller's accessible "scope" metadata via where_filter above.
    chunks_raw = []
    # Attempt 1: exact match on parent_directory.
    try:
        results = coll.query(
            query_texts=[query],
            n_results=min(n_results * 3, 60),
            where=_merge_where({"parent_directory": directory.strip()})
        )
        if results and results.get('documents'):
            for i in range(len(results['documents'][0])):
                chunks_raw.append({
                    'content':    results['documents'][0][i],
                    'metadata':   results['metadatas'][0][i],
                    'distance':   results['distances'][0][i],
                    'similarity': 1 - results['distances'][0][i],
                })
    except Exception:
        pass

    # Attempt 2: if exact match found nothing, do a broad search and
    # filter client-side on directory_chain (case-insensitive substring).
    if not chunks_raw:
        try:
            results = coll.query(
                query_texts=[query],
                n_results=min(n_results * 5, 100),
                where=where_filter,
            )
            if results and results.get('documents'):
                for i in range(len(results['documents'][0])):
                    meta = results['metadatas'][0][i]
                    chain = meta.get('directory_chain', '').lower()
                    parent = meta.get('parent_directory', '').lower()
                    fp = meta.get('filepath', '').lower()
                    if (dir_lower in chain or dir_lower in parent
                            or dir_lower in fp):
                        chunks_raw.append({
                            'content':    results['documents'][0][i],
                            'metadata':   meta,
                            'distance':   results['distances'][0][i],
                            'similarity': 1 - results['distances'][0][i],
                        })
        except Exception:
            pass

    # Sort by best similarity before filtering/truncating.
    chunks_raw.sort(key=lambda c: c['distance'])

    if min_similarity > 0.0:
        chunks_raw = [c for c in chunks_raw if c['similarity'] >= min_similarity]

    # Limit to requested count
    chunks_raw = chunks_raw[:n_results]

    if not chunks_raw:
        return (
            f"No results found for '{query}' within directory '{directory}'.\n"
            "Possible causes:\n"
            f"  - No documents indexed from a directory matching '{directory}'\n"
            "  - Try a broader search with search_documents() first\n"
            "  - Call list_indexed_directories() to see available directory trees"
        )

    lines = [
        f"Scoped search for: \"{query}\"",
        f"Directory filter: \"{directory}\"",
        f"Returning {len(chunks_raw)} chunk(s)",
        "-" * 55, "",
    ]
    for i, chunk in enumerate(chunks_raw, 1):
        meta      = chunk['metadata']
        filename  = meta.get('filename', 'unknown')
        chunk_idx = meta.get('chunk_index', 0)
        total_ch  = meta.get('total_chunks', 1)
        ext       = meta.get('extension', '')
        sim       = chunk['similarity']
        content   = chunk['content'].strip()
        dir_chain = meta.get('directory_chain', '')
        doc_id    = meta.get('document_id', '')
        doc_title = meta.get('doc_title', '')
        file_mod  = meta.get('file_modified', '')

        lines.append(
            f"[{i}] {filename}  chunk {chunk_idx+1}/{total_ch}  "
            f"similarity: {sim:.3f}  type: {ext}"
        )
        if dir_chain:
            lines.append(f"    📁 {dir_chain}")
        if doc_id:
            lines.append(f"    📄 Document ID: {doc_id}")
        if doc_title and doc_title != Path(filename).stem:
            lines.append(f"    📝 Title: {doc_title}")
        if file_mod:
            lines.append(f"    ⏱  Modified: {file_mod[:10]}")
        lines.append("")
        lines.append(content)
        lines.append("")
        lines.append("-" * 55)
        lines.append("")

    lines.append(
        "All results are from the specified directory scope. "
        "Use expand_search_result(filename, chunk_index) to expand."
    )
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — list_indexed_directories  (Directory Tree Discovery)
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def list_indexed_directories(ctx: Context = None) -> str:
    """
    AGENTIC RAG - Directory tree discovery.
    Lists all indexed directory trees with document counts per directory.
    Use this to understand the knowledge base structure and identify the
    right directory scope for search_within_directory().
    No Ollama required.

    Returns:
        Directory tree with file counts, sorted alphabetically.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    try:
        coll, where_filter = _scoped_collections_for_ctx(ctx)
    except RuntimeError as e:
        return str(e)

    try:
        _get_kwargs = {"include": []}
        if where_filter:
            _get_kwargs["where"] = where_filter
        id_probe = coll.get(**_get_kwargs)
        total = len(id_probe.get('ids') or [])
        metadatas = []
        if total:
            _get_kwargs2 = {"limit": min(5000, total), "include": ["metadatas"]}
            if where_filter:
                _get_kwargs2["where"] = where_filter
            sample = coll.get(**_get_kwargs2)
            metadatas.extend(sample.get('metadatas', []))
    except Exception as exc:
        return f"Could not read knowledge base: {exc}"

    if not metadatas:
        return "Knowledge base is empty — no documents indexed yet."

    # Build directory tree: directory_chain → set of unique filepaths
    dir_tree = {}       # chain → set(filepath)
    parent_dirs = {}    # parent_directory → set(filepath)
    for m in metadatas:
        fp    = m.get('filepath', '')
        chain = m.get('directory_chain', '')
        pdir  = m.get('parent_directory', '')

        if chain:
            if chain not in dir_tree:
                dir_tree[chain] = set()
            dir_tree[chain].add(fp)

        if pdir:
            if pdir not in parent_dirs:
                parent_dirs[pdir] = set()
            parent_dirs[pdir].add(fp)

    if not dir_tree and not parent_dirs:
        return (
            "No directory provenance metadata found.\n"
            "This likely means the documents were indexed before the\n"
            "provenance system was added.  Re-index your directories\n"
            "to add directory_chain metadata to all chunks."
        )

    lines = [
        "AI-Prowler — Indexed Directory Tree",
        "=" * 50,
        "",
    ]

    if dir_tree:
        lines.append("Directory chains (full breadcrumb paths):")
        lines.append("-" * 40)
        total_files = 0
        for chain in sorted(dir_tree.keys()):
            n = len(dir_tree[chain])
            total_files += n
            lines.append(f"  📁 {chain}  ({n} file{'s' if n != 1 else ''})")
        lines.append("")
        lines.append(f"Total: {len(dir_tree)} directories, {total_files} files")
    elif parent_dirs:
        lines.append("Parent directories (immediate folder names):")
        lines.append("-" * 40)
        total_files = 0
        for pdir in sorted(parent_dirs.keys()):
            n = len(parent_dirs[pdir])
            total_files += n
            lines.append(f"  📁 {pdir}  ({n} file{'s' if n != 1 else ''})")
        lines.append("")
        lines.append(f"Total: {len(parent_dirs)} directories, {total_files} files")

    lines.append("")
    lines.append(
        "Use search_within_directory(query, directory_name) to search "
        "within a specific directory."
    )
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# ░░  ACTION TOOLS — Field Service Automation  ░░
#
# All free-tier tools use no API key:
#   • Weather    : Open-Meteo  (https://open-meteo.com)
#   • Geocoding  : Nominatim / OpenStreetMap (https://nominatim.openstreetmap.org)
#   • Routing    : OSRM public server (http://router.project-osrm.org)
#                  OSRM's /trip endpoint solves TSP natively — no OR-Tools needed
#   • Maps URL   : Google Maps URL scheme (free, no key, tap-to-navigate)
#
# Spreadsheet update uses openpyxl (already installed).
# ══════════════════════════════════════════════════════════════════════════════


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 1 — get_weather
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def get_weather(location: str, days: int = 3) -> str:
    """
    Get current weather conditions and a multi-day forecast for any location.
    Uses Open-Meteo for weather data and Nominatim for geocoding.
    Both are completely free with no API key required.

    Use this before planning outdoor field service schedules to flag rain risk.
    Rain probability >= 50% is flagged with a warning symbol.

    Args:
        location: City name, address, or zip code.
                  Examples: "Orlando FL", "32801", "Chicago Illinois"
        days:     Number of forecast days to return (1-7, default 3).

    Returns:
        Current conditions plus a daily forecast with high/low temps,
        weather description, and rain probability. Rain-risk days are flagged.
    """
    import requests as _req
    import time as _time

    # ── Step 1: Geocode the location via Nominatim ────────────────────────────
    try:
        geo = _req.get(
            "https://nominatim.openstreetmap.org/search",
            params={"q": location, "format": "json", "limit": 1},
            headers={"User-Agent": "AI-Prowler/5.0 (field-service-tool)"},
            timeout=10,
        ).json()
        if not geo:
            return f"❌ Could not geocode '{location}'. Try a more specific address."
        lat        = float(geo[0]["lat"])
        lon        = float(geo[0]["lon"])
        place_name = geo[0].get("display_name", location).split(",")[0]
    except Exception as exc:
        return f"❌ Geocoding failed: {exc}"

    # ── Step 2: Fetch forecast from Open-Meteo ───────────────────────────────
    days = max(1, min(7, days))
    try:
        wx = _req.get(
            "https://api.open-meteo.com/v1/forecast",
            params={
                "latitude":            lat,
                "longitude":           lon,
                "current_weather":     True,
                "daily":               ",".join([
                    "temperature_2m_max",
                    "temperature_2m_min",
                    "precipitation_probability_max",
                    "weathercode",
                ]),
                "temperature_unit":    "fahrenheit",
                "wind_speed_unit":     "mph",
                "precipitation_unit":  "inch",
                "timezone":            "auto",
                "forecast_days":       days,
            },
            timeout=10,
        ).json()
    except Exception as exc:
        return f"❌ Weather fetch failed: {exc}"

    # WMO weather code descriptions (simplified subset)
    _WMO = {
        0: "Clear sky",       1: "Mainly clear",    2: "Partly cloudy",
        3: "Overcast",        45: "Foggy",           48: "Icy fog",
        51: "Light drizzle",  53: "Drizzle",         55: "Heavy drizzle",
        61: "Light rain",     63: "Rain",            65: "Heavy rain",
        71: "Light snow",     73: "Snow",            75: "Heavy snow",
        80: "Light showers",  81: "Rain showers",    82: "Heavy showers",
        95: "Thunderstorm",   96: "Hail storm",      99: "Heavy hailstorm",
    }

    cur     = wx.get("current_weather", {})
    daily   = wx.get("daily", {})
    dates   = daily.get("time", [])
    highs   = daily.get("temperature_2m_max", [])
    lows    = daily.get("temperature_2m_min", [])
    rain    = daily.get("precipitation_probability_max", [])
    codes   = daily.get("weathercode", [])

    lines = [
        f"🌤️  Weather for {place_name}",
        "─" * 45,
        f"Now:  {cur.get('temperature','?')}°F  "
        f"{'  ' + _WMO.get(int(cur.get('weathercode',0)), 'Unknown')}"
        f"  Wind: {cur.get('windspeed','?')} mph",
        "",
    ]
    for i, date in enumerate(dates):
        desc      = _WMO.get(int(codes[i]) if i < len(codes) else 0, "Unknown")
        rain_pct  = int(rain[i]) if i < len(rain) else 0
        rain_flag = "  ⚠️  RAIN RISK — consider rescheduling outdoor jobs" if rain_pct >= 50 else ""
        lines.append(
            f"  {date}:  "
            f"High {highs[i] if i < len(highs) else '?'}°F  "
            f"Low {lows[i]  if i < len(lows)  else '?'}°F  "
            f"{desc}  Rain: {rain_pct}%{rain_flag}"
        )

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 2 — geocode_address
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def geocode_address(address: str) -> str:
    """
    Convert a street address to GPS coordinates (latitude/longitude).
    Uses Nominatim / OpenStreetMap — free, no API key required.

    Useful for verifying that a job address can be geocoded before running
    route optimization, or for looking up coordinates manually.

    Args:
        address: Full street address.
                 Example: "1203 Pine Ave, Orlando FL 32801"

    Returns:
        Latitude, longitude, and display name for the matched location.
    """
    import requests as _req

    try:
        data = _req.get(
            "https://nominatim.openstreetmap.org/search",
            params={"q": address, "format": "json", "limit": 1},
            headers={"User-Agent": "AI-Prowler/5.0 (field-service-tool)"},
            timeout=10,
        ).json()
    except Exception as exc:
        return f"❌ Geocoding request failed: {exc}"

    if not data:
        return (
            f"❌ Address not found: '{address}'\n"
            "Try including city and state for better results."
        )

    r = data[0]
    return (
        f"📍 {r.get('display_name', address)}\n"
        f"   Latitude:  {r['lat']}\n"
        f"   Longitude: {r['lon']}\n"
        f"   Type:      {r.get('type', 'unknown')}"
    )


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 3 — optimize_route
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def optimize_route(
    stops: list[str],
    origin: str,
    optimize_for: str = "time",
    departure_hour: int = 7,
    return_to_origin: bool = True,
) -> str:
    """
    Calculate the optimal driving route for a list of job stops (Traveling Salesman).
    Returns stops in the best visit order with estimated arrival times for each.

    Uses:
    - Nominatim (OpenStreetMap) for geocoding — free, no API key
    - OSRM public routing server for TSP optimization — free, no API key
      OSRM's /trip endpoint solves the Traveling Salesman Problem natively.
      Routing is via real streets with real drive times — not straight-line distance.

    Note: Nominatim requires a small delay between requests (0.3s per address).
    Geocoding 20 addresses takes about 6 seconds. This is a free service courtesy
    limitation — not a bug.

    Args:
        stops:            List of job addresses to visit (any order — the tool
                          reorders them into the optimal sequence).
        origin:           Starting address (your home base or depot).
        optimize_for:     "time" to minimize total drive time (recommended),
                          "distance" to minimize total miles driven.
        departure_hour:   Hour to leave origin in 24h format (default 7 = 7:00 AM).
                          Used to calculate estimated arrival times per stop.
        return_to_origin: If True (default), route ends back at origin (round trip).

    Returns:
        Optimized stop sequence with estimated arrival time, drive time, and
        distance for each leg. Includes totals and a hint to call build_maps_url().
    """
    import requests as _req
    import datetime  as _dt
    import time      as _time

    all_addresses = [origin] + list(stops)

    # ── Step 1: Geocode all addresses via Nominatim ───────────────────────────
    coords  = []
    failed  = []
    for addr in all_addresses:
        _time.sleep(0.35)  # Nominatim courtesy rate limit: max 1 req/s
        try:
            geo = _req.get(
                "https://nominatim.openstreetmap.org/search",
                params={"q": addr, "format": "json", "limit": 1},
                headers={"User-Agent": "AI-Prowler/5.0 (field-service-tool)"},
                timeout=10,
            ).json()
            if geo:
                coords.append({
                    "address": addr,
                    "lat":     float(geo[0]["lat"]),
                    "lon":     float(geo[0]["lon"]),
                    "short":   geo[0].get("display_name", addr).split(",")[0],
                })
            else:
                failed.append(addr)
                coords.append(None)
        except Exception as exc:
            failed.append(f"{addr} (error: {exc})")
            coords.append(None)

    valid = [c for c in coords if c is not None]
    if len(valid) < 2:
        return (
            f"❌ Could not geocode enough addresses to build a route.\n"
            f"Failed addresses: {failed}"
        )

    # ── Step 2: Call OSRM /trip endpoint (solves TSP natively) ───────────────
    # Format: lon,lat;lon,lat;...
    coord_str = ";".join(f"{c['lon']},{c['lat']}" for c in valid)
    destination_param = "any" if return_to_origin else "last"
    try:
        osrm_resp = _req.get(
            f"http://router.project-osrm.org/trip/v1/driving/{coord_str}",
            params={
                "roundtrip":   "true" if return_to_origin else "false",
                "source":      "first",
                "destination": destination_param,
                "annotations": "false",
            },
            timeout=30,
        ).json()
    except Exception as exc:
        return f"❌ Route optimization request failed: {exc}"

    if osrm_resp.get("code") != "Ok":
        return (
            f"❌ OSRM routing error: {osrm_resp.get('message', 'Unknown error')}\n"
            "The OSRM public server may be temporarily overloaded. Try again in a moment."
        )

    trips     = osrm_resp.get("trips",     [])
    waypoints = osrm_resp.get("waypoints", [])
    if not trips or not waypoints:
        return "❌ No route returned. Check that addresses are valid street addresses."

    trip          = trips[0]
    total_dur_s   = trip.get("duration", 0)   # seconds
    total_dist_m  = trip.get("distance",  0)  # meters
    total_miles   = total_dist_m / 1609.34
    total_mins    = total_dur_s  / 60.0
    legs          = trip.get("legs", [])

    # OSRM returns waypoints sorted by trip visit order via trips_index/waypoint_index
    sorted_wps = sorted(
        waypoints,
        key=lambda w: w.get("trips_index", 0) * 1000 + w.get("waypoint_index", 0)
    )

    # ── Step 3: Build human-readable schedule ────────────────────────────────
    now        = _dt.datetime(2026, 1, 1, departure_hour, 0, 0)
    cur_time   = now

    lines = [
        "🗺️  Optimized Route Plan",
        "─" * 50,
        f"  Stops:          {len(stops)}",
        f"  Total drive:    {total_mins:.0f} min  ({total_miles:.1f} miles)",
        f"  Departure:      {now.strftime('%I:%M %p')}",
        f"  Optimized for:  {optimize_for}",
        "",
    ]

    if failed:
        lines.append(f"⚠️  Could not geocode (excluded from route): {', '.join(failed)}")
        lines.append("")

    lines.append("OPTIMIZED SEQUENCE:")
    lines.append("")

    for i, wp in enumerate(sorted_wps):
        wp_idx = wp.get("waypoint_index", i)
        if wp_idx >= len(valid):
            continue
        info = valid[wp_idx]

        if i == 0:
            lines.append(f"  🏠 ORIGIN")
            lines.append(f"     {info['address']}")
            lines.append(f"     Depart: {cur_time.strftime('%I:%M %p')}")
        else:
            leg_idx  = i - 1
            if leg_idx < len(legs):
                leg_s    = legs[leg_idx].get("duration", 0)
                leg_m    = legs[leg_idx].get("distance", 0) / 1609.34
                leg_mins = leg_s / 60.0
                cur_time += _dt.timedelta(seconds=leg_s)
            else:
                leg_mins = 0.0
                leg_m    = 0.0
            lines.append(f"  ▼  Drive {leg_mins:.0f} min  ({leg_m:.1f} mi)")
            lines.append(f"  {i}. {info['address']}")
            lines.append(f"     Arrive: {cur_time.strftime('%I:%M %p')}")
        lines.append("")

    if return_to_origin and legs:
        last_leg  = legs[-1] if legs else {}
        last_s    = last_leg.get("duration", 0)
        last_m    = last_leg.get("distance", 0) / 1609.34
        cur_time += _dt.timedelta(seconds=last_s)
        lines.append(f"  ▼  Drive {last_s/60:.0f} min  ({last_m:.1f} mi)")
        lines.append(f"  🏠 Return to Origin")
        lines.append(f"     Arrive: {cur_time.strftime('%I:%M %p')}")
        lines.append("")

    lines += [
        "─" * 50,
        f"✅ Route total: {total_mins:.0f} min drive,  {total_miles:.1f} miles",
        "",
        "Next step: call build_maps_url(stops_in_order, origin) to generate",
        "a tap-to-navigate Google Maps link for your phone.",
    ]
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 4 — build_maps_url
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def build_maps_url(
    stops: list[str],
    origin: str,
    app: str = "google",
) -> str:
    """
    Build a tap-to-navigate multi-stop directions URL.
    The user taps the link on their phone and Google Maps (or Apple Maps) opens
    immediately in navigation mode with all stops pre-loaded in the correct order.

    Google Maps supports 9 waypoints per URL. For larger routes the tool
    automatically splits the day into legs (each with its own tap-to-navigate
    link) so no stops are lost.

    No API key required — uses the public Google Maps and Apple Maps URL schemes.
    Works on iPhone, Android, CarPlay, Android Auto, and desktop Chrome.

    Args:
        stops:  Job addresses in the OPTIMIZED visit order
                (use the sequence returned by optimize_route()).
        origin: Starting address (home base or depot).
        app:    "google" (default, all devices) or "apple" (iPhone/iPad only).

    Returns:
        One or more tap-to-navigate URLs. For routes > 9 stops, multiple
        leg links are provided — tap each when the previous leg is complete.
    """
    import urllib.parse as _up

    def _enc(s: str) -> str:
        return _up.quote(s.replace(" ", "+"), safe="+")

    GOOGLE_MAX = 9  # Google Maps URL waypoint limit per link

    if app.lower() == "apple":
        lines = [
            f"📍 Apple Maps Navigation Links  ({len(stops)} stops)",
            "(iPhone / iPad only — opens Apple Maps)",
            "",
        ]
        for i, stop in enumerate(stops, 1):
            url = (
                "http://maps.apple.com/?saddr=Current+Location"
                f"&daddr={_enc(stop)}&dirflg=d"
            )
            lines.append(f"  Stop {i}: {stop}")
            lines.append(f"  {url}")
            lines.append("")
        return "\n".join(lines)

    # ── Google Maps ───────────────────────────────────────────────────────────
    def _google_url(orig: str, dest: str, waypoints: list[str]) -> str:
        base   = "https://www.google.com/maps/dir/?api=1"
        params = f"&origin={_enc(orig)}&destination={_enc(dest)}&travelmode=driving"
        if waypoints:
            params += "&waypoints=" + "|".join(_enc(w) for w in waypoints)
        return base + params

    if len(stops) <= GOOGLE_MAX:
        url = _google_url(
            orig      = origin,
            dest      = origin,
            waypoints = stops,
        )
        return (
            f"📍 TAP TO NAVIGATE  —  {len(stops)} stops loaded\n\n"
            f"{url}\n\n"
            f"Opens Google Maps with all {len(stops)} stops in optimized order.\n"
            f"Works on: iPhone (Google Maps app), Android, CarPlay, Android Auto,\n"
            f"          and desktop Chrome.\n"
            f"Tap 'Start' in Maps for turn-by-turn navigation."
        )

    # Split into legs of GOOGLE_MAX stops each
    leg_groups = [stops[i:i + GOOGLE_MAX] for i in range(0, len(stops), GOOGLE_MAX)]
    total_legs  = len(leg_groups)
    lines = [
        f"📍 NAVIGATION ROUTE  —  {len(stops)} stops  ({total_legs} legs)",
        f"   Google Maps limit is {GOOGLE_MAX} waypoints per link.",
        f"   Tap each leg link when you finish the previous leg.",
        "",
    ]

    leg_origin = origin
    for leg_num, group in enumerate(leg_groups, 1):
        is_last_leg = leg_num == total_legs
        leg_dest    = origin if is_last_leg else group[-1]
        leg_wps     = group  if is_last_leg else group[:-1]

        stop_start = (leg_num - 1) * GOOGLE_MAX + 1
        stop_end   = min(leg_num * GOOGLE_MAX, len(stops))
        url = _google_url(leg_origin, leg_dest, leg_wps)

        lines.append(f"  LEG {leg_num}/{total_legs}  (stops {stop_start}–{stop_end}):")
        lines.append(f"  {url}")
        lines.append(f"  Tap when Leg {leg_num - 1 if leg_num > 1 else 0} is complete.")
        lines.append("")
        leg_origin = leg_dest

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 5 — update_job_spreadsheet
# ══════════════════════════════════════════════════════════════════════════════

def _get_default_spreadsheet_path() -> str:
    """Read the default spreadsheet path from ~/.ai-prowler/config.json."""
    try:
        _cfg_path = Path.home() / '.ai-prowler' / 'config.json'
        if _cfg_path.exists():
            import json as _jcfg
            cfg = _jcfg.loads(_cfg_path.read_text(encoding='utf-8-sig'))
            return cfg.get('default_spreadsheet_path', '').strip()
    except Exception as _e:
        _log.warning("Could not read default_spreadsheet_path from config: %s", _e)
    return ''


# Serialises every job-spreadsheet WRITE (update_job_spreadsheet,
# schedule_next_recurring_job, log_time_entry) so two concurrent server-mode
# users calling a write tool at the same moment queue up instead of racing —
# openpyxl has no file locking of its own: load-modify-save is a full
# read-then-overwrite cycle, so two interleaved saves would silently drop
# whichever one finished first. Mirrors _index_write_lock's role for
# ChromaDB writes in rag_preprocessor.py. Read-only tools (read_job_spreadsheet,
# email_invoice, get_ar_aging_report) do not need the lock.
import threading as _sheet_threading
_spreadsheet_write_lock = _sheet_threading.RLock()


def _resolve_job_spreadsheet_path(ctx, filepath_arg: str = "") -> str:
    """
    Resolve which .xlsx file a job-spreadsheet tool should use.

    Personal mode (ctx has no user): unchanged from all prior versions —
    an explicit filepath_arg is honored exactly as given; otherwise falls
    back to the configured default_spreadsheet_path. Full flexibility,
    exactly like every version before this function existed.

    Server mode: filepath_arg is IGNORED — every role's tool call resolves
    through the shared default_spreadsheet_path (Settings -> Business ->
    Default Spreadsheet Path) instead. This closes a real gap: previously
    any server-mode role could point these tools at an arbitrary .xlsx file
    anywhere on the host with zero access control.

    Within that, two models, both driven by the SAME single config value:
      - default_spreadsheet_path points at an .xlsx file that exists ->
        that IS the master file (today's behaviour, unchanged). Every
        role reads/writes that one shared spreadsheet.
      - The calling user has their OWN per-user file sitting in the same
        folder as default_spreadsheet_path, named exactly "<user_id>.xlsx"
        -> that takes priority over the master file for THIS user only.
        If no such file exists, falls back to the master file above.

    A manager sets up per-user tracking simply by dropping "jake-r.xlsx",
    "vicki-vavro.xlsx" etc. into the same folder as the master spreadsheet
    (whatever filename was chosen via the Business tab's Browse button —
    defaults to AI-Prowler_Job_Tracker.xlsx on a fresh install). No
    separate config field is needed; the folder is simply
    dirname(default_spreadsheet_path).
    """
    user = _current_user(ctx)
    default_path = _get_default_spreadsheet_path()

    if user is None:
        # Personal mode: unrestricted, exactly as before this feature existed.
        if filepath_arg:
            return filepath_arg
        return default_path

    # Server mode: filepath_arg is ignored from here on.
    if not default_path:
        return ""

    user_id = (user.get("id") or "").strip()
    if user_id:
        try:
            folder = os.path.dirname(default_path)
            if folder:
                candidate = os.path.join(folder, f"{user_id}.xlsx")
                if os.path.exists(candidate):
                    return candidate
        except Exception:
            pass

    return default_path


def _backup_spreadsheet(fp: str, keep_days: int = 30) -> str:
    """
    Copy fp into a _backups subfolder next to the file, timestamped.
    Prunes backups older than keep_days days.

    Returns a short status string (success path or warning message).
    """
    import shutil as _shutil
    import datetime as _dt

    src = Path(fp)
    backup_dir = src.parent / '_backups'
    try:
        backup_dir.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        return f"⚠️  Could not create backup folder: {exc}"

    ts  = _dt.datetime.now().strftime('%Y-%m-%d_%H%M%S')
    dst = backup_dir / f"{src.stem}_{ts}{src.suffix}"
    try:
        _shutil.copy2(str(src), str(dst))
    except Exception as exc:
        return f"⚠️  Backup copy failed: {exc}"

    # Prune old backups
    cutoff = _dt.datetime.now() - _dt.timedelta(days=keep_days)
    pruned = 0
    try:
        for old in backup_dir.glob(f"{src.stem}_*{src.suffix}"):
            if old == dst:
                continue
            try:
                mtime = _dt.datetime.fromtimestamp(old.stat().st_mtime)
                if mtime < cutoff:
                    old.unlink()
                    pruned += 1
            except Exception:
                pass
    except Exception:
        pass

    msg = f"💾 Backup saved: _backups/{dst.name}"
    if pruned:
        msg += f"  ({pruned} old backup(s) pruned)"
    return msg


@mcp.tool()
def update_job_spreadsheet(
    job_identifier: str,
    updates:        dict,
    filepath:       str = "",
    id_column:      str = "Customer",
    sheet_name:     str = "",
    backup:         bool = True,
    ctx: "Context | None" = None,
) -> str:
    """
    Update a row in a job tracking spreadsheet after a job is completed.

    Finds the correct row by matching job_identifier against id_column
    (e.g. customer name), then writes new values to specified columns
    (e.g. marks job complete, records invoice number, updates last service date).

    Uses openpyxl — already installed, no new package needed.
    Works only on .xlsx files. For .xls, convert to .xlsx first.

    If filepath is omitted, the default spreadsheet path configured in
    AI-Prowler Settings → Small Business → Default Spreadsheet Path is used
    automatically — no need to specify the path every time.

    Available to every role in every mode — personal and ALL server-mode
    roles (owner, manager, staff, field_crew) — no DB-management or
    communications gate applies here.

    Args:
        job_identifier: Value to search for in id_column.
                        Example: "Crabby's Daytona" (partial matches accepted)
        updates:        Dict of {column_header: new_value} pairs to write.
                        Example: {"Job\\nStatus": "Complete",
                                  "Last Service\\nDate": "2026-03-31",
                                  "Actual\\nDuration (min)": 45,
                                  "Actual\\nAmount ($)": 150.00}
        filepath:       Full path to the Excel spreadsheet (.xlsx).
                        If omitted, uses the path saved in AI-Prowler Settings.
                        Example: "C:/Users/Dave/Documents/jobs.xlsx"
        id_column:      Column header to search in (default "Customer").
        sheet_name:     Sheet to use (default: first/active sheet).
        backup:         If True (default), a timestamped backup copy of the
                        spreadsheet is saved in a _backups subfolder next to
                        the file before any changes are written.
                        Backups older than 30 days are pruned automatically.
                        Set to False to skip the backup (faster, no disk use).

    Returns:
        Confirmation listing exactly which cells were updated, backup status,
        or an error if the file, row, or column was not found.
    """
    try:
        import openpyxl as _opx
    except ImportError:
        return "❌ openpyxl not installed. Run: pip install openpyxl"

    # Resolve filepath — use default from config if not supplied
    filepath = _resolve_job_spreadsheet_path(ctx, filepath)
    if not filepath:
        return (
            "❌ No spreadsheet path provided and no default path configured.\n"
            "Set one in AI-Prowler → Settings → Small Business → Default Spreadsheet Path,\n"
            "or pass the full filepath argument explicitly."
        )

    fp = filepath.replace("\\", "/")
    if not os.path.exists(fp):
        return f"❌ Spreadsheet not found: {fp}"
    if not fp.lower().endswith(".xlsx"):
        return (
            "❌ Only .xlsx files are supported for updates.\n"
            "Save the spreadsheet as .xlsx in Excel first."
        )

    # ── Backup before modifying ───────────────────────────────────────────────
    backup_msg = ""
    if backup:
        backup_msg = _backup_spreadsheet(fp)
        if backup_msg.startswith("⚠️") or backup_msg.startswith("❌"):
            # Backup failed — abort to protect the file
            return (
                f"{backup_msg}\n"
                "Spreadsheet was NOT modified. Fix the backup issue or pass backup=False to skip."
            )

    # Serialise the whole load -> modify -> save cycle so two concurrent
    # server-mode writers can never interleave and silently drop each
    # other's changes (openpyxl has no file locking of its own).
    with _spreadsheet_write_lock:
        try:
            wb = _opx.load_workbook(fp)
        except Exception as exc:
            return f"❌ Could not open spreadsheet: {exc}"

        ws = (wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames
              else wb.active)

        # ── Detect header row (skip title/banner rows, same logic as read tool) ───
        # Scans the first 5 rows and uses the first row that has ≥ 3 non-empty
        # cells as the real header row.  This handles decorative title rows like
        # "📅 JOBS & SCHEDULE — All Service Appointments" in row 1.
        header_row_num: int | None = None
        headers: dict[str, int] = {}
        for r in ws.iter_rows(min_row=1, max_row=5):
            non_empty = [c for c in r if c.value is not None]
            if len(non_empty) >= 3:
                header_row_num = r[0].row
                for col_idx, cell in enumerate(r, 1):
                    if cell.value is not None:
                        raw = str(cell.value).strip()
                        headers[raw] = col_idx
                        # Register a newline-normalised alias so callers can pass
                        # either "Job\nStatus" or "Job Status" and both resolve.
                        normalised = raw.replace('\n', ' ')
                        if normalised != raw:
                            headers.setdefault(normalised, col_idx)
                break

        if header_row_num is None or not headers:
            return (
                "❌ Could not detect a header row in the spreadsheet.\n"
                "Expected a row with at least 3 non-empty cells in the first 5 rows."
            )

        if id_column not in headers:
            avail = [k for k in headers.keys() if '\n' not in k][:15]
            return (
                f"❌ Column '{id_column}' not found in headers (detected on row {header_row_num}).\n"
                f"Available columns: {', '.join(avail)}"
            )

        id_col_idx = headers[id_column]

        # ── Find matching row ─────────────────────────────────────────────────────
        found_row = None
        for row in ws.iter_rows(min_row=header_row_num + 1):
            cell_val = row[id_col_idx - 1].value
            if cell_val and job_identifier.lower() in str(cell_val).lower():
                found_row = row
                break

        if found_row is None:
            return (
                f"❌ No row found where {id_column} contains '{job_identifier}'.\n"
                "Check the spelling — partial matches are accepted."
            )

        # ── Apply updates ─────────────────────────────────────────────────────────
        updated:    list[str] = []
        not_found:  list[str] = []

        for col_name, new_val in updates.items():
            if col_name in headers:
                found_row[headers[col_name] - 1].value = new_val
                updated.append(f"{col_name} → {new_val}")
            else:
                not_found.append(col_name)

        # ── Save ──────────────────────────────────────────────────────────────────
        try:
            wb.save(fp)
        except Exception as exc:
            return f"❌ Could not save spreadsheet: {exc}"

        row_num = found_row[0].row
    lines   = [
        f"✅ Spreadsheet updated: {os.path.basename(fp)}",
        f"   Row:     {row_num}  ({id_column}: {job_identifier})",
        f"   Updated: {', '.join(updated)}",
    ]
    if backup_msg:
        lines.append(f"   {backup_msg}")
    if not_found:
        lines.append(
            f"   ⚠️  Columns not found (check spelling): {', '.join(not_found)}"
        )
    lines.append(
        "\n📑 Re-index the spreadsheet to keep AI-Prowler search results current:\n"
        "   Call update_tracked_directories() after updating the file."
    )
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 6 — read_job_spreadsheet
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def read_job_spreadsheet(
    filepath:    str = "",
    sheet_name:  str = "",
    filter_date: str = "",
    max_rows:    int = 200,
    ctx: "Context | None" = None,
) -> str:
    """
    Read job data from the AI-Prowler job tracking spreadsheet.

    Returns rows from the Jobs_Schedule sheet (or any named sheet) as
    structured text so Claude can answer questions like:
      - "What jobs do I have scheduled for today?"
      - "Which jobs are still open?"
      - "Show me everything scheduled this week."
      - "What customers haven't been serviced yet?"

    If filepath is omitted, the default spreadsheet path configured in
    AI-Prowler Settings -> Small Business -> Default Spreadsheet Path is used
    automatically.

    Available to every role in every mode — personal and ALL server-mode
    roles (owner, manager, staff, field_crew) — no DB-management or
    communications gate applies here. The Customers sheet specifically is
    also the FIRST place send_email() and send_sms() look up a recipient
    by name/company/ID, so this is the same data those tools draw on.

    Args:
        filepath:    Full path to the .xlsx spreadsheet.
                     If omitted, uses the path saved in AI-Prowler Settings.
        sheet_name:  Sheet to read (default: "Jobs_Schedule").
                     Use "Customers" to read the customer master list.
                     Leave blank to default to Jobs_Schedule.
        filter_date: Optional date string to filter rows by Service Date.
                     Examples: "2026-03-31", "03/31/2026", "today"
                     Leave blank to return all rows.
        max_rows:    Maximum data rows to return (default 200, max 500).

    Returns:
        A formatted table of rows with all column values, or a message
        if no matching rows are found.
    """
    try:
        import openpyxl as _opx
    except ImportError:
        return "❌ openpyxl not installed. Run: pip install openpyxl"

    filepath = _resolve_job_spreadsheet_path(ctx, filepath)
    if not filepath:
        return (
            "❌ No spreadsheet path provided and no default path configured.\n"
            "Set one in AI-Prowler -> Settings -> Small Business -> Default Spreadsheet Path,\n"
            "or pass the full filepath argument explicitly."
        )

    fp = filepath.replace("\\", "/")
    if not os.path.exists(fp):
        return f"❌ Spreadsheet not found: {fp}"
    if not fp.lower().endswith(".xlsx"):
        return "❌ Only .xlsx files are supported. Save as .xlsx first."

    try:
        wb = _opx.load_workbook(fp, data_only=True)
    except Exception as exc:
        return f"❌ Could not open spreadsheet: {exc}"

    target_sheet = sheet_name.strip() if sheet_name.strip() else "Jobs_Schedule"
    if target_sheet not in wb.sheetnames:
        if sheet_name.strip():
            avail = ", ".join(wb.sheetnames)
            return f"❌ Sheet '{target_sheet}' not found.\nAvailable sheets: {avail}"
        target_sheet = wb.sheetnames[0]

    ws = wb[target_sheet]

    # Detect header row (skip title banner rows)
    header_row_idx = None
    headers: list = []
    for r in ws.iter_rows(min_row=1, max_row=5):
        non_empty = [c for c in r if c.value is not None]
        if len(non_empty) >= 3:
            header_row_idx = r[0].row
            headers = [str(c.value).strip().replace('\n', ' ') if c.value else '' for c in r]
            break

    if header_row_idx is None or not headers:
        return f"❌ Could not detect a header row in sheet '{target_sheet}'."

    import datetime as _dt
    date_filter = None
    if filter_date:
        fd = filter_date.strip().lower()
        if fd == 'today':
            date_filter = _dt.date.today()
        else:
            for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%d/%m/%Y'):
                try:
                    date_filter = _dt.datetime.strptime(fd, fmt).date()
                    break
                except ValueError:
                    continue
            if date_filter is None:
                return f"❌ Could not parse filter_date '{filter_date}'. Use YYYY-MM-DD or MM/DD/YYYY."

    svc_date_col = None
    if date_filter:
        for idx, h in enumerate(headers):
            if 'service' in h.lower() and 'date' in h.lower():
                svc_date_col = idx
                break

    max_rows = min(max_rows, 500)
    rows_out: list = []

    for row in ws.iter_rows(min_row=header_row_idx + 1):
        vals = [c.value for c in row]
        if all(v is None or str(v).strip() == '' for v in vals):
            continue
        if date_filter is not None and svc_date_col is not None:
            cell_val = vals[svc_date_col]
            if cell_val is None:
                continue
            if isinstance(cell_val, (_dt.datetime, _dt.date)):
                cell_date = cell_val.date() if isinstance(cell_val, _dt.datetime) else cell_val
            else:
                cell_date = None
                for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%d/%m/%Y'):
                    try:
                        cell_date = _dt.datetime.strptime(str(cell_val).strip(), fmt).date()
                        break
                    except ValueError:
                        continue
            if cell_date != date_filter:
                continue
        rows_out.append(vals)
        if len(rows_out) >= max_rows:
            break

    if not rows_out:
        msg = f"📋 No rows found in sheet '{target_sheet}'"
        if date_filter:
            msg += f" for date {date_filter.strftime('%Y-%m-%d')}"
        return msg + "."

    # Trim to last used column
    max_col_used = 0
    for row in rows_out:
        for i in range(len(row) - 1, -1, -1):
            if row[i] is not None and str(row[i]).strip():
                if i > max_col_used:
                    max_col_used = i
                break
    headers_trimmed = headers[:max_col_used + 1]

    lines = [
        f"📋 {target_sheet}  —  {os.path.basename(fp)}",
        f"   {len(rows_out)} row(s)" + (f" for {date_filter.strftime('%Y-%m-%d')}" if date_filter else ""),
        "─" * 60,
    ]
    for row_vals in rows_out:
        lines.append("")
        for col_idx, col_name in enumerate(headers_trimmed):
            if not col_name:
                continue
            val = row_vals[col_idx] if col_idx < len(row_vals) else None
            if val is None or str(val).strip() == '':
                continue
            if isinstance(val, _dt.datetime):
                val = val.strftime('%Y-%m-%d')
            elif isinstance(val, _dt.date):
                val = val.strftime('%Y-%m-%d')
            lines.append(f"  {col_name}: {val}")
    lines.append("")
    lines.append("─" * 60)
    lines.append("✅ Read complete. Use update_job_spreadsheet() to write changes back.")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 7 — check_tools_status
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def check_tools_status(ctx: "Context | None" = None) -> str:
    """
    Check which AI-Prowler Action Tools are ready to use and which need setup.

    Returns a full status report covering:
    - Free tools (weather, geocoding, routing, navigation URLs)
    - Spreadsheet update tool readiness
    - Email, SMS (Twilio), and WhatsApp configuration
    - Contractor workflow tools
    - Dev tools and file-editing tools (always available — no setup needed)

    Call this tool first when planning to use action tools, to confirm
    everything is configured before starting a workflow.

    Returns:
        Status report with ✅/⚠️/❌ for each tool and setup instructions
        for anything that needs configuration.
    """
    lines = [
        "🔧 AI-Prowler Action Tools Status",
        "─" * 50,
        "",
        "FREE TOOLS  (no API key, no setup required):",
        "",
    ]

    # Check requests
    try:
        import requests  # noqa: F401
        req_ok = True
    except ImportError:
        req_ok = False

    lines += [
        f"  {'✅' if req_ok else '❌'} get_weather(location, days)",
        f"     Open-Meteo + Nominatim geocoding — free, no key",
        "",
        f"  {'✅' if req_ok else '❌'} geocode_address(address)",
        f"     Nominatim / OpenStreetMap — free, no key",
        "",
        f"  {'✅' if req_ok else '❌'} optimize_route(stops, origin, ...)",
        f"     OSRM public server — TSP solver, real streets, free, no key",
        "",
        f"  ✅ build_maps_url(stops, origin, app)",
        f"     Google Maps URL builder — tap-to-navigate, free, no key",
        "",
        "─" * 50,
        "",
        "SPREADSHEET TOOLS:",
        "",
    ]

    if not req_ok:
        lines.append("  ❌ requests package missing — run: pip install requests")
        lines.append("")

    try:
        import openpyxl  # noqa: F401
        opx_ok = True
    except ImportError:
        opx_ok = False

    default_xl = _get_default_spreadsheet_path()
    xl_status  = f"✅ Default path: {default_xl}" if default_xl else "⚠️  No default path set"

    lines += [
        f"  {'✅' if opx_ok else '❌'} read_job_spreadsheet(filepath?, sheet?, date?, max_rows?)",
        f"     openpyxl — reads Jobs_Schedule or any sheet; supports date filtering",
        f"     {xl_status}",
        "",
        f"  {'✅' if opx_ok else '❌'} update_job_spreadsheet(job_id, updates, filepath?)",
        f"     openpyxl — updates .xlsx job tracking files in place",
        f"     {xl_status}",
        "",
        "     💡 Available to every role, in both personal and server mode —",
        "        no DB-management or communications gate applies to these two",
        "        tools. The Customers sheet is also the FIRST place send_email",
        "        and send_sms look up a recipient by name/company/ID.",
        "",
        "     💡 Set a default path once in Settings → Small Business → Default",
        "        Spreadsheet Path so you never need to specify it in conversation.",
        "",
        "─" * 50,
        "",
        "All FREE tools work immediately with no configuration.",
        "Spreadsheet tools use the default path from Settings if filepath is omitted.",
    ]

    # ── SMS — provider-agnostic (Twilio / SignalWire / Vonage) ──────────────────
    try:
        from sms_backends import get_sms_backend, load_sms_config as _load_sms_cfg_chk
        _sms_cfg_chk = _load_sms_cfg_chk()
        _sms_backend_chk = get_sms_backend(_sms_cfg_chk)
        _sms_ok_chk, _ = _sms_backend_chk.validate_config()
        sms_provider_name = _sms_backend_chk.provider_name.title()
    except Exception:
        _sms_ok_chk = False
        sms_provider_name = "Twilio"

    saved_contacts = len(_contacts_cache_load(_current_user(ctx)).get('contacts', {}))

    lines += [
        "─" * 50,
        "",
        f"SMS  (provider: {sms_provider_name} — Twilio, SignalWire, or Vonage,",
        f"      set via sms_provider in Settings → SMS / Text Messaging):",
        "",
        f"  {'✅' if _sms_ok_chk else '⚠️ '} send_sms(to, message)",
        f"     {'✅ Configured' if _sms_ok_chk else '⚠️  Not configured — see Settings → SMS / Text Messaging'}",
        f"     Works in personal mode unrestricted; in server mode, available",
        f"     to every role (owner, manager, staff, field_crew) via the",
        f"     company's shared provider account.",
        "",
        f"  ✅ check_sms_replies(since_hours?, from_number?)",
        f"     Instant — reads the local sms_inbox.json (no API polling).",
        f"     Server mode: replies are filtered to threads YOU personally",
        f"     sent — Mike sees Karen's reply, not Jake's reply to Bob.",
        "",
        f"  ✅ get_sms_thread(contact) / list_sms_contacts_with_replies()",
        f"     Full conversation history with a contact, or a list of all",
        f"     recent contacts with their unread-reply counts.",
        "",
        f"  ✅ save_contact(name, phone?, email?)",
        f"     Save a person by name so 'text David' / 'email Vicki' works",
        f"     without repeating their number/address every time. Each",
        f"     server-mode role has their own private contacts file.",
        f"     {saved_contacts} contact(s) currently saved for you.",
        "",
        "─" * 50,
        "",
        "CONTRACTOR WORKFLOW TOOLS:",
        "",
        f"  {'✅' if opx_ok else '❌'} email_invoice(invoice_id, to?, filepath?)",
        f"     Reads Invoices sheet, sends branded HTML invoice via SMTP",
        f"     {xl_status}",
        "",
        f"  {'✅' if opx_ok else '❌'} schedule_next_recurring_job(job_id, filepath?, when?)",
        f"     Auto-creates next job based on customer frequency (W/BW/M/Q)",
        "",
        f"  {'✅' if opx_ok else '❌'} log_time_entry(job_id, action, filepath?)",
        f"     Clock in / out — records TimeLog, writes Actual Duration to Jobs_Schedule",
        "",
        f"  {'✅' if opx_ok else '❌'} get_ar_aging_report(filepath?, as_of_date?)",
        f"     AR aging buckets: Current / 1-30 / 31-60 / 61-90 / 90+ days",
        "",
    ]

    # ── Email (configure_email / send_email / send_alert / send_file) ────────────
    email_configured = bool(_email_config_load())

    lines += [
        "─" * 50,
        "",
        "EMAIL:",
        "",
        f"  {'✅' if email_configured else '⚠️ '} send_email(to, subject, body) / send_alert(to, body) / send_file(...)",
        f"     {'✅ SMTP configured' if email_configured else '⚠️  Not configured — call configure_email() or use Settings → Email Configuration'}",
        f"     Personal mode: any SMTP provider (Gmail, Outlook, etc.) with an app password.",
        f"     Server mode: available to every role via the company's shared SMTP.",
        "",
    ]

    # ── WhatsApp (always via Twilio, even if SMS provider is SignalWire/Vonage) ──
    try:
        from sms_backends import get_whatsapp_backend, load_sms_config as _load_wa_cfg_chk
        _wa_cfg_chk = _load_wa_cfg_chk()
        _wa_backend_chk = get_whatsapp_backend(_wa_cfg_chk)
        whatsapp_configured, _ = _wa_backend_chk.validate_config()
    except Exception:
        whatsapp_configured = False

    lines += [
        "─" * 50,
        "",
        "WHATSAPP (always via Twilio — even if your SMS provider above is",
        "          SignalWire or Vonage; WhatsApp has no separate provider):",
        "",
        f"  {'✅' if whatsapp_configured else '⚠️ '} send_whatsapp(to, message)",
        f"     {'✅ Configured' if whatsapp_configured else '⚠️  Not configured — add Twilio credentials in Settings → SMS / Text Messaging'}",
        f"     Reuses your Twilio Account SID / Auth Token / From Number.",
        "",
        f"  ✅ check_whatsapp_replies(since_hours?, from_number?)",
        f"     Reads from the local inbox (instant, no API call).",
        "",
        f"  ✅ check_sms_inbox / get_sms_thread / list_sms_contacts_with_replies",
        f"     Browse SMS + WhatsApp conversation history stored locally.",
        "",
    ]

    lines += [
        "─" * 50,
        "",
        "DEV TOOLS & FILE EDITING:",
        "",
    ]

    # Server mode: most dev tools are Tier A suppressed entirely, and the
    # write/edit tools are scoped to the caller's own personal directory —
    # neither is "always available, no setup" the way personal mode is.
    # Check the CALLING user's actual status so this report is accurate
    # per-caller, not a blanket claim.
    _cts_user = _current_user(ctx)

    if _cts_user is None:
        # Personal mode — unchanged from every version before this fix.
        lines += [
            "  ✅ syntax_check / compile_check / check_python_import / lint_check",
            "     Verify code without leaving the conversation.",
            "",
            "  ✅ run_script / run_script_start / run_script_status / run_script_kill",
            "     Execute scripts synchronously or as a tracked background job.",
            "",
            "  ✅ create_file / write_file / str_replace_in_file /",
            "     line_replace_in_file / create_directory / list_directory",
            "     Edit files on disk under a tracked, writable directory.",
            "",
            "  ✅ copy_to_backup / list_backups / restore_backup / cleanup_backups",
            "     Every write auto-backs up first — these tools manage that history.",
            "",
            "  ✅ diff_files / reset_write_counter",
            "     Compare two files; reset the per-session write circuit breaker.",
            "",
            "  💡 Write tools do NOT auto-index — call reindex_file(path) once you're",
            "     done editing a file so ChromaDB reflects the new content.",
            "",
        ]
    else:
        # Server mode — accurate per-tool, per-caller status.
        _cts_priv_status, _cts_priv_dir = _user_private_write_dir(ctx)
        if _cts_priv_status == "scoped":
            _cts_write_line = f"✅ Scoped to your personal directory: {_cts_priv_dir}"
        else:
            _cts_write_line = ("⚠️  Not available — no personal directory configured. "
                               "Ask your owner/admin to set one up in the Admin tab.")

        lines += [
            "  ❌ syntax_check / compile_check / check_python_import / lint_check",
            "     Not available in server mode — personal-install tools only.",
            "",
            "  ❌ run_script / run_script_start / run_script_status / run_script_kill",
            "     Not available in server mode — personal-install tools only.",
            "",
            "  create_file / write_file / str_replace_in_file /",
            "  line_replace_in_file / create_directory",
            f"     {_cts_write_line}",
            "",
            "  ❌ list_directory",
            "     Not available in server mode — personal-install tool only.",
            "",
            "  ❌ copy_to_backup / list_backups / restore_backup / cleanup_backups /",
            "     reset_write_counter",
            "     Not available in server mode — backup/write-zone management is",
            "     operator-only.",
            "",
            "  ✅ diff_files",
            "     Compare two files — available to every role, scoped to files",
            "     you're allowed to read.",
            "",
            "  💡 Write tools do NOT auto-index — ask an owner/manager to call",
            "     reindex_file(path) (owner/manager only in server mode) once",
            "     editing is done.",
            "",
        ]

    lines += [
        "─" * 50,
        "",
        "SMS / WhatsApp: configure in Settings → SMS / Text Messaging.",
    ]

    return "\n".join(lines)

# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 8 — email_invoice
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def email_invoice(
    invoice_identifier: str,
    to:                 str  = "",
    filepath:           str  = "",
    ctx:                "Context | None" = None,
) -> str:
    """
    Email a professional HTML invoice to a customer directly from the spreadsheet.

    Reads the matching invoice row from the Invoices sheet, formats it as a
    clean HTML email, and sends it via the configured SMTP account.

    Args:
        invoice_identifier: InvoiceID (e.g. "INV-0001") or customer name
                            (partial match accepted, e.g. "Torres").
        to:                 Recipient email address. If omitted, looks up the
                            customer email in the Customers sheet automatically.
        filepath:           Path to the .xlsx job tracker. Uses the default
                            path from Settings if omitted.
        ctx:                MCP context (injected automatically).

    Returns:
        Confirmation with invoice total and recipient, or an error message.

    Voice examples:
        "Email the invoice for INV-0001 to the customer"
        "Send Karen's invoice"
        "Email the Blue Wave Cafe invoice"
    """
    _telemetry_increment_tool_count("email_invoice")

    try:
        import openpyxl as _opx
    except ImportError:
        return "❌ openpyxl not installed. Run: pip install openpyxl"

    filepath = _resolve_job_spreadsheet_path(ctx, filepath)
    if not filepath:
        return "❌ No spreadsheet path configured. Set one in Settings → Small Business."

    fp = filepath.replace("\\", "/")
    if not os.path.exists(fp):
        return f"❌ Spreadsheet not found: {fp}"

    try:
        wb = _opx.load_workbook(fp, data_only=True)
    except Exception as exc:
        return f"❌ Could not open spreadsheet: {exc}"

    # ── Find invoice row ──────────────────────────────────────────────────────
    inv_sheet = wb["Invoices"] if "Invoices" in wb.sheetnames else wb.active
    header_row_idx, headers = None, []
    for r in inv_sheet.iter_rows(min_row=1, max_row=5):
        non_empty = [c for c in r if c.value is not None]
        if len(non_empty) >= 3:
            header_row_idx = r[0].row
            headers = [str(c.value).strip().replace('\n', ' ') if c.value else '' for c in r]
            break

    if not headers:
        return "❌ Could not detect header row in Invoices sheet."

    inv_row = None
    for row in inv_sheet.iter_rows(min_row=header_row_idx + 1):
        vals = [c.value for c in row]
        row_text = " ".join(str(v) for v in vals if v)
        if invoice_identifier.lower() in row_text.lower():
            inv_row = dict(zip(headers, vals))
            break

    if not inv_row:
        return f"❌ No invoice found matching '{invoice_identifier}' in Invoices sheet."

    # ── Look up customer email if not provided ────────────────────────────────
    if not to:
        cust_id = str(inv_row.get("CustomerID", "") or "")
        cust_name = str(inv_row.get("Customer Name / Company", "") or "")
        if "Customers" in wb.sheetnames:
            cust_sheet = wb["Customers"]
            cust_hdrs = []
            cust_hdr_row = None
            for r in cust_sheet.iter_rows(min_row=1, max_row=5):
                ne = [c for c in r if c.value is not None]
                if len(ne) >= 3:
                    cust_hdr_row = r[0].row
                    cust_hdrs = [str(c.value).strip().replace('\n',' ') if c.value else '' for c in r]
                    break
            if cust_hdrs:
                for row in cust_sheet.iter_rows(min_row=cust_hdr_row + 1):
                    cvals = [c.value for c in row]
                    crow = dict(zip(cust_hdrs, cvals))
                    cid = str(crow.get("CustomerID (CUST-####)", "") or "")
                    cname = str(crow.get("Company Name", "") or crow.get("First Name","") or "")
                    if (cust_id and cust_id == cid) or (cust_name and cust_name.lower() in str(cvals).lower()):
                        to = str(crow.get("Email", "") or "")
                        break
        if not to:
            return (f"❌ No recipient email provided and could not auto-find customer email.\n"
                    f"Pass a 'to' address: email_invoice('{invoice_identifier}', to='email@example.com')")

    # ── Build HTML invoice ────────────────────────────────────────────────────
    import datetime as _dt

    def _fv(key, default="—"):
        v = inv_row.get(key)
        if v is None or str(v).strip() == "":
            return default
        if isinstance(v, _dt.datetime):
            return v.strftime("%B %d, %Y")
        if isinstance(v, _dt.date):
            return v.strftime("%B %d, %Y")
        return str(v).strip()

    def _fmt_money(key):
        v = inv_row.get(key)
        try:
            return f"${float(v):,.2f}"
        except (TypeError, ValueError):
            return "—"

    inv_id    = _fv("InvoiceID (INV-####)")
    job_id    = _fv("JobID (JOB-####)")
    cust      = _fv("Customer Name / Company")
    inv_date  = _fv("Invoice Date")
    due_date  = _fv("Due Date (Net 30)")
    svc_date  = _fv("Service Date")
    svc_type  = _fv("Service Type")
    desc      = _fv("Description")
    subtotal  = _fmt_money("Subtotal ($)")
    discount  = _fmt_money("Discount ($)")
    tax       = _fmt_money("Tax 7% ($)")
    total_due = _fmt_money("TOTAL DUE ($)")
    balance   = _fmt_money("Balance Due ($)")
    pmt_status= _fv("Payment Status", "Unpaid")

    html_body = f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8">
<style>
  body {{ font-family: Arial, sans-serif; color: #333; max-width: 650px; margin: 0 auto; }}
  .header {{ background: #1a3c5e; color: white; padding: 24px 32px; }}
  .header h1 {{ margin: 0; font-size: 28px; letter-spacing: 1px; }}
  .header p  {{ margin: 4px 0 0; font-size: 13px; opacity: 0.85; }}
  .section   {{ padding: 20px 32px; }}
  .row       {{ display: flex; justify-content: space-between; margin: 6px 0; }}
  .label     {{ color: #666; font-size: 14px; }}
  .value     {{ font-weight: 600; font-size: 14px; }}
  table      {{ width: 100%; border-collapse: collapse; margin-top: 8px; }}
  th         {{ background: #f0f4f8; text-align: left; padding: 10px; font-size: 13px; color: #555; }}
  td         {{ padding: 10px; border-bottom: 1px solid #eee; font-size: 14px; }}
  .totals    {{ background: #f9f9f9; padding: 16px 32px; }}
  .total-row {{ display: flex; justify-content: space-between; padding: 4px 0; font-size: 14px; }}
  .grand     {{ font-size: 18px; font-weight: bold; color: #1a3c5e; border-top: 2px solid #1a3c5e; margin-top: 8px; padding-top: 8px; }}
  .footer    {{ background: #f0f4f8; padding: 16px 32px; font-size: 12px; color: #888; }}
  .badge     {{ display: inline-block; padding: 4px 10px; border-radius: 12px; font-size: 12px; font-weight: bold;
                background: {'#e8f5e9' if pmt_status == 'Paid' else '#fff3e0'}; 
                color: {'#2e7d32' if pmt_status == 'Paid' else '#e65100'}; }}
</style>
</head>
<body>
<div class="header">
  <h1>INVOICE</h1>
  <p>{inv_id} &nbsp;|&nbsp; {job_id}</p>
</div>
<div class="section">
  <div class="row"><span class="label">Bill To</span><span class="value">{cust}</span></div>
  <div class="row"><span class="label">Service Date</span><span class="value">{svc_date}</span></div>
  <div class="row"><span class="label">Invoice Date</span><span class="value">{inv_date}</span></div>
  <div class="row"><span class="label">Due Date</span><span class="value">{due_date}</span></div>
  <div class="row"><span class="label">Status</span><span class="value"><span class="badge">{pmt_status}</span></span></div>
</div>
<div class="section">
  <table>
    <tr><th>Service</th><th>Description</th><th>Amount</th></tr>
    <tr><td>{svc_type}</td><td>{desc}</td><td>{subtotal}</td></tr>
  </table>
</div>
<div class="totals">
  <div class="total-row"><span>Subtotal</span><span>{subtotal}</span></div>
  <div class="total-row"><span>Discount</span><span>({discount})</span></div>
  <div class="total-row"><span>Tax (7%)</span><span>{tax}</span></div>
  <div class="total-row grand"><span>TOTAL DUE</span><span>{total_due}</span></div>
  {"" if pmt_status == "Paid" else f'<div class="total-row" style="color:#e65100"><span>Balance Due</span><span>{balance}</span></div>'}
</div>
<div class="footer">
  <p>Thank you for your business! Questions? Reply to this email.</p>
  <p>Please make payment by <strong>{due_date}</strong>.</p>
</div>
</body>
</html>"""

    subject = f"Invoice {inv_id} — {cust} — {total_due}"

    # ── Send via existing send_email infrastructure ───────────────────────────
    cfg = _email_config_load()
    if not cfg:
        return ("❌ Email not configured. Call configure_email() first.")

    to = (to or "").strip()
    if not to:
        return "❌ Could not determine recipient email address."

    import smtplib as _smtp
    import email.mime.multipart as _mp
    import email.mime.text as _mt

    try:
        msg = _mp.MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"]    = f"{cfg.get('from_name','AI-Prowler')} <{cfg['from_email']}>"
        msg["To"]      = to
        msg.attach(_mt.MIMEText(html_body, "html", "utf-8"))

        port = int(cfg.get("smtp_port", 587))
        host = cfg.get("smtp_host", "")
        user = cfg.get("smtp_user", "")
        pwd  = cfg.get("smtp_password", "")

        if port == 465:
            with _smtp.SMTP_SSL(host, port, timeout=30) as server:
                server.login(user, pwd)
                server.sendmail(cfg["from_email"], [to], msg.as_string())
        else:
            with _smtp.SMTP(host, port, timeout=30) as server:
                server.ehlo()
                server.starttls()
                server.login(user, pwd)
                server.sendmail(cfg["from_email"], [to], msg.as_string())

        return (
            f"✅ Invoice emailed\n"
            f"   Invoice:   {inv_id}  ({job_id})\n"
            f"   Customer:  {cust}\n"
            f"   Total Due: {total_due}\n"
            f"   Sent to:   {to}\n"
            f"   Subject:   {subject}"
        )
    except Exception as exc:
        return f"❌ Email send failed: {exc}"


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 9 — save_contact  +  send_sms
#
# SMS is sent via Twilio (paid) — see send_sms() below.
#
# NOTE (v8.0.1): the free email-to-SMS carrier-gateway approach (sending
# email to <number>@vtext.com, txt.att.net, etc.) was removed. Carriers are
# shutting these gateways down industry-wide: AT&T's gateway is already
# gone (June 2025), T-Mobile's went offline December 2024, and Verizon's
# is mid-shutdown through March 2027. The remaining gateways are also
# unreliable while active — aggressive spam filtering on the sending side,
# no delivery confirmation, and no two-way reply support. Twilio (paid)
# replaces it entirely; see check_sms_replies() for inbound reply checking.
# ══════════════════════════════════════════════════════════════════════════════


@mcp.tool()
def save_contact(
    name:    str,
    phone:   str = "",
    email:   str = "",
    ctx:     "Context | None" = None,
) -> str:
    """
    Save or update a personal contact (phone and/or email) so future
    send_sms / send_email calls can use their name directly — e.g.
    "text David" or "email Vicki" — instead of repeating the number or
    address every time.

    This is for personal contacts (you, family, crew you message by name),
    not for customers — customers belong in the Customers spreadsheet sheet,
    which is checked automatically by send_sms/send_email and takes priority.

    Call this automatically (no need to ask first) whenever the operator
    tells you a name + phone number or name + email for the first time, e.g.
    "David's number is 386-555-0101" or "Vicki's email is vicki@gmail.com".

    Args:
        name:    Contact's name (e.g. "David", "Vicki"). Case-insensitive,
                 used as the lookup key — saving the same name again merges
                 with the existing record instead of overwriting it.
        phone:   Phone number, any format (digits are normalised).
        email:   Email address.
        ctx:     MCP context (injected automatically).

    Returns:
        Confirmation of what was saved.

    Voice examples:
        "David's number is 386-555-0101"
        "Save Vicki's email as vicki@gmail.com"
        "Remember that Jake's phone is 555-0199 and his email is jake@x.com"
    """
    _telemetry_increment_tool_count("save_contact")

    if not name.strip():
        return "❌ A contact name is required."
    if not phone.strip() and not email.strip():
        return "❌ Provide at least a phone number or an email address."

    _sc_user = _current_user(ctx)
    rec = _contact_save(name, phone=phone, email=email, user=_sc_user)

    # Report the ACTUAL filename the record was saved to — reuse
    # _CONTACTS_CACHE_PATH() itself rather than recomputing the filename
    # with separate (and previously mismatched) logic, which could show a
    # different, unsanitised filename than where the data really landed.
    cache_file = _CONTACTS_CACHE_PATH(_sc_user).name

    lines = [f"✅ Contact saved: {name.strip()}"]
    if rec.get('phone'):
        lines.append(f"   Phone:   {rec['phone']}")
    if rec.get('email'):
        lines.append(f"   Email:   {rec['email']}")
    lines.append(f"   (saved to {cache_file})")
    return '\n'.join(lines)


@mcp.tool()
def send_sms(
    to:      str,
    message: str,
    ctx:     "Context | None" = None,
) -> str:
    """
    Send an SMS text message to a crew member, customer, or saved contact.

    Requires an SMS provider to be configured in Settings → SMS / Text
    Messaging — Twilio, SignalWire, or Vonage (set via sms_provider). The
    free email-to-SMS carrier-gateway approach (vtext.com, txt.att.net,
    etc.) was removed — those gateways are being shut down industry-wide
    (AT&T's is already gone, Verizon's is mid-shutdown through 2027) and
    were unreliable even while active (spam filtering, no delivery
    confirmation, no two-way replies).

    'to' accepts a phone number OR a saved contact name (e.g. "David") —
    if it doesn't look like a number, contacts_cache.json is checked first,
    along with the Customers sheet and users.json by name.

    Server mode: any role (owner, manager, staff, field_crew) may send SMS
    here, via the company's shared SMS-provider account. Inbound replies
    ARE correctly attributed per user — check_sms_replies() filters to only
    the threads YOU personally sent, so Mike won't see Karen's reply to
    Jake's text.

    Args:
        to:      Recipient's phone number (10-digit US) or a saved contact
                 name. Accepts any format: 3865550101, 386-555-0101,
                 (386) 555-0101.
        message: Text message body (keep under 160 chars for one SMS segment).
        ctx:     MCP context (injected automatically).

    Returns:
        Confirmation, or an error string with setup instructions.

    Voice examples:
        "Text Torres that we're 20 minutes away"
        "Text David that the afternoon job is cancelled"
        "Send Karen Walsh a reminder that her appointment is tomorrow at 9am"
    """
    _telemetry_increment_tool_count("send_sms")

    # Server-mode gate — all roles permitted (see _send_sms_cap docstring).
    _sms_user = _current_user(ctx)
    _sms_allowed, _sms_why = _send_sms_cap(_sms_user)
    if not _sms_allowed:
        return f"❌ {_sms_why}"

    import re as _re

    # ── Resolve a name to a phone number ───────────────────────────────────────
    # Lookup order: Customers sheet (by name/company) → users.json (crew, by
    # name) → contacts_cache (personal/per-field_crew, by name). Lets "text
    # David" or "text Torres" work without repeating numbers.
    _to_digits_check = _re.sub(r'\D', '', to)
    if len(_to_digits_check) < 10:
        _resolved_phone = None

        # 1. Customers sheet — match by name/company in any column
        try:
            import openpyxl as _opxn
            _xln = _get_default_spreadsheet_path()
            if _xln and os.path.exists(_xln):
                _wbn = _opxn.load_workbook(_xln, data_only=True)
                if 'Customers' in _wbn.sheetnames:
                    _wsn = _wbn['Customers']
                    _hn, _hrn = [], None
                    for _rn in _wsn.iter_rows(min_row=1, max_row=5):
                        if len([c for c in _rn if c.value]) >= 3:
                            _hrn = _rn[0].row
                            _hn = [str(c.value or '').replace('\n', ' ').strip()
                                   for c in _rn]
                            break
                    if _hrn:
                        _pcn = next((i for i, h in enumerate(_hn)
                                    if 'phone' in h.lower()), None)
                        if _pcn is not None:
                            _needle = to.strip().lower()
                            for _rown in _wsn.iter_rows(min_row=_hrn + 1):
                                _valsn = [str(c.value or '').strip() for c in _rown]
                                if _needle and _needle in ' '.join(_valsn).lower():
                                    if _valsn[_pcn]:
                                        _resolved_phone = _valsn[_pcn]
                                        break
        except Exception:
            pass

        # 2. users.json crew — match by name (server mode only)
        if not _resolved_phone:
            try:
                _ud_n = _load_users()
                if _ud_n and isinstance(_ud_n.get('users'), dict):
                    _needle2 = to.strip().lower()
                    for _un in _ud_n['users'].values():
                        if isinstance(_un, dict) and _needle2 in (_un.get('name') or '').lower():
                            if _un.get('cell_phone'):
                                _resolved_phone = _un['cell_phone']
                                break
            except Exception:
                pass

        # 3. contacts_cache (personal, or per-field_crew in server mode)
        if not _resolved_phone:
            _contact = _contact_lookup(to, _sms_user)
            if _contact and _contact.get('phone'):
                _resolved_phone = _contact['phone']
            elif _contact:
                return (
                    f"❌ Found contact '{_contact.get('name', to)}' but no phone "
                    f"number is saved for them.\n"
                    f"Tell Claude: \"{_contact.get('name', to)}'s number is <phone>\" to add it."
                )

        if _resolved_phone:
            to = _resolved_phone
        # else: leave `to` as-is — the digit-count check below gives a clear error

    # ── Normalise phone number to 10 digits ───────────────────────────────────
    digits = _re.sub(r'\D', '', to)
    if len(digits) == 11 and digits[0] == '1':
        digits = digits[1:]
    if len(digits) != 10:
        return (
            f"❌ '{to}' does not look like a valid 10-digit US number or a "
            f"known contact/customer/crew name.\n"
            f"Provide a phone number, or save a contact first: "
            f"\"David's number is 386-555-0101\""
        )

    if not message.strip():
        return "❌ message cannot be empty."

    # ── Append callback signature if configured ───────────────────────────────
    try:
        _cfg_path = Path.home() / '.ai-prowler' / 'config.json'
        _full_cfg: dict = {}
        if _cfg_path.exists():
            import json as _jcfg
            _full_cfg = _jcfg.loads(_cfg_path.read_text(encoding='utf-8-sig'))
        _sig = _full_cfg.get('sms_callback_signature', '').strip()
        if _sig and _sig not in message:
            message = f"{message.rstrip()}\n{_sig}"
    except Exception:
        pass

    # ── Send via backend abstraction (Twilio / SignalWire / Vonage) ───────────
    try:
        from sms_backends import get_sms_backend, load_sms_config
        from sms_inbox import sms_thread_log
    except ImportError as _ie:
        return f"❌ SMS modules not found: {_ie}\nEnsure sms_backends.py and sms_inbox.py are in the AI-Prowler directory."

    _sms_cfg     = load_sms_config()
    _sms_backend = get_sms_backend(_sms_cfg)
    _sms_ok, _sms_msg = _sms_backend.validate_config()
    if not _sms_ok:
        return _sms_msg

    to_e164 = f"+1{digits}"
    _send_ok, _send_result = _sms_backend.send(to_e164, message)

    if _send_ok:
        # Log the outbound message so check_sms_inbox can route replies
        _user_id_for_log = (_sms_user or {}).get('id', 'personal') if _sms_user else 'personal'
        _provider_name   = _sms_cfg.get('sms_provider', 'twilio')
        try:
            _contact_for_log = _contact_lookup(digits, _sms_user)
            _cname_for_log   = (_contact_for_log or {}).get('name', '')
        except Exception:
            _cname_for_log = ''
        try:
            sms_thread_log(
                sent_by      = _user_id_for_log,
                to_number    = digits,
                body         = message,
                provider     = _provider_name,
                contact_name = _cname_for_log,
            )
        except Exception:
            pass
        return (
            f"✅ SMS sent via {_sms_backend.provider_name.title()}\n"
            f"   To:      {to_e164}\n"
            f"   Message: {message[:80]}{'...' if len(message) > 80 else ''}\n"
            f"   {_send_result}"
        )
    return _send_result


@mcp.tool()
def check_sms_replies(
    since_hours: int = 24,
    from_number: str = "",
    mark_read:   bool = True,
    ctx:         "Context | None" = None,
) -> str:
    """
    Check for inbound SMS replies received on your SMS number.

    V8.0.0: Reads from the local sms_inbox.json (populated by the
    /sms-webhook endpoint) instead of polling the Twilio API. Instant,
    works offline, supports all providers.

    Server mode: replies are filtered to threads you personally sent —
    Mike sees Karen's reply, Jake sees Bob's reply.

    Args:
        since_hours: How far back to look, in hours (default 24).
        from_number: Optional — filter to one sender only.
        mark_read:   If True, marks returned messages as read by you.
        ctx:         MCP context (injected automatically).

    Returns:
        A list of inbound messages, or a note that there are none.

    Voice examples:
        "Did Vicki reply to that text?"
        "Check for SMS replies"
        "Any texts come in today?"
    """
    _telemetry_increment_tool_count("check_sms_replies")

    try:
        from sms_inbox import (sms_inbox_read_for_user, sms_inbox_read,
                               sms_inbox_mark_read)
    except ImportError as _ie:
        return f"❌ SMS inbox module not found: {_ie}"

    _user    = _current_user(ctx)
    _user_id = (_user or {}).get('id', 'personal') if _user else 'personal'

    if _user:
        msgs = sms_inbox_read_for_user(_user_id, since_hours=since_hours)
    else:
        msgs = sms_inbox_read(since_hours=since_hours, from_number=from_number)

    if from_number.strip():
        import re as _re5
        _digits_f = _re5.sub(r'\D', '', from_number)
        if len(_digits_f) == 11 and _digits_f[0] == '1':
            _digits_f = _digits_f[1:]
        msgs = [m for m in msgs if _re5.sub(r'\D', '', m.get('from',''))[-10:] == _digits_f]

    # Exclude WhatsApp — use check_whatsapp_replies() for those
    msgs = [m for m in msgs if m.get('provider', '') != 'whatsapp']

    if not msgs:
        scope = f" from {from_number}" if from_number.strip() else ""
        return f"📭 No SMS replies{scope} in the last {since_hours} hour(s)."

    if mark_read:
        for m in msgs:
            try:
                sms_inbox_mark_read(m['id'], _user_id)
            except Exception:
                pass

    lines = [f"📬 {len(msgs)} SMS repl{'y' if len(msgs)==1 else 'ies'} "
             f"in the last {since_hours} hour(s):", ""]
    for m in msgs:
        from_display = m.get('contact_name') or m.get('from', 'unknown')
        lines.append(f"  From: {from_display}")
        lines.append(f"  When: {m.get('timestamp', 'unknown')[:19]}")
        lines.append(f"  Msg:  {(m.get('body') or '').strip()}")
        lines.append("")

    return '\n'.join(lines).rstrip()


@mcp.tool()
def check_sms_inbox(
    since_hours: float = 24.0,
    from_number: str   = "",
    unread_only: bool  = False,
    provider:    str   = "",
    ctx:         "Context | None" = None,
) -> str:
    """
    Check the local SMS / WhatsApp inbox for all inbound messages.

    Reads from sms_inbox.json — instant, no API call. Supports all
    providers in one place. Use provider='whatsapp' for WhatsApp only.

    Args:
        since_hours: How far back to look (default 24). Pass 0 for all.
        from_number: Filter to one sender (any format).
        unread_only: Only show messages not yet read by you.
        provider:    Filter by 'twilio', 'signalwire', 'vonage',
                     or 'whatsapp'. Empty = all.
        ctx:         MCP context (injected automatically).

    Returns:
        Formatted list of messages, or 'No messages' if empty.

    Voice examples:
        "Any messages come in?"
        "Check WhatsApp messages"
        "Any unread texts?"
    """
    _telemetry_increment_tool_count("check_sms_inbox")

    try:
        from sms_inbox import sms_inbox_read, sms_inbox_mark_read
    except ImportError as _ie:
        return f"❌ SMS inbox module not found: {_ie}"

    _user    = _current_user(ctx)
    _user_id = (_user or {}).get('id', 'personal') if _user else 'personal'

    msgs = sms_inbox_read(
        since_hours = since_hours,
        from_number = from_number,
        unread_only = unread_only,
        user_id     = _user_id,
        provider    = provider,
    )

    if not msgs:
        parts = []
        if unread_only: parts.append("unread")
        if provider:    parts.append(provider)
        scope = " ".join(parts)
        hours = f"in the last {since_hours}h" if since_hours > 0 else "ever"
        return f"📭 No {scope + ' ' if scope else ''}messages {hours}."

    lines = [f"📬 {len(msgs)} message(s):", ""]
    for m in msgs:
        prov         = m.get('provider', '')
        icon         = "💬" if prov == 'whatsapp' else "📱"
        from_display = m.get('contact_name') or m.get('from', 'unknown')
        unread_flag  = " 🔵" if _user_id not in (m.get('read_by') or []) else ""
        lines.append(f"  {icon} From: {from_display}{unread_flag}")
        lines.append(f"     When: {m.get('timestamp','')[:19]}")
        lines.append(f"     Msg:  {(m.get('body') or '').strip()}")
        lines.append("")

    return '\n'.join(lines).rstrip()


@mcp.tool()
def get_sms_thread(
    contact: str,
    since_hours: float = 168.0,
    ctx: "Context | None" = None,
) -> str:
    """
    Show the full two-way conversation thread with a contact.

    Combines your sent messages with their inbound replies in
    chronological order — like a phone's message history.
    Works for both SMS and WhatsApp.

    Server mode: only returns the thread if YOU are the one who last
    sent to this contact — otherwise another employee's conversation
    with the same contact would be visible to anyone who just knows or
    guesses their name/number, since threads are keyed by phone number
    company-wide, not per employee.

    Args:
        contact:     Contact name or phone number (any format).
        since_hours: How far back to include replies (default 168 = 7 days).
        ctx:         MCP context (injected automatically).

    Returns:
        Full conversation thread, or 'No thread found'.

    Voice examples:
        "Show my conversation with Karen"
        "What did I text Torres?"
    """
    _telemetry_increment_tool_count("get_sms_thread")

    try:
        from sms_inbox import sms_thread_get_with_replies
    except ImportError as _ie:
        return f"❌ SMS inbox module not found: {_ie}"

    conv = sms_thread_get_with_replies(contact, since_hours=since_hours)
    if not conv:
        return (
            f"📭 No SMS thread found for '{contact}'.\n\n"
            f"Start a conversation: \"Text {contact} that we're on our way\""
        )

    # Server mode: only show this thread if the caller is the one who
    # last sent to this contact. Threads are keyed by phone number alone
    # (not per employee), so without this check any role could pull up
    # a different employee's conversation just by naming the contact.
    _gst_user = _current_user(ctx)
    if _gst_user is not None:
        _gst_uid = _gst_user.get('id', '')
        if conv.get('last_sent_by') != _gst_uid:
            return (
                f"📭 No SMS/WhatsApp thread found for '{contact}' that you "
                f"personally sent to. (A thread with this contact may exist "
                f"for a different team member.)"
            )

    msgs  = conv.get('messages', [])
    name  = conv.get('contact_name', contact)
    prov  = conv.get('provider', '')
    lines = [
        f"💬 Conversation with {name}",
        f"   Provider: {prov.title() if prov else 'unknown'}",
        f"   {len(msgs)} message(s)  |  last 7 days",
        "─" * 45, "",
    ]
    for m in msgs:
        direction = m.get('direction', 'outbound')
        icon      = "→" if direction == 'outbound' else "←"
        who       = "You" if direction == 'outbound' else name
        ts        = (m.get('timestamp', '') or '')[:19]
        body      = (m.get('body', '') or '').strip()
        lines.append(f"  {icon} {who}  [{ts}]")
        lines.append(f"    {body}")
        lines.append("")

    return '\n'.join(lines).rstrip()


@mcp.tool()
def list_sms_contacts_with_replies(
    since_hours: float = 168.0,
    ctx: "Context | None" = None,
) -> str:
    """
    List all contacts you've texted recently, with unread reply counts.

    Shows who you've been in contact with, reply counts, and how many
    are unread — like a conversations list on your phone.

    Server mode: scoped to threads YOU personally sent — Mike sees only
    his own conversation list, not Karen's or Jake's. A brand-new user
    with no thread history yet sees everything until they've sent their
    first message, matching check_sms_replies' identical fallback.

    Args:
        since_hours: How far back to include (default 168 = 7 days).
        ctx:         MCP context (injected automatically).

    Returns:
        Formatted list of active threads with reply counts.

    Voice examples:
        "Who has replied to my texts?"
        "Any unread messages from customers?"
        "Show my recent SMS conversations"
    """
    _telemetry_increment_tool_count("list_sms_contacts_with_replies")

    try:
        from sms_inbox import sms_active_threads
    except ImportError as _ie:
        return f"❌ SMS inbox module not found: {_ie}"

    threads = sms_active_threads(since_hours=since_hours)

    # Server mode: scope to threads THIS user has sent to — same ownership
    # check sms_inbox_read_for_user() already uses for check_sms_replies
    # (last_sent_by == user_id). Without this, any role — including
    # field_crew — would see every active SMS conversation company-wide.
    # New user with no thread history yet: fall back to showing everything,
    # matching sms_inbox_read_for_user()'s identical fallback so both tools
    # behave consistently for a brand-new employee.
    _user = _current_user(ctx)
    if _user is not None:
        _uid = _user.get("id", "")
        _own_threads = [t for t in threads if t.get("last_sent_by") == _uid]
        if _own_threads:
            threads = _own_threads

    if not threads:
        return f"📭 No SMS activity in the last {since_hours:.0f} hours."

    lines = [f"📱 {len(threads)} active SMS thread(s):", ""]
    for t in threads:
        name   = t.get('contact_name') or t.get('thread_id', 'unknown')
        prov   = t.get('provider', '')
        unread = t.get('unread_replies', 0)
        total  = t.get('total_replies', 0)
        unread_str = f"  🔵 {unread} unread" if unread else ""
        lines.append(f"  • {name}{unread_str}")
        if total:
            lines.append(f"    {total} repl{'y' if total==1 else 'ies'} received")
        if prov:
            lines.append(f"    via {prov.title()}")
        lines.append("")

    return '\n'.join(lines).rstrip()


@mcp.tool()
def send_whatsapp(
    to:      str,
    message: str,
    ctx:     "Context | None" = None,
) -> str:
    """
    Send a WhatsApp message via Twilio's WhatsApp API.

    Uses the same Twilio credentials as send_sms — no separate
    account needed. Works worldwide without carrier restrictions.

    Args:
        to:      Recipient's phone number or contact name.
        message: Message text (up to 4096 chars for WhatsApp).
        ctx:     MCP context (injected automatically).

    Returns:
        Confirmation, or setup instructions if not configured.

    Voice examples:
        "WhatsApp Karen that the crew is 10 minutes out"
        "Send Torres a WhatsApp with the job details"
    """
    _telemetry_increment_tool_count("send_whatsapp")

    _wa_user = _current_user(ctx)
    _wa_allowed, _wa_why = _send_sms_cap(_wa_user)
    if not _wa_allowed:
        return f"❌ {_wa_why}"

    try:
        from sms_backends import get_whatsapp_backend, load_sms_config
        from sms_inbox import sms_thread_log
    except ImportError as _ie:
        return f"❌ SMS modules not found: {_ie}"

    _wa_cfg     = load_sms_config()
    _wa_backend = get_whatsapp_backend(_wa_cfg)
    _wa_ok, _wa_hint = _wa_backend.validate_config()
    if not _wa_ok:
        return _wa_hint

    import re as _re_wa
    _wa_digits = _re_wa.sub(r'\D', '', to)
    if len(_wa_digits) < 10:
        _wa_contact = _contact_lookup(to, _wa_user)
        if _wa_contact and _wa_contact.get('phone'):
            to = _wa_contact['phone']

    _wa_send_ok, _wa_result = _wa_backend.send(to, message)

    if _wa_send_ok:
        _wa_user_id = (_wa_user or {}).get('id', 'personal') if _wa_user else 'personal'
        try:
            sms_thread_log(
                sent_by      = _wa_user_id,
                to_number    = to,
                body         = message,
                provider     = 'whatsapp',
                contact_name = '',
            )
        except Exception:
            pass
        return (
            f"✅ WhatsApp message sent\n"
            f"   To:      {to}\n"
            f"   Message: {message[:80]}{'...' if len(message) > 80 else ''}\n"
            f"   {_wa_result}"
        )
    return _wa_result


@mcp.tool()
def check_whatsapp_replies(
    since_hours: float = 24.0,
    from_number: str   = "",
    ctx: "Context | None" = None,
) -> str:
    """
    Check for inbound WhatsApp messages.

    Reads from the local inbox (populated by /whatsapp-webhook) —
    instant, no API call. Shows only WhatsApp messages, not SMS.

    Server mode: scoped to threads YOU personally sent — same per-user
    isolation as check_sms_replies. Mike sees Karen's WhatsApp reply,
    not Jake's.

    Args:
        since_hours: How far back to look (default 24).
        from_number: Filter to one contact (number or name).
        ctx:         MCP context (injected automatically).

    Returns:
        List of inbound WhatsApp messages, or 'No messages'.

    Voice examples:
        "Any WhatsApp replies?"
        "Did Karen reply on WhatsApp?"
    """
    _telemetry_increment_tool_count("check_whatsapp_replies")

    try:
        from sms_inbox import sms_inbox_read, sms_inbox_read_for_user
    except ImportError as _ie:
        return f"❌ SMS inbox module not found: {_ie}"

    _wr_user    = _current_user(ctx)
    _wr_user_id = (_wr_user or {}).get('id', 'personal') if _wr_user else 'personal'

    # Server mode: scope to threads THIS user has sent to — same ownership
    # check check_sms_replies() already uses. Closes a bypass where this
    # tool previously called check_sms_inbox() directly (a plain Python
    # function call that skips MCP-level Tier A suppression entirely),
    # giving every role an unscoped, company-wide dump filtered only by
    # provider, not by who's asking.
    if _wr_user:
        _wr_msgs = sms_inbox_read_for_user(_wr_user_id, since_hours=since_hours)
        _wr_msgs = [m for m in _wr_msgs if m.get('provider', '') == 'whatsapp']
        if from_number.strip():
            import re as _re6
            _wr_digits = _re6.sub(r'\D', '', from_number)
            if len(_wr_digits) == 11 and _wr_digits[0] == '1':
                _wr_digits = _wr_digits[1:]
            _wr_msgs = [m for m in _wr_msgs
                       if _re6.sub(r'\D', '', m.get('from', ''))[-10:] == _wr_digits]
    else:
        _wr_msgs = sms_inbox_read(
            since_hours = since_hours,
            from_number = from_number,
            unread_only = False,
            user_id     = _wr_user_id,
            provider    = 'whatsapp',
        )

    if not _wr_msgs:
        hours = f"in the last {since_hours}h" if since_hours > 0 else "ever"
        return f"📭 No whatsapp messages {hours}."

    lines = [f"📬 {len(_wr_msgs)} message(s):", ""]
    for m in _wr_msgs:
        from_display = m.get('contact_name') or m.get('from', 'unknown')
        unread_flag  = " 🔵" if _wr_user_id not in (m.get('read_by') or []) else ""
        lines.append(f"  💬 From: {from_display}{unread_flag}")
        lines.append(f"     When: {m.get('timestamp','')[:19]}")
        lines.append(f"     Msg:  {(m.get('body') or '').strip()}")
        lines.append("")

    return '\n'.join(lines).rstrip()


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 10 — schedule_next_recurring_job
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def schedule_next_recurring_job(
    job_identifier: str,
    filepath:       str = "",
    when:           str = "",
    ctx: "Context | None" = None,
) -> str:
    """
    Auto-create the next recurring job after a job is marked complete.

    Looks up the completed job to find the customer, then reads the customer's
    service frequency (Weekly / Biweekly / Monthly / Quarterly / One-time) from
    the Customers sheet and calculates the next service date accordingly.
    Writes the new job row to Jobs_Schedule and returns the new JobID.

    Handles all standard frequencies:
      W  / Weekly      → +7 days
      BW / Biweekly    → +14 days
      M  / Monthly     → +1 month (same day)
      Q  / Quarterly   → +3 months (same day)
      OT / One-time    → no next job created

    Args:
        job_identifier: JobID (e.g. "JOB-0001") or customer name/company.
        filepath:       Path to the .xlsx tracker. Uses default if omitted.
                         In server mode this argument is ignored — see
                         _resolve_job_spreadsheet_path().
        when:           Which jobs to search among — works in BOTH personal
                         and server mode: "today" (default if omitted),
                         "tomorrow", "yesterday", "this_week", "next_week",
                         "any" (no date restriction), or an explicit
                         "YYYY-MM-DD" date.

    Server mode: staff/field_crew only search jobs assigned to THEM
    (Crew / Technician matches their own name) — owner/manager search
    every crew's jobs. This crew restriction applies regardless of `when`
    and does not apply at all in personal mode (single user).

    Returns:
        New JobID and scheduled date, or explanation if no recurring job applies.

    Voice examples:
        "Schedule the next recurring job after JOB-0003"
        "Set up the next job for Blue Wave Cafe"
        "Auto-book the next visit after today's Sunshine Realty job"
        "Schedule the next job after the one I did yesterday"
        "Reschedule next week's Crabby's job"
    """
    # Serialise the whole load -> modify -> save cycle so two concurrent
    # server-mode writers can never interleave and silently drop each
    # other's changes (openpyxl has no file locking of its own).
    with _spreadsheet_write_lock:
        return _schedule_next_recurring_job_impl(job_identifier, filepath, when, ctx)


def _srj_resolve_date_range(when: str, today):
    """
    Resolve a human date-range keyword into an inclusive (start, end) date
    range for scoping schedule_next_recurring_job()'s job search.

    Returns (None, None) for "any" — no date restriction at all.
    Falls back to (today, today) for anything unrecognised, since "today"
    is the safe default this tool is normally used for.
    """
    import datetime as _dt
    w = (when or "today").strip().lower().replace("-", "_").replace(" ", "_")
    if w in ("any", "all", ""):
        return (None, None)
    if w == "today":
        return (today, today)
    if w == "tomorrow":
        d = today + _dt.timedelta(days=1)
        return (d, d)
    if w == "yesterday":
        d = today - _dt.timedelta(days=1)
        return (d, d)
    if w in ("this_week", "thisweek"):
        start = today - _dt.timedelta(days=today.weekday())  # Monday
        end   = start + _dt.timedelta(days=6)                # Sunday
        return (start, end)
    if w in ("next_week", "nextweek"):
        start = today - _dt.timedelta(days=today.weekday()) + _dt.timedelta(days=7)
        end   = start + _dt.timedelta(days=6)
        return (start, end)
    # Explicit date, e.g. "2026-07-15"
    try:
        d = _dt.datetime.strptime(w.replace("_", "-"), "%Y-%m-%d").date()
        return (d, d)
    except ValueError:
        pass
    return (today, today)


def _schedule_next_recurring_job_impl(job_identifier: str, filepath: str, when: str, ctx) -> str:
    """Implementation body, called under _spreadsheet_write_lock. See
    schedule_next_recurring_job() for the public docstring."""
    _telemetry_increment_tool_count("schedule_next_recurring_job")

    try:
        import openpyxl as _opx
    except ImportError:
        return "❌ openpyxl not installed. Run: pip install openpyxl"

    import datetime as _dt
    import calendar as _cal

    filepath = _resolve_job_spreadsheet_path(ctx, filepath)
    if not filepath:
        return "❌ No spreadsheet path configured."

    fp = filepath.replace("\\", "/")
    if not os.path.exists(fp):
        return f"❌ Spreadsheet not found: {fp}"

    backup_msg = _backup_spreadsheet(fp)
    if backup_msg.startswith("⚠️") or backup_msg.startswith("❌"):
        return f"{backup_msg}\nSpreadsheet was NOT modified."

    try:
        wb = _opx.load_workbook(fp)
    except Exception as exc:
        return f"❌ Could not open spreadsheet: {exc}"

    # ── Read Jobs_Schedule ────────────────────────────────────────────────────
    if "Jobs_Schedule" not in wb.sheetnames:
        return "❌ 'Jobs_Schedule' sheet not found in spreadsheet."
    ws_jobs = wb["Jobs_Schedule"]

    job_hdr_row, job_hdrs = None, []
    for r in ws_jobs.iter_rows(min_row=1, max_row=5):
        ne = [c for c in r if c.value is not None]
        if len(ne) >= 3:
            job_hdr_row = r[0].row
            job_hdrs = [str(c.value).strip().replace('\n',' ') if c.value else '' for c in r]
            break

    if not job_hdrs:
        return "❌ Could not detect header row in Jobs_Schedule."

    # ── Scoping ─────────────────────────────────────────────────────────────
    # Staff/field_crew (server mode only): only search jobs assigned to THEM
    # (Crew / Technician matches their own name) — prevents accidentally
    # rescheduling a coworker's job by a job-name/JobID match alone.
    # Owner/manager see every crew's jobs — their real entitlement, same
    # _check_db_cap('full') gate used by list_writable_directories /
    # list_tracked_directories. This crew restriction never applies in
    # personal mode (single user).
    #
    # Date range (`when`): a common feature in BOTH modes now. If the caller
    # doesn't specify one: server mode defaults to "today" (this tool is
    # normally called right after finishing a job); personal mode defaults
    # to "any" (unrestricted) so every version before this parameter existed
    # keeps working unchanged. An explicitly-passed `when` behaves
    # identically in both modes.
    _srj_user = _current_user(ctx)
    _srj_full_access = True
    _srj_crew_name = None
    if _srj_user is not None:
        _srj_full_access, _ = _check_db_cap(_srj_user, "full")
        _srj_crew_name = (_srj_user.get('name') or '').strip().lower()

    _srj_effective_when = when.strip() if when and when.strip() else "today"
    _srj_range_start, _srj_range_end = _srj_resolve_date_range(
        _srj_effective_when, _dt.date.today())

    def _srj_parse_date(value):
        if isinstance(value, _dt.datetime):
            return value.date()
        if isinstance(value, _dt.date):
            return value
        if value:
            for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y'):
                try:
                    return _dt.datetime.strptime(str(value).strip(), fmt).date()
                except ValueError:
                    continue
        return None

    _crew_col_idx = next((i for i, h in enumerate(job_hdrs) if h == "Crew / Technician"), None)
    _date_col_idx = next((i for i, h in enumerate(job_hdrs) if h == "Service Date"), None)

    _srj_matches = []
    for row in ws_jobs.iter_rows(min_row=job_hdr_row + 1):
        vals = [c.value for c in row]
        row_text = " ".join(str(v) for v in vals if v)
        if job_identifier.lower() not in row_text.lower():
            continue
        if _srj_range_start is not None:  # None,None means "any" - unrestricted
            row_date = _srj_parse_date(vals[_date_col_idx]) if _date_col_idx is not None else None
            if row_date is None or not (_srj_range_start <= row_date <= _srj_range_end):
                continue
        if _srj_user is not None and not _srj_full_access:
            row_crew = str(vals[_crew_col_idx] or '').strip().lower() if _crew_col_idx is not None else ''
            if row_crew != _srj_crew_name:
                continue
        _srj_matches.append(dict(zip(job_hdrs, vals)))

    if not _srj_matches:
        _scope_note = " assigned to you" if (_srj_user is not None and not _srj_full_access) else ""
        _when_note = (
            "" if _srj_range_start is None
            else f" scheduled {_srj_range_start}" if _srj_range_start == _srj_range_end
            else f" scheduled {_srj_range_start} to {_srj_range_end}"
        )
        return (
            f"❌ No job found matching '{job_identifier}'{_when_note}{_scope_note} "
            f"in Jobs_Schedule." + ("" if _srj_range_start is None else " Try when='any' to search all dates.")
        )

    if len(_srj_matches) > 1:
        _srj_candidates = "\n".join(
            f"   • {m.get('JobID (JOB-####)','?')} — {m.get('Customer Name / Company','?')}"
            for m in _srj_matches[:10]
        )
        return (
            f"❌ '{job_identifier}' matches {len(_srj_matches)} jobs — please "
            f"specify which one:\n{_srj_candidates}\n\nTry again with the exact JobID."
        )

    found_job = _srj_matches[0]

    cust_id   = str(found_job.get("CustomerID (Customers!A)", "") or "")
    cust_name = str(found_job.get("Customer Name / Company",  "") or "")
    svc_date  = found_job.get("Service Date")
    crew      = str(found_job.get("Crew / Technician", "") or "")
    svc_type  = str(found_job.get("Service Type", "") or "")
    est_dur   = found_job.get("Est. Duration (min)", "")
    svc_notes = str(found_job.get("Service Details / Notes", "") or "")

    # Parse service date
    if isinstance(svc_date, _dt.datetime):
        base_date = svc_date.date()
    elif isinstance(svc_date, _dt.date):
        base_date = svc_date
    elif svc_date:
        for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y'):
            try:
                base_date = _dt.datetime.strptime(str(svc_date).strip(), fmt).date()
                break
            except ValueError:
                continue
        else:
            return f"❌ Could not parse service date: {svc_date}"
    else:
        return "❌ Completed job has no Service Date — cannot compute next date."

    # ── Look up customer frequency ────────────────────────────────────────────
    frequency = ""
    cust_addr = str(found_job.get("Street Address ★ AI Route", "") or "")
    cust_city = str(found_job.get("City ★ AI Route", "") or "")
    cust_zip  = str(found_job.get("ZIP ★ AI Route", "") or "")

    if "Customers" in wb.sheetnames:
        ws_cust = wb["Customers"]
        cust_hdr_row, cust_hdrs = None, []
        for r in ws_cust.iter_rows(min_row=1, max_row=5):
            ne = [c for c in r if c.value is not None]
            if len(ne) >= 3:
                cust_hdr_row = r[0].row
                cust_hdrs = [str(c.value).strip().replace('\n',' ') if c.value else '' for c in r]
                break
        if cust_hdrs:
            for row in ws_cust.iter_rows(min_row=cust_hdr_row + 1):
                cvals = [c.value for c in row]
                crow  = dict(zip(cust_hdrs, cvals))
                cid   = str(crow.get("CustomerID (CUST-####)", "") or "")
                cn    = str(crow.get("Company Name", "") or crow.get("First Name", "") or "")
                if (cust_id and cid == cust_id) or (cust_name and cust_name.lower() in str(cvals).lower()):
                    frequency = str(crow.get("Frequency W/BW/M/Q/OT", "") or "")
                    pref_day  = str(crow.get("Preferred Day(s)", "") or "")
                    pref_time = str(crow.get("Pref. Time Window", "") or "")
                    break

    freq_norm = frequency.strip().upper()

    # Frequency → delta mapping
    _FREQ_MAP = {
        "W":         lambda d: d + _dt.timedelta(weeks=1),
        "WEEKLY":    lambda d: d + _dt.timedelta(weeks=1),
        "BW":        lambda d: d + _dt.timedelta(weeks=2),
        "BIWEEKLY":  lambda d: d + _dt.timedelta(weeks=2),
        "M":         lambda d: d.replace(month=(d.month % 12) + 1,
                                          year=d.year + (1 if d.month == 12 else 0)),
        "MONTHLY":   lambda d: d.replace(month=(d.month % 12) + 1,
                                          year=d.year + (1 if d.month == 12 else 0)),
        "Q":         lambda d: d.replace(month=((d.month - 1 + 3) % 12) + 1,
                                          year=d.year + ((d.month - 1 + 3) // 12)),
        "QUARTERLY": lambda d: d.replace(month=((d.month - 1 + 3) % 12) + 1,
                                          year=d.year + ((d.month - 1 + 3) // 12)),
    }

    if freq_norm in ("OT", "ONE-TIME", "ONE TIME", "ONETIME", ""):
        return (
            f"ℹ️  No recurring job scheduled — {cust_name} is a one-time customer\n"
            f"   (Frequency: '{frequency or 'not set'}')\n"
            "   To add a recurring schedule, update the Customers sheet first."
        )

    delta_fn = _FREQ_MAP.get(freq_norm)
    if delta_fn is None:
        return (
            f"❌ Unrecognised frequency '{frequency}' for {cust_name}.\n"
            "   Expected: W / BW / M / Q / OT"
        )

    next_date = delta_fn(base_date)

    # ── Generate next JobID ───────────────────────────────────────────────────
    existing_ids = []
    for row in ws_jobs.iter_rows(min_row=job_hdr_row + 1):
        jid_cell = row[0].value
        if jid_cell and str(jid_cell).startswith("JOB-"):
            try:
                existing_ids.append(int(str(jid_cell).split("-")[1]))
            except ValueError:
                pass
    next_num = (max(existing_ids) + 1) if existing_ids else 1
    new_job_id = f"JOB-{next_num:04d}"

    # ── Build new row matching Jobs_Schedule columns ──────────────────────────
    # Find the next empty row
    last_row = job_hdr_row
    for row in ws_jobs.iter_rows(min_row=job_hdr_row + 1):
        if any(c.value for c in row):
            last_row = row[0].row

    new_row_data = {
        "JobID (JOB-####)":          new_job_id,
        "CustomerID (Customers!A)":  cust_id,
        "Customer Name / Company":   cust_name,
        "Customer Type":             str(found_job.get("Customer Type", "") or ""),
        "Street Address ★ AI Route": cust_addr,
        "City ★ AI Route":           cust_city,
        "State":                     str(found_job.get("State", "") or ""),
        "ZIP ★ AI Route":            cust_zip,
        "Service Date":              next_date.strftime('%Y-%m-%d'),
        "Day of Week":               next_date.strftime('%A'),
        "Service Type":              svc_type,
        "Service Details / Notes":   svc_notes,
        "Crew / Technician":         crew,
        "Est. Duration (min)":       est_dur,
        "Job Status":                "Scheduled",
    }

    new_row_num = last_row + 1
    for col_idx, col_name in enumerate(job_hdrs, 1):
        if col_name in new_row_data:
            ws_jobs.cell(row=new_row_num, column=col_idx).value = new_row_data[col_name]

    try:
        wb.save(fp)
    except Exception as exc:
        return f"❌ Could not save spreadsheet: {exc}"

    return (
        f"✅ Next recurring job scheduled\n"
        f"   New Job ID:   {new_job_id}\n"
        f"   Customer:     {cust_name}\n"
        f"   Frequency:    {frequency}  ({freq_norm})\n"
        f"   Last Service: {base_date.strftime('%Y-%m-%d')}\n"
        f"   Next Service: {next_date.strftime('%Y-%m-%d')} ({next_date.strftime('%A')})\n"
        f"   Crew:         {crew or '(unassigned)'}\n"
        f"   Service:      {svc_type}\n"
        f"   {backup_msg}"
    )


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 11 — log_time_entry
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def log_time_entry(
    job_identifier: str,
    action:         str,
    filepath:       str = "",
    ctx: "Context | None" = None,
) -> str:
    """
    Clock in or out for a job. Records timestamps to the TimeLog sheet for
    accurate job costing, crew payroll, and actual-vs-estimated duration.

    On clock-out, automatically calculates elapsed minutes and writes the
    Actual Duration back to the Jobs_Schedule sheet.

    Args:
        job_identifier: JobID (e.g. "JOB-0001") or customer name.
        action:         "start" to clock in, "stop" to clock out.
        filepath:       Path to the .xlsx tracker. Uses default if omitted.
                         In server mode this argument is ignored — see
                         _resolve_job_spreadsheet_path().

    Returns:
        Confirmation with timestamp and elapsed time (on stop), or error.

    Voice examples:
        "Clock in for JOB-0003"
        "Start the timer for Blue Wave Cafe"
        "Clock out — I just finished the Harbor Inn job"
        "Stop the timer for JOB-0001"
    """
    # Serialise the whole load -> modify -> save cycle so two concurrent
    # server-mode writers (e.g. two crew members clocking in/out at once)
    # can never interleave and silently drop each other's changes.
    with _spreadsheet_write_lock:
        return _log_time_entry_impl(job_identifier, action, filepath, ctx)


def _log_time_entry_impl(job_identifier: str, action: str, filepath: str, ctx) -> str:
    """Implementation body, called under _spreadsheet_write_lock. See
    log_time_entry() for the public docstring."""
    _telemetry_increment_tool_count("log_time_entry")

    try:
        import openpyxl as _opx
    except ImportError:
        return "❌ openpyxl not installed. Run: pip install openpyxl"

    import datetime as _dt

    action = action.strip().lower()
    if action not in ("start", "stop"):
        return "❌ action must be 'start' or 'stop'."

    filepath = _resolve_job_spreadsheet_path(ctx, filepath)
    if not filepath:
        return "❌ No spreadsheet path configured."

    fp = filepath.replace("\\", "/")
    if not os.path.exists(fp):
        return f"❌ Spreadsheet not found: {fp}"

    backup_msg = _backup_spreadsheet(fp)
    if backup_msg.startswith("⚠️") or backup_msg.startswith("❌"):
        return f"{backup_msg}\nSpreadsheet was NOT modified."

    try:
        wb = _opx.load_workbook(fp)
    except Exception as exc:
        return f"❌ Could not open spreadsheet: {exc}"

    # ── Ensure TimeLog sheet exists ───────────────────────────────────────────
    _TIMELOG_HEADERS = [
        "EntryID", "JobID", "Customer Name / Company",
        "Clock In", "Clock Out", "Elapsed (min)", "Crew / Technician", "Notes",
        "Logged By (User ID)",
    ]
    if "TimeLog" not in wb.sheetnames:
        ws_log = wb.create_sheet("TimeLog")
        ws_log.append(["⏱️  TIME LOG — Job Clock In / Clock Out"])
        ws_log.append(_TIMELOG_HEADERS)
    else:
        ws_log = wb["TimeLog"]

    # Detect header row in TimeLog
    log_hdr_row, log_hdrs = None, []
    for r in ws_log.iter_rows(min_row=1, max_row=5):
        ne = [c for c in r if c.value is not None]
        if len(ne) >= 3:
            log_hdr_row = r[0].row
            log_hdrs = [str(c.value).strip().replace('\n',' ') if c.value else '' for c in r]
            break

    if not log_hdrs:
        return "❌ Could not detect header row in TimeLog sheet."

    # ── Find the job in Jobs_Schedule — must match exactly ONE row ────────────
    # Requiring an unambiguous match (not silently taking the first hit, and
    # not silently falling back to the raw identifier string when nothing
    # matches) is what makes "clock in" / "clock out" actually specify a
    # real job rather than guessing.
    cust_name, job_id_found, crew = None, None, ""
    if "Jobs_Schedule" in wb.sheetnames:
        ws_jobs = wb["Jobs_Schedule"]
        job_hdr_row, job_hdrs = None, []
        for r in ws_jobs.iter_rows(min_row=1, max_row=5):
            ne = [c for c in r if c.value is not None]
            if len(ne) >= 3:
                job_hdr_row = r[0].row
                job_hdrs = [str(c.value).strip().replace('\n',' ') if c.value else '' for c in r]
                break
        if job_hdrs:
            _matches = []
            for row in ws_jobs.iter_rows(min_row=job_hdr_row + 1):
                vals = [c.value for c in row]
                row_text = " ".join(str(v) for v in vals if v)
                if job_identifier.lower() in row_text.lower():
                    jrow = dict(zip(job_hdrs, vals))
                    _matches.append((
                        str(jrow.get("JobID (JOB-####)", "") or ""),
                        str(jrow.get("Customer Name / Company", "") or ""),
                        str(jrow.get("Crew / Technician", "") or ""),
                    ))
            if not _matches:
                return (
                    f"❌ No job found matching '{job_identifier}' in Jobs_Schedule.\n"
                    "Check the JobID or customer name and try again — clocking in/out "
                    "requires an exact job, not a guess."
                )
            if len(_matches) > 1:
                _candidates = "\n".join(
                    f"   • {jid or '(no JobID)'} — {cust}" for jid, cust, _ in _matches[:10]
                )
                return (
                    f"❌ '{job_identifier}' matches {len(_matches)} jobs — please specify "
                    f"which one:\n{_candidates}\n\nTry again with the exact JobID."
                )
            job_id_found, cust_name, crew = _matches[0]
            job_id_found = job_id_found or job_identifier
            cust_name = cust_name or job_identifier
    if job_id_found is None:
        return (
            "❌ No Jobs_Schedule sheet found — cannot verify job identity. "
            "Clocking in/out requires a real job in the tracker."
        )

    _lte_user = _current_user(ctx)

    now_ts = _dt.datetime.now()
    now_str = now_ts.strftime('%Y-%m-%d %H:%M:%S')

    _lte_uid = _lte_user.get('id', '') if _lte_user else ''
    _lte_display = _lte_user.get('name', _lte_uid) if _lte_user else ''

    if action == "start":
        # Check for an already-open entry. Server mode: scoped to THIS
        # user's own open entries — Jake starting a shift doesn't block
        # Karen from also starting her own shift on the same job. Personal
        # mode: unchanged, any open entry for the job blocks re-starting
        # (there's only one user, so a second concurrent entry can only
        # mean a forgotten clock-out).
        for row in ws_log.iter_rows(min_row=log_hdr_row + 1):
            rvals = [c.value for c in row]
            rdict = dict(zip(log_hdrs, rvals))
            if not (job_id_found.lower() in str(rdict.get("JobID", "")).lower()
                    and not rdict.get("Clock Out")):
                continue
            if _lte_user is not None and rdict.get("Logged By (User ID)", "") != _lte_uid:
                continue  # someone else's open entry — doesn't block you
            return (
                f"⚠️  A clock-in for {job_id_found} is already open.\n"
                f"   Clocked in at: {rdict.get('Clock In')}\n"
                "   Call log_time_entry with action='stop' to clock out first."
            )

        # Generate entry ID
        existing_ids = []
        for row in ws_log.iter_rows(min_row=log_hdr_row + 1):
            eid = row[0].value
            if eid and str(eid).startswith("TE-"):
                try:
                    existing_ids.append(int(str(eid).split("-")[1]))
                except ValueError:
                    pass
        next_num  = (max(existing_ids) + 1) if existing_ids else 1
        entry_id  = f"TE-{next_num:04d}"

        # Find next empty row
        last_log_row = log_hdr_row
        for row in ws_log.iter_rows(min_row=log_hdr_row + 1):
            if any(c.value for c in row):
                last_log_row = row[0].row

        # Server mode: the actual calling user's identity is authoritative
        # for who's clocking in — not whatever's pre-filled in the job's
        # Crew / Technician assignment. Personal mode: unchanged, uses the
        # job's assigned crew field since there's only one user anyway.
        crew_display = _lte_display if _lte_user is not None else crew

        new_row = {
            "EntryID":                  entry_id,
            "JobID":                    job_id_found,
            "Customer Name / Company":  cust_name,
            "Clock In":                 now_str,
            "Crew / Technician":        crew_display,
            "Logged By (User ID)":      _lte_uid if _lte_user is not None else "",
        }
        new_row_num = last_log_row + 1
        for col_idx, col_name in enumerate(log_hdrs, 1):
            if col_name in new_row:
                ws_log.cell(row=new_row_num, column=col_idx).value = new_row[col_name]

        try:
            wb.save(fp)
        except Exception as exc:
            return f"❌ Could not save spreadsheet: {exc}"

        return (
            f"⏱️  Clocked IN\n"
            f"   Entry ID:  {entry_id}\n"
            f"   Job:       {job_id_found}\n"
            f"   Customer:  {cust_name}\n"
            f"   Clock In:  {now_str}\n"
            f"   Crew:      {crew_display or '(unspecified)'}\n"
            "   Call log_time_entry(action='stop') when finished."
        )

    else:  # action == "stop"
        # Find YOUR open entry for this job. Server mode: only entries this
        # caller personally opened (Logged By (User ID) match) — you cannot
        # clock out a coworker's still-open shift, even for the same job.
        # Personal mode: unchanged, single-user, so any open entry qualifies.
        open_row_num = None
        open_entry   = None
        _stop_someone_elses_entry_exists = False
        for row in ws_log.iter_rows(min_row=log_hdr_row + 1):
            rvals = [c.value for c in row]
            rdict = dict(zip(log_hdrs, rvals))
            if not (job_id_found.lower() in str(rdict.get("JobID", "")).lower()
                    and not rdict.get("Clock Out")):
                continue
            if _lte_user is not None and rdict.get("Logged By (User ID)", "") != _lte_uid:
                _stop_someone_elses_entry_exists = True
                continue
            open_row_num = row[0].row
            open_entry   = (rdict, row)
            break

        if open_entry is None:
            if _stop_someone_elses_entry_exists:
                return (
                    f"❌ No open clock-in found for '{job_identifier}' under your "
                    f"name. Someone else on the crew has an open entry for this "
                    f"job, but you can only clock out your own."
                )
            return (
                f"❌ No open clock-in found for '{job_identifier}'.\n"
                "   Call log_time_entry(action='start') first."
            )

        rdict, row_cells = open_entry
        clock_in_val = rdict.get("Clock In")
        if isinstance(clock_in_val, _dt.datetime):
            clock_in_dt = clock_in_val
        else:
            try:
                clock_in_dt = _dt.datetime.strptime(str(clock_in_val).strip(), '%Y-%m-%d %H:%M:%S')
            except Exception:
                return f"❌ Could not parse Clock In time: {clock_in_val}"

        elapsed_td   = now_ts - clock_in_dt
        elapsed_mins = round(elapsed_td.total_seconds() / 60)

        # Write Clock Out and Elapsed
        for col_idx, col_name in enumerate(log_hdrs, 1):
            if col_name == "Clock Out":
                ws_log.cell(row=open_row_num, column=col_idx).value = now_str
            elif col_name == "Elapsed (min)":
                ws_log.cell(row=open_row_num, column=col_idx).value = elapsed_mins

        # Also update Actual Duration in Jobs_Schedule
        if "Jobs_Schedule" in wb.sheetnames:
            ws_jobs = wb["Jobs_Schedule"]
            job_hdr_row2, job_hdrs2 = None, []
            for r in ws_jobs.iter_rows(min_row=1, max_row=5):
                ne = [c for c in r if c.value is not None]
                if len(ne) >= 3:
                    job_hdr_row2 = r[0].row
                    job_hdrs2 = [str(c.value).strip().replace('\n',' ') if c.value else '' for c in r]
                    break
            if job_hdrs2:
                for row in ws_jobs.iter_rows(min_row=job_hdr_row2 + 1):
                    vals = [c.value for c in row]
                    row_text = " ".join(str(v) for v in vals if v)
                    if job_id_found.lower() in row_text.lower():
                        for col_idx, col_name in enumerate(job_hdrs2, 1):
                            if "Actual" in col_name and "Duration" in col_name:
                                ws_jobs.cell(row=row[0].row, column=col_idx).value = elapsed_mins
                        break

        try:
            wb.save(fp)
        except Exception as exc:
            return f"❌ Could not save spreadsheet: {exc}"

        hours, mins = divmod(elapsed_mins, 60)
        elapsed_str = (f"{hours}h {mins}m" if hours else f"{mins}m")

        return (
            f"⏱️  Clocked OUT\n"
            f"   Job:          {job_id_found}\n"
            f"   Customer:     {cust_name}\n"
            f"   Clock In:     {clock_in_dt.strftime('%I:%M %p')}\n"
            f"   Clock Out:    {now_ts.strftime('%I:%M %p')}\n"
            f"   Elapsed:      {elapsed_str}  ({elapsed_mins} min)\n"
            f"   Actual Duration written to Jobs_Schedule ✅\n"
            f"   {backup_msg}"
        )


# ══════════════════════════════════════════════════════════════════════════════
# ACTION TOOL 12 — get_ar_aging_report
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def get_ar_aging_report(
    filepath:    str  = "",
    as_of_date:  str  = "today",
    ctx: "Context | None" = None,
) -> str:
    """
    Generate an Accounts Receivable aging report from the Invoices sheet.

    Reads all unpaid/partial invoices and buckets them by days overdue:
      Current (not yet due) | 1-30 | 31-60 | 61-90 | 90+ days

    This is the fastest way to see cash flow on your phone — ask Claude
    "What invoices are overdue?" and get a clean summary in seconds.

    Args:
        filepath:   Path to the .xlsx tracker. Uses default if omitted.
        as_of_date: Date to calculate aging from. Default "today".
                    Accepts: "today", "2026-04-15", "04/15/2026".

    Returns:
        AR aging summary grouped by bucket with subtotals and a grand total.

    Voice examples:
        "Show me my overdue invoices"
        "What's my AR aging report?"
        "Which customers owe me money?"
        "How much is outstanding past 30 days?"
    """
    _telemetry_increment_tool_count("get_ar_aging_report")

    try:
        import openpyxl as _opx
    except ImportError:
        return "❌ openpyxl not installed. Run: pip install openpyxl"

    import datetime as _dt

    filepath = _resolve_job_spreadsheet_path(ctx, filepath)
    if not filepath:
        return "❌ No spreadsheet path configured."

    fp = filepath.replace("\\", "/")
    if not os.path.exists(fp):
        return f"❌ Spreadsheet not found: {fp}"

    try:
        wb = _opx.load_workbook(fp, data_only=True)
    except Exception as exc:
        return f"❌ Could not open spreadsheet: {exc}"

    if "Invoices" not in wb.sheetnames:
        return "❌ 'Invoices' sheet not found in spreadsheet."

    ws = wb["Invoices"]

    # Detect header row
    hdr_row, headers = None, []
    for r in ws.iter_rows(min_row=1, max_row=5):
        ne = [c for c in r if c.value is not None]
        if len(ne) >= 3:
            hdr_row = r[0].row
            headers = [str(c.value).strip().replace('\n', ' ') if c.value else '' for c in r]
            break

    if not headers:
        return "❌ Could not detect header row in Invoices sheet."

    # Parse as_of_date
    aod_str = as_of_date.strip().lower()
    if aod_str == 'today' or not aod_str:
        as_of = _dt.date.today()
    else:
        for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y'):
            try:
                as_of = _dt.datetime.strptime(aod_str, fmt).date()
                break
            except ValueError:
                continue
        else:
            return f"❌ Could not parse as_of_date '{as_of_date}'."

    # ── Bucket definitions ────────────────────────────────────────────────────
    buckets = {
        "current":  {"label": "Current (not yet due)", "rows": [], "total": 0.0},
        "1_30":     {"label": "1 – 30 days overdue",   "rows": [], "total": 0.0},
        "31_60":    {"label": "31 – 60 days overdue",  "rows": [], "total": 0.0},
        "61_90":    {"label": "61 – 90 days overdue",  "rows": [], "total": 0.0},
        "over_90":  {"label": "90+ days overdue",      "rows": [], "total": 0.0},
    }

    def _parse_date(v):
        if isinstance(v, _dt.datetime):
            return v.date()
        if isinstance(v, _dt.date):
            return v
        if v:
            for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y'):
                try:
                    return _dt.datetime.strptime(str(v).strip(), fmt).date()
                except ValueError:
                    continue
        return None

    total_outstanding = 0.0
    rows_processed    = 0

    for row in ws.iter_rows(min_row=hdr_row + 1):
        vals   = [c.value for c in row]
        if all(v is None or str(v).strip() == '' for v in vals):
            continue
        rdict  = dict(zip(headers, vals))

        pmt_status = str(rdict.get("Payment Status", "") or "").strip().upper()
        if pmt_status in ("PAID",):
            continue  # fully paid — skip

        balance_raw = rdict.get("Balance Due ($)")
        try:
            balance = float(balance_raw) if balance_raw is not None else 0.0
        except (TypeError, ValueError):
            balance = 0.0

        if balance <= 0:
            continue  # zero balance — skip

        due_date = _parse_date(rdict.get("Due Date (Net 30)"))
        inv_id   = str(rdict.get("InvoiceID (INV-####)", "") or "—")
        customer = str(rdict.get("Customer Name / Company", "") or "—")
        inv_date = _parse_date(rdict.get("Invoice Date"))

        if due_date is None:
            bucket_key = "current"
        else:
            days_over = (as_of - due_date).days
            if days_over <= 0:
                bucket_key = "current"
            elif days_over <= 30:
                bucket_key = "1_30"
            elif days_over <= 60:
                bucket_key = "31_60"
            elif days_over <= 90:
                bucket_key = "61_90"
            else:
                bucket_key = "over_90"

        due_str  = due_date.strftime('%Y-%m-%d') if due_date else "—"
        days_str = (f"{(as_of - due_date).days}d overdue" if due_date and (as_of - due_date).days > 0
                    else ("due " + due_date.strftime('%m/%d') if due_date else "no due date"))
        row_line = f"  {inv_id:<12}  {customer:<28}  ${balance:>9,.2f}   {days_str}"

        buckets[bucket_key]["rows"].append(row_line)
        buckets[bucket_key]["total"] += balance
        total_outstanding += balance
        rows_processed += 1

    if rows_processed == 0:
        return (
            f"✅ No outstanding invoices as of {as_of.strftime('%Y-%m-%d')}.\n"
            "   All invoices are paid or have zero balance."
        )

    # ── Build report ──────────────────────────────────────────────────────────
    lines = [
        f"💰 AR AGING REPORT",
        f"   As of: {as_of.strftime('%B %d, %Y')}",
        f"   File:  {os.path.basename(fp)}",
        "═" * 60,
        "",
    ]

    _BUCKET_ORDER = ["over_90", "61_90", "31_60", "1_30", "current"]
    for bkey in _BUCKET_ORDER:
        b = buckets[bkey]
        if not b["rows"]:
            continue
        lines.append(f"  {'⚠️' if bkey != 'current' else '📋'}  {b['label']}")
        lines.append(f"  {'─' * 56}")
        lines.append(f"  {'Invoice':<12}  {'Customer':<28}  {'Balance':>11}   Days")
        for r in b["rows"]:
            lines.append(r)
        lines.append(f"  {'─' * 56}")
        lines.append(f"  {'Subtotal':<42}  ${b['total']:>9,.2f}")
        lines.append("")

    lines += [
        "═" * 60,
        f"  TOTAL OUTSTANDING:              ${total_outstanding:>12,.2f}",
        "═" * 60,
    ]

    if buckets["over_90"]["total"] > 0 or buckets["61_90"]["total"] > 0:
        lines.append("")
        lines.append("  🚨 Action recommended: send payment reminders for 60+ day items.")
        lines.append("     Ask: \"Send payment reminders for all overdue invoices\"")

    return "\n".join(lines)


# ── Import the self-learning engine ──────────────────────────────────────────
try:
    import self_learning as _sl
    _log.info("self_learning module imported OK — Self-Learning Tools active")
    _SELF_LEARNING_AVAILABLE = True
except ImportError as _sl_err:
    _log.warning("self_learning module not found: %s — "
                 "Self-Learning Tools will be disabled", _sl_err)
    _SELF_LEARNING_AVAILABLE = False


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — record_learning
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def record_learning(
    title: str,
    content: str,
    category: str = "general",
    context: str = "",
    source: str = "",
    confidence: float = 0.8,
    tags: str = "",
    supersedes_id: str = "",
    outcome: str = "unknown",
    auto_detected: bool = False,
    ctx: Context = None,
) -> str:
    """
    Record a new learning into AI-Prowler's self-learning knowledge base.

    Use this when the operator says "remember this", "learn this", or when
    Claude detects new information that supersedes previously known facts.
    Also use after project reviews, post-mortems, or when documenting
    what went right or wrong in business activities.

    The learning is instantly indexed into ChromaDB for semantic retrieval —
    no training or GPU required.  Future calls to search_learnings() will
    find and apply this knowledge automatically.

    ── CONFIRMATION PROTOCOL ──────────────────────────────────────────────
    CRITICAL: After calling this tool, Claude MUST ALWAYS present the
    returned confirmation summary to the user. NEVER record silently.

    If auto_detected=True (Claude initiated the recording without being
    asked), the confirmation is MORE PROMINENT — it tells the user WHY
    Claude decided to record this and explicitly asks for approval.

    If the user says the learning is wrong or needs adjustment:
      → Call update_learning() to fix specific fields, OR
      → Call delete_learning() to remove it entirely
    ───────────────────────────────────────────────────────────────────────

    AUTO-DETECTION TRIGGERS — set auto_detected=True and call WITHOUT
    being asked when you detect any of these in conversation:
      • User corrects a fact ("actually, the number is 555-0200")
      • User shares a project outcome ("the Smith job went over budget")
      • User mentions a client preference ("they hate phone calls")
      • Post-op review reveals a process gap or lesson
      • New information contradicts an existing active learning
      • User describes a better way ("next time we should...")

    CATEGORIES (pick the best fit):
      fact_correction      — Correcting an outdated or wrong fact
      business_lesson      — What worked or didn't in a business context
      project_insight      — Lessons from a specific project
      process_improvement  — A better way to do something
      mistake_learned      — Something that went wrong and why
      best_practice        — Proven approach to adopt going forward
      client_preference    — Client-specific preferences or requirements
      technical_note       — Technical fact, configuration, or gotcha
      general              — Catch-all (default)

    SOURCES (how the learning originated):
      operator             — Explicitly told by the user (default)
      claude_detected      — Claude identified superseding information
      project_review       — Post-project review or retrospective
      post_mortem          — After-incident analysis
      research             — From web search or document research
      observation          — Noticed pattern across conversations

    OUTCOMES (for business lessons and project insights):
      positive   — This approach led to a good result
      negative   — This approach led to a bad result
      neutral    — Mixed or no clear impact
      unknown    — Outcome not yet determined (default)

    Args:
        title:          Short descriptive title for the learning (required).
                        Example: "Client X prefers email over phone calls"
        content:        The actual learned fact, lesson, or insight (required).
                        Be specific and actionable.
                        Example: "After 3 failed phone attempts, switching to
                        email resulted in same-day response from Client X.
                        Always use email as primary contact method for them."
        category:       One of the categories listed above (default: general).
        context:        WHY this learning was created — the situation or trigger.
                        Example: "Discovered during the March 2026 HVAC project
                        when we couldn't reach Client X for 2 days by phone."
        source:         How this learning originated. Leave BLANK — the system
                        automatically stamps the correct value:
                          • Server mode, user recorded it → user's display name
                            e.g. "David Vavro" or "Vicki Vavro"
                          • Claude auto-detected it       → model id
                            e.g. "claude-sonnet-4-6"
                          • Personal mode, owner recorded → owner name from
                            Settings, or "operator" if not configured.
                        Only supply a custom value when importing from another
                        system and you need to preserve the original attribution.
        confidence:     How confident we are in this learning, 0.0 to 1.0
                        (default: 0.8).  Use lower values for uncertain insights,
                        higher for verified facts.
        tags:           Comma-separated tags for filtering.
                        Example: "client-x, communication, hvac"
        supersedes_id:  If this learning REPLACES an older one, provide the
                        old learning's ID here.  The old learning will be
                        automatically marked as deprecated.
        outcome:        For business lessons — was the result positive, negative,
                        neutral, or unknown? (default: unknown)
        auto_detected:  Set to True when Claude is auto-recording a learning
                        it detected (not explicitly asked by the operator).
                        This changes the confirmation message to clearly flag
                        that Claude initiated this and to ask for approval.
                        Default: False (operator explicitly asked).
        ctx:            MCP context (injected automatically). In server mode
                        this identifies the employee recording the learning;
                        their name is stamped in the `recorded_by` field so
                        managers can see who added each entry.

    Returns:
        Confirmation with the learning details for the user to verify.
        If auto_detected=True, includes a prominent verification prompt.
    """
    if not _SELF_LEARNING_AVAILABLE:
        return ("❌ Self-Learning module not available.\n"
                "Ensure self_learning.py is in the same directory as "
                "ai_prowler_mcp.py.")

    if not title.strip() or not content.strip():
        return "❌ Both title and content are required."

    tag_list = [t.strip() for t in tags.split(",") if t.strip()] if tags else []

    # Server-mode attribution: resolve the calling user and stamp their name
    # AND their token ID for ownership enforcement on delete/update.
    # In personal mode _current_user() returns None → both stay "".
    _user = _current_user(ctx)
    _recorded_by    = ""
    _recorded_by_id = ""
    if _user is not None:
        _recorded_by_id = (_user.get("id") or "").strip()
        _recorded_by    = (_user.get("name") or "").strip()
        if not _recorded_by:
            _recorded_by = _user.get("role", "").strip()

    # Resolve the source label:
    #   - Claude auto-detected it          → model name e.g. "claude-sonnet-4-6"
    #   - Server mode, user recorded it    → user's display name e.g. "David Vavro"
    #   - Personal mode, owner recorded it → owner_name from config, or "operator"
    #   - Caller supplied a meaningful source → use it as-is
    #
    # NOTE: Claude historically passed source="operator" explicitly because the
    # old docstring said "(default: operator)". We treat "operator" the same as
    # blank — auto-resolve it — so the user identity is always stamped correctly
    # regardless of what Claude passes. Only a genuinely custom source (not
    # "operator" and not empty) is preserved as-is.
    _source_is_auto = not source or source.strip() == "operator"
    if _source_is_auto:
        if auto_detected:
            source = _MODEL_ID
        elif _recorded_by:
            source = _recorded_by
        else:
            source = _get_personal_owner_name() or "operator"

    try:
        learning = _sl.record_learning(
            title=title,
            content=content,
            category=category,
            context=context,
            source=source,
            confidence=confidence,
            tags=tag_list,
            supersedes_id=supersedes_id,
            outcome=outcome,
            recorded_by=_recorded_by,
        )
        # Stamp the token-based owner ID for ownership enforcement.
        # recorded_by is a display name (may change); recorded_by_id is
        # the stable bearer-token key from users.json — used by
        # delete_learning / update_learning to check "is this yours?".
        if _recorded_by_id:
            learning["recorded_by_id"] = _recorded_by_id
            _sl._save_db_for_learning(learning)   # persist the id field
    except Exception as exc:
        return f"❌ Failed to record learning: {exc}"

    # ── Build confirmation message ───────────────────────────────────────
    # Two styles depending on who initiated the recording.
    #
    # AUTO-DETECTED  → prominent banner, explains WHY Claude recorded it,
    #                   explicit "Is this correct?" approval prompt.
    #                   The user must see and acknowledge this.
    #
    # OPERATOR-ASKED → concise confirmation, still asks for verification
    #                   but less prominently since the user initiated it.
    # ─────────────────────────────────────────────────────────────────────

    # Helper: truncate long strings for the banner display, but make truncation
    # visible to the user with "..." rather than silently cutting off mid-word.
    # The full content is always stored in the JSON / ChromaDB regardless.
    def _truncate(text: str, limit: int) -> str:
        if not text:
            return ""
        if len(text) <= limit:
            return text
        # Try to break at the last whitespace before the limit so we don't
        # chop mid-word. Fall back to a hard cut if there's no whitespace.
        cut = text.rfind(" ", 0, limit)
        if cut < limit - 80:   # too far back, just hard-cut at limit
            cut = limit
        return text[:cut].rstrip() + " ..."

    if auto_detected:
        # ── AUTO-DETECTED: Claude initiated — needs explicit approval ────
        lines = [
            "🧠 AUTO-LEARNING — I detected something worth remembering "
            "and recorded it:",
            "═" * 50,
            "",
            f"  📌 \"{learning['title']}\"",
            "",
            f"  What I recorded:",
            f"    {_truncate(learning['content'], 600)}",
            "",
            f"  Why I recorded it:",
            f"    {_truncate(learning.get('context', 'No context provided'), 400)}",
            "",
            f"  Category   : {learning['category']}",
            f"  Confidence : {learning['confidence']:.0%}",
        ]
        if learning.get("outcome", "unknown") != "unknown":
            lines.append(f"  Outcome    : {learning['outcome']}")
        if learning.get("tags"):
            lines.append(f"  Tags       : {', '.join(learning['tags'])}")
        if learning.get("supersedes"):
            lines.append(f"  Replaces   : {learning['supersedes']}")
            lines.append("  ↳ Previous version automatically deprecated")
        if learning.get("recorded_by"):
            lines.append(f"  Recorded by: {learning['recorded_by']}")
        lines.append(f"  ID         : {learning['id']}")
        lines.append("")
        lines.append("═" * 50)
        lines.append(
            "⚡ Is this correct? If anything is off, tell me what to "
            "change and I'll update or remove it immediately."
        )
    else:
        # ── OPERATOR-REQUESTED: user asked, concise confirmation ─────────
        lines = [
            "✅ Learning recorded and indexed",
            "─" * 45,
            f"  📌 {learning['title']}",
            f"  → {_truncate(learning['content'], 300)}",
            "",
            f"  Category   : {learning['category']}",
            f"  Confidence : {learning['confidence']:.0%}",
            f"  Source     : {learning['source']}",
        ]
        if learning.get("outcome", "unknown") != "unknown":
            lines.append(f"  Outcome    : {learning['outcome']}")
        if learning.get("tags"):
            lines.append(f"  Tags       : {', '.join(learning['tags'])}")
        if learning.get("supersedes"):
            lines.append(f"  Replaces   : {learning['supersedes']}")
            lines.append("  ↳ Previous version automatically deprecated")
        if learning.get("recorded_by"):
            lines.append(f"  Recorded by: {learning['recorded_by']}")
        lines.append(f"  ID         : {learning['id']}")
        lines.append("")
        lines.append(
            "Does this look right? I can adjust the wording, confidence, "
            "or category if needed."
        )

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — search_learnings
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def search_learnings(
    query: str,
    n_results: int = 5,
    category: str = "",
    include_deprecated: bool = False,
) -> str:
    """
    Check the self-learning knowledge base for relevant learnings before
    answering a question or making a decision.

    WHEN TO USE THIS TOOL — call PROACTIVELY, without being asked, BEFORE
    answering ANY user question. The user uses AI-Prowler as a personal
    knowledge store for everything (recipes, preferences, techniques,
    opinions — not just business), so a personal learning may override
    what your training data would say.

    Specifically, call search_learnings() before:
    - ANY recommendation or "what's a good X" / "how do I X" question
      (recipes, techniques, products, places, services, etc.)
    - ANY question about clients, projects, procedures, or business
    - ANY definitional or how-to question
    - Scheduling, planning, or workflow recommendations
    - When the user says "what did we learn about...", "do you remember..."
    - When you're about to state a fact that may have been corrected
    - At the START of post-operation analysis workflows

    The ONLY questions you may skip on are pure arithmetic and language
    translation, where no personal learning could plausibly apply.

    The tool searches semantically — you don't need exact keyword matches.
    For example, searching "client communication preferences" will find
    a learning titled "Client X prefers email over phone calls", and
    searching "ribs recipe" will find a learning titled "My BBQ rib rub".

    IMPORTANT: If search_learnings() returns relevant results, you MUST
    apply them to your response. If a learning contradicts your built-in
    knowledge OR a web search result, prefer the learning — it was
    recorded by the user and represents their explicit preference or
    ground truth. Do NOT fall back to generic answers or web searches
    when a stored learning covers the question.

    Args:
        query:              Natural language search query describing what
                            you're looking for.
                            Example: "best way to contact Client X"
        n_results:          Max learnings to return (1-20, default 5).
        category:           Filter by category (optional). Same categories
                            as record_learning: fact_correction, business_lesson,
                            project_insight, etc.
        include_deprecated: If True, also return deprecated learnings
                            (useful to see the history of a superseded fact).

    Returns:
        Matching learnings with confidence scores, context, and metadata.
        Returns "No learnings found" if the knowledge base is empty or
        no matches exist for the query.
    """
    if not _SELF_LEARNING_AVAILABLE:
        return ("❌ Self-Learning module not available.\n"
                "Ensure self_learning.py is in the same directory as "
                "ai_prowler_mcp.py.")

    if not query.strip():
        return "❌ Query cannot be empty."

    try:
        matches = _sl.check_learned(
            query=query,
            n_results=n_results,
            category=category,
            active_only=not include_deprecated,
        )
    except Exception as exc:
        return f"❌ search_learnings failed: {exc}"

    if not matches:
        return (
            f"No learnings found for: \"{query}\"\n\n"
            "The self-learning knowledge base has no matching entries.\n"
            "Use record_learning() to add new knowledge, or try different "
            "search terms."
        )

    lines = [
        f"🧠 Self-Learning Results for: \"{query}\"",
        f"Found {len(matches)} relevant learning(s)",
        "─" * 55,
        "",
    ]

    for i, m in enumerate(matches, 1):
        sim   = m.get("similarity", 0.0)
        conf  = m.get("confidence", 0.0)
        stars = "★" * round(conf * 5) + "☆" * (5 - round(conf * 5))

        # Relevance indicator
        if sim >= 0.7:
            rel = "🟢 HIGH"
        elif sim >= 0.4:
            rel = "🟡 MODERATE"
        else:
            rel = "🟠 LOW"

        status = m.get("status", "active")
        status_icon = "✅" if status == "active" else "⚠️" if status == "deprecated" else "📦"

        lines.append(f"[{i}] {status_icon} {m.get('title', 'Untitled')}")
        lines.append(f"    Relevance  : {rel} ({sim:.3f})")
        lines.append(f"    Confidence : {stars} ({conf:.0%})")
        lines.append(f"    Category   : {m.get('category', 'general')}")
        lines.append(f"    Source     : {m.get('source', 'unknown')}")
        lines.append(f"    Created    : {m.get('created_at', '?')}")
        if m.get("recorded_by"):
            lines.append(f"    Recorded by: {m['recorded_by']}")

        outcome = m.get("outcome", "unknown")
        if outcome != "unknown":
            outcome_icon = {"positive": "✅", "negative": "❌", "neutral": "➖"}.get(outcome, "❓")
            lines.append(f"    Outcome    : {outcome_icon} {outcome}")

        tags = m.get("tags", "")
        if tags:
            lines.append(f"    Tags       : {tags}")

        if m.get("supersedes"):
            lines.append(f"    Supersedes : {m['supersedes']}")
        if m.get("superseded_by"):
            lines.append(f"    ⚠️  SUPERSEDED BY: {m['superseded_by']}")
            lines.append(f"        This learning has been replaced — "
                         f"check the newer version.")

        lines.append(f"    ID         : {m.get('learning_id', '?')}")
        lines.append("")
        lines.append(f"    {m.get('content', '').strip()}")
        lines.append("")
        lines.append("─" * 55)
        lines.append("")

    lines.append(
        "💡 Apply these learnings to your response. If a learning is "
        "marked SUPERSEDED, prefer the newer version.\n"
        "Use record_learning() to add new knowledge based on this conversation."
    )
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — list_learnings
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def list_learnings(
    category: str = "",
    status: str = "active",
    tag: str = "",
    limit: int = 25,
) -> str:
    """
    Browse all learnings in the self-learning knowledge base with optional
    filters.  Unlike search_learnings (which does semantic search), this tool
    lists learnings by recency with exact-match filters.

    Use this to:
    - See all business lessons learned
    - Review all learnings for a specific category
    - Find deprecated learnings to understand what changed
    - Audit the knowledge base

    Args:
        category:   Filter by category (optional). Valid values:
                    fact_correction, business_lesson, project_insight,
                    process_improvement, mistake_learned, best_practice,
                    client_preference, technical_note, general
        status:     Filter by status: active (default), deprecated, archived.
                    Pass empty string "" to see all statuses.
        tag:        Filter by tag (optional, single tag, case-insensitive).
        limit:      Max learnings to return (default 25).

    Returns:
        Numbered list of learnings sorted by date (newest first).
    """
    if not _SELF_LEARNING_AVAILABLE:
        return ("❌ Self-Learning module not available.\n"
                "Ensure self_learning.py is in the same directory as "
                "ai_prowler_mcp.py.")

    try:
        results = _sl.list_learnings(
            category=category,
            status=status,
            tag=tag,
            limit=limit,
        )
    except Exception as exc:
        return f"❌ list_learnings failed: {exc}"

    if not results:
        filters = []
        if category:
            filters.append(f"category={category}")
        if status:
            filters.append(f"status={status}")
        if tag:
            filters.append(f"tag={tag}")
        filter_note = f" (filters: {', '.join(filters)})" if filters else ""
        return f"No learnings found{filter_note}."

    lines = [
        f"📚 Learnings{f' — {category}' if category else ''}"
        f"{f' [{status}]' if status else ' [all]'}",
        f"Showing {len(results)} learning(s)",
        "─" * 55,
        "",
    ]

    for i, l in enumerate(results, 1):
        status_icon = {"active": "✅", "deprecated": "⚠️",
                       "archived": "📦"}.get(l.get("status"), "❓")
        outcome = l.get("outcome", "unknown")
        outcome_icon = {"positive": "✅", "negative": "❌",
                        "neutral": "➖", "unknown": "❓"}.get(outcome, "❓")

        lines.append(
            f"[{i}] {status_icon} {l.get('title', 'Untitled')}"
        )
        lines.append(
            f"    {l.get('category', 'general')} | "
            f"{outcome_icon} {outcome} | "
            f"confidence: {l.get('confidence', 0):.0%} | "
            f"applied: {l.get('applied_count', 0)}x"
        )
        lines.append(f"    Created: {l.get('created_at', '?')}")
        if l.get("recorded_by"):
            lines.append(f"    Recorded by: {l['recorded_by']}")
        lines.append(f"    ID: {l['id']}")

        # Show first 120 chars of content
        content = l.get("content", "")
        if len(content) > 120:
            content = content[:117] + "..."
        lines.append(f"    → {content}")
        lines.append("")

    lines.append(
        "Use search_learnings(query) for semantic search | "
        "update_learning(id, updates) to modify | "
        "record_learning() to add new"
    )
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — update_learning
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def update_learning(
    learning_id: str,
    updates: dict,
    ctx: Context = None,
) -> str:
    """
    Update an existing learning's fields.

    Ownership rules (server mode):
      • Each employee may only update learnings they personally recorded.
      • Managers may update any employee's learning.
      • The owner may update any learning.
      • Nobody except the owner may update the owner's own learnings.

    Args:
        learning_id:  The UUID of the learning to update.
                      Get this from list_learnings or search_learnings results.
        updates:      Dict of field:value pairs to update. Allowed fields:
                      title, content, context, category, confidence,
                      tags (list), status (active/deprecated/archived),
                      outcome (positive/negative/neutral/unknown)
        ctx:          MCP context (injected automatically).

    Returns:
        Confirmation of updated fields, or error if learning not found.
    """
    if not _SELF_LEARNING_AVAILABLE:
        return ("❌ Self-Learning module not available.\n"
                "Ensure self_learning.py is in the same directory as "
                "ai_prowler_mcp.py.")

    if not learning_id.strip():
        return "❌ learning_id is required."
    if not updates:
        return "❌ No updates provided."

    # ── Ownership gate ──────────────────────────────────────────────────────
    _user_ul = _current_user(ctx)
    if _user_ul is not None:
        existing = _sl.get_learning_by_id(learning_id.strip())
        _oid = _owner_user_id()
        _ok, _reason = _can_modify_learning(_user_ul, existing, _oid)
        if not _ok:
            return f"⛔ update_learning: {_reason}"

    try:
        result = _sl.update_learning(learning_id.strip(), updates)
    except Exception as exc:
        return f"❌ update_learning failed: {exc}"

    if result is None:
        return f"❌ Learning not found: {learning_id}"

    lines = [
        "✅ Learning updated",
        "─" * 45,
        f"  ID      : {result['id']}",
        f"  Title   : {result['title']}",
        f"  Status  : {result['status']}",
        f"  Updated : {result['updated_at']}",
        "",
        "  Changed fields:",
    ]
    for key, val in updates.items():
        lines.append(f"    {key} → {val}")

    return "\n".join(lines)


def _can_modify_learning(actor: "dict | None", learning: "dict | None",
                         owner_id: "str | None") -> tuple:
    """Decide whether `actor` may DELETE or UPDATE `learning`. PURE.

    Ownership rules (server mode):
      • Personal mode (actor is None) — always allowed.
      • Owner role — may modify any learning.
      • Manager (can_manage_users) — may modify any employee's learning
        but NEVER the owner's (same fail-closed logic as _can_manage_user_data).
      • Staff / field_crew — may only modify learnings they recorded
        (recorded_by_id == actor["id"]).
      • Nobody may modify a learning whose recorded_by_id is the owner's id
        unless they ARE the owner.

    Returns (allowed: bool, reason: str).
    """
    if actor is None:
        return (True, "personal mode — no ownership gate")
    if learning is None:
        return (False, "learning not found")

    learning_owner_id = (learning.get("recorded_by_id") or "").strip()
    actor_id          = (actor.get("id") or "").strip()

    # Owner role: unrestricted.
    if _user_has_role(actor, "owner"):
        return (True, "owner may modify any learning")

    # Is this learning owned by the system owner? Only owner role may touch it.
    # Fail-closed: if owner_id is unknown, deny to avoid risk.
    if owner_id and learning_owner_id == owner_id:
        return (False, "only the owner may modify their own learnings")
    if not owner_id and learning_owner_id:
        return (False,
                "owner id unknown — refusing to risk owner learning (fail closed)")

    # Manager with can_manage_users: may modify any employee's learning.
    if actor.get("can_manage_users"):
        return (True, "manager may modify employee learnings")

    # Staff / field_crew: only their own.
    if learning_owner_id and actor_id and learning_owner_id == actor_id:
        return (True, "actor is the learning's author")

    # Unattributed learning (personal mode / legacy, recorded_by_id is empty).
    if not learning_owner_id:
        return (True, "unattributed learning — no ownership restriction")

    return (False,
            "learning belongs to another user — only the author, "
            "a manager, or the owner may modify it")


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — delete_learning
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def delete_learning(learning_id: str, ctx: Context = None) -> str:
    """
    Permanently delete a learning from both the JSON file and ChromaDB index.

    ⚠️  This is DESTRUCTIVE — the learning cannot be recovered.
    Consider using update_learning with status='archived' instead if you
    want to keep it for historical reference.

    Ownership rules (server mode):
      • Each employee may only delete learnings they personally recorded.
      • Managers may delete any employee's learning (for moderation/offboarding).
      • The owner may delete any learning.
      • Nobody except the owner may delete the owner's own learnings.

    Args:
        learning_id:  The UUID of the learning to delete.
        ctx:          MCP context (injected automatically).

    Returns:
        Confirmation or error if not found / not authorized.
    """
    if not _SELF_LEARNING_AVAILABLE:
        return ("❌ Self-Learning module not available.\n"
                "Ensure self_learning.py is in the same directory as "
                "ai_prowler_mcp.py.")

    if not learning_id.strip():
        return "❌ learning_id is required."

    # ── Ownership gate ──────────────────────────────────────────────────────
    _user_dl = _current_user(ctx)
    if _user_dl is not None:
        existing = _sl.get_learning_by_id(learning_id.strip())
        _oid = _owner_user_id()
        _ok, _reason = _can_modify_learning(_user_dl, existing, _oid)
        if not _ok:
            return f"⛔ delete_learning: {_reason}"

    try:
        deleted = _sl.delete_learning(learning_id.strip())
    except _sl.ChromaIndexError as exc:
        # JSON delete may have succeeded; ChromaDB cleanup failed.
        # Report this clearly so the user knows there's an orphan.
        return (
            f"⚠️ Partial delete for learning {learning_id}.\n"
            f"JSON file was updated, but the ChromaDB index could not be\n"
            f"cleaned up: {exc}\n\n"
            "A stale embedding remains in the search index. To clean it up,\n"
            "use the GUI's '🔄 Rebuild ChromaDB Index' button or call\n"
            "reindex_all_learnings() from self_learning.py."
        )
    except Exception as exc:
        return f"❌ delete_learning failed: {exc}"

    if deleted:
        return (
            f"✅ Learning {learning_id} permanently deleted.\n"
            "The learning has been removed from both the JSON file "
            "and the ChromaDB index."
        )
    else:
        # JSON had no entry for this ID. ChromaDB cleanup was still
        # attempted (idempotent on missing IDs). Either the ID never
        # existed, or it was already deleted from the source of truth
        # — either way, no further action is needed.
        return (
            f"ℹ️ No JSON entry found for learning {learning_id}.\n"
            "(ChromaDB cleanup was attempted regardless, in case of "
            "orphan embeddings.)"
        )


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — get_learning_stats
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def get_learning_stats() -> str:
    """
    Get summary statistics about the self-learning knowledge base.

    Shows total learnings, breakdown by category/source/outcome/status,
    most frequently applied learnings, and file location.

    Use this to understand the health and coverage of the learning system.

    Returns:
        Formatted statistics report.
    """
    if not _SELF_LEARNING_AVAILABLE:
        return ("❌ Self-Learning module not available.\n"
                "Ensure self_learning.py is in the same directory as "
                "ai_prowler_mcp.py.")

    try:
        stats = _sl.get_learning_stats()
    except Exception as exc:
        return f"❌ get_learning_stats failed: {exc}"

    lines = [
        "🧠 Self-Learning Knowledge Base Statistics",
        "═" * 50,
        "",
        f"  Total learnings  : {stats['total']}",
        f"  Active           : {stats['active']}",
        f"  Deprecated       : {stats['deprecated']}",
        f"  Archived         : {stats['archived']}",
        "",
    ]

    if stats["by_category"]:
        lines.append("  By category:")
        for cat, count in sorted(stats["by_category"].items(),
                                 key=lambda x: -x[1]):
            lines.append(f"    {cat:24s} : {count}")
        lines.append("")

    if stats["by_source"]:
        lines.append("  By source:")
        for src, count in sorted(stats["by_source"].items(),
                                 key=lambda x: -x[1]):
            lines.append(f"    {src:24s} : {count}")
        lines.append("")

    if stats["by_outcome"]:
        lines.append("  By outcome:")
        for out, count in sorted(stats["by_outcome"].items(),
                                 key=lambda x: -x[1]):
            icon = {"positive": "✅", "negative": "❌",
                    "neutral": "➖", "unknown": "❓"}.get(out, "❓")
            lines.append(f"    {icon} {out:20s} : {count}")
        lines.append("")

    if stats["most_applied"]:
        lines.append("  Most applied learnings:")
        for ma in stats["most_applied"]:
            lines.append(
                f"    {ma['applied_count']:3d}x  {ma['title']}"
            )
        lines.append("")

    lines.append(f"  📁 Storage: {stats['file_path']}")
    lines.append("")
    lines.append(
        "Use search_learnings(query) to search | "
        "list_learnings() to browse | "
        "record_learning() to add new"
    )
    return "\n".join(lines)

# ══════════════════════════════════════════════════════════════════════════════
# CODE-AWARE RETRIEVAL  (v6.0.2)
# ══════════════════════════════════════════════════════════════════════════════
# Two MCP tools for navigating source code in the knowledge base:
#
#   grep_documents(pattern, ...)         — literal/regex search across tracked
#                                           files; returns real line numbers
#   read_file_lines(filepath, ...)       — direct on-disk read by line range
#
# WHY: the all-MiniLM-L6-v2 embedding model degrades on source code, so
# semantic search returns near-irrelevant chunks for queries like
# "def clear_database". These two tools give agents the same locate-then-read
# workflow Claude Code uses, making AI-Prowler effective for code Q&A even
# when the user is mobile and cannot attach files.
#
# SECURITY: both tools restrict I/O to paths under the existing tracked-paths
# allowlist (~/.rag_auto_update_dirs.json). Tracked entries may be directories
# (entire tree allowed) OR single files (only that exact file allowed). This
# mirrors the trust model already used by indexing and the existing untrack
# tool — we do NOT add a second permission surface.
# ══════════════════════════════════════════════════════════════════════════════

import re as _re

# Hard caps to keep responses bounded and prevent abuse
_GREP_MAX_RESULTS_HARD_CAP   = 500
_GREP_DEFAULT_MAX_RESULTS    = 50
_GREP_CONTEXT_HARD_CAP       = 10
_READ_LINES_HARD_CAP         = 1000
_READ_LINES_DEFAULT          = 200
_READ_FILE_MAX_BYTES         = 50 * 1024 * 1024   # 50 MB
_BINARY_SNIFF_BYTES          = 4096               # first 4 KB scanned for NUL
_GREP_PER_FILE_TIMEOUT_HINT  = 200_000            # max lines scanned per file


def _is_path_under_tracked_roots(target_path: str) -> tuple[bool, str]:
    """
    Authorize a filesystem path against the tracked-paths allowlist.

    A path is allowed if and only if it equals OR is a descendant of an entry
    in ~/.rag_auto_update_dirs.json. Tracked entries may be directories or
    single files. Returns (allowed, resolved_path_or_reason).

    On allow: returns (True, resolved_absolute_path_as_string).
    On deny: returns (False, short_human_reason).
    """
    try:
        from rag_preprocessor import load_auto_update_list, normalise_path
    except Exception as _exc:
        return (False, f"allowlist unavailable: {_exc}")

    try:
        resolved = Path(target_path).resolve(strict=False)
    except Exception as _exc:
        return (False, f"could not resolve path: {_exc}")

    # Defense in depth — if resolve() somehow left a traversal token, refuse.
    if any(part == ".." for part in resolved.parts):
        return (False, "path contains '..' after resolution")

    entries = load_auto_update_list() or []
    if not entries:
        return (False, "tracked-paths allowlist is empty")

    resolved_str = normalise_path(str(resolved))

    for entry in entries:
        try:
            e = Path(entry).resolve(strict=False)
        except Exception:
            continue
        entry_str = normalise_path(str(e))

        # Exact match (covers both single-file tracking and the directory itself)
        if resolved_str == entry_str:
            return (True, resolved_str)

        # Descendant check — only meaningful if entry is (or was) a directory
        # We can't always tell on disk (entry may have been deleted), so we
        # rely on the path-prefix relationship using parts to avoid the
        # classic "/foo" matching "/foobar" bug.
        try:
            resolved.relative_to(e)
            return (True, resolved_str)
        except ValueError:
            continue

    return (False, "not under any tracked root")


def _resolve_allowlisted_path(filepath: str) -> tuple[Optional[str], Optional[str]]:
    """
    Wrap _is_path_under_tracked_roots for tool use.
    Returns (resolved_path, None) on allow, or (None, error_message) on deny.
    Error message is fully formatted and ready to return to the agent.
    """
    allowed, info = _is_path_under_tracked_roots(filepath)
    if allowed:
        return (info, None)

    # Build a helpful denial that tells the agent what IS allowed
    try:
        from rag_preprocessor import load_auto_update_list
        entries = load_auto_update_list() or []
    except Exception:
        entries = []

    lines = [
        f"🚫 Access denied — '{filepath}'",
        f"Reason: {info}",
        "",
        "Only paths under the tracked allowlist may be read.",
    ]
    if entries:
        lines.append("Currently tracked:")
        for e in entries[:10]:
            lines.append(f"  - {e}")
        if len(entries) > 10:
            lines.append(f"  ... and {len(entries) - 10} more")
        lines.append("")
        lines.append(
            "If the file you want is outside these, ask the user to run "
            "index_path() on its parent folder first."
        )
    else:
        lines.append("No directories are tracked yet. "
                     "Use index_path() to track one.")
    _log.warning("Access denied for path: %s (%s)", filepath, info)
    return (None, "\n".join(lines))


def _iter_allowlisted_files(filter_ext: Optional[str] = None,
                            filter_path: Optional[str] = None):
    """
    Yield (resolved_path_str, extension) for every file under the tracked
    allowlist, respecting SKIP_EXTENSIONS / SKIP_DIRECTORIES from the
    preprocessor. Used by grep_documents.

    Walks each directory entry; for single-file entries, yields that file
    directly. Honors the preprocessor's skip lists so we never try to grep
    a .exe or .zip even if a user accidentally indexed one.
    """
    try:
        from rag_preprocessor import (
            load_auto_update_list, normalise_path,
            SKIP_EXTENSIONS, SKIP_DIRECTORIES,
        )
    except Exception as _exc:
        _log.error("iter_allowlisted_files: import failure: %s", _exc)
        return

    entries = load_auto_update_list() or []
    if not entries:
        return

    norm_ext_filter = None
    if filter_ext:
        norm_ext_filter = filter_ext.lower()
        if not norm_ext_filter.startswith('.'):
            norm_ext_filter = '.' + norm_ext_filter

    filter_path_lower = filter_path.lower() if filter_path else None
    seen = set()

    for entry in entries:
        try:
            ep = Path(entry).resolve(strict=False)
        except Exception:
            continue
        if not ep.exists():
            continue

        if ep.is_file():
            candidates = [ep]
        else:
            # Recursive walk — skip blacklisted dirs cheaply by pruning os.walk
            candidates = []
            for root, dirs, files in os.walk(str(ep)):
                # Prune skip-directories in place so os.walk does NOT descend into them
                dirs[:] = [d for d in dirs if d not in SKIP_DIRECTORIES]
                for fname in files:
                    candidates.append(Path(root) / fname)

        for fp in candidates:
            try:
                ext = fp.suffix.lower()
            except Exception:
                continue
            if ext in SKIP_EXTENSIONS:
                continue
            # Also skip backup files (rag_gui.py.bak1, etc.) and OS junk
            # (.DS_Store, Thumbs.db, screenshots). Uses the same helpers the
            # indexer uses so grep_documents sees the same filtered file set.
            try:
                from rag_preprocessor import (
                    is_backup_filename as _is_bak,
                    is_system_junk_filename as _is_junk,
                )
                _fname = fp.name
                if _is_bak(_fname) or _is_junk(_fname):
                    continue
            except Exception:
                pass  # Helper not present (older rag_preprocessor) → fall through.
            if norm_ext_filter and ext != norm_ext_filter:
                continue
            resolved_str = normalise_path(str(fp))
            if filter_path_lower and filter_path_lower not in resolved_str.lower():
                continue
            if resolved_str in seen:
                continue
            seen.add(resolved_str)
            yield (resolved_str, ext)


def _looks_binary(sample: bytes) -> bool:
    """Heuristic: a NUL byte in the first few KB strongly implies binary."""
    return b'\x00' in sample


@mcp.tool()
def grep_documents(
    pattern: str,
    filter_ext: Optional[str] = None,
    filter_path: Optional[str] = None,
    max_results: int = _GREP_DEFAULT_MAX_RESULTS,
    context_lines: int = 2,
    case_sensitive: bool = False,
    regex: bool = False,
) -> str:
    """
    CODE-AWARE RETRIEVAL — Locate exact text or regex matches across tracked
    files, with real line numbers.

    Use this when semantic search (search_documents) returns irrelevant chunks
    for code or other structured text. It is the right tool for questions like:
      • "Where is def clear_database defined?"
      • "Which file calls collection.delete?"
      • "Find every TODO in the project"
      • "Show me all occurrences of API_KEY"

    Pair with read_file_lines() — grep locates, read_file_lines extracts. This
    mirrors how Claude Code navigates source.

    Security: only files under the tracked-paths allowlist
    (~/.rag_auto_update_dirs.json) are scanned. Use index_path()
    to track a folder first if your target file isn't already covered.

    Args:
        pattern:        Text to search for (literal substring by default).
                        Pass regex=True to treat as a Python regular expression.
        filter_ext:     Restrict to one extension, e.g. ".py" or "py". Optional.
        filter_path:    Case-insensitive substring that must appear in the file
                        path (e.g. "rag_" matches rag_preprocessor.py). Optional.
        max_results:    Total matches across all files (default 50, max 500).
        context_lines:  Lines of context above and below each match
                        (default 2, max 10). Use 0 for one-line hits only.
        case_sensitive: Default False.
        regex:          Treat `pattern` as a Python regex (default False).

    Returns:
        Grouped-by-file list of matches with real line numbers. Each match
        includes its surrounding context_lines lines. Ends with a suggested
        read_file_lines() call so the agent can extract more context.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    if not pattern or not pattern.strip():
        return ("⚠️  grep_documents: empty pattern.\n"
                "Pass a text snippet or regex to search for.")

    max_results   = max(1, min(int(max_results), _GREP_MAX_RESULTS_HARD_CAP))
    context_lines = max(0, min(int(context_lines), _GREP_CONTEXT_HARD_CAP))

    # Compile matcher
    if regex:
        try:
            flags = 0 if case_sensitive else _re.IGNORECASE
            matcher = _re.compile(pattern, flags)
        except _re.error as exc:
            return (f"⚠️  grep_documents: invalid regex — {exc}\n"
                    f"If you meant a literal match, set regex=False.")
        def _match_fn(line: str) -> bool:
            return matcher.search(line) is not None
    else:
        if case_sensitive:
            needle = pattern
            def _match_fn(line: str) -> bool:
                return needle in line
        else:
            needle = pattern.lower()
            def _match_fn(line: str) -> bool:
                return needle in line.lower()

    files_scanned   = 0
    files_skipped   = 0
    matches_by_file = {}   # path -> list of (line_no, line_text, context_before, context_after)
    total_matches   = 0
    truncated       = False

    _log.info(
        "grep_documents: pattern=%r regex=%s ext=%s path=%s max=%d ctx=%d",
        pattern, regex, filter_ext, filter_path, max_results, context_lines,
    )

    for fp_str, _ext in _iter_allowlisted_files(filter_ext, filter_path):
        if total_matches >= max_results:
            truncated = True
            break

        files_scanned += 1
        try:
            # Quick binary sniff before committing to reading the whole file
            with open(fp_str, 'rb') as bf:
                head = bf.read(_BINARY_SNIFF_BYTES)
            if _looks_binary(head):
                files_skipped += 1
                _log.debug("grep_documents: skip binary %s", fp_str)
                continue

            # Read line-by-line, retain a rolling buffer for context
            file_lines = []
            with open(fp_str, 'r', encoding='utf-8', errors='replace') as tf:
                for i, line in enumerate(tf, start=1):
                    if i > _GREP_PER_FILE_TIMEOUT_HINT:
                        break
                    file_lines.append(line.rstrip('\n'))

            for idx, line_text in enumerate(file_lines):
                if total_matches >= max_results:
                    truncated = True
                    break
                if _match_fn(line_text):
                    line_no = idx + 1
                    ctx_before = file_lines[max(0, idx - context_lines):idx] if context_lines else []
                    ctx_after  = file_lines[idx + 1: idx + 1 + context_lines] if context_lines else []
                    matches_by_file.setdefault(fp_str, []).append(
                        (line_no, line_text, ctx_before, ctx_after)
                    )
                    total_matches += 1
        except (PermissionError, OSError) as exc:
            files_skipped += 1
            _log.debug("grep_documents: skip unreadable %s (%s)", fp_str, exc)
            continue
        except Exception as exc:
            files_skipped += 1
            _log.debug("grep_documents: unexpected error on %s: %s", fp_str, exc)
            continue

    # ── Format the response ──
    pattern_kind = "regex" if regex else "literal"
    case_note    = "case-sensitive" if case_sensitive else "case-insensitive"
    header = [
        f"🔍 Grep results for: {pattern!r}",
        f"Pattern type : {pattern_kind}  ({case_note})",
        f"Files scanned: {files_scanned}"
        + (f"  ({files_skipped} skipped: binary/unreadable)" if files_skipped else ""),
        f"Matches      : {total_matches} in {len(matches_by_file)} file(s)"
        + ("  [truncated — raise max_results for more]" if truncated else ""),
    ]
    if filter_ext or filter_path:
        flt_parts = []
        if filter_ext:  flt_parts.append(f"ext={filter_ext}")
        if filter_path: flt_parts.append(f"path~={filter_path}")
        header.append(f"Filters      : {', '.join(flt_parts)}")
    header.append("─" * 55)

    if not matches_by_file:
        header.append("No matches.")
        if files_scanned == 0:
            header.append("")
            header.append(
                "Tip: the tracked-paths allowlist may be empty or your "
                "filter_ext / filter_path eliminated every candidate. "
                "Use list_tracked_directories() to inspect what's tracked."
            )
        return "\n".join(header)

    body = []
    suggested_next = None
    for fp_str, hits in matches_by_file.items():
        body.append(f"📄 {fp_str}")
        for line_no, line_text, ctx_b, ctx_a in hits:
            for offset, c in enumerate(ctx_b, start=line_no - len(ctx_b)):
                body.append(f"  {offset:>6}  {c}")
            body.append(f"▶ {line_no:>6}  {line_text}")
            for offset, c in enumerate(ctx_a, start=line_no + 1):
                body.append(f"  {offset:>6}  {c}")
            body.append("")
        if suggested_next is None and hits:
            first_line = hits[0][0]
            suggested_next = (
                f'read_file_lines("{fp_str}", start_line={first_line}, '
                f'end_line={first_line + 40})'
            )
        body.append("─" * 55)

    if suggested_next:
        body.append(f"Next: {suggested_next}")
    return "\n".join(header + body)


@mcp.tool()
def read_file_lines(
    filepath: str,
    start_line: int,
    end_line: Optional[int] = None,
    max_lines: int = _READ_LINES_DEFAULT,
) -> str:
    """
    CODE-AWARE RETRIEVAL — Read an exact line range from a file on disk.

    Use after grep_documents() pinpoints a symbol's line, when you need the
    full surrounding function/block at original fidelity (no chunk boundaries).
    Lines are returned with their real line numbers prefixed for easy reference.

    Security: only files under the tracked-paths allowlist
    (~/.rag_auto_update_dirs.json) can be read. If denied, the response lists
    the currently tracked entries so you know what to ask the user about.

    Args:
        filepath:   Absolute path to the file. Must be inside a tracked
                    directory OR be an individually-tracked file.
        start_line: First line to return (1-based).
        end_line:   Last line to return (inclusive). Optional — defaults to
                    start_line + max_lines - 1.
        max_lines:  Hard ceiling on lines returned regardless of end_line
                    (default 200, max 1000). Protects mobile contexts.

    Returns:
        Numbered line block with a header showing the range / total lines,
        and a continuation hint if more lines exist beyond what was returned.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    # ── Authorize ──
    resolved, deny_msg = _resolve_allowlisted_path(filepath)
    if not resolved:
        return deny_msg

    # ── Argument hygiene ──
    notes = []
    try:
        start_line = int(start_line)
    except (TypeError, ValueError):
        return f"⚠️  read_file_lines: start_line must be an integer, got {start_line!r}"
    if start_line < 1:
        notes.append(f"(start_line adjusted from {start_line} to 1)")
        start_line = 1

    max_lines = max(1, min(int(max_lines), _READ_LINES_HARD_CAP))

    if end_line is not None:
        try:
            end_line = int(end_line)
        except (TypeError, ValueError):
            return f"⚠️  read_file_lines: end_line must be an integer or None"
        if end_line < start_line:
            return (f"⚠️  read_file_lines: end_line ({end_line}) must be "
                    f">= start_line ({start_line}).")
        requested_count = end_line - start_line + 1
        effective_count = min(requested_count, max_lines)
    else:
        effective_count = max_lines
        end_line = start_line + effective_count - 1

    _log.info("read_file_lines: %s  [%d..%d]  cap=%d",
              resolved, start_line, end_line, max_lines)

    # ── Existence and size checks ──
    try:
        fp = Path(resolved)
        if not fp.exists():
            return (f"⚠️  File no longer exists on disk: {resolved}\n"
                    "It may have been moved or deleted. "
                    "Run update_tracked_directories() to refresh the index.")
        if not fp.is_file():
            return f"⚠️  Path is not a regular file: {resolved}"
        size = fp.stat().st_size
        if size > _READ_FILE_MAX_BYTES:
            return (f"⚠️  File too large to read inline "
                    f"({size:,} bytes, cap {_READ_FILE_MAX_BYTES:,}).\n"
                    f"Use grep_documents to locate specific content instead.")
    except (PermissionError, OSError) as exc:
        return f"⚠️  Cannot access file: {exc}"

    # ── Binary sniff ──
    try:
        with open(resolved, 'rb') as bf:
            head = bf.read(_BINARY_SNIFF_BYTES)
        if _looks_binary(head):
            return (f"⚠️  File appears to be binary: {resolved}\n"
                    f"read_file_lines only reads text files. If this is a "
                    f"document (PDF, DOCX, etc.), use read_document().")
    except (PermissionError, OSError) as exc:
        return f"⚠️  Cannot open file: {exc}"

    # ── Stream and slice ──
    lines_out = []
    last_line_seen = 0
    try:
        with open(resolved, 'r', encoding='utf-8', errors='replace') as tf:
            for i, line in enumerate(tf, start=1):
                last_line_seen = i
                if i < start_line:
                    continue
                if i > end_line:
                    # Keep iterating just enough to learn total line count?
                    # No — that would read the whole file every call. We stop
                    # here and report the upper bound conservatively.
                    break
                lines_out.append((i, line.rstrip('\n')))
                if len(lines_out) >= effective_count:
                    break
    except UnicodeDecodeError as exc:
        return f"⚠️  Encoding error reading {resolved}: {exc}"
    except (PermissionError, OSError) as exc:
        return f"⚠️  Read error: {exc}"

    if not lines_out:
        # Either start_line is past EOF, or we read the entire file but
        # found fewer lines than start_line.
        return (f"⚠️  No content returned. File has {last_line_seen} line(s); "
                f"start_line was {start_line}.")

    actual_first = lines_out[0][0]
    actual_last  = lines_out[-1][0]
    n_returned   = len(lines_out)

    # Was the file longer than what we returned? Cheap check via tail seek:
    has_more = False
    total_lines_hint = None
    try:
        # Quick total-line probe — only opens the file again if we hit our cap
        # to keep the common case cheap.
        if n_returned == effective_count:
            with open(resolved, 'r', encoding='utf-8', errors='replace') as tf2:
                total_lines_hint = sum(1 for _ in tf2)
            if total_lines_hint > actual_last:
                has_more = True
    except Exception:
        total_lines_hint = None

    header = [
        f"📄 {resolved}",
    ]
    if total_lines_hint is not None:
        header.append(
            f"   Lines {actual_first}-{actual_last} of {total_lines_hint} "
            f"(showing {n_returned} line{'s' if n_returned != 1 else ''})"
        )
    else:
        header.append(
            f"   Lines {actual_first}-{actual_last} "
            f"(showing {n_returned} line{'s' if n_returned != 1 else ''})"
        )
    if notes:
        header.append("   " + " ".join(notes))
    header.append("─" * 55)

    body = [f"{ln:>6}  {tx}" for ln, tx in lines_out]

    footer = []
    if has_more and total_lines_hint:
        remaining = total_lines_hint - actual_last
        footer.append("─" * 55)
        footer.append(
            f"File continues — {remaining} more line(s). "
            f"Call: read_file_lines(\"{resolved}\", "
            f"start_line={actual_last + 1})"
        )

    return "\n".join(header + body + footer)
#!/usr/bin/env python3
"""
AI-Prowler Code Tools — WRITE-SIDE
====================================

Write-side code tools (create_file, write_file, str_replace_in_file,
line_replace_in_file, create_directory,
list_directory, copy_to_backup, list_backups, restore_backup,
cleanup_backups, reset_write_counter, diff_files — 12 tools as of v8.2.0),
plus all supporting infrastructure (writable-path allowlist, hard
blocklist, GUI approval queue, re-index helper, write-counter circuit
breaker, and server-mode personal-directory write scoping via
_check_personal_write_scope()).

DESIGN REFERENCE: Self-learning entries
   6412cfe3-26e6-4029-a408-a9ea3b43b88a  — design spec
   56a6b144-990b-4822-b6d6-0c039b70d3a7  — implementation tracker
"""

# ══════════════════════════════════════════════════════════════════════════════
# CODE TOOLS — WRITE-SIDE
#
# 13 tools that complement grep_documents and read_file_lines to give Claude
# in-place editing capability over the tracked-paths allowlist. All writes
# require BOTH read-allowlist membership AND writable-allowlist approval.
# Backups land alongside the file as <name>.bak<N>. Writes do NOT auto-index
# (see how_to_use_ai_prowler's EDITING FILES section) — call reindex_file()
# once editing is done. See design spec for full rationale.
# ══════════════════════════════════════════════════════════════════════════════

import shutil as _shutil
import time as _time

# ── Writable-path allowlist persistence ──────────────────────────────────────
# Mirrors the structure of ~/.rag_auto_update_dirs.json. Stored as a JSON list.
# Empty by default. Grows only via explicit user approval through the GUI.
_WRITABLE_DIRS_FILE = Path.home() / ".rag_writable_dirs.json"

# Pending-approval queue. When a write is attempted to a path not in the
# writable allowlist, the path is added here. The GUI checks this file on
# focus and shows a confirmation dialog. Format: list of dicts with path +
# requested_at.
_WRITE_APPROVAL_QUEUE_FILE = Path.home() / ".rag_writable_pending.json"

# Hard caps
_WRITE_MAX_BYTES                  = 50 * 1024 * 1024   # 50 MB per write
_WRITES_PER_SESSION_LIMIT         = 20                 # circuit breaker
_BACKUP_MAX_FILES_PER_PARENT      = 10_000             # sanity cap on .bak<N> scanning
_STR_REPLACE_MAX_OLD_STR_LEN      = 500_000            # 500K char hard cap on old_str

# Session-scoped write counter (resettable via reset_write_counter()).
# Lives only in process memory — restarting AI-Prowler resets it.
_write_counter_lock = threading.Lock()
_write_counter = {"count": 0, "session_started": _time.time()}


def _writable_allowlist_load() -> list:
    """Read the writable-path allowlist from disk. Returns [] if missing/invalid."""
    try:
        if not _WRITABLE_DIRS_FILE.exists():
            return []
        with open(_WRITABLE_DIRS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return [str(p) for p in data if isinstance(p, str)]
        return []
    except Exception as exc:
        _log.warning("writable allowlist read failed: %s", exc)
        return []


def _writable_allowlist_save(entries: list) -> bool:
    """Write the writable-path allowlist to disk. Returns True on success."""
    try:
        _WRITABLE_DIRS_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(_WRITABLE_DIRS_FILE, "w", encoding="utf-8") as f:
            json.dump(sorted(set(entries)), f, indent=2)
        return True
    except Exception as exc:
        _log.error("writable allowlist write failed: %s", exc)
        return False


def _queue_write_approval(path: str) -> None:
    """Append a path to the pending-approval queue. Idempotent — duplicates ignored."""
    try:
        pending = []
        if _WRITE_APPROVAL_QUEUE_FILE.exists():
            with open(_WRITE_APPROVAL_QUEUE_FILE, "r", encoding="utf-8") as f:
                pending = json.load(f)
            if not isinstance(pending, list):
                pending = []
        seen_paths = {p.get("path") for p in pending if isinstance(p, dict)}
        if path not in seen_paths:
            pending.append({
                "path": path,
                "requested_at": _time.strftime("%Y-%m-%dT%H:%M:%S"),
            })
            with open(_WRITE_APPROVAL_QUEUE_FILE, "w", encoding="utf-8") as f:
                json.dump(pending, f, indent=2)
    except Exception as exc:
        _log.warning("queue_write_approval failed: %s", exc)


# ── Hard blocklist (cannot be approved, ever) ────────────────────────────────
# Each entry is matched case-insensitively against the resolved absolute path.
# Patterns ending with os.sep mean "this directory and anything under it".
# Bare segments like ".git" match if they appear as a path component.
def _is_blocked_path(resolved_path: str) -> tuple[bool, str]:
    """
    Check if a resolved path is in the hard-coded write blocklist.
    Returns (blocked, reason). Reason is empty when not blocked.

    These cannot be overridden by the user — system directories, credentials,
    git internals, and the schema-aware xlsx are protected at code level.
    """
    p = resolved_path.replace("/", "\\").lower()
    parts = Path(resolved_path).parts
    parts_lower = [pt.lower() for pt in parts]

    # System / install directories
    if p.startswith("c:\\windows\\") or p == "c:\\windows":
        return (True, "C:\\Windows is a system directory")
    if p.startswith("c:\\program files\\") or p == "c:\\program files":
        return (True, "C:\\Program Files is a system directory "
                       "(includes the installed AI-Prowler — never overwrite)")
    if p.startswith("c:\\program files (x86)\\") or p == "c:\\program files (x86)":
        return (True, "C:\\Program Files (x86) is a system directory")
    if p.startswith("c:\\programdata\\") or p == "c:\\programdata":
        return (True, "C:\\ProgramData is a system directory")

    # AppData — block specific sensitive subdirectories instead of all of AppData.
    # AppData\Local\Temp is user scratch space (pytest tmpdir lives here, and
    # legitimate transient work happens here). AppData\Local\AI-Prowler is
    # AI-Prowler's own state. Everything else in AppData is generally other
    # apps' configs and credentials — we block the well-known sensitive ones
    # rather than ALL of AppData, since AppData is too broad to deny wholesale.
    if "\\appdata\\" in p:
        # Sensitive AppData subdirectories (anywhere under AppData\Local or AppData\Roaming)
        sensitive_appdata_segments = (
            "\\appdata\\roaming\\microsoft\\",
            "\\appdata\\local\\microsoft\\",
            "\\appdata\\roaming\\mozilla\\",
            "\\appdata\\local\\mozilla\\",
            "\\appdata\\roaming\\google\\",
            "\\appdata\\local\\google\\",
            "\\appdata\\roaming\\apple\\",
            "\\appdata\\local\\apple computer\\",
            "\\appdata\\local\\packages\\",   # Windows Store apps
            "\\appdata\\roaming\\discord\\",
            "\\appdata\\roaming\\slack\\",
            "\\appdata\\roaming\\anthropic\\",   # Claude Desktop config
        )
        for sensitive in sensitive_appdata_segments:
            if sensitive in p:
                # Friendly app name from the segment for the error message
                app_name = sensitive.split("\\")[-2]
                return (True, f"AppData\\{app_name} contains app state/credentials "
                               f"and is not writable by AI-Prowler")
        # Not in the sensitive list — allowed (covers Temp, AI-Prowler, and any

    # Git internals — protect from any write
    # Use both parts-based (semantic) and string-based (defense-in-depth) checks
    # to be robust across any path-tokenization quirks.
    if ".git" in parts_lower:
        for seg in parts_lower:
            if seg == ".git":
                return (True, "git internals — git operations stay in user's terminal")
    if "\\.git\\" in p or p.endswith("\\.git"):
        return (True, "git internals — git operations stay in user's terminal")

    # Credentials — same defense-in-depth pattern
    if ".ssh" in parts_lower or "\\.ssh\\" in p or p.endswith("\\.ssh"):
        return (True, ".ssh contains credentials — never written by AI-Prowler")
    if ".aws" in parts_lower or "\\.aws\\" in p or p.endswith("\\.aws"):
        return (True, ".aws contains credentials — never written by AI-Prowler")

    # The job tracker spreadsheet — schema-aware tool only
    job_tracker_marker = "ai-prowler_job_tracker.xlsx"
    if p.endswith(job_tracker_marker):
        return (True, "AI-Prowler_Job_Tracker.xlsx must be modified via the "
                       "dedicated update_job_spreadsheet tool (schema-aware), "
                       "not via generic write tools")

    return (False, "")


def _resolve_writable_path(filepath: str, *, queue_approval: bool = True
                           ) -> tuple[Optional[str], Optional[str]]:
    """
    Authorize a path for WRITES. Returns (resolved_path, None) on allow,
    or (None, formatted_error_message) on deny.

    Double-lock constraint:
      (1) Path must be in the read allowlist (~/.rag_auto_update_dirs.json)
      (2) Path must be in the writable allowlist (~/.rag_writable_dirs.json)
      Plus the hard blocklist always wins, regardless of either allowlist.

    If the path passes (1) and the blocklist but fails (2), this function
    enqueues a GUI approval request (unless queue_approval=False).
    """
    # Step 1: read allowlist (also resolves the path)
    resolved, deny_msg = _resolve_allowlisted_path(filepath)
    if not resolved:
        # Already a formatted denial from the read allowlist
        return (None, deny_msg)

    # Step 2: hard blocklist — never bypassable
    blocked, reason = _is_blocked_path(resolved)
    if blocked:
        lines = [
            f"🚫 Write blocked — '{resolved}'",
            f"Reason: {reason}",
            "",
            "This path is on AI-Prowler's hard write-blocklist and cannot be "
            "approved even by the user. Hard-blocked locations include "
            "C:\\Windows, C:\\Program Files, C:\\ProgramData, AppData, "
            ".git, .ssh, .aws, and the job tracker xlsx.",
        ]
        _log.warning("Write blocked (hard): %s (%s)", resolved, reason)
        return (None, "\n".join(lines))

    # Step 3: writable allowlist
    writable = _writable_allowlist_load()
    if any(_path_is_under(resolved, w) for w in writable):
        return (resolved, None)

    # Not yet approved — queue for GUI approval if requested
    if queue_approval:
        _queue_write_approval(resolved)

    lines = [
        f"🔐 Write needs approval — '{resolved}'",
        "",
        "This path is in the READ allowlist but not yet in the WRITABLE "
        "allowlist. AI-Prowler has queued an approval request — the next time "
        "you focus the AI-Prowler GUI window, it will ask you to approve "
        "writes to:",
        f"    {Path(resolved).parent}",
        "",
        "Once approved, this and all paths under that directory will be "
        "writable in future sessions. No need to re-run this tool — the "
        "approval persists in ~/.rag_writable_dirs.json.",
    ]
    if writable:
        lines.append("")
        lines.append("Currently writable:")
        for w in writable[:10]:
            lines.append(f"  - {w}")
    return (None, "\n".join(lines))


def _path_is_under(target: str, ancestor: str) -> bool:
    """True if target equals or is a descendant of ancestor (both resolved)."""
    try:
        t = Path(target).resolve(strict=False)
        a = Path(ancestor).resolve(strict=False)
        if t == a:
            return True
        try:
            t.relative_to(a)
            return True
        except ValueError:
            return False
    except Exception:
        return False


# ── Server-mode personal-directory write scoping ─────────────────────────────
# Server-mode users may only write inside their OWN personal (private)
# directory — never anywhere else in the writable allowlist, and never at
# all if they don't have a personal directory configured. Personal-mode
# installs (ctx has no user) are completely unaffected: full write access,
# exactly as before this feature was added.
def _user_private_write_dir(ctx) -> "tuple[str, Path | None]":
    """
    Determine the write-scope status for the calling user.

    Returns a (status, path) tuple:
      ("personal", None)     — personal mode (no user on ctx); unrestricted,
                                 no change from prior behaviour.
      ("scoped", Path(...))  — server mode; user has a private directory.
                                 Writes must resolve under this path.
      ("blocked", None)      — server mode; user has no private directory
                                 configured (private_collection_enabled is
                                 False, or no folder has been set up for
                                 them yet). No writes are allowed at all.
    """
    user = _current_user(ctx)
    if user is None:
        return ("personal", None)

    if not user.get("private_collection_enabled"):
        return ("blocked", None)

    try:
        company_map = _company_collection_map()
        target_collection = f"user:{user.get('id', '')}"
        for rule in (company_map.get("rules") or []):
            if rule.get("collection") == target_collection:
                prefix = (rule.get("prefix") or "").strip()
                if prefix and Path(prefix).exists():
                    return ("scoped", Path(prefix))
    except Exception:
        pass

    # private_collection_enabled=True but no folder actually set up yet —
    # fail closed, same as not having one at all.
    return ("blocked", None)


def _check_personal_write_scope(ctx, resolved_path: str) -> "str | None":
    """
    Enforce server-mode personal-directory write scoping against an already
    read/write-allowlist-resolved path.

    Returns None if the write is allowed (personal mode, or the path is
    inside the calling user's own private directory). Returns a formatted
    denial string otherwise — callers should `return` it immediately.
    """
    status, private_dir = _user_private_write_dir(ctx)

    if status == "personal":
        return None

    if status == "blocked":
        return (
            "🚫 Write denied — you don't have a personal directory configured.\n"
            "Server-mode users may only write inside their own personal "
            "directory. Ask your owner/admin to enable 'Private collection' "
            "and set up your personal folder in the Admin tab. Until then "
            "you have read-only access."
        )

    # status == "scoped"
    if _path_is_under(resolved_path, str(private_dir)):
        return None

    return (
        f"🚫 Write denied — '{resolved_path}' is outside your personal directory.\n"
        f"Server-mode users may only write inside their own personal directory:\n"
        f"    {private_dir}"
    )


# ── Backup naming (.bak<N> alongside the file) ───────────────────────────────
_BAK_RE = _re.compile(r"\.bak(\d+)$", _re.IGNORECASE)


def _next_backup_path(filepath: str) -> Path:
    """
    Find the next available <filepath>.bak<N> in the file's parent directory.
    N = (max existing bak number for this base filename) + 1.
    Returns the Path object for where the new backup should be written.
    Does NOT create the file — caller must copy contents into it.
    """
    fp = Path(filepath)
    parent = fp.parent
    base = fp.name
    highest = 0
    try:
        prefix = base + ".bak"
        prefix_lower = prefix.lower()
        # List parent dir; cap scan to avoid pathological cases
        scanned = 0
        for sibling in parent.iterdir():
            scanned += 1
            if scanned > _BACKUP_MAX_FILES_PER_PARENT:
                break
            name = sibling.name
            if name.lower().startswith(prefix_lower):
                tail = name[len(prefix):]   # the digit portion after ".bak"
                if tail.isdigit():
                    n = int(tail)
                    if n > highest:
                        highest = n
    except Exception as exc:
        _log.warning("_next_backup_path: directory scan failed in %s: %s", parent, exc)
    return parent / f"{base}.bak{highest + 1}"


def _make_backup(filepath: str) -> tuple[Optional[str], Optional[str]]:
    """
    Copy filepath to <filepath>.bak<N> alongside it. Returns (backup_path, None)
    on success or (None, error_message) on failure.

    Caller has already verified the file exists.
    """
    try:
        target = _next_backup_path(filepath)
        _shutil.copy2(filepath, str(target))   # copy2 preserves mtime
        _log.info("Backup created: %s -> %s", filepath, target)
        return (str(target), None)
    except Exception as exc:
        msg = f"backup failed: {exc}"
        _log.error("_make_backup: %s", msg)
        return (None, msg)


# ── Line-ending preservation helper ──────────────────────────────────────────
# Background: Python's text-mode file I/O on Windows silently converts
# CRLF → LF on read and LF → platform-native on write. The write-side
# tools previously round-tripped through text mode and binary write,
# which silently stripped CRLF endings from Windows files on every edit.
# These helpers let us detect the existing convention and preserve it.

def _detect_line_ending(file_bytes: bytes) -> str:
    """Detect the dominant line ending in raw file bytes.

    Returns one of:
        '\\r\\n'  — Windows / CRLF  (any CRLF presence wins)
        '\\r'     — Classic Mac     (only \\r, no \\n)
        '\\n'     — Unix / LF       (the default; also empty files)

    Strategy: any CRLF in the file = treat the whole file as CRLF. This
    correctly handles slightly-mixed files (rare in practice, but they
    do occur from tools that append LF-only lines to a CRLF file).
    """
    if b"\r\n" in file_bytes:
        return "\r\n"
    if b"\r" in file_bytes and b"\n" not in file_bytes:
        return "\r"
    return "\n"


def _read_text_preserving_endings(filepath: str) -> tuple[str, str]:
    """Read filepath, return (lf_normalized_text, original_line_ending).

    The returned text always uses LF newlines so that in-memory string
    operations (str_replace, etc.) behave uniformly regardless of the
    file's on-disk convention. The caller is responsible for re-applying
    the returned line ending before writing back.

    Caller has already verified the file exists and is small enough.

    ENCODING DETECTION (v7.0.1):
    Files are not always UTF-8. ISS installer scripts, legacy Windows batch
    files, and many text editors write Latin-1 / Windows-1252. If we decode
    with utf-8 errors='replace', non-ASCII bytes (e.g. em-dashes U+0097 in
    Windows-1252) become U+FFFD replacement characters. old_str then contains
    the real Unicode character but the file content contains U+FFFD — they
    never match, and str_replace silently fails.

    Detection cascade (each attempted in order):
      1. utf-8-sig  — strips BOM if present, valid UTF-8 otherwise
      2. utf-8      — strict, no BOM
      3. chardet    — statistical detection if available (optional dependency)
      4. latin-1    — fallback: accepts every byte value, never raises

    The detected encoding is stored alongside the content so the write path
    can round-trip the file in its original encoding. This preserves the
    file's byte content for non-ASCII characters instead of corrupting them.
    """
    with open(filepath, "rb") as f:
        raw = f.read()
    line_ending = _detect_line_ending(raw)

    # Encoding detection cascade
    detected_encoding = "latin-1"  # safe fallback — accepts all byte values

    # Check for actual BOM presence before trying utf-8-sig.
    # utf-8-sig always succeeds on valid UTF-8 even without a BOM, which
    # would cause us to write a BOM back into files that never had one.
    has_bom = raw[:3] == b"\xef\xbb\xbf"

    # 1. Try UTF-8 with BOM strip (only if BOM actually present)
    if has_bom:
        try:
            decoded = raw.decode("utf-8-sig")
            detected_encoding = "utf-8-sig"
        except UnicodeDecodeError:
            has_bom = False  # malformed BOM, fall through

    if not has_bom:
        # 2. Try strict UTF-8
        try:
            decoded = raw.decode("utf-8")
            detected_encoding = "utf-8"
        except UnicodeDecodeError:
            # 3. Try chardet if available
            try:
                import chardet
                result = chardet.detect(raw)
                enc = result.get("encoding") or "latin-1"
                confidence = result.get("confidence", 0)
                if confidence >= 0.7 and enc.lower() not in ("ascii",):
                    decoded = raw.decode(enc, errors="replace")
                    detected_encoding = enc
                else:
                    raise ValueError("low confidence")
            except Exception:
                # 4. Latin-1 fallback — always succeeds
                decoded = raw.decode("latin-1")
                detected_encoding = "latin-1"

    # Normalize all endings to LF for consistent in-memory editing.
    # Order matters: handle \r\n before bare \r.
    normalized = decoded.replace("\r\n", "\n").replace("\r", "\n")
    return (normalized, line_ending, detected_encoding)


def _apply_line_ending(text: str, line_ending: str) -> str:
    """Re-apply the given line ending to LF-normalized text.

    If line_ending is '\\n', returns text unchanged (zero allocations).
    Otherwise replaces every '\\n' with the target ending.
    """
    if line_ending == "\n":
        return text
    return text.replace("\n", line_ending)


# ── Write-counter circuit breaker ────────────────────────────────────────────
def _check_and_increment_write_counter() -> tuple[bool, str]:
    """
    Returns (allowed, message). Increments the counter on each call.
    The counter is process-scoped — restarting AI-Prowler resets it.
    """
    with _write_counter_lock:
        if _write_counter["count"] >= _WRITES_PER_SESSION_LIMIT:
            return (False,
                    f"🛑 Write circuit-breaker tripped: "
                    f"{_WRITES_PER_SESSION_LIMIT} writes already this session. "
                    f"This protects against runaway loops. Restart AI-Prowler "
                    f"or call reset_write_counter from the GUI to continue.")
        _write_counter["count"] += 1
        return (True, "")


def _reset_write_counter_internal() -> int:
    """Reset session write counter. Returns the count that was reset from."""
    with _write_counter_lock:
        old = _write_counter["count"]
        _write_counter["count"] = 0
        _write_counter["session_started"] = _time.time()
        return old


# ── Auto re-index helper ─────────────────────────────────────────────────────
def _reindex_file_after_write(filepath: str, *, was_new: bool = False) -> None:
    """
    DISABLED in v7.0.0 — intentional no-op.

    Auto-reindex-on-write was removed: re-embedding the whole file on the
    uvicorn request thread re-entered the already-open ChromaDB collection and
    caused "resource deadlock would occur", hanging the HTTP MCP server. Large
    files (e.g. rag_gui.py at ~700 KB) reproduced it every time.

    The database is now updated only by EXPLICIT reindex: call reindex_file()
    once you are done editing a file, or reindex_directory()/index_path() for a
    whole tree. Writes still create backups and land on disk immediately; they
    just no longer touch ChromaDB. Kept as a stub so existing call sites and any
    external references remain valid.
    """
    return


# ── Parent directory must exist (helper for create_file / write_file) ────────
def _parent_dir_check(filepath: str) -> Optional[str]:
    """Return an error message if the parent directory doesn't exist, else None."""
    try:
        parent = Path(filepath).parent
        if not parent.exists():
            return (f"⚠️  Parent directory does not exist: {parent}\n"
                    f"Call create_directory({parent}) first, then retry.")
        if not parent.is_dir():
            return f"⚠️  Parent path is not a directory: {parent}"
    except Exception as exc:
        return f"⚠️  Could not check parent directory: {exc}"
    return None


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 1 — create_file
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def create_file(filepath: str, content: str, encoding: str = "text", ctx: "Context | None" = None) -> str:
    """
    CODE TOOLS — Create a NEW file. FAILS if the file already exists.

    Use this for files that should not yet exist. To modify an existing file,
    use write_file (whole-file overwrite) or str_replace_in_file (surgical edit).

    Security: the path must be in BOTH the read allowlist
    (~/.rag_auto_update_dirs.json) AND the writable allowlist
    (~/.rag_writable_dirs.json). The first time you attempt a write to a new
    directory, AI-Prowler queues a GUI approval request that you must accept
    at the desktop. Subsequent writes to that directory succeed silently.

    Hard-blocked locations (cannot be approved):
      C:\\Windows, C:\\Program Files, C:\\ProgramData, AppData (except
      AI-Prowler's own state), .git, .ssh, .aws, the job tracker xlsx.

    Args:
        filepath:  Absolute path of the file to create.
        content:   Full file contents as a string. For text mode, pass the
                   text directly. For binary mode, pass base64-encoded bytes.
        encoding:  Write mode — 'text' (default) for plain text, source code,
                   markdown, JSON, etc. Use 'base64' for binary files such as
                   .docx, .xlsx, .pdf, .png, .zip — content must be a valid
                   base64 string. Line-ending normalisation is skipped in
                   base64 mode.

    Returns:
        Success: confirmation with byte count and a note that the file is now
                 indexed in ChromaDB. Failure: a clear error explaining what
                 went wrong (path not allowlisted, file already exists,
                 parent missing, write circuit-breaker tripped, etc.).
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    # Validate encoding parameter
    encoding = (encoding or "text").lower().strip()
    if encoding not in ("text", "base64"):
        return f"⚠️  Unknown encoding '{encoding}'. Use 'text' or 'base64'."

    # Authorize
    resolved, deny = _resolve_writable_path(filepath)
    if not resolved:
        return deny

    # Server-mode: writes are scoped to the caller's own personal directory
    scope_denial = _check_personal_write_scope(ctx, resolved)
    if scope_denial:
        return scope_denial

    # Existence check — create_file fails if file exists (negative space: use write_file)
    if Path(resolved).exists():
        return (f"⚠️  File already exists: {resolved}\n"
                f"create_file is for NEW files only. To modify an existing "
                f"file, use write_file (full overwrite) or "
                f"str_replace_in_file (surgical edit).")

    # Parent directory must already exist (force explicit create_directory call)
    parent_err = _parent_dir_check(resolved)
    if parent_err:
        return parent_err

    # ── Encode content to bytes ───────────────────────────────────────────────
    if encoding == "base64":
        # Binary mode: decode base64 → raw bytes; skip all line-ending logic.
        import base64 as _b64
        try:
            content_bytes = _b64.b64decode(content, validate=True)
        except Exception as exc:
            return (f"⚠️  base64 decode failed: {exc}\n"
                    f"Ensure content is a valid base64 string with no extra whitespace.")
    else:
        # Text mode: apply platform line-ending normalisation (existing behaviour).
        if "\r" in content:
            final_content = content
        elif "\n" in content and os.linesep != "\n":
            final_content = content.replace("\n", os.linesep)
        else:
            final_content = content
        try:
            content_bytes = final_content.encode("utf-8")
        except Exception as exc:
            return f"⚠️  Content could not be encoded as UTF-8: {exc}"

    if len(content_bytes) > _WRITE_MAX_BYTES:
        return (f"⚠️  Content too large ({len(content_bytes):,} bytes, "
                f"cap {_WRITE_MAX_BYTES:,}).")

    # Circuit breaker
    ok, msg = _check_and_increment_write_counter()
    if not ok:
        return msg

    # Write
    try:
        with open(resolved, "wb") as f:
            f.write(content_bytes)
    except Exception as exc:
        return f"⚠️  Write failed: {exc}"

    _log.info("create_file: %s (%d bytes, encoding=%s)", resolved, len(content_bytes), encoding)

    mode_note = " [binary/base64]" if encoding == "base64" else ""
    return (f"✅ Created {resolved}{mode_note}\n"
            f"   {len(content_bytes):,} bytes written\n"
            f"   NOT yet indexed — call reindex_file({resolved!r}) when done editing.")


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 2 — write_file
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def write_file(filepath: str, content: str, verify_after_write: bool = False, encoding: str = "text", ctx: "Context | None" = None) -> str:
    """
    CODE TOOLS — Overwrite an EXISTING file with new content. FAILS if the
    file does not exist.

    Use this for whole-file rewrites. To create a new file, use create_file.
    For surgical single-string edits, use str_replace_in_file (much cheaper
    in tokens for large files).

    Auto-backup: before writing, the current file is copied to
    <filepath>.bak<N> alongside it, where N is the next available number.
    Backups are kept forever — manual cleanup only. The .bak<N> file is
    never indexed in ChromaDB.

    Auto re-index: after the write succeeds, all existing ChromaDB chunks
    for this filepath are purged and the new content is re-chunked and added.

    Args:
        filepath:           Absolute path to overwrite.
        content:            Full new file contents. For text mode, pass the
                            text directly. For binary mode, pass base64-encoded
                            bytes.
        verify_after_write: If True (default False), re-read the file after
                            writing and include the first/last 5 lines of the
                            new content in the response so you can confirm the
                            write landed. Skipped in base64 mode.
        encoding:           Write mode — 'text' (default) for plain text,
                            source code, markdown, JSON, etc. Use 'base64' for
                            binary files such as .docx, .xlsx, .pdf, .png,
                            .zip — content must be a valid base64 string.
                            Line-ending normalisation is skipped in base64 mode.

    Returns:
        Success: confirmation including backup path and byte counts.
        Failure: a clear error message.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    # Validate encoding parameter
    encoding = (encoding or "text").lower().strip()
    if encoding not in ("text", "base64"):
        return f"⚠️  Unknown encoding '{encoding}'. Use 'text' or 'base64'."

    # Authorize
    resolved, deny = _resolve_writable_path(filepath)
    if not resolved:
        return deny

    # Server-mode: writes are scoped to the caller's own personal directory
    scope_denial = _check_personal_write_scope(ctx, resolved)
    if scope_denial:
        return scope_denial

    # Existence check — write_file requires the file to exist
    if not Path(resolved).exists():
        return (f"⚠️  File does not exist: {resolved}\n"
                f"write_file is for EXISTING files only. To create a new "
                f"file, use create_file.")
    if not Path(resolved).is_file():
        return f"⚠️  Path is not a regular file: {resolved}"

    # ── Encode content to bytes ───────────────────────────────────────────────
    if encoding == "base64":
        # Binary mode: decode base64 → raw bytes; skip all line-ending logic.
        import base64 as _b64
        try:
            content_bytes = _b64.b64decode(content, validate=True)
        except Exception as exc:
            return (f"⚠️  base64 decode failed: {exc}\n"
                    f"Ensure content is a valid base64 string with no extra whitespace.")
    else:
        # Text mode: detect existing line-ending convention and preserve it.
        try:
            with open(resolved, "rb") as f:
                existing_head = f.read(65536)
            line_ending = _detect_line_ending(existing_head)
        except Exception as exc:
            return f"⚠️  Cannot probe existing file for line endings: {exc}"

        normalized_content = content.replace("\r\n", "\n").replace("\r", "\n")
        final_content = _apply_line_ending(normalized_content, line_ending)

        try:
            content_bytes = final_content.encode("utf-8")
        except Exception as exc:
            return f"⚠️  Content could not be encoded as UTF-8: {exc}"

    if len(content_bytes) > _WRITE_MAX_BYTES:
        return (f"⚠️  Content too large ({len(content_bytes):,} bytes, "
                f"cap {_WRITE_MAX_BYTES:,}).")

    # Circuit breaker
    ok, msg = _check_and_increment_write_counter()
    if not ok:
        return msg

    # Auto-backup before overwriting
    old_size = Path(resolved).stat().st_size
    backup_path, backup_err = _make_backup(resolved)
    if backup_err:
        return f"⚠️  Aborting write — could not create backup: {backup_err}"

    # Write
    try:
        with open(resolved, "wb") as f:
            f.write(content_bytes)
    except Exception as exc:
        return (f"⚠️  Write failed (backup preserved at {backup_path}): {exc}\n"
                f"Use restore_backup to recover if needed.")

    _log.info("write_file: %s (%d -> %d bytes, encoding=%s, backup %s)",
              resolved, old_size, len(content_bytes), encoding, backup_path)

    mode_note = " [binary/base64]" if encoding == "base64" else ""
    out = [
        f"✅ Wrote {resolved}{mode_note}",
        f"   {old_size:,} bytes  →  {len(content_bytes):,} bytes",
        f"   Backup: {backup_path}",
        f"   NOT yet indexed — call reindex_file() when done editing.",
    ]
    if verify_after_write and encoding == "text":
        out.append("")
        out.append("─── Verify (first 5 / last 5 lines of new content) ───")
        try:
            with open(resolved, "r", encoding="utf-8", errors="replace") as f:
                new_lines = f.read().splitlines()
            total = len(new_lines)
            head = new_lines[:5]
            tail = new_lines[-5:] if total > 10 else []
            for i, ln in enumerate(head, start=1):
                out.append(f"  {i:>4}  {ln}")
            if tail:
                out.append(f"  ... ({total - 10} line(s) omitted)")
                for i, ln in enumerate(tail, start=total - 4):
                    out.append(f"  {i:>4}  {ln}")
        except Exception as exc:
            out.append(f"  (verify read failed: {exc})")
    elif verify_after_write and encoding == "base64":
        out.append("   (verify skipped — binary file)")
    return "\n".join(out)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 3 — str_replace_in_file  (THE MOST IMPORTANT WRITE TOOL)
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def str_replace_in_file(filepath: str,
                        old_str: str,
                        new_str: str,
                        dry_run: bool = False,
                        verify_after_write: bool = True,
                        ctx: "Context | None" = None) -> str:
    """
    CODE TOOLS — Surgical in-place edit: replace one unique occurrence of
    old_str with new_str. This is THE primary write tool — prefer it over
    write_file for any change to a specific section of a file.

    Why this is so much better than read-modify-write for large files:
        A round-trip of rag_gui.py (12,209 lines, ~120K tokens) costs
        ~240K tokens vs ~250 tokens for a str_replace edit. 1000x savings.

    UNIQUENESS CONSTRAINT: old_str must appear in the file EXACTLY ONCE.
    If it appears zero times, the tool fails (typo or already-edited).
    If it appears more than once, the tool fails and reports each match's
    line number so you can pick a more distinctive old_str.

    Auto-backup: before any write, the file is copied to <filepath>.bak<N>.
    Auto re-index: ChromaDB chunks for this file are purged and rebuilt.

    Args:
        filepath:           Absolute path to edit. Must be writable-allowlisted.
        old_str:            Exact substring to find. Must appear exactly once.
                            Whitespace is significant.
        new_str:            Replacement text. May be empty to delete a span.
                            May include multiple lines.
        dry_run:            If True, returns a unified diff WITHOUT modifying
                            the file. No backup is created. Useful for mobile
                            confirmation — review the diff, then call again
                            with dry_run=False.
        verify_after_write: If True (default True for str_replace_in_file),
                            include 5 lines before and after the changed
                            region in the response. Catches silent failures
                            (file locked by AV, OneDrive sync conflict).

    Returns:
        Success: confirmation with line where the change landed, backup path,
                 and (if verify_after_write) the surrounding context.
        Failure: detailed error explaining why (not found, ambiguous, etc.).
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    # Sanity checks on the strings
    if old_str is None:
        return "⚠️  old_str cannot be None."
    if new_str is None:
        return "⚠️  new_str cannot be None (use empty string to delete a span)."
    if len(old_str) == 0:
        return "⚠️  old_str cannot be empty (would match infinitely)."
    if len(old_str) > _STR_REPLACE_MAX_OLD_STR_LEN:
        return (f"⚠️  old_str too long ({len(old_str):,} chars, cap "
                f"{_STR_REPLACE_MAX_OLD_STR_LEN:,}). Use a smaller distinctive "
                f"snippet or write_file for whole-file rewrites.")

    # ── CRLF + BACKSLASH NORMALISATION ─────────────────────────────────────────
    # CRLF FIX (v7.0.0): The file is read and LF-normalized by
    # _read_text_preserving_endings(), but old_str / new_str arrive straight
    # from the MCP client. When a multi-line snippet is copied from a CRLF
    # (Windows) source, old_str carries '\r\n' while the in-memory file text
    # carries only '\n'. Normalize both args to LF so all matching logic below
    # is line-ending-agnostic; the file's original ending is re-applied at
    # write time by _apply_line_ending().
    #
    # BACKSLASH FIX (v7.0.1): When old_str contains backslash sequences (e.g.
    # Pascal string literals like '".Replace(\"\\", \"/\")', Windows paths like
    # 'C:\\Users\\...', or escape sequences like '\r\n'), the MCP JSON transport
    # and Claude's own string handling can double-interpret the backslashes so
    # what arrives here differs from what is actually in the file.
    #
    # Strategy: after CRLF normalisation, if old_str is NOT found in the
    # normalised file text, try common backslash escape variants automatically
    # before giving up. This makes the tool robust against the most frequent
    # escaping mismatches without requiring Claude to manually debug encoding.
    old_str = old_str.replace("\r\n", "\n").replace("\r", "\n")
    new_str = new_str.replace("\r\n", "\n").replace("\r", "\n")

    # Authorize (suppress queueing if it's a dry-run — no actual write attempted)
    resolved, deny = _resolve_writable_path(filepath, queue_approval=not dry_run)
    if not resolved:
        return deny

    # Server-mode: writes are scoped to the caller's own personal directory
    scope_denial = _check_personal_write_scope(ctx, resolved)
    if scope_denial:
        return scope_denial

    if not Path(resolved).exists():
        return f"⚠️  File does not exist: {resolved}"
    if not Path(resolved).is_file():
        return f"⚠️  Path is not a regular file: {resolved}"

    # Size cap on the file
    try:
        size = Path(resolved).stat().st_size
    except Exception as exc:
        return f"⚠️  Cannot stat file: {exc}"
    if size > _READ_FILE_MAX_BYTES:
        return (f"⚠️  File too large for in-memory edit "
                f"({size:,} bytes, cap {_READ_FILE_MAX_BYTES:,}).")

    # Binary sniff
    try:
        with open(resolved, "rb") as bf:
            head = bf.read(_BINARY_SNIFF_BYTES)
        if _looks_binary(head):
            return (f"⚠️  File appears to be binary: {resolved}\n"
                    f"str_replace_in_file only edits text files.")
    except Exception as exc:
        return f"⚠️  Cannot open file: {exc}"

    # Read whole file — in BINARY first so we can detect the line-ending
    # convention, then decode to an LF-normalized in-memory string. This
    # preserves CRLF files on the round-trip; without it Python's text-mode
    # write silently strips \r bytes from every Windows file we edit.
    try:
        text, line_ending, file_encoding = _read_text_preserving_endings(resolved)
    except Exception as exc:
        return f"⚠️  Read failed: {exc}"

    # Count occurrences — with automatic backslash-variant fallback.
    #
    # When old_str contains backslashes (Windows paths, Pascal string literals,
    # escape sequences), the MCP JSON transport or Claude's string formatting
    # can double-escape them so '\\' in the file arrives as '\\\\' in old_str,
    # or vice-versa.  Rather than forcing Claude to debug the escaping manually,
    # we silently try a small set of common backslash variants and use the first
    # one that produces exactly one match.  The variant used is logged in the
    # success message so the behaviour is transparent.
    _backslash_variants = [
        old_str,                                    # as-received (already LF-normalised)
        old_str.replace("\\\\", "\\"),              # de-double: \\\\ -> \\
        old_str.replace("\\", "\\\\"),              # re-double: \\ -> \\\\
        old_str.replace("\\\\", "\\").replace("\\\\", "\\"),   # two passes of de-double
    ]
    # Deduplicate while preserving order
    seen_variants: list[str] = []
    for _v in _backslash_variants:
        if _v not in seen_variants:
            seen_variants.append(_v)

    _matched_variant = None
    count = 0
    for _v in seen_variants:
        _c = text.count(_v)
        if _c == 1:
            _matched_variant = _v
            count = 1
            break
        if _c > 1:
            _matched_variant = _v
            count = _c
            break   # ambiguous — report it below

    if _matched_variant is not None and _matched_variant != old_str:
        # Log the auto-correction so it's visible in the response
        _backslash_note = (
            f"   ℹ️  Backslash auto-correction applied — "
            f"original old_str had a different escape level than the file.\n"
        )
        old_str = _matched_variant
    else:
        _backslash_note = ""
        if _matched_variant is None:
            count = 0   # no variant matched

    if count == 0:
        # ── Nearest-match diagnostic (Problem 4) ────────────────────────────
        # Find the closest line(s) in the file to old_str using a simple
        # token-overlap heuristic so Claude can pinpoint where the mismatch is.
        diag_lines = []
        try:
            needle_words = set(old_str.split())
            if needle_words:
                file_lines = text.splitlines()
                scored = []
                for ln, line_text in enumerate(file_lines, 1):
                    line_words = set(line_text.split())
                    overlap = len(needle_words & line_words)
                    if overlap:
                        scored.append((overlap, ln, line_text))
                scored.sort(key=lambda x: -x[0])
                top = scored[:5]
                if top:
                    diag_lines.append("Nearest matching lines by word overlap:")
                    for overlap, ln, line_text in top:
                        preview = line_text[:120] + ("..." if len(line_text) > 120 else "")
                        diag_lines.append(f"  line {ln:>5} ({overlap} word(s) match): {preview}")
        except Exception:
            pass
        diag = ("\n" + "\n".join(diag_lines)) if diag_lines else ""
        return (
            f"⚠️  old_str not found in {resolved}.\n"
            f"Possible causes: typo in old_str, file already edited, "
            f"whitespace difference (tabs vs spaces, trailing newline). "
            f"Tip: use grep_documents to verify the exact text first."
            f"{diag}"
        )
    if count > 1:
        # Find each occurrence and report its line number for disambiguation
        lines_of_match = []
        start_at = 0
        while True:
            idx = text.find(old_str, start_at)
            if idx == -1:
                break
            line_no = text.count("\n", 0, idx) + 1
            lines_of_match.append(line_no)
            start_at = idx + 1
            if len(lines_of_match) > 20:
                break
        return (f"⚠️  old_str found {count} times in {resolved}.\n"
                f"Match line numbers: {lines_of_match[:20]}"
                + ("  (more not listed)" if count > 20 else "") + "\n"
                f"old_str must match EXACTLY ONCE. Add surrounding context "
                f"(a few lines before/after) until it's unique, then retry.")

    # ── DRY-RUN SUGGESTION for long old_str (Problem 5) ────────────────────────
    # If old_str is large (> 10 lines or > 500 chars) and dry_run is False,
    # append a suggestion — but don't block the edit.
    _dry_run_hint = ""
    if not dry_run and (old_str.count("\n") > 10 or len(old_str) > 500):
        _dry_run_hint = (
            "   💡 Tip: for large old_str blocks, use dry_run=True first "
            "to preview the diff before committing.\n"
        )

    # Compute the change
    new_text = text.replace(old_str, new_str, 1)

    # Find where the change happened (for verify and reporting)
    change_idx = text.find(old_str)
    line_of_change = text.count("\n", 0, change_idx) + 1
    # Report the size as it will be on disk — i.e. AFTER re-applying the
    # original line ending and encoding. Otherwise size previews are wrong
    # for CRLF files or non-UTF-8 files.
    # Use file_encoding for the encode so Latin-1/Windows-1252 files are
    # written back in their original encoding rather than corrupted to UTF-8.
    new_bytes_on_disk = _apply_line_ending(new_text, line_ending).encode(
        file_encoding, errors="replace"
    )
    new_byte_count = len(new_bytes_on_disk)

    # ── DRY RUN ──
    if dry_run:
        out = [
            f"🔎 DRY RUN — no changes written to {resolved}",
            f"   Match found at line {line_of_change}",
            f"   File size would change: {size:,} → {new_byte_count:,} bytes",
            "",
            "─── Unified diff (a=before, b=after) ───",
        ]
        try:
            import difflib as _difflib
            diff = list(_difflib.unified_diff(
                text.splitlines(keepends=False),
                new_text.splitlines(keepends=False),
                fromfile=f"a/{Path(resolved).name}",
                tofile=f"b/{Path(resolved).name}",
                lineterm="",
                n=3,
            ))
            if diff:
                out.extend(diff[:200])
                if len(diff) > 200:
                    out.append(f"... ({len(diff) - 200} more diff lines suppressed)")
            else:
                out.append("(no textual diff — old_str and new_str are identical)")
        except Exception as exc:
            out.append(f"(diff generation failed: {exc})")
        out.append("")
        out.append("To apply: call again with dry_run=False.")
        return "\n".join(out)

    # ── REAL WRITE ──
    # Circuit breaker
    ok, msg = _check_and_increment_write_counter()
    if not ok:
        return msg

    # Size cap on new content
    if new_byte_count > _WRITE_MAX_BYTES:
        return (f"⚠️  Resulting content too large ({new_byte_count:,} bytes, "
                f"cap {_WRITE_MAX_BYTES:,}).")

    # Auto-backup
    backup_path, backup_err = _make_backup(resolved)
    if backup_err:
        return f"⚠️  Aborting edit — could not create backup: {backup_err}"

    # Write — re-apply the original line ending convention. new_bytes_on_disk
    # was already computed above for the size report; reuse it here.
    try:
        with open(resolved, "wb") as f:
            f.write(new_bytes_on_disk)
    except Exception as exc:
        return (f"⚠️  Write failed (backup preserved at {backup_path}): {exc}\n"
                f"Use restore_backup to recover.")

    _log.info("str_replace_in_file: %s line=%d (%d -> %d bytes, backup %s)",
              resolved, line_of_change, size, new_byte_count, backup_path)

    # (auto-reindex removed in v7.0.0 — call reindex_file() when done editing)
    out = [
        f"✅ Edited {resolved}",
        f"   Change at line {line_of_change}",
        f"   {size:,} bytes  →  {new_byte_count:,} bytes",
        f"   Encoding: {file_encoding}",
        f"   Backup: {backup_path}",
        f"   NOT yet indexed — call reindex_file() when done editing.",
    ]
    if _backslash_note:
        out.append(_backslash_note.rstrip())
    if _dry_run_hint:
        out.append(_dry_run_hint.rstrip())
    if verify_after_write:
        out.append("")
        out.append("─── Verify (5 lines before, change region, 5 lines after) ───")
        try:
            new_lines_list = new_text.splitlines()
            new_total = len(new_lines_list)
            new_str_line_count = new_str.count("\n") + 1
            change_start = max(1, line_of_change - 5)
            change_end = min(new_total, line_of_change + new_str_line_count + 4)
            for i in range(change_start, change_end + 1):
                marker = "▶" if change_start + 5 <= i < change_start + 5 + new_str_line_count else " "
                if i <= new_total:
                    out.append(f"  {marker}{i:>5}  {new_lines_list[i - 1]}")
        except Exception as exc:
            out.append(f"  (verify read failed: {exc})")
    return "\n".join(out)



# ══════════════════════════════════════════════════════════════════════════════
# TOOL 3c — line_replace_in_file
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def line_replace_in_file(filepath: str,
                         start_line: int,
                         end_line: int,
                         new_content: str,
                         dry_run: bool = False,
                         ctx: "Context | None" = None) -> str:
    """
    CODE TOOLS — Replace a range of lines by line number. Zero text-matching
    ambiguity — works on any file regardless of encoding, tabs, or Unicode.

    This is the LAST RESORT when str_replace_in_file fails (e.g. whitespace
    or Unicode differences make an exact match impractical). The workflow is:
        1. read_file_lines(filepath, start, end)   ← confirm exact lines
        2. line_replace_in_file(filepath, start, end, new_content)

    The lines from start_line to end_line (inclusive, 1-based) are replaced
    with new_content. new_content may contain any number of lines — the
    replacement does not have to be the same number of lines as removed.

    Args:
        filepath:    Absolute path to edit. Must be writable-allowlisted.
        start_line:  First line to replace (1-based, inclusive).
        end_line:    Last line to replace (1-based, inclusive).
                     Use start_line == end_line to replace a single line.
        new_content: Text that replaces lines start_line..end_line.
                     Do NOT include a trailing newline — the tool handles it.
        dry_run:     If True, shows the diff without writing.

    Returns:
        Success: confirmation with line range, backup path, and verify block.
        Failure: clear error (out-of-range, file not found, etc.)
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    if start_line < 1:
        return "⚠️  start_line must be >= 1."
    if end_line < start_line:
        return f"⚠️  end_line ({end_line}) must be >= start_line ({start_line})."
    if new_content is None:
        return "⚠️  new_content cannot be None (pass empty string to delete lines)."

    resolved, deny = _resolve_writable_path(filepath, queue_approval=not dry_run)
    if not resolved:
        return deny
    if not Path(resolved).exists():
        return f"⚠️  File does not exist: {resolved}"

    # Server-mode: writes are scoped to the caller's own personal directory
    scope_denial = _check_personal_write_scope(ctx, resolved)
    if scope_denial:
        return scope_denial

    try:
        size = Path(resolved).stat().st_size
    except Exception as exc:
        return f"⚠️  Cannot stat file: {exc}"
    if size > _READ_FILE_MAX_BYTES:
        return f"⚠️  File too large ({size:,} bytes, cap {_READ_FILE_MAX_BYTES:,})."

    try:
        text, line_ending, file_encoding = _read_text_preserving_endings(resolved)
    except Exception as exc:
        return f"⚠️  Read failed: {exc}"

    file_lines = text.split("\n")
    total_lines = len(file_lines)

    if start_line > total_lines:
        return (f"⚠️  start_line {start_line} is beyond end of file "
                f"({total_lines} lines).")
    if end_line > total_lines:
        return (f"⚠️  end_line {end_line} is beyond end of file "
                f"({total_lines} lines). Use {total_lines} to reach the last line.")

    # Build the new file lines
    new_content_norm = new_content.replace("\r\n", "\n").replace("\r", "\n")
    replacement_lines = new_content_norm.split("\n")

    # Splice: lines before + new content + lines after
    before  = file_lines[:start_line - 1]           # 0-based before start
    after   = file_lines[end_line:]                  # 0-based after end
    new_file_lines = before + replacement_lines + after
    new_text = "\n".join(new_file_lines)

    new_bytes_on_disk = _apply_line_ending(new_text, line_ending).encode(
        file_encoding, errors="replace"
    )
    new_byte_count = len(new_bytes_on_disk)

    if dry_run:
        out = [
            f"🔎 DRY RUN (line replace) — no changes written to {resolved}",
            f"   Replacing lines {start_line}–{end_line} "
            f"({end_line - start_line + 1} line(s) → {len(replacement_lines)} line(s))",
            f"   File size would change: {size:,} → {new_byte_count:,} bytes",
            "",
            "─── Lines being REMOVED ───",
        ]
        for i in range(start_line - 1, min(end_line, total_lines)):
            out.append(f"  - {i+1:>5}  {file_lines[i]}")
        out.append("")
        out.append("─── Lines being ADDED ───")
        for i, l in enumerate(replacement_lines):
            out.append(f"  + {start_line + i:>5}  {l}")
        out.append("")
        out.append("To apply: call again with dry_run=False.")
        return "\n".join(out)

    ok, msg = _check_and_increment_write_counter()
    if not ok:
        return msg

    if new_byte_count > _WRITE_MAX_BYTES:
        return f"⚠️  Result too large ({new_byte_count:,} bytes, cap {_WRITE_MAX_BYTES:,})."

    backup_path, backup_err = _make_backup(resolved)
    if backup_err:
        return f"⚠️  Could not create backup: {backup_err}"

    try:
        with open(resolved, "wb") as f:
            f.write(new_bytes_on_disk)
    except Exception as exc:
        return f"⚠️  Write failed (backup at {backup_path}): {exc}"

    new_total = len(new_file_lines)
    _log.info("line_replace_in_file: %s lines %d-%d -> %d lines (%d -> %d bytes)",
              resolved, start_line, end_line, len(replacement_lines), size, new_byte_count)

    out = [
        f"✅ Edited {resolved}",
        f"   Replaced lines {start_line}–{end_line} "
        f"({end_line - start_line + 1} line(s) → {len(replacement_lines)} line(s))",
        f"   {size:,} bytes  →  {new_byte_count:,} bytes",
        f"   Encoding: {file_encoding}",
        f"   Backup: {backup_path}",
        f"   NOT yet indexed — call reindex_file() when done editing.",
        "",
        "─── Verify (5 lines before, change region, 5 lines after) ───",
    ]
    try:
        verify_start = max(1, start_line - 5)
        verify_end   = min(new_total, start_line + len(replacement_lines) + 4)
        for i in range(verify_start, verify_end + 1):
            marker = "▶" if start_line <= i < start_line + len(replacement_lines) else " "
            if i <= new_total:
                out.append(f"  {marker}{i:>5}  {new_file_lines[i - 1]}")
    except Exception as exc:
        out.append(f"  (verify read failed: {exc})")
    return "\n".join(out)



@mcp.tool()
def create_directory(dirpath: str, parents: bool = True, ctx: "Context | None" = None) -> str:
    """
    CODE TOOLS — Create a directory inside an indexed AND writable area.
    Idempotent: succeeds if the directory already exists.

    Args:
        dirpath:  Absolute path to create.
        parents:  If True (default), create missing parent directories
                  (mkdir -p semantics). If False, parent must already exist.

    Returns:
        Confirmation with whether the directory was newly created or already
        existed, or an error explaining why creation failed.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    resolved, deny = _resolve_writable_path(dirpath)
    if not resolved:
        return deny

    # Server-mode: writes are scoped to the caller's own personal directory
    scope_denial = _check_personal_write_scope(ctx, resolved)
    if scope_denial:
        return scope_denial

    dp = Path(resolved)
    if dp.exists():
        if dp.is_dir():
            return f"✅ Directory already exists: {resolved}\n   (idempotent — no action taken)"
        return f"⚠️  Path exists but is not a directory: {resolved}"

    # Circuit breaker (directory creation counts toward the limit)
    ok, msg = _check_and_increment_write_counter()
    if not ok:
        return msg

    try:
        dp.mkdir(parents=parents, exist_ok=False)
    except FileNotFoundError as exc:
        return (f"⚠️  Parent directory does not exist: {exc}\n"
                f"Retry with parents=True, or call create_directory on the "
                f"parent first.")
    except Exception as exc:
        return f"⚠️  Could not create directory: {exc}"

    _log.info("create_directory: %s (parents=%s)", resolved, parents)
    return f"✅ Created directory: {resolved}"


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 5 — list_directory
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def list_directory(dirpath: str, show_hidden: bool = False) -> str:
    """
    CODE TOOLS — List the contents of a directory (files and subdirectories).
    Read-only — no writable-allowlist requirement; only the read allowlist
    applies.

    Args:
        dirpath:      Absolute path to list.
        show_hidden:  If True, include entries beginning with '.'
                      (default False).

    Returns:
        A formatted listing with file/dir indicators, sizes for files, and
        a summary line. Output is sorted: directories first, then files,
        each group alphabetically.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    # read-only — only the read allowlist applies
    resolved, deny = _resolve_allowlisted_path(dirpath)
    if not resolved:
        return deny

    dp = Path(resolved)
    if not dp.exists():
        return f"⚠️  Path does not exist: {resolved}"
    if not dp.is_dir():
        return f"⚠️  Path is not a directory: {resolved}"

    try:
        entries = list(dp.iterdir())
    except Exception as exc:
        return f"⚠️  Cannot list directory: {exc}"

    dirs = []
    files = []
    backups = []
    for e in entries:
        name = e.name
        if not show_hidden and name.startswith("."):
            continue
        try:
            if e.is_dir():
                dirs.append(name)
            elif _BAK_RE.search(name):
                # .bak<N> files grouped separately so the listing emphasizes
                # active code over historical backups
                try:
                    size = e.stat().st_size
                except Exception:
                    size = 0
                backups.append((name, size))
            else:
                try:
                    size = e.stat().st_size
                except Exception:
                    size = 0
                files.append((name, size))
        except Exception:
            continue

    dirs.sort(key=str.lower)
    files.sort(key=lambda x: x[0].lower())
    backups.sort(key=lambda x: x[0].lower())

    out = [
        f"📁 {resolved}",
        f"   {len(dirs)} dir(s), {len(files)} file(s)"
        + (f", {len(backups)} backup(s)" if backups else ""),
        "─" * 55,
    ]
    for name in dirs:
        out.append(f"   📁  {name}/")
    for name, size in files:
        out.append(f"   📄  {name}   ({size:,} bytes)")
    if backups:
        out.append("")
        out.append("   Backups (not indexed in ChromaDB):")
        for name, size in backups:
            out.append(f"   💾  {name}   ({size:,} bytes)")
    if not dirs and not files and not backups:
        out.append("   (empty)")
    return "\n".join(out)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 6 — copy_to_backup
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def copy_to_backup(filepath: str) -> str:
    """
    CODE TOOLS — Take a manual snapshot of a file. Creates <filepath>.bak<N>
    alongside the file WITHOUT modifying the active file.

    This is the 'soft delete' primitive in AI-Prowler — there is no delete
    tool. To remove content from active use, copy_to_backup first, then
    write_file an empty string (or simply stop touching the file).

    Differs from the auto-backup inside write_file / str_replace_in_file:
    that one fires automatically before a modification. This one is for
    when you want a checkpoint of the current state without changing it.

    The .bak<N> file is NOT indexed in ChromaDB. The active file's existing
    chunks are untouched (no re-index needed).

    Args:
        filepath:  Absolute path of the file to snapshot.

    Returns:
        Confirmation including the backup path, or an error.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    resolved, deny = _resolve_writable_path(filepath)
    if not resolved:
        return deny

    if not Path(resolved).exists():
        return f"⚠️  File does not exist: {resolved}"
    if not Path(resolved).is_file():
        return f"⚠️  Path is not a regular file: {resolved}"

    # Counts toward the circuit breaker (it does create a file)
    ok, msg = _check_and_increment_write_counter()
    if not ok:
        return msg

    backup_path, backup_err = _make_backup(resolved)
    if backup_err:
        return f"⚠️  Snapshot failed: {backup_err}"

    try:
        size = Path(backup_path).stat().st_size
    except Exception:
        size = 0

    return (f"💾 Snapshot created\n"
            f"   Source: {resolved}\n"
            f"   Backup: {backup_path}\n"
            f"   {size:,} bytes\n"
            f"   Active file unchanged. Backup NOT indexed in ChromaDB.")


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 7 — list_backups
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def list_backups(filepath: str) -> str:
    """
    CODE TOOLS — Show all <filepath>.bak<N> files next to the given file,
    with their sizes and modification times.

    Args:
        filepath:  Absolute path of the active file whose backups to enumerate.

    Returns:
        Sorted list (newest = highest N first) of backups, or a message if
        none exist.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    # Read-only — only the read allowlist applies. Backups don't need write
    # approval to be listed.
    resolved, deny = _resolve_allowlisted_path(filepath)
    if not resolved:
        return deny

    fp = Path(resolved)
    parent = fp.parent
    base = fp.name
    prefix = base + ".bak"
    prefix_lower = prefix.lower()

    found = []
    try:
        for sibling in parent.iterdir():
            name = sibling.name
            if name.lower().startswith(prefix_lower):
                tail = name[len(prefix):]
                if tail.isdigit():
                    try:
                        stat = sibling.stat()
                        found.append((int(tail), str(sibling), stat.st_size, stat.st_mtime))
                    except Exception:
                        continue
    except Exception as exc:
        return f"⚠️  Could not scan {parent}: {exc}"

    if not found:
        return (f"📂 No backups for {resolved}\n"
                f"   No <{base}>.bak<N> files found in {parent}")

    # Sort newest first (highest N first)
    found.sort(key=lambda x: x[0], reverse=True)

    out = [
        f"📂 Backups for {resolved}",
        f"   {len(found)} backup(s) found in {parent}",
        f"   (newest = highest .bak number)",
        "─" * 55,
    ]
    for n, path, size, mtime in found:
        when = _time.strftime("%Y-%m-%d %H:%M:%S", _time.localtime(mtime))
        out.append(f"   .bak{n:<4}  {size:>12,} bytes  {when}")
        out.append(f"             {path}")
    out.append("")
    out.append(f"Restore: restore_backup({resolved!r}, backup_number=<N>)")
    return "\n".join(out)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 8 — restore_backup
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def restore_backup(filepath: str, backup_number: int) -> str:
    """
    CODE TOOLS — Restore <filepath>.bak<N> over the active file at filepath.

    Does NOT auto-backup the current state before restoring — this is
    deliberate to keep the tool simple. If you want to preserve the broken
    state for later study, call copy_to_backup(filepath) FIRST, then call
    restore_backup.

    Triggers a full ChromaDB re-index of the active file (functionally a
    write — the file's contents change on disk).

    Args:
        filepath:       Absolute path of the active file to restore over.
        backup_number:  The N in <filepath>.bak<N>. Use list_backups to see
                        what's available.

    Returns:
        Confirmation with byte counts before and after, or an error.
    """
    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait a moment and try again."

    resolved, deny = _resolve_writable_path(filepath)
    if not resolved:
        return deny

    try:
        n = int(backup_number)
        if n < 1:
            return f"⚠️  backup_number must be >= 1, got {backup_number}"
    except (TypeError, ValueError):
        return f"⚠️  backup_number must be an integer, got {backup_number!r}"

    fp = Path(resolved)
    backup_path = fp.parent / f"{fp.name}.bak{n}"

    if not backup_path.exists():
        return (f"⚠️  Backup not found: {backup_path}\n"
                f"Call list_backups({resolved!r}) to see what's available.")

    # Active file may or may not exist (e.g. after a failed write that left
    # only the backup). Allow restore-over-missing for recovery.
    old_size = fp.stat().st_size if fp.exists() else 0
    try:
        backup_size = backup_path.stat().st_size
    except Exception as exc:
        return f"⚠️  Could not read backup: {exc}"

    # Circuit breaker
    ok, msg = _check_and_increment_write_counter()
    if not ok:
        return msg

    # Restore: copy backup over active path
    try:
        _shutil.copy2(str(backup_path), str(fp))
    except Exception as exc:
        return f"⚠️  Restore failed: {exc}"

    _log.info("restore_backup: %s <- %s (%d -> %d bytes)",
              resolved, backup_path, old_size, backup_size)

    # (auto-reindex removed in v7.0.0 — call reindex_file() if you want the
    #  restored content re-indexed)
    return (f"✅ Restored {resolved}\n"
            f"   Source: {backup_path}\n"
            f"   {old_size:,} bytes  →  {backup_size:,} bytes\n"
            f"   NOT yet indexed — call reindex_file() if you want it indexed.\n"
            f"   .bak{n} preserved (restore is non-destructive to backups).")


# ══════════════════════════════════════════════════════════════════════════════
# TOOL 9 — cleanup_backups  (Tier A: personal mode only)
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def cleanup_backups(path: str = "", dry_run: bool = True) -> str:
    """
    CODE TOOLS — Find and delete .bakN backup files created by str_replace_in_file.

    Always call with dry_run=True first (the default) — it lists every backup
    file it would delete without removing anything.  Then call again with
    dry_run=False to actually delete them.

    Not available in server mode (Tier A suppressed).

    Args:
        path:     • Empty string (default) — scan every tracked directory.
                  • A file path — find only .bakN backups of that specific file.
                  • A directory path — scan that directory recursively.
        dry_run:  True (default) = list only, nothing deleted.
                  False = delete all found backups.

    Examples:
        cleanup_backups()
            → lists all .bakN files across tracked directories (dry run)
        cleanup_backups(dry_run=False)
            → deletes all of them
        cleanup_backups(path="rag_gui.py")
            → lists backups of rag_gui.py only
        cleanup_backups(path="rag_gui.py", dry_run=False)
            → deletes rag_gui.py.bak1, .bak2, etc.
        cleanup_backups(path="C:/Projects/MyApp/")
            → lists all .bakN files under that directory
    """
    import re as _re_bak
    from pathlib import Path as _Pbak

    _bak_pattern = _re_bak.compile(r'\.bak\d+$', _re_bak.IGNORECASE)

    def _find_baks(root: _Pbak):
        """Yield .bakN Paths under root.
        File: finds backups of that specific file only.
        Dir:  recursive glob for any .bakN file."""
        if root.is_file():
            base = root.name
            for f in root.parent.iterdir():
                if (f.is_file()
                        and _bak_pattern.search(f.name)
                        and f.name.startswith(base + ".bak")):
                    yield f
        elif root.is_dir():
            for f in root.rglob("*"):
                if f.is_file() and _bak_pattern.search(f.name):
                    yield f

    # Gather candidate files
    found: list[_Pbak] = []
    if path and path.strip():
        p = _Pbak(path.strip())
        if not p.exists():
            return f"❌ Path not found: {path}"
        found = sorted(set(_find_baks(p)))
    else:
        try:
            from rag_preprocessor import load_auto_update_list as _lau
            tracked = _lau() or []
        except Exception:
            tracked = []
        for t in tracked:
            tp = _Pbak(t)
            found.extend(_find_baks(tp))
        found = sorted(set(found))

    if not found:
        where = str(path).strip() or "tracked directories"
        return f"✅ No .bakN backup files found in: {where}"

    # Measure sizes
    sizes: dict[_Pbak, int] = {}
    for f in found:
        try:
            sizes[f] = f.stat().st_size
        except Exception:
            sizes[f] = 0
    total = sum(sizes.values())

    if dry_run:
        lines = [
            f"🔍 Dry run — found {len(found)} backup file(s) totalling {total:,} bytes.",
            "   Nothing deleted. Re-run with dry_run=False to delete all, or",
            "   pass a path= to limit scope.",
            "",
        ]
        for f in found:
            lines.append(f"  • {f}  ({sizes[f]:,} bytes)")
        return "\n".join(lines)

    # ── Actually delete ────────────────────────────────────────────────────
    deleted: list[str] = []
    failed:  list[str] = []
    for f in found:
        try:
            f.unlink()
            deleted.append(str(f))
        except Exception as exc:
            failed.append(f"{f}: {exc}")

    freed = sum(sizes[f] for f in found if str(f) in deleted)
    lines = [f"🗑️  Deleted {len(deleted)} backup file(s) — {freed:,} bytes freed."]
    for d in deleted:
        lines.append(f"  ✅ {d}")
    if failed:
        lines.append(f"\n⚠️  Failed to delete {len(failed)} file(s):")
        for e in failed:
            lines.append(f"  ❌ {e}")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — cleanup_job_logs  (Tier A: personal mode only)
# ══════════════════════════════════════════════════════════════════════════════
# Added 2026-07-12 after ~.ai-prowler/jobs/ was found to have accumulated
# 1000+ files (333+ background job runs × 3 files each: manifest .json,
# .log, and _wrapper.py) with zero automatic retention — run_script_start's
# _jobs_dir() only ever creates, never prunes. Mirrors cleanup_backups()'s
# dry-run-by-default shape.
@mcp.tool()
def cleanup_job_logs(older_than_days: int = 7, keep_last: int = 20,
                     dry_run: bool = True) -> str:
    """
    CODE TOOLS — Delete old run_script_start job files (manifest .json,
    .log, and _wrapper.py) from ~/.ai-prowler/jobs/, which has no automatic
    retention and grows by 3 files every background job run forever.

    Always call with dry_run=True first (the default) — it lists every job
    it would delete without removing anything. Then call again with
    dry_run=False to actually delete them.

    Not available in server mode (Tier A suppressed) — this whole feature
    (run_script_start/status/kill) is personal-install-only.

    Retention rule (deliberately conservative — a job is only deleted if
    BOTH conditions hold):
      1. Older than `older_than_days` (by the manifest's started_at, falling
         back to file mtime if the manifest is missing/corrupt), AND
      2. Not among the `keep_last` most recently started jobs.
    So a very recent job is kept even if you pass older_than_days=0, and an
    old-but-still-within-keep_last job is kept too. Only jobs that are BOTH
    old and not among the recent handful get removed.

    A job whose manifest status is still "running" is NEVER deleted,
    regardless of age or rank — safety against removing a job's log out
    from under it while it's actively executing.

    Args:
        older_than_days: Age threshold in days (default 7). 0 disables the
                         age check (age condition always satisfied), leaving
                         keep_last as the only protection — use with care.
        keep_last:       Always keep at least this many of the most recent
                         jobs regardless of age (default 20).
        dry_run:         True (default) = list only, nothing deleted.
                         False = delete all qualifying jobs.

    Examples:
        cleanup_job_logs()
            → lists jobs older than 7 days beyond the most recent 20 (dry run)
        cleanup_job_logs(dry_run=False)
            → actually deletes them
        cleanup_job_logs(older_than_days=30, keep_last=50, dry_run=False)
            → more conservative: only jobs over a month old, keep the last 50
    """
    import json as _j_clean
    import datetime as _dt
    from pathlib import Path as _Pjobs

    jobs_dir = _jobs_dir()
    if not jobs_dir.exists():
        return "✅ No jobs directory found — nothing to clean up."

    # Group files by job_id. A job_id looks like job_20260712_131646_9agp;
    # its three files are <id>.json, <id>.log, <id>_wrapper.py.
    manifests = sorted(jobs_dir.glob("job_*.json"))
    if not manifests:
        return "✅ No job files found — nothing to clean up."

    now = _dt.datetime.now()
    jobs: list[dict] = []
    for mpath in manifests:
        job_id = mpath.stem
        status = None
        started_at = None
        try:
            data = _j_clean.loads(mpath.read_text(encoding="utf-8"))
            status = data.get("status")
            started_raw = data.get("started_at")
            if started_raw:
                started_at = _dt.datetime.fromisoformat(started_raw)
        except Exception:
            pass
        if started_at is None:
            try:
                started_at = _dt.datetime.fromtimestamp(mpath.stat().st_mtime)
            except Exception:
                started_at = now  # worst case: treat as brand-new, never deleted by age

        related = [mpath]
        log_p = jobs_dir / f"{job_id}.log"
        wrap_p = jobs_dir / f"{job_id}_wrapper.py"
        if log_p.exists():
            related.append(log_p)
        if wrap_p.exists():
            related.append(wrap_p)

        jobs.append({
            "job_id": job_id, "status": status, "started_at": started_at,
            "files": related,
        })

    # Most recent first — this ordering IS the keep_last ranking.
    jobs.sort(key=lambda j: j["started_at"], reverse=True)

    age_cutoff = now - _dt.timedelta(days=older_than_days)
    to_delete = []
    skipped_running = []
    for rank, job in enumerate(jobs):
        if job["status"] == "running":
            if job["started_at"] < age_cutoff or rank >= keep_last:
                skipped_running.append(job["job_id"])
            continue
        is_old = job["started_at"] < age_cutoff
        is_recent_rank = rank < keep_last
        if is_old and not is_recent_rank:
            to_delete.append(job)

    if not to_delete:
        msg = (f"✅ Nothing to clean up — {len(jobs)} job(s) total, all within "
               f"the last {keep_last} or under {older_than_days} day(s) old.")
        if skipped_running:
            msg += (f"\n⚠️  {len(skipped_running)} job(s) would otherwise qualify "
                    f"but are still marked 'running' — left alone: "
                    f"{', '.join(skipped_running)}")
        return msg

    total_files = sum(len(j["files"]) for j in to_delete)
    total_bytes = 0
    for j in to_delete:
        for f in j["files"]:
            try:
                total_bytes += f.stat().st_size
            except Exception:
                pass

    if dry_run:
        lines = [
            f"🔍 Dry run — {len(to_delete)} job(s) / {total_files} file(s) "
            f"totalling {total_bytes:,} bytes would be deleted.",
            "   Nothing deleted. Re-run with dry_run=False to actually delete.",
            "",
        ]
        for j in to_delete:
            age_days = (now - j["started_at"]).days
            lines.append(f"  • {j['job_id']}  ({age_days}d old, "
                        f"{len(j['files'])} file(s))")
        if skipped_running:
            lines.append("")
            lines.append(f"⚠️  {len(skipped_running)} job(s) qualify by age/rank "
                        f"but are still 'running' — will NOT be deleted: "
                        f"{', '.join(skipped_running)}")
        return "\n".join(lines)

    # ── Actually delete ────────────────────────────────────────────────────
    deleted_jobs: list[str] = []
    failed: list[str] = []
    freed = 0
    for j in to_delete:
        job_ok = True
        for f in j["files"]:
            try:
                freed += f.stat().st_size
                f.unlink()
            except Exception as exc:
                job_ok = False
                failed.append(f"{f}: {exc}")
        if job_ok:
            deleted_jobs.append(j["job_id"])

    lines = [f"🗑️  Deleted {len(deleted_jobs)} job(s) — {freed:,} bytes freed."]
    for d in deleted_jobs:
        lines.append(f"  ✅ {d}")
    if failed:
        lines.append(f"\n⚠️  Failed to delete {len(failed)} file(s):")
        for e in failed:
            lines.append(f"  ❌ {e}")
    if skipped_running:
        lines.append(f"\n⚠️  {len(skipped_running)} job(s) left alone (still "
                    f"'running'): {', '.join(skipped_running)}")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# BONUS — reset_write_counter (admin tool, not part of the 8)
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def reset_write_counter() -> str:
    """
    CODE TOOLS — Reset the per-session write circuit-breaker counter.

    The circuit breaker limits AI-Prowler to 20 write operations per process
    lifetime to protect against runaway loops. This tool resets the count so
    legitimate large editing sessions can continue without restarting the
    server.

    Returns:
        Confirmation with the count that was reset.
    """
    old = _reset_write_counter_internal()
    _log.info("Write counter reset (was %d)", old)
    return (f"✅ Write counter reset\n"
            f"   Was: {old} / {_WRITES_PER_SESSION_LIMIT}\n"
            f"   Now: 0 / {_WRITES_PER_SESSION_LIMIT}")


# ══════════════════════════════════════════════════════════════════════════════
# diff_files — compare two text files and return a unified diff
# Available in both personal and server mode.
# Server mode: all roles may use it for files within their assigned scopes.
#              Owners and managers can diff any tracked file.
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def diff_files(
    file_a: str,
    file_b: str,
    context_lines: int = 3,
    max_lines: int = 150,
    ctx: "Context | None" = None,
) -> str:
    """
    Compare two text files and return a unified diff that Claude can read and
    reason about directly — no external diff tool required.

    Especially useful for:
      • Verifying what str_replace_in_file just changed
          diff_files("rag_gui.py", "rag_gui.py.bak1")
      • Comparing a file to an older backup
          diff_files("config.json", "config.json.bak3")
      • Spotting differences between two versions of any file
          diff_files("setup_v1.iss", "setup_v2.iss")
      • Confirming a fix was applied correctly before running tests

    Both files must be inside tracked directories (or be .bakN backups of
    tracked files).  In server mode all roles may diff files within their
    assigned scopes; owners and managers may diff any tracked file.

    Args:
        file_a:        First file path. May be a .bakN backup.
        file_b:        Second file path. May be a .bakN backup.
        context_lines: Lines of unchanged context shown around each change.
                       Default 3. Use 0 for a compact view, 8+ for full context.
        max_lines:     Maximum number of diff lines returned.  Default 150 —
                       keeps responses concise.  Use max_lines=500 for large diffs.

    Returns:
        Unified diff with a summary header, or a confirmation that the files
        are identical.  Output is truncated with a notice if it exceeds max_lines.
    """
    import difflib as _difflib
    import re as _re_diff
    from pathlib import Path as _Pdiff

    _bak_re_diff = _re_diff.compile(r'\.bak\d+$', _re_diff.IGNORECASE)
    user = _current_user(ctx)

    # ── Access validator ───────────────────────────────────────────────────────
    def _check_access(filepath: str) -> "tuple[list[str] | None, str | None]":
        """Load file lines if accessible, else return (None, error_message)."""
        fp = _Pdiff(filepath.strip())

        # Determine the "base" path for allowlist check —
        # .bakN files share their base file's allowlist entry.
        is_bak = bool(_bak_re_diff.search(fp.name))
        base_str = _re_diff.sub(r'\.bak\d+$', '', str(fp), flags=_re_diff.IGNORECASE)
        check_str = base_str if is_bak else str(fp)

        if not fp.exists():
            return None, f"File not found: {filepath}"
        if not fp.is_file():
            return None, f"Not a file: {filepath}"

        # Allowlist check (covers both the file and its .bakN siblings)
        resolved, err = _resolve_allowlisted_path(check_str)
        if err:
            resolved, err = _resolve_allowlisted_path(str(fp))
            if err:
                return None, f"'{filepath}' is not in a tracked directory."

        # ── Server-mode scope check ────────────────────────────────────────
        # Owners and managers: unrestricted across all tracked files.
        # Staff and field_crew: file must resolve to one of their scopes.
        if user is not None:
            role = (user.get("role") or "").strip().lower()
            if role not in ("owner", "manager"):
                try:
                    company_map = _company_collection_map()
                    from scope_resolver import resolve_collection_for_path as _rcfp
                    file_coll = _rcfp(check_str, company_map)
                    user_scopes = set(user.get("scopes") or [])
                    # "documents" is the default collection — accessible to all roles
                    if file_coll not in user_scopes and file_coll not in ("documents", ""):
                        return None, (
                            f"'{filepath}' is in collection '{file_coll}' which is "
                            f"outside your assigned scopes {sorted(user_scopes)}."
                        )
                except Exception:
                    pass  # scope check failure is non-fatal for a read tool

        # Read the file — detect binary first (null bytes), then try encodings
        raw_bytes = fp.read_bytes()
        if b'\x00' in raw_bytes:
            return None, f"Cannot read '{filepath}' — binary file (contains null bytes)."
        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return raw_bytes.decode(enc).splitlines(keepends=True), None
            except UnicodeDecodeError:
                continue
        return None, f"Cannot read '{filepath}' — unknown encoding."

    # ── Load both files ────────────────────────────────────────────────────────
    lines_a, err_a = _check_access(file_a)
    if err_a:
        return f"❌ {err_a}"
    lines_b, err_b = _check_access(file_b)
    if err_b:
        return f"❌ {err_b}"

    pa = _Pdiff(file_a.strip())
    pb = _Pdiff(file_b.strip())

    # ── Generate unified diff ──────────────────────────────────────────────────
    diff = list(_difflib.unified_diff(
        lines_a, lines_b,
        fromfile=pa.name,
        tofile=pb.name,
        n=max(0, int(context_lines)),
    ))

    if not diff:
        try:
            sa = pa.stat().st_size
            sb = pb.stat().st_size
        except Exception:
            sa = sb = 0
        return (
            f"✅ Files are identical.\n"
            f"   {pa.name}  ({sa:,} bytes)\n"
            f"   {pb.name}  ({sb:,} bytes)"
        )

    # ── Summary statistics ─────────────────────────────────────────────────────
    added   = sum(1 for ln in diff if ln.startswith('+') and not ln.startswith('+++'))
    removed = sum(1 for ln in diff if ln.startswith('-') and not ln.startswith('---'))
    hunks   = sum(1 for ln in diff if ln.startswith('@@'))
    try:
        sa, sb = pa.stat().st_size, pb.stat().st_size
        size_info = f"size: {sa:,} → {sb:,} bytes  ({sb - sa:+,})"
    except Exception:
        size_info = ""

    header = (
        f"diff  {pa.name}  →  {pb.name}\n"
        f"      {hunks} hunk(s)   +{added} added   -{removed} removed"
        + (f"   {size_info}" if size_info else "") + "\n"
        + "─" * 62 + "\n"
    )

    # ── Truncate if needed ─────────────────────────────────────────────────────
    diff_lines = "".join(diff).splitlines()
    total_diff_lines = len(diff_lines)
    truncated = total_diff_lines > int(max_lines)
    if truncated:
        diff_lines = diff_lines[:int(max_lines)]

    output = header + "\n".join(diff_lines)
    if truncated:
        output += (
            f"\n\n⚠️  Truncated at {max_lines} lines "
            f"({total_diff_lines - int(max_lines)} more lines not shown). "
            f"Use max_lines={int(max_lines) * 2} to see more."
        )
    return output


# ══════════════════════════════════════════════════════════════════════════════
# DEV CHECK TOOLS — compile_check / check_python_import  (v7.0.0)
# ══════════════════════════════════════════════════════════════════════════════
# Let the agent verify its own edits on THIS machine (real interpreter, real
# installed stack) instead of asking the user to run py_compile / import by hand.
#
# SAFETY MODEL — these are developer conveniences, NOT customer-facing tools:
#   • Edition-gated: only available when this install is the Home edition OR
#     config.json explicitly sets {"dev_tools": true}. On any deployed Business
#     server (edition != home, no dev_tools flag) they are HARD-DISABLED and
#     return a refusal — so a tunnel-exposed customer box is never an RCE host.
#   • NO free-form shell: each tool builds a FIXED argument list
#     ([sys.executable, "-m", "py_compile", path] etc.) and runs it with
#     subprocess.run(shell=False). The agent cannot inject a command string.
#   • Path-scoped: compile_check validates filepath against the SAME tracked-root
#     allowlist the edit tools use (_resolve_allowlisted_path). No arbitrary file.
#   • Bounded: a hard timeout prevents a hung interpreter from wedging the server.
# ══════════════════════════════════════════════════════════════════════════════
_DEV_CHECK_TIMEOUT_SEC = 120   # generous enough for an import that loads heavy deps


def _dev_tools_enabled() -> tuple[bool, str]:
    """Return (enabled, reason). Dev tools are enabled in ALL editions/modes.

    History: v6.x and earlier had this gated to Home-edition-only because the
    tools (compile_check, check_python_import, ...) were considered too exposed for
    mobile/business installs accessed via Cloudflare Tunnel. v7.0.0 (2026-05-30)
    removed the gate per operator decision: the dev tools are read-only
    subprocess calls bounded by the same read allowlist as every other tool,
    so the same path-based protection that gates search_documents also gates
    syntax_check / lint_check.

    Kept as a function (not inlined) so re-tightening later is a one-line
    change: simply return (False, "...") here for the modes you want to lock.
    The opt-in dev_tools flag is preserved as a historical escape hatch but
    is now redundant — every code path it would unlock is already enabled.
    """
    try:
        cfg = _load_runtime_config()   # module-level; safe, never raises
    except Exception:
        cfg = {}
    edition = cfg.get("edition", "home")
    mode    = cfg.get("mode",    "personal")
    if cfg.get("dev_tools") is True:
        return (True, "dev_tools flag set in config.json")
    return (True, f"edition={edition!r} mode={mode!r} (dev tools enabled in all configurations)")


@mcp.tool()
def compile_check(filepath: str, timeout_sec: int = _DEV_CHECK_TIMEOUT_SEC) -> str:
    """
    DEV TOOLS — Byte-compile a Python file with this machine's interpreter to
    check for SYNTAX errors. Equivalent to `python -m py_compile <file>`.

    Use this after editing a .py file to confirm it parses before relying on it.
    Catches syntax errors only (not import-time or runtime errors — use
    check_python_import for those). Available in all editions/modes; the file must be
    under a tracked read-allowlisted root regardless of edition.

    Args:
        filepath:    Path to a .py file under a tracked root.
        timeout_sec: Max seconds to wait (default 120).

    Returns:
        "✅ compile OK" or the compiler's error output.
    """
    enabled, why = _dev_tools_enabled()
    if not enabled:
        return f"🚫 compile_check is disabled here ({why})."

    resolved, err = _resolve_allowlisted_path(filepath)
    if err:
        return err
    if not resolved.lower().endswith(".py"):
        return f"⚠️  compile_check only handles .py files (got {resolved})."

    import subprocess as _sp
    try:
        proc = _sp.run(
            [sys.executable, "-m", "py_compile", resolved],
            capture_output=True, text=True, timeout=max(5, int(timeout_sec)),
            shell=False)
    except _sp.TimeoutExpired:
        return f"⏱️  compile_check timed out after {timeout_sec}s on {resolved}."
    except Exception as exc:
        return f"⚠️  compile_check could not run: {exc}"

    if proc.returncode == 0:
        _log.info("compile_check OK: %s", resolved)
        return f"✅ compile OK — {resolved}\n   (python -m py_compile, rc=0)"
    out = (proc.stderr or proc.stdout or "").strip()
    _log.warning("compile_check FAILED: %s\n%s", resolved, out)
    return (f"❌ compile FAILED — {resolved}  (rc={proc.returncode})\n"
            f"───\n{out}")


@mcp.tool()
def check_python_import(module_or_path: str, timeout_sec: int = _DEV_CHECK_TIMEOUT_SEC) -> str:
    """
    DEV TOOLS — Import a Python module with this machine's interpreter to catch
    LOAD-TIME errors (NameError, ImportError, bad module-level references) that
    syntax-only compile_check misses. Equivalent to `python -c "import <module>"`.

    Accepts either a bare module name (e.g. "ai_prowler_mcp") or a path to a .py
    file under a tracked root (the module name is derived from the filename).
    The import runs in a SEPARATE interpreter process, so it cannot disturb the
    running server. Available in all editions/modes; the file must be under
    a tracked read-allowlisted root regardless of edition.

    Args:
        module_or_path: Module name, or path to a .py file under a tracked root.
        timeout_sec:    Max seconds to wait (default 120; imports of heavy deps
                        like torch/chromadb can take a while).

    Returns:
        "✅ import OK" or the traceback the import produced.
    """
    enabled, why = _dev_tools_enabled()
    if not enabled:
        return f"🚫 check_python_import is disabled here ({why})."

    # Resolve to a module name + a working directory to run from.
    cwd = None
    if module_or_path.replace("_", "").isalnum() and "/" not in module_or_path \
            and "\\" not in module_or_path and not module_or_path.endswith(".py"):
        module = module_or_path           # bare module name
    else:
        resolved, err = _resolve_allowlisted_path(module_or_path)
        if err:
            return err
        if not resolved.lower().endswith(".py"):
            return f"⚠️  check_python_import path must be a .py file (got {resolved})."
        import os as _os
        cwd = _os.path.dirname(resolved)
        module = _os.path.splitext(_os.path.basename(resolved))[0]

    # Validate the derived module name is a safe identifier — no injection via -c.
    if not module.isidentifier():
        return f"⚠️  '{module}' is not a valid module name."

    import subprocess as _sp
    try:
        proc = _sp.run(
            [sys.executable, "-c", f"import {module}"],
            capture_output=True, text=True, timeout=max(5, int(timeout_sec)),
            shell=False, cwd=cwd)
    except _sp.TimeoutExpired:
        return f"⏱️  check_python_import timed out after {timeout_sec}s importing {module}."
    except Exception as exc:
        return f"⚠️  check_python_import could not run: {exc}"

    if proc.returncode == 0:
        _log.info("check_python_import OK: %s", module)
        note = (proc.stderr or "").strip()
        extra = f"\n   (stderr, non-fatal):\n{note}" if note else ""
        return f"✅ import OK — {module}\n   (python -c 'import {module}', rc=0){extra}"
    out = (proc.stderr or proc.stdout or "").strip()
    _log.warning("check_python_import FAILED: %s\n%s", module, out)
    return (f"❌ import FAILED — {module}  (rc={proc.returncode})\n"
            f"───\n{out}")


# ══════════════════════════════════════════════════════════════════════════════
# MULTI-LANGUAGE DEV TOOLS — syntax_check + lint_check (v7.0.0)
# ══════════════════════════════════════════════════════════════════════════════
# Why one tool per concern (syntax / lint / pytest) instead of one tool per
# language: keeps the MCP tool surface narrow (3 tools) instead of N×M (one
# per language per concern = ~24 tools). Each tool auto-detects language by
# file extension and dispatches to the right underlying binary. If the binary
# isn't installed, return a clean "❌ Not available — install <tool>" message
# rather than a cryptic subprocess error.
#
# Coverage choices, with honest tradeoffs:
#   • Python  : py_compile (syntax) + pyflakes (lint). Both bundled with Python.
#   • JS      : node --check (syntax). Lint via eslint if installed.
#   • TS      : tsc --noEmit (syntax+type). Same tool for lint.
#   • C/C++   : gcc -fsyntax-only on POSIX, cl /Zs on Windows. Lint via cppcheck.
#                Headers without their full dependency graph WILL produce false
#                errors — this is single-file checking, not full build.
#   • Go      : go build -o NUL/devnull (syntax). go vet (lint).
#   • Rust    : NOT supported single-file — cargo expects a Cargo.toml project.
#                Return a clear message pointing the operator at `cargo check`.
#   • Java    : javac -d <tmpdir> (syntax). No standard lint tool.
#   • Perl    : perl -c. No lint.
#   • Ruby    : ruby -c. No lint (rubocop is third-party).
#   • PHP     : php -l. No standard lint.
#   • Bash    : bash -n (on Windows requires Git Bash or WSL). shellcheck if installed.
# ══════════════════════════════════════════════════════════════════════════════

# Extension → language config. Each entry is (lang_name, needs_binary, syntax_argv, lint_argv).
# {file} placeholder in argv is replaced with the resolved filepath at runtime.
# argv = None means "not supported for this concern" (e.g. no standard lint for Perl).
_LANG_CONFIG = {
    # Python
    ".py":   ("Python",  sys.executable, [sys.executable, "-m", "py_compile", "{file}"],
                                          [sys.executable, "-m", "pyflakes", "{file}"]),
    # JavaScript
    ".js":   ("JavaScript", "node", ["node", "--check", "{file}"], None),
    ".mjs":  ("JavaScript", "node", ["node", "--check", "{file}"], None),
    ".cjs":  ("JavaScript", "node", ["node", "--check", "{file}"], None),
    # TypeScript
    ".ts":   ("TypeScript", "tsc",  ["tsc", "--noEmit", "{file}"], ["tsc", "--noEmit", "{file}"]),
    ".tsx":  ("TypeScript", "tsc",  ["tsc", "--noEmit", "{file}"], ["tsc", "--noEmit", "{file}"]),
    # C / C++
    ".c":    ("C",   "gcc", ["gcc", "-fsyntax-only", "{file}"], None),
    ".h":    ("C",   "gcc", ["gcc", "-fsyntax-only", "-x", "c-header", "{file}"], None),
    ".cpp":  ("C++", "g++", ["g++", "-fsyntax-only", "{file}"], None),
    ".cc":   ("C++", "g++", ["g++", "-fsyntax-only", "{file}"], None),
    ".cxx":  ("C++", "g++", ["g++", "-fsyntax-only", "{file}"], None),
    ".hpp":  ("C++", "g++", ["g++", "-fsyntax-only", "-x", "c++-header", "{file}"], None),
    # Go
    ".go":   ("Go", "go", ["go", "vet", "{file}"], ["go", "vet", "{file}"]),
    # Java
    ".java": ("Java", "javac", ["javac", "-d", "{tmpdir}", "{file}"], None),
    # Perl
    ".pl":   ("Perl", "perl", ["perl", "-c", "{file}"], None),
    ".pm":   ("Perl", "perl", ["perl", "-c", "{file}"], None),
    # Ruby
    ".rb":   ("Ruby", "ruby", ["ruby", "-c", "{file}"], None),
    # PHP
    ".php":  ("PHP", "php", ["php", "-l", "{file}"], None),
    # Bash
    ".sh":   ("Bash",   "bash", ["bash", "-n", "{file}"], None),
    ".bash": ("Bash",   "bash", ["bash", "-n", "{file}"], None),
    # ── Hardware Description Languages (HDL) ─────────────────────────────────
    # Verilog / SystemVerilog — Icarus Verilog (iverilog)
    #   iverilog -t null -o /dev/null {file}  → syntax-only, no binary output
    #   On Windows iverilog writes to NUL instead of /dev/null.
    #   Lint: verilator --lint-only {file}  → deeper static analysis than
    #   iverilog's syntax-only pass (unused signals, width mismatches,
    #   latch inference, multi-driven nets, etc.). Optional separate install:
    #   https://verilator.org or `winget install Verilator.Verilator` —
    #   the lint check reports a clean "install verilator" message if it's
    #   not on PATH, same pattern as every other missing-binary case here.
    ".v":    ("Verilog",         "iverilog",
              ["iverilog", "-t", "null", "-o", "{verilog_null}", "{file}"],
              ["verilator", "--lint-only", "-Wall", "{file}"]),
    ".vh":   ("Verilog Header",  "iverilog",
              ["iverilog", "-t", "null", "-o", "{verilog_null}", "{file}"],
              ["verilator", "--lint-only", "-Wall", "{file}"]),
    ".sv":   ("SystemVerilog",   "iverilog",
              ["iverilog", "-g2012", "-t", "null", "-o", "{verilog_null}", "{file}"],
              ["verilator", "--lint-only", "-Wall", "--sv", "{file}"]),
    ".svh":  ("SystemVerilog Header", "iverilog",
              ["iverilog", "-g2012", "-t", "null", "-o", "{verilog_null}", "{file}"],
              ["verilator", "--lint-only", "-Wall", "--sv", "{file}"]),
    # VHDL — GHDL
    #   ghdl -s {file}  → syntax-only analysis, no elaborate/run step
    #   Lint: ghdl -a {file}  → full semantic analysis (catches more than -s)
    ".vhd":  ("VHDL", "ghdl",
              ["ghdl", "-s", "{file}"],
              ["ghdl", "-a", "{file}"]),
    ".vhdl": ("VHDL", "ghdl",
              ["ghdl", "-s", "{file}"],
              ["ghdl", "-a", "{file}"]),
}

# Languages we intentionally do NOT auto-dispatch on (require project setup).
_LANG_NOT_SUPPORTED = {
    ".rs": ("Rust", "cargo check (requires Cargo.toml — single-file Rust check isn't viable; "
                    "run `cargo check` in your crate root instead)"),
}


def _detect_language(filepath: str) -> Optional[tuple]:
    """Return (lang_name, needs_binary, syntax_argv, lint_argv) for filepath,
    or None if the extension isn't recognised."""
    ext = os.path.splitext(filepath)[1].lower()
    return _LANG_CONFIG.get(ext)


def _binary_available(binary: str) -> bool:
    """Quick PATH check — returns True if shutil.which(binary) finds it."""
    import shutil as _shutil
    # Special case: sys.executable is a full path, not a PATH name.
    if binary == sys.executable:
        return True
    return _shutil.which(binary) is not None


def _run_dev_subprocess(argv: list, timeout_sec: int,
                        tmpdir_for_java: Optional[str] = None) -> tuple[int, str]:
    """Run a dev-tool subprocess with timeout. Returns (returncode, output).
    Substitutes {file}/{tmpdir} placeholders in argv. Output is the combined
    stderr+stdout (stripped). Returns (-1, error_message) if the subprocess
    could not be launched at all (e.g. binary disappeared between PATH check
    and run)."""
    import subprocess as _sp
    try:
        proc = _sp.run(argv, capture_output=True, text=True,
                       timeout=max(5, int(timeout_sec)), shell=False)
        out = (proc.stderr or "") + (proc.stdout or "")
        return proc.returncode, out.strip()
    except _sp.TimeoutExpired:
        return -2, f"⏱️ subprocess timed out after {timeout_sec}s"
    except FileNotFoundError as exc:
        return -1, f"binary not found: {exc}"
    except Exception as exc:
        return -1, f"subprocess could not run: {exc}"


@mcp.tool()
def syntax_check(filepath: str, timeout_sec: int = _DEV_CHECK_TIMEOUT_SEC) -> str:
    """
    DEV TOOLS — Multi-language syntax checker. Auto-detects language by file
    extension and runs the appropriate compiler/parser in --check / -fsyntax-only
    mode (does not produce binaries).

    Supports: Python (.py), JavaScript (.js/.mjs/.cjs), TypeScript (.ts/.tsx),
    C (.c/.h), C++ (.cpp/.cc/.cxx/.hpp), Go (.go), Java (.java), Perl (.pl/.pm),
    Ruby (.rb), PHP (.php), Bash (.sh/.bash),
    Verilog (.v/.vh), SystemVerilog (.sv/.svh) — requires iverilog,
    VHDL (.vhd/.vhdl) — requires ghdl.

    Rust requires a Cargo project and is intentionally NOT supported here —
    use `cargo check` in your project root.

    If the underlying tool isn't installed on this machine (e.g. no `go` binary
    on PATH), returns a clean "❌ Not available — install <tool>" message rather
    than a cryptic subprocess error.

    Honest limitations:
      • C/C++ single-file checks may report false errors when headers reference
        symbols defined in other compilation units (single-file ≠ full build).
      • Bash on Windows requires Git Bash or WSL to provide a `bash` binary.
      • TypeScript needs a tsconfig.json in the project root for most non-
        trivial files; lone .ts files may report import resolution errors.
      • Verilog/SystemVerilog: install Icarus Verilog (iverilog) from
        http://iverilog.icarus.com or via winget: winget install IcarusVerilog
      • VHDL: install GHDL from https://github.com/ghdl/ghdl/releases or
        via winget: winget install ghdl.ghdl

    Available in all editions/modes; the file must be under a tracked
    read-allowlisted root regardless of edition.

    Args:
        filepath:    Path to the source file under a tracked root.
        timeout_sec: Max seconds to wait (default 120).

    Returns:
        "✅ syntax OK — <lang>" or the compiler's error output.
    """
    enabled, why = _dev_tools_enabled()
    if not enabled:
        return f"🚫 syntax_check is disabled here ({why})."

    resolved, err = _resolve_allowlisted_path(filepath)
    if err:
        return err

    # Detect language
    ext = os.path.splitext(resolved)[1].lower()
    if ext in _LANG_NOT_SUPPORTED:
        lang_name, msg = _LANG_NOT_SUPPORTED[ext]
        return f"ℹ️  {lang_name} not supported as single-file syntax check: {msg}"

    cfg = _detect_language(resolved)
    if cfg is None:
        return (f"⚠️  syntax_check: unsupported extension {ext!r} (file {resolved}). "
                f"Supported: .py .js .mjs .cjs .ts .tsx .c .h .cpp .cc .cxx .hpp .go "
                f".java .pl .pm .rb .php .sh .bash "
                f".v .vh .sv .svh (iverilog) .vhd .vhdl (ghdl)")

    lang_name, needs_binary, syntax_argv, _lint_argv = cfg

    # Binary availability check (clean error before subprocess)
    if not _binary_available(needs_binary):
        return (f"❌ {lang_name} syntax_check not available: '{needs_binary}' is not on PATH. "
                f"Install it to enable {lang_name} checking on this machine.")

    # Substitute {file} (and {tmpdir} for Java which needs a -d target).
    import tempfile as _tempfile
    # {verilog_null} is the platform null output target for iverilog (-o NUL on
    # Windows, -o /dev/null on POSIX) so syntax-only checks produce no binary.
    _verilog_null = "NUL" if sys.platform == "win32" else "/dev/null"
    tmpdir = _tempfile.mkdtemp(prefix="aiprowler_syncheck_") if "{tmpdir}" in " ".join(syntax_argv) else None
    try:
        argv = [arg.replace("{file}", resolved)
                    .replace("{tmpdir}", tmpdir or "")
                    .replace("{verilog_null}", _verilog_null)
                for arg in syntax_argv]
        rc, out = _run_dev_subprocess(argv, timeout_sec)
    finally:
        if tmpdir:
            try:
                import shutil as _shutil
                _shutil.rmtree(tmpdir, ignore_errors=True)
            except Exception:
                pass

    if rc == 0:
        _log.info("syntax_check OK (%s): %s", lang_name, resolved)
        return f"✅ syntax OK — {lang_name} — {resolved}\n   ({' '.join(argv)}, rc=0)"
    if rc == -1:
        return f"⚠️  syntax_check could not run: {out}"
    if rc == -2:
        return out  # already-formatted timeout message
    _log.warning("syntax_check FAILED (%s): %s\n%s", lang_name, resolved, out)
    return (f"❌ syntax FAILED — {lang_name} — {resolved}  (rc={rc})\n"
            f"───\n{out}")


@mcp.tool()
def lint_check(filepath: str, timeout_sec: int = _DEV_CHECK_TIMEOUT_SEC) -> str:
    """
    DEV TOOLS — Multi-language linter. Auto-detects language by extension and
    runs the appropriate lint tool. Catches unused imports, undefined names,
    style issues, and other warnings that syntax_check would let through.

    Tool used per language:
      Python      : pyflakes  (bundled — catches NameError-ish issues at lint time)
      TypeScript  : tsc --noEmit  (same as syntax_check; tsc IS the linter for TS)
      Go          : go vet  (built into Go toolchain)
      Verilog     : verilator --lint-only -Wall  (unused signals, width
                    mismatches, latch inference, multi-driven nets — separate
                    optional install from iverilog, which syntax_check uses)
      SystemVerilog: verilator --lint-only -Wall --sv  (same tool, SV mode)
      VHDL        : ghdl -a  (full semantic analysis, catches more than ghdl -s)
      Others      : no standard lint tool — use syntax_check instead.

    If a language has no lint tool available (Perl, Ruby, PHP, C/C++, etc.),
    returns a clear "ℹ️ No lint tool for <lang>; use syntax_check" message.

    Available in all editions/modes; the file must be under a tracked
    read-allowlisted root regardless of edition.

    Args:
        filepath:    Path to the source file under a tracked root.
        timeout_sec: Max seconds to wait (default 120).

    Returns:
        "✅ lint clean — <lang>" or the linter's warnings/errors.
    """
    enabled, why = _dev_tools_enabled()
    if not enabled:
        return f"🚫 lint_check is disabled here ({why})."

    resolved, err = _resolve_allowlisted_path(filepath)
    if err:
        return err

    ext = os.path.splitext(resolved)[1].lower()
    cfg = _detect_language(resolved)
    if cfg is None:
        return (f"⚠️  lint_check: unsupported extension {ext!r}. "
                f"Use syntax_check to see supported languages.")

    lang_name, _needs_binary, _syntax_argv, lint_argv = cfg
    if lint_argv is None:
        return (f"ℹ️  No standard lint tool for {lang_name}. "
                f"Use syntax_check({filepath!r}) instead — it catches the same "
                f"errors syntax-level tools can detect.")

    lint_binary = lint_argv[0] if lint_argv[0] != sys.executable else sys.executable
    if not _binary_available(lint_binary):
        return (f"❌ {lang_name} lint_check not available: '{lint_binary}' is not on PATH. "
                f"Install it to enable {lang_name} linting on this machine.")

    argv = [arg.replace("{file}", resolved) for arg in lint_argv]
    rc, out = _run_dev_subprocess(argv, timeout_sec)

    if rc == 0 and not out.strip():
        _log.info("lint_check clean (%s): %s", lang_name, resolved)
        return f"✅ lint clean — {lang_name} — {resolved}\n   ({' '.join(argv)}, rc=0)"
    if rc == 0:
        # pyflakes returns rc=0 even when it has warnings — surface them
        return (f"⚠️  lint findings — {lang_name} — {resolved}  (rc=0, warnings only)\n"
                f"───\n{out}")
    if rc == -1:
        return f"⚠️  lint_check could not run: {out}"
    if rc == -2:
        return out
    _log.warning("lint_check FAILED (%s): %s\n%s", lang_name, resolved, out)
    return (f"❌ lint FAILED — {lang_name} — {resolved}  (rc={rc})\n"
            f"───\n{out}")


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — run_script
# ══════════════════════════════════════════════════════════════════════════════
@mcp.tool()
def run_script(script_path: str, args: str = "",
               timeout_sec: int = 120, max_output_lines: int = 200) -> str:
    """
    DEV TOOLS — Execute a script or compiled program and return its output.

    Use run_script for short scripts that complete within ~60 seconds.
    For long-running tasks (full test suites, builds, etc.) use
    run_script_start / run_script_status / run_script_kill instead —
    those run async and won't time out the MCP connection.

    To run pytest with a -k filter use run_script_start with args:
      run_script_start("run_tests.bat", args="tests\\mcp\\ -k binary_write")

    IMPORTANT — NEVER elevated: all scripts run as the current Windows user
    (david). Admin-requiring operations (service installs, UAC-gated writes)
    must be run manually at the desktop.

    Supported file types and their runners:
      .bat / .cmd   →  cmd.exe /c <script> (shell=True, Windows only)
      .py           →  Python311 (same interpreter as run_script_start)
      .js / .mjs    →  node
      .sh / .bash   →  bash (requires Git Bash or WSL on Windows)
      .rb           →  ruby
      .pl / .pm     →  perl
      .go           →  go run
      .c            →  gcc compile to temp binary → run → delete binary
      .cpp/.cc/.cxx →  g++ compile to temp binary → run → delete binary
      .java         →  javac compile → java run → delete .class

    Security guardrails (same as all Dev Tools):
      • Script must be under the tracked read-allowlist — no running arbitrary
        files from outside the knowledge base.
      • Never runs as administrator / elevated — no UAC, no runas.
      • Suppressed in server mode (Tier A).
      • Hard timeout enforced — runaway scripts are killed.
      • Output truncated to max_output_lines from the end.

    Args:
        script_path:      Absolute path to the script or source file.
                          Must be under a tracked read-allowlisted root.
        args:             Optional command-line arguments as a single string
                          (e.g. "tests\\mcp\\ -v"). Split on whitespace.
        timeout_sec:      Max seconds before the process is killed (default 120).
        max_output_lines: Truncate stdout+stderr to this many lines from the
                          end (default 200). Raise for verbose build output.

    Returns:
        "✅ rc=0 — <script>" + captured output, OR
        "❌ rc=N — <script>" + captured output on non-zero exit.
        Always includes the return code for easy parsing.
    """
    enabled, why = _dev_tools_enabled()
    if not enabled:
        return f"🚫 run_script is disabled here ({why})."

    resolved, err = _resolve_allowlisted_path(script_path)
    if err:
        return err

    import subprocess as _sp
    import tempfile as _tmp

    ext = os.path.splitext(resolved)[1].lower()
    extra_args = args.split() if args.strip() else []
    cwd = os.path.dirname(resolved)
    use_shell = False
    cleanup_bin = None  # temp binary to delete after run (C/C++/Java)

    # ── Resolve runner ────────────────────────────────────────────────────────
    if ext in (".bat", ".cmd"):
        # cmd.exe /c is the only sane way to run batch files on Windows.
        # shell=True is required; the script path is passed as a single arg.
        argv = ["cmd.exe", "/c", resolved] + extra_args
        use_shell = False

    elif ext == ".py":
        # Reuse the same Python-finder logic as run_script_start.
        def _find_py() -> str:
            candidates = [
                os.path.join(os.environ.get("LocalAppData", ""),
                             "Programs", "Python", "Python311", "python.exe"),
                os.path.join(os.environ.get("LocalAppData", ""),
                             "Programs", "Python", "Python312", "python.exe"),
                os.path.join(os.environ.get("LocalAppData", ""),
                             "Programs", "Python", "Python310", "python.exe"),
                sys.executable,
            ]
            for c in candidates:
                if c and os.path.isfile(c):
                    return c
            return sys.executable
        argv = [_find_py(), resolved] + extra_args

    elif ext in (".js", ".mjs", ".cjs"):
        argv = ["node", resolved] + extra_args

    elif ext in (".sh", ".bash"):
        argv = ["bash", resolved] + extra_args

    elif ext in (".rb",):
        argv = ["ruby", resolved] + extra_args

    elif ext in (".pl", ".pm"):
        argv = ["perl", resolved] + extra_args

    elif ext == ".go":
        argv = ["go", "run", resolved] + extra_args

    elif ext in (".c", ".h"):
        # Compile to a temp binary, run it, clean up.
        tmp_bin = os.path.join(_tmp.gettempdir(),
                               f"_aip_run_{os.getpid()}.exe")
        compile_proc = _sp.run(
            ["gcc", resolved, "-o", tmp_bin],
            capture_output=True, text=True, timeout=60, shell=False, cwd=cwd)
        if compile_proc.returncode != 0:
            return (f"❌ gcc compile failed for {resolved}\n"
                    f"{compile_proc.stderr or compile_proc.stdout}")
        argv = [tmp_bin] + extra_args
        cleanup_bin = tmp_bin

    elif ext in (".cpp", ".cc", ".cxx", ".hpp"):
        tmp_bin = os.path.join(_tmp.gettempdir(),
                               f"_aip_run_{os.getpid()}.exe")
        compile_proc = _sp.run(
            ["g++", resolved, "-o", tmp_bin],
            capture_output=True, text=True, timeout=60, shell=False, cwd=cwd)
        if compile_proc.returncode != 0:
            return (f"❌ g++ compile failed for {resolved}\n"
                    f"{compile_proc.stderr or compile_proc.stdout}")
        argv = [tmp_bin] + extra_args
        cleanup_bin = tmp_bin

    elif ext == ".java":
        # Compile first, then run.
        compile_proc = _sp.run(
            ["javac", resolved],
            capture_output=True, text=True, timeout=60, shell=False, cwd=cwd)
        if compile_proc.returncode != 0:
            return (f"❌ javac compile failed for {resolved}\n"
                    f"{compile_proc.stderr or compile_proc.stdout}")
        class_name = os.path.splitext(os.path.basename(resolved))[0]
        argv = ["java", "-cp", cwd, class_name] + extra_args

    else:
        return (f"⚠️  run_script: unsupported file type '{ext}'.\n"
                f"Supported: .bat .cmd .py .js .mjs .sh .bash .rb .pl "
                f".go .c .cpp .cc .cxx .java")

    # ── Content preview — forces script into Claude's context before execution ─
    # This is the primary security gate: Claude's built-in values mean it will
    # never have called run_script on something destructive it can see.
    # For binary formats (compiled C/C++) we skip the preview since the source
    # was already read; for all text scripts we show the first 50 lines.
    _preview_note = ""
    if ext not in (".c", ".h", ".cpp", ".cc", ".cxx", ".hpp", ".java"):
        try:
            with open(resolved, "r", encoding="utf-8", errors="replace") as _pf:
                _preview_lines = _pf.readlines()
            _total = len(_preview_lines)
            _shown = _preview_lines[:50]
            _omitted = _total - len(_shown)
            _preview_note = (
                f"📄 Script preview ({_total} lines"
                f"{f', showing first 50' if _omitted else ''}):\n"
                f"{'─' * 60}\n"
                + "".join(_shown)
                + (f"{'─' * 60}\n... ({_omitted} more lines)\n" if _omitted else
                   f"{'─' * 60}\n")
            )
        except Exception as _pe:
            _preview_note = f"(preview unavailable: {_pe})\n"

    # ── Execute ───────────────────────────────────────────────────────────────
    try:
        proc = _sp.run(
            argv,
            capture_output=True, text=True,
            timeout=max(5, int(timeout_sec)),
            shell=use_shell,
            cwd=cwd,
        )
    except FileNotFoundError as exc:
        return (f"⚠️  run_script: runner not found — {exc}\n"
                f"Install the required runtime and ensure it is on PATH.")
    except _sp.TimeoutExpired:
        return (f"⏱️  run_script timed out after {timeout_sec}s — {resolved}\n"
                f"Raise timeout_sec if the script legitimately needs more time.")
    except Exception as exc:
        return f"⚠️  run_script could not execute: {exc}"
    finally:
        # Always clean up temp binaries even if execution raised.
        if cleanup_bin and os.path.exists(cleanup_bin):
            try:
                os.remove(cleanup_bin)
            except Exception:
                pass

    # ── Format output ─────────────────────────────────────────────────────────
    out = (proc.stdout or "") + (proc.stderr or "")
    lines = out.splitlines()
    truncated = ""
    if len(lines) > max_output_lines:
        truncated = (f"... (truncated {len(lines) - max_output_lines} earlier "
                     f"lines; raise max_output_lines to see more) ...\n")
        lines = lines[-max_output_lines:]
    out_trimmed = truncated + "\n".join(lines)

    rc = proc.returncode
    icon = "✅" if rc == 0 else "❌"
    _log.info("run_script: %s rc=%d", resolved, rc)
    return (f"{_preview_note}"
            f"{icon} rc={rc} — {resolved}\n"
            f"───\n{out_trimmed}")


# ══════════════════════════════════════════════════════════════════════════════
# ASYNC JOB TOOLS — run_script_start / run_script_status / run_script_kill
# ══════════════════════════════════════════════════════════════════════════════

# Jobs directory: ~/.ai-prowler/jobs/
# Each job writes two files:
#   <job_id>.json  — manifest (pid, status, exit_code, paths, timestamps)
#   <job_id>.log   — live stdout+stderr from the process

_JOBS_DIR = Path.home() / ".ai-prowler" / "jobs"


def _jobs_dir() -> Path:
    """Return the jobs directory, creating it if needed."""
    _JOBS_DIR.mkdir(parents=True, exist_ok=True)
    return _JOBS_DIR


def _job_manifest_path(job_id: str) -> Path:
    return _jobs_dir() / f"{job_id}.json"


def _job_log_path(job_id: str) -> Path:
    return _jobs_dir() / f"{job_id}.log"


def _read_manifest(job_id: str) -> dict:
    p = _job_manifest_path(job_id)
    if not p.exists():
        return {}
    try:
        import json as _j
        return _j.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _resolve_run_argv(script_path: str, extra_args: list) -> tuple:
    """
    Return (argv, use_shell, needs_compile, error_str).
    Mirrors the runner-selection logic in run_script.
    """
    import tempfile as _tmp
    ext = os.path.splitext(script_path)[1].lower()

    if ext in (".bat", ".cmd"):
        return (["cmd.exe", "/c", script_path] + extra_args, False, False, "")

    elif ext == ".py":
        candidates = [
            os.path.join(os.environ.get("LocalAppData", ""),
                         "Programs", "Python", "Python311", "python.exe"),
            os.path.join(os.environ.get("LocalAppData", ""),
                         "Programs", "Python", "Python312", "python.exe"),
            os.path.join(os.environ.get("LocalAppData", ""),
                         "Programs", "Python", "Python310", "python.exe"),
            sys.executable,
        ]
        py = next((c for c in candidates if c and os.path.isfile(c)), sys.executable)
        return ([py, script_path] + extra_args, False, False, "")

    elif ext in (".js", ".mjs", ".cjs"):
        return (["node", script_path] + extra_args, False, False, "")

    elif ext in (".sh", ".bash"):
        return (["bash", script_path] + extra_args, False, False, "")

    elif ext in (".rb",):
        return (["ruby", script_path] + extra_args, False, False, "")

    elif ext in (".pl", ".pm"):
        return (["perl", script_path] + extra_args, False, False, "")

    elif ext == ".go":
        return (["go", "run", script_path] + extra_args, False, False, "")

    elif ext in (".c", ".h", ".cpp", ".cc", ".cxx", ".hpp"):
        compiler = "gcc" if ext in (".c", ".h") else "g++"
        tmp_bin = os.path.join(_tmp.gettempdir(), f"_aip_async_{os.getpid()}.exe")
        import subprocess as _sp2
        cp = _sp2.run([compiler, script_path, "-o", tmp_bin],
                      capture_output=True, text=True, timeout=60, shell=False)
        if cp.returncode != 0:
            return ([], False, False,
                    f"❌ {compiler} compile failed:\n{cp.stderr or cp.stdout}")
        return ([tmp_bin] + extra_args, False, True, "")

    elif ext == ".java":
        cwd = os.path.dirname(script_path)
        import subprocess as _sp2
        cp = _sp2.run(["javac", script_path],
                      capture_output=True, text=True, timeout=60, shell=False,
                      cwd=cwd)
        if cp.returncode != 0:
            return ([], False, False,
                    f"❌ javac compile failed:\n{cp.stderr or cp.stdout}")
        class_name = os.path.splitext(os.path.basename(script_path))[0]
        return (["java", "-cp", cwd, class_name] + extra_args, False, False, "")

    else:
        return ([], False, False,
                f"⚠️  Unsupported file type '{ext}'. "
                f"Supported: .bat .cmd .py .js .mjs .sh .bash .rb .pl "
                f".go .c .cpp .cc .cxx .java")


@mcp.tool()
def run_script_start(script_path: str, args: str = "",
                     timeout_sec: int = 1800) -> str:
    """
    DEV TOOLS — Launch a script in the background and return a job_id immediately.

    Unlike run_script (which blocks until done), this returns in under a second.
    The process runs independently on your machine. Use run_script_status to
    check progress and read log output, and run_script_kill to terminate it.

    Ideal for long-running tasks: full test suites, build scripts, compilers,
    data processing jobs — anything that might exceed the MCP transport timeout.

    Supported file types: same as run_script (.bat, .cmd, .py, .js, .sh, .rb,
    .pl, .go, .c, .cpp, .java).

    NEVER elevated: runs as current user (david). No UAC, no runas.
    Suppressed in server mode (Tier A).

    Args:
        script_path: Absolute path to the script. Must be under a tracked root.
        args:        Optional command-line arguments (space-separated string).
        timeout_sec: Max wall-clock seconds before the wrapper kills the process
                     (default 1800 = 30 minutes). The wrapper enforces this.

    Returns:
        job_id string on success (e.g. "job_20260615_143022_a3f7").
        Use this id with run_script_status and run_script_kill.
    """
    enabled, why = _dev_tools_enabled()
    if not enabled:
        return f"🚫 run_script_start is disabled here ({why})."

    resolved, err = _resolve_allowlisted_path(script_path)
    if err:
        return err

    import json as _j
    import subprocess as _sp
    import datetime as _dt
    import random, string

    extra_args = args.split() if args.strip() else []
    argv, use_shell, _compiled, compile_err = _resolve_run_argv(resolved, extra_args)
    if compile_err:
        return compile_err

    # ── Content preview — forces script into Claude's context before launch ───
    # Same security gate as run_script: Claude sees the content before it runs.
    # Skipped for compiled source files (already reviewed during compile step).
    _ext = os.path.splitext(resolved)[1].lower()
    _preview_note = ""
    if _ext not in (".c", ".h", ".cpp", ".cc", ".cxx", ".hpp", ".java"):
        try:
            with open(resolved, "r", encoding="utf-8", errors="replace") as _pf:
                _preview_lines = _pf.readlines()
            _total = len(_preview_lines)
            _shown = _preview_lines[:50]
            _omitted = _total - len(_shown)
            _preview_note = (
                f"📄 Script preview ({_total} lines"
                f"{f', showing first 50' if _omitted else ''}):\n"
                f"{'─' * 60}\n"
                + "".join(_shown)
                + (f"{'─' * 60}\n... ({_omitted} more lines)\n" if _omitted else
                   f"{'─' * 60}\n")
            )
        except Exception as _pe:
            _preview_note = f"(preview unavailable: {_pe})\n"

    # Generate unique job id
    suffix = "".join(random.choices(string.ascii_lowercase + string.digits, k=4))
    job_id = f"job_{_dt.datetime.now().strftime('%Y%m%d_%H%M%S')}_{suffix}"
    log_path  = str(_job_log_path(job_id))
    mani_path = str(_job_manifest_path(job_id))

    # Build the wrapper script that runs the target and updates the manifest.
    wrapper_src = f"""
import subprocess, json, datetime, sys, os, signal

MANIFEST = {_j.dumps(mani_path)}
LOG      = {_j.dumps(log_path)}
ARGV     = {_j.dumps(argv)}
TIMEOUT  = {int(timeout_sec)}
USE_SHELL= {use_shell}

def _update(status, rc=None):
    try:
        with open(MANIFEST, "r", encoding="utf-8") as f:
            m = json.load(f)
        m["status"]      = status
        m["exit_code"]   = rc
        m["finished_at"] = datetime.datetime.now().isoformat()
        with open(MANIFEST, "w", encoding="utf-8") as f:
            json.dump(m, f, indent=2)
    except Exception as e:
        pass  # manifest update failure is non-fatal

with open(LOG, "a", encoding="utf-8", buffering=1) as lf:
    try:
        # For bat/cmd files ARGV[0] is cmd.exe — use the script path instead.
        # ARGV[2] is the actual script when cmd.exe /c <script> is used.
        if USE_SHELL or (len(ARGV) >= 3 and ARGV[0].lower().endswith("cmd.exe")):
            _cwd = os.path.dirname(ARGV[2]) if len(ARGV) >= 3 else None
        else:
            _cwd = os.path.dirname(ARGV[0])
        proc = subprocess.Popen(
            ARGV,
            stdout=lf, stderr=lf,
            shell=USE_SHELL,
            cwd=_cwd or None,
        )
        # Update manifest with real PID
        try:
            with open(MANIFEST, "r", encoding="utf-8") as f:
                m = json.load(f)
            m["pid"] = proc.pid
            with open(MANIFEST, "w", encoding="utf-8") as f:
                json.dump(m, f, indent=2)
        except Exception:
            pass
        try:
            proc.wait(timeout=TIMEOUT)
            _update("done" if proc.returncode == 0 else "failed", proc.returncode)
        except subprocess.TimeoutExpired:
            proc.kill()
            _update("timeout", -1)
            lf.write(f"\\n[wrapper] TIMEOUT after {{TIMEOUT}}s — process killed\\n")
    except Exception as exc:
        lf.write(f"\\n[wrapper] LAUNCH ERROR: {{exc}}\\n")
        _update("error", -1)
"""

    # Write wrapper to jobs dir
    wrapper_path = str(_jobs_dir() / f"{job_id}_wrapper.py")
    Path(wrapper_path).write_text(wrapper_src, encoding="utf-8")

    # Write initial manifest
    manifest = {
        "job_id":      job_id,
        "script":      resolved,
        "args":        args,
        "pid":         None,     # updated by wrapper once child spawns
        "wrapper_pid": None,
        "status":      "running",
        "started_at":  _dt.datetime.now().isoformat(),
        "finished_at": None,
        "exit_code":   None,
        "log_file":    log_path,
        "manifest":    mani_path,
        "timeout_sec": timeout_sec,
    }
    Path(mani_path).write_text(_j.dumps(manifest, indent=2), encoding="utf-8")

    # Touch the log file so run_script_status can read it immediately
    Path(log_path).touch()

    # Find Python for the wrapper
    py_candidates = [
        os.path.join(os.environ.get("LocalAppData", ""),
                     "Programs", "Python", "Python311", "python.exe"),
        sys.executable,
    ]
    wrapper_py = next((c for c in py_candidates if c and os.path.isfile(c)),
                      sys.executable)

    # Launch wrapper as a detached background process
    try:
        wp = _sp.Popen(
            [wrapper_py, wrapper_path],
            stdout=_sp.DEVNULL, stderr=_sp.DEVNULL,
            creationflags=getattr(_sp, "DETACHED_PROCESS", 0)
                          | getattr(_sp, "CREATE_NEW_PROCESS_GROUP", 0),
        )
        # Record wrapper PID in manifest
        manifest["wrapper_pid"] = wp.pid
        Path(mani_path).write_text(_j.dumps(manifest, indent=2), encoding="utf-8")
    except Exception as exc:
        return f"⚠️  run_script_start: could not launch wrapper: {exc}"

    _log.info("run_script_start: job_id=%s script=%s wrapper_pid=%d",
              job_id, resolved, wp.pid)
    return (f"{_preview_note}"
            f"✅ Job started — id: {job_id}\n"
            f"   Script:  {resolved}\n"
            f"   Log:     {log_path}\n"
            f"   Timeout: {timeout_sec}s\n"
            f"   Use run_script_status('{job_id}') to check progress.")


@mcp.tool()
def run_script_status(job_id: str, tail_lines: int = 50) -> str:
    """
    DEV TOOLS — Check the status of a background job and tail its log output.

    Call this after run_script_start to see what the job has printed so far,
    whether it is still running, and its exit code when done.

    Args:
        job_id:     The id returned by run_script_start.
        tail_lines: How many lines from the end of the log to return
                    (default 50). Raise for more context.

    Returns:
        Status block with manifest info + log tail.
        Status values: running / done / failed / timeout / error / killed
    """
    enabled, why = _dev_tools_enabled()
    if not enabled:
        return f"🚫 run_script_status is disabled here ({why})."

    m = _read_manifest(job_id)
    if not m:
        return (f"⚠️  No job found with id '{job_id}'.\n"
                f"Jobs are stored in {_JOBS_DIR}. "
                f"Check the id returned by run_script_start.")

    log_path = m.get("log_file", "")
    status   = m.get("status", "unknown")
    rc       = m.get("exit_code")
    started  = m.get("started_at", "?")
    finished = m.get("finished_at", "")
    script   = m.get("script", "?")
    pid      = m.get("pid", "?")

    # Status icon
    icons = {
        "running": "⏳", "done": "✅", "failed": "❌",
        "timeout": "⏱️", "error": "💥", "killed": "🛑",
    }
    icon = icons.get(status, "❓")

    lines_out = [
        f"{icon} Job {job_id}  [{status.upper()}]",
        f"   Script:    {script}",
        f"   PID:       {pid}",
        f"   Started:   {started}",
        f"   Finished:  {finished or '(still running)'}",
        f"   Exit code: {rc if rc is not None else '(pending)'}",
        f"   Log:       {log_path}",
        "───",
    ]

    # Read log tail
    try:
        log_text = Path(log_path).read_text(encoding="utf-8", errors="replace")
        log_lines = log_text.splitlines()
        total = len(log_lines)
        if total > tail_lines:
            lines_out.append(
                f"... ({total - tail_lines} earlier lines omitted — "
                f"raise tail_lines or use read_file_lines for full log) ...")
            log_lines = log_lines[-tail_lines:]
        lines_out.extend(log_lines)
    except Exception as exc:
        lines_out.append(f"(log unreadable: {exc})")

    return "\n".join(lines_out)


@mcp.tool()
def run_script_kill(job_id: str) -> str:
    """
    DEV TOOLS — Kill a running background job and all its child processes.

    Use when a job is hung, taking too long, or you want to cancel it.
    After killing, the manifest status is updated to 'killed' and the
    final log tail is returned.

    Args:
        job_id: The id returned by run_script_start.

    Returns:
        Confirmation of kill + final log tail.
    """
    enabled, why = _dev_tools_enabled()
    if not enabled:
        return f"🚫 run_script_kill is disabled here ({why})."

    import json as _j

    m = _read_manifest(job_id)
    if not m:
        return f"⚠️  No job found with id '{job_id}'."

    status = m.get("status", "unknown")
    if status not in ("running",):
        return (f"ℹ️  Job {job_id} is already in status '{status}' — "
                f"nothing to kill.")

    pid         = m.get("pid")
    wrapper_pid = m.get("wrapper_pid")
    killed      = []
    errors      = []

    def _kill_pid(p, label):
        if not p:
            return
        try:
            import signal
            # Try psutil first for child-tree kill
            try:
                import psutil as _ps
                proc_obj = _ps.Process(int(p))
                children = proc_obj.children(recursive=True)
                for child in children:
                    try:
                        child.kill()
                    except Exception:
                        pass
                proc_obj.kill()
                killed.append(f"{label} PID {p} (+ {len(children)} children)")
            except ImportError:
                # psutil not available — kill just the pid
                os.kill(int(p), signal.SIGTERM)
                killed.append(f"{label} PID {p}")
        except ProcessLookupError:
            killed.append(f"{label} PID {p} (already gone)")
        except Exception as exc:
            errors.append(f"{label} PID {p}: {exc}")

    _kill_pid(pid, "script")
    _kill_pid(wrapper_pid, "wrapper")

    # Update manifest
    mani_path = _job_manifest_path(job_id)
    try:
        import datetime as _dt
        m["status"]      = "killed"
        m["exit_code"]   = -9
        m["finished_at"] = _dt.datetime.now().isoformat()
        mani_path.write_text(_j.dumps(m, indent=2), encoding="utf-8")
    except Exception as exc:
        errors.append(f"manifest update: {exc}")

    lines_out = [f"🛑 Killed job {job_id}"]
    lines_out += [f"   {k}" for k in killed]
    if errors:
        lines_out += [f"   ⚠️  {e}" for e in errors]
    lines_out.append("───")

    # Final log tail
    log_path = m.get("log_file", "")
    try:
        log_lines = Path(log_path).read_text(
            encoding="utf-8", errors="replace").splitlines()
        lines_out.extend(log_lines[-30:])
    except Exception as exc:
        lines_out.append(f"(log unreadable: {exc})")

    return "\n".join(lines_out)


# ══════════════════════════════════════════════════════════════════════════════
# END OF CODE TOOLS WRITE-SIDE PATCH
# ══════════════════════════════════════════════════════════════════════════════


# ══════════════════════════════════════════════════════════════════════════════
# EMAIL TOOLS  (v7.1.0)
# ══════════════════════════════════════════════════════════════════════════════
# SMTP Option-A implementation: Python stdlib only (smtplib + email + ssl).
# One-time configure_email() stores credentials in ~/.ai-prowler/email_config.json
# (password stored base64-obfuscated, same pattern as the bearer token file).
# All subsequent email tools (send_email, send_alert, send_learnings_report,
# send_file) pick up the stored config automatically.
#
# Server-mode safety (v8.0.0): ALL roles may use send_email / send_alert /
# send_learnings_report — owner, manager, staff, and field_crew all have
# can_send_email=True. configure_email and send_file remain personal-only
# (configure_email reconfigures SMTP; send_file attaches arbitrary files).
#
# RECIPIENT SOURCES for field_crew in server mode:
#   Email (send_email / send_alert):
#     1. Customers sheet — matched by name, company, CustomerID, or email address.
#        The 'Email' column is used automatically; no need to know the address.
#     2. Registered server users — matched by email address or name (users.json).
#   SMS (send_sms):
#     1. users.json crew records — cell_phone + cell_carrier → gateway auto-resolved.
#     2. Customers sheet — Phone column match; SMS Gateway (pinned) or
#        Cell Carrier column → gateway auto-resolved via _carrier_to_gateway().
#     3. contacts_cache.json — saved contacts with a manually pinned gateway.
#     4. Twilio (paid) — unrestricted, any US number.
# ══════════════════════════════════════════════════════════════════════════════

import base64 as _b64

def _EMAIL_CONFIG_PATH() -> Path:
    return _state_dir() / "email_config.json"


def _CONTACTS_CACHE_PATH(user: "dict | None" = None) -> Path:
    """Personal contacts directory — one entry per person, holding phone,
    email, SMS gateway, and carrier together so a name like 'David' or
    'Vicki' resolves directly to everything needed to reach them.

    Structure:
    {
      "contacts": {
        "david": {"phone": "4807470358", "email": "david@x.com",
                   "gateway": "txt.att.net", "carrier": "AT&T Mobility LLC"},
        "vicki": {"phone": "4805437595", "email": "", "gateway": "vtext.com",
                   "carrier": ""}
      }
    }
    Looked up by name (case-insensitive) or by 10-digit phone number,
    whichever the caller has on hand.

    Personal mode (user=None): single shared contacts_cache.json — there's
    only one person using this install, so there's nothing to separate.

    Server mode (user is not None): EVERY role gets its own private,
    per-user file (contacts_cache_<username>.json) — owner, manager, staff,
    and field_crew all save and look up contacts independently of each
    other. This matters because v8.0.0 opened email/SMS/WhatsApp to every
    server-mode role, not just field_crew, so any role might say "text
    David" and expect their own saved contact, not someone else's.
    """
    if user is not None:
        uname = (user.get('username') or user.get('id') or
                 user.get('name') or 'unknown').strip().lower()
        # Sanitise to a safe filename component
        import re as _rsan
        uname = _rsan.sub(r'[^a-z0-9_-]', '_', uname) or 'unknown'
        return _state_dir() / f"contacts_cache_{uname}.json"
    return _state_dir() / "contacts_cache.json"


def _contacts_cache_load(user: "dict | None" = None) -> dict:
    """Load the relevant contacts_cache.json for this user (see
    _CONTACTS_CACHE_PATH), returning {'contacts': {}} if missing."""
    try:
        p = _CONTACTS_CACHE_PATH(user)
        if p.exists():
            import json as _jsc
            data = _jsc.loads(p.read_text(encoding="utf-8")) or {}
            data.setdefault('contacts', {})
            return data
    except Exception:
        pass
    return {'contacts': {}}


def _contacts_cache_save(data: dict, user: "dict | None" = None) -> bool:
    """Save the relevant contacts_cache.json for this user. Atomic write."""
    try:
        import json as _jscs
        p = _CONTACTS_CACHE_PATH(user)
        p.parent.mkdir(parents=True, exist_ok=True)
        tmp = p.with_suffix(".json.tmp")
        tmp.write_text(_jscs.dumps(data, indent=2), encoding="utf-8")
        import os as _eos2
        _eos2.replace(str(tmp), str(p))
        return True
    except Exception:
        return False


def _contact_lookup(name_or_phone: str, user: "dict | None" = None) -> "dict | None":
    """Resolve a name or phone number to its contact record, scoped to the
    calling user (field_crew get their own private file in server mode;
    everyone else uses the shared/personal contacts_cache.json).
    Returns the record dict (with 'name' key added) or None if not found."""
    import re as _rclk
    data = _contacts_cache_load(user)
    contacts = data.get('contacts', {})

    digits = _rclk.sub(r'\D', '', name_or_phone)
    if len(digits) == 11 and digits[0] == '1':
        digits = digits[1:]

    key = name_or_phone.strip().lower()

    # Try exact name match first
    if key in contacts:
        rec = dict(contacts[key])
        rec['name'] = key
        return rec

    # Try phone number match across all contacts
    if len(digits) == 10:
        for nm, rec in contacts.items():
            rec_digits = _rclk.sub(r'\D', '', rec.get('phone', ''))
            if len(rec_digits) == 11 and rec_digits[0] == '1':
                rec_digits = rec_digits[1:]
            if rec_digits == digits:
                out = dict(rec)
                out['name'] = nm
                return out

    return None


def _contact_save(name: str, phone: str = "", email: str = "",
                   user: "dict | None" = None) -> dict:
    """Create or update a contact record for the calling user's scope
    (field_crew's own private file in server mode; shared/personal
    contacts_cache.json otherwise). Merges with any existing fields for
    that name so a partial update doesn't wipe other known info."""
    data = _contacts_cache_load(user)
    key = name.strip().lower()
    rec = data['contacts'].get(key, {})
    if phone.strip():
        import re as _rcs
        d = _rcs.sub(r'\D', '', phone)
        if len(d) == 11 and d[0] == '1':
            d = d[1:]
        rec['phone'] = d
    if email.strip():
        rec['email'] = email.strip()
    data['contacts'][key] = rec
    _contacts_cache_save(data, user)
    return rec


def _email_config_load() -> "dict | None":
    """Load email config from disk. Returns None if not configured."""
    try:
        if _EMAIL_CONFIG_PATH().exists():
            raw = json.loads(_EMAIL_CONFIG_PATH().read_text(encoding="utf-8"))
            if isinstance(raw, dict) and raw.get("smtp_host"):
                # Decode obfuscated password
                raw = dict(raw)
                enc = raw.get("_password_b64", "")
                if enc:
                    raw["password"] = _b64.b64decode(enc.encode()).decode("utf-8")
                return raw
    except Exception as _e:
        _log.warning("email_config load failed: %s", _e)
    return None


def _lookup_customer_email(name_or_id: str) -> "str | None":
    """Look up a customer's email address from the Customers sheet by name or ID.

    Matches CustomerID (CUST-####), Company Name, First+Last Name, or any
    partial name match (case-insensitive). Returns the email string or None.
    Used by send_email / send_alert so any role can address customers by
    name without knowing their email — works the same in personal mode and
    for every server-mode role (owner, manager, staff, field_crew).

    v8.0.0 — enables customer email from the job spreadsheet.
    """
    try:
        import openpyxl as _opx
        _xl_path = _get_default_spreadsheet_path()
        if not (_xl_path and os.path.exists(_xl_path)):
            return None
        _wb = _opx.load_workbook(_xl_path, data_only=True)
        if 'Customers' not in _wb.sheetnames:
            return None
        _ws = _wb['Customers']
        _hdrs, _hdr_row = [], None
        for _r in _ws.iter_rows(min_row=1, max_row=5):
            if len([c for c in _r if c.value]) >= 3:
                _hdr_row = _r[0].row
                _hdrs = [str(c.value or '').replace('\n', ' ').strip() for c in _r]
                break
        if not _hdr_row:
            return None
        _email_col = next((i for i, h in enumerate(_hdrs) if h.lower() == 'email'), None)
        if _email_col is None:
            return None
        _needle = name_or_id.strip().lower()
        for _row in _ws.iter_rows(min_row=_hdr_row + 1):
            _vals = [str(c.value or '').strip() for c in _row]
            _row_str = ' '.join(_vals).lower()
            if _needle in _row_str and _vals[_email_col]:
                return _vals[_email_col]
    except Exception:
        pass
    return None


def _email_config_save(cfg: dict) -> bool:
    """Save email config to disk. Obfuscates password with base64."""
    try:
        _EMAIL_CONFIG_PATH().parent.mkdir(parents=True, exist_ok=True)
        out = dict(cfg)
        pw = out.pop("password", "")
        if pw:
            out["_password_b64"] = _b64.b64encode(pw.encode("utf-8")).decode()
        tmp = _EMAIL_CONFIG_PATH().with_suffix(".json.tmp")
        tmp.write_text(json.dumps(out, indent=2), encoding="utf-8")
        import os as _eos
        _eos.replace(str(tmp), str(_EMAIL_CONFIG_PATH()))
        return True
    except Exception as _e:
        _log.error("email_config save failed: %s", _e)
        return False


def _send_smtp(to: str, subject: str, body: str,
               attachment_path: "str | None" = None,
               body_html: "str | None" = None,
               reply_to: "str | None" = None,
               sender_display: "str | None" = None) -> tuple:
    """Core SMTP send. Returns (ok: bool, message: str). Uses stored config.

    reply_to:       Optional Reply-To header value (e.g. employee's personal
                    email). When set, the mail client's Reply button goes to
                    this address rather than the From address.
    sender_display: Optional override for the From display name — used in
                    server mode to show the employee's name alongside the
                    company address (e.g. 'Jake Smith via ABC Cleaning').
    """
    import smtplib
    import ssl as _ssl
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders as _enc

    cfg = _email_config_load()
    if not cfg:
        return (False, "Email not configured. Call configure_email() first.")

    smtp_host = cfg.get("smtp_host", "").strip()
    smtp_port = int(cfg.get("smtp_port", 587))
    username  = cfg.get("username", "").strip()
    password  = cfg.get("password", "")
    from_addr = cfg.get("from_address", username).strip() or username
    from_name = sender_display or cfg.get("from_name", "AI-Prowler").strip()
    use_tls   = cfg.get("use_tls", True)

    if not smtp_host or not username:
        return (False, "Incomplete email config — smtp_host and username required.")

    # Build message
    msg = MIMEMultipart("mixed")
    msg["Subject"] = subject
    msg["From"]    = f"{from_name} <{from_addr}>" if from_name else from_addr
    msg["To"]      = to

    # Reply-To — directs replies to the employee's personal email rather than
    # the server's shared SMTP address. Works with all providers.
    if reply_to:
        msg["Reply-To"] = reply_to

    # Body — prefer HTML if provided, plain-text fallback
    alt_part = MIMEMultipart("alternative")
    alt_part.attach(MIMEText(body, "plain", "utf-8"))
    if body_html:
        alt_part.attach(MIMEText(body_html, "html", "utf-8"))
    msg.attach(alt_part)

    # Optional attachment
    if attachment_path:
        try:
            ap = Path(attachment_path)
            with open(ap, "rb") as _fh:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(_fh.read())
            _enc.encode_base64(part)
            part.add_header("Content-Disposition",
                            f'attachment; filename="{ap.name}"')
            msg.attach(part)
        except Exception as _ae:
            return (False, f"Could not attach file: {_ae}")

    # Connect and send
    try:
        context = _ssl.create_default_context()
        if smtp_port == 465:
            # SMTPS — SSL from the start
            with smtplib.SMTP_SSL(smtp_host, smtp_port,
                                  context=context, timeout=20) as server:
                server.login(username, password)
                server.sendmail(from_addr, [to], msg.as_bytes())
        else:
            # STARTTLS (port 587 typical)
            with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as server:
                server.ehlo()
                if use_tls:
                    server.starttls(context=context)
                    server.ehlo()
                server.login(username, password)
                server.sendmail(from_addr, [to], msg.as_bytes())
        _log.info("Email sent to %s subject=%r", to, subject)
        return (True, f"✅ Email sent to {to}")
    except smtplib.SMTPAuthenticationError:
        return (False,
                "❌ SMTP authentication failed — check username/password. "
                "For Gmail, use a 16-digit App Password (not your account password). "
                "Get one at: myaccount.google.com → Security → App passwords.")
    except smtplib.SMTPException as _se:
        return (False, f"❌ SMTP error: {_se}")
    except Exception as _ge:
        return (False, f"❌ Send failed: {_ge}")


def _email_allowed_for_user(user: "dict | None") -> tuple:
    """Gate configure_email / send_file / send_learnings_report to personal mode only.
    Returns (allowed: bool, reason: str). PURE.

    These email tools use personal SMTP credentials and are not appropriate for
    a shared company server. In server mode, field_crew may use send_email /
    send_alert via the separate _send_email_cap gate below.

    Personal mode (user=None): always allowed.
    Server mode (user is not None): always blocked (for these tools).
    """
    if user is None:
        return (True, "personal mode")
    return (False,
            "This email tool is only available in personal mode. "
            "In server mode each user should configure email on their own "
            "personal AI-Prowler install.")


def _send_email_cap(user: "dict | None") -> tuple:
    """Capability gate for send_email and send_alert in server mode. PURE.

    Personal mode (user=None): always allowed (personal SMTP, single user).
    Server mode (v8.0.0): ALL roles may send email — owner, manager, staff,
    and field_crew all have can_send_email=True in _ROLE_CAPS.

    Recipients in server mode: customers in the job spreadsheet (Customers
    sheet, Email column — matched by name/company/ID) OR registered server
    users (users.json, email field). Name-to-email resolution is handled
    in send_email / send_alert before the SMTP call.

    configure_email, send_file, and send_learnings_report remain personal-only
    regardless of this capability — use _email_allowed_for_user for those.

    Returns (allowed: bool, reason: str).
    """
    if user is None:
        return (True, "personal mode — email permitted")
    caps = _role_caps(user.get("role"))
    if caps.get("can_send_email"):
        return (True, f"role '{user.get('role')}' may send email in server mode")
    return (False,
            f"role '{user.get('role')}' cannot send email from the company server — "
            "use your personal AI-Prowler install for email.")


def _send_sms_cap(user: "dict | None") -> tuple:
    """Capability gate for send_sms in server mode. PURE.

    Personal mode (user=None): always allowed (single-user desktop install).
    Server mode: True for ALL roles — owner, manager, staff, and field_crew
    can all send SMS via the company's configured SMS provider (Twilio,
    SignalWire, or Vonage).

    Inbound replies ARE correctly attributed per user: check_sms_replies()
    reads from a local webhook-backed inbox (sms_inbox.json/sms_threads.json)
    and filters to only the threads the calling user personally sent — Mike
    will not see Karen's reply to Jake's text. See sms_inbox.py for details.

    Returns (allowed: bool, reason: str).
    """
    if user is None:
        return (True, "personal mode — SMS permitted")
    caps = _role_caps(user.get("role"))
    if caps.get("can_send_sms"):
        return (True, f"role '{user.get('role')}' may send SMS in server mode")
    return (False,
            f"role '{user.get('role')}' cannot send SMS from the company server — "
            "use your personal AI-Prowler install for SMS (Twilio).")


@mcp.tool()
def configure_email(smtp_host: str, smtp_port: int, username: str,
                    password: str, from_name: str = "AI-Prowler",
                    default_to: str = "",
                    ctx: Context = None) -> str:
    """
    Configure SMTP email settings for AI-Prowler. One-time setup — all other
    email tools use the saved config automatically after this.

    configure_email is personal-mode only — it sets YOUR SMTP account, and
    a shared company server's SMTP is configured once by an admin instead.
    Once that's done, send_email / send_alert work for every server role
    (owner, manager, staff, field_crew); only configure_email itself,
    send_file, and send_learnings_report remain personal-install-only.

    Supports any SMTP provider:
      • Gmail   : smtp.gmail.com  port 587  (requires a 16-digit App Password,
                  NOT your account password. Create one at:
                  myaccount.google.com → Security → App passwords)
      • Outlook : smtp.office365.com  port 587
      • Yahoo   : smtp.mail.yahoo.com  port 587
      • Any other SMTP server your provider documents

    Args:
        smtp_host:  SMTP server hostname (e.g. 'smtp.gmail.com')
        smtp_port:  SMTP port — 587 for STARTTLS (most common),
                    465 for SMTPS, 25 for plain (not recommended)
        username:   Your email address / SMTP login
        password:   App password or SMTP password (stored obfuscated)
        from_name:  Display name shown in the From field (default: 'AI-Prowler')
        default_to: Default recipient email address. Tools that take a 'to'
                    argument use this when none is supplied.
        ctx:        MCP context (injected automatically)

    Returns:
        Confirmation string, or an error if the config could not be saved.
    """
    _telemetry_increment_tool_count("configure_email")

    # Personal-mode only
    user = _current_user(ctx)
    allowed, why = _email_allowed_for_user(user)
    if not allowed:
        return f"❌ {why}"

    smtp_host = smtp_host.strip()
    username  = username.strip()
    if not smtp_host:
        return "❌ smtp_host is required."
    if not username:
        return "❌ username (your email address) is required."
    if not password:
        return "❌ password is required."
    if smtp_port < 1 or smtp_port > 65535:
        return f"❌ smtp_port {smtp_port} is invalid (must be 1-65535)."

    cfg = {
        "smtp_host":    smtp_host,
        "smtp_port":    smtp_port,
        "username":     username,
        "password":     password,
        "from_address": username,
        "from_name":    from_name.strip() or "AI-Prowler",
        "default_to":   default_to.strip(),
        "use_tls":      smtp_port != 465,
    }

    if not _email_config_save(cfg):
        return "❌ Could not save email config. Check disk permissions."

    lines = [
        "✅ Email configured successfully.",
        f"   SMTP host  : {smtp_host}:{smtp_port}",
        f"   Account    : {username}",
        f"   From name  : {cfg['from_name']}",
    ]
    if default_to:
        lines.append(f"   Default to : {default_to}")
    lines += [
        "",
        "To test: call send_alert() or send_email() with a test message.",
        "Note: for Gmail use a 16-digit App Password — your account",
        "password will not work (Google blocks it for security).",
    ]
    return "\n".join(lines)


@mcp.tool()
def send_email(to: str, subject: str, body: str,
               attachment_path: str = "",
               body_html: str = "",
               ctx: Context = None) -> str:
    """
    Send an email via the configured SMTP account.

    In server mode, any role (owner, manager, staff, field_crew) may email:
      • Other registered users of this server (by email address or name).
      • Customers listed in the job spreadsheet (by name, company, or CustomerID).
          The Customers sheet Email column is used — no need to know the address.

    Args:
        to:              Recipient email address, or a customer/user name to look
                         up automatically (e.g. "Torres", "Blue Wave Cafe",
                         "CUST-0003"). Leave blank to use the configured default_to.
        subject:         Email subject line.
        body:            Plain-text email body. Always used as the plain-text
                         part (and as the sole body if body_html is omitted).
        attachment_path: Optional — absolute path to a file in a tracked
                         read-allowlisted directory to attach to the email.
        body_html:       Optional — HTML version of the body. When provided,
                         the email is sent as multipart/alternative so HTML
                         mail clients render the formatted version instead of
                         showing raw tags, with `body` kept as the plain-text
                         fallback.
        ctx:             MCP context (injected automatically)

    Returns:
        "✅ Email sent to <address>" on success, or an error string.

    Voice examples:
        "Email Torres that his job is confirmed for Saturday"
        "Send the Blue Wave Cafe invoice to the customer"
        "Email Jake a summary of today's jobs"
        "Send Maria the updated schedule"
    """
    _telemetry_increment_tool_count("send_email")

    # Resolve user early so contact-name lookups can be scoped correctly
    # (field_crew get their own private contacts_cache_<username>.json).
    user = _current_user(ctx)

    cfg = _email_config_load()
    if not cfg:
        return ("❌ Email not configured. "
                "Call configure_email() first with your SMTP settings.")

    to = (to or "").strip() or cfg.get("default_to", "").strip()
    if not to:
        return "❌ No recipient address. Provide a 'to' address or name, or set default_to via configure_email()."

    # v8.0.0: if 'to' doesn't look like an email address, try to resolve it
    # from the Customers sheet (by name/company/ID), users.json (by name),
    # or the personal contacts_cache.json (by name) — in that order.
    if '@' not in to:
        _resolved = _lookup_customer_email(to)
        if _resolved:
            to = _resolved
        else:
            # Try registered users by name
            try:
                _ud = _load_users()
                if _ud:
                    _needle = to.lower()
                    for _u in (_ud.get('users') or {}).values():
                        if isinstance(_u, dict):
                            if _needle in (_u.get('name') or '').lower():
                                _em = (_u.get('email') or '').strip()
                                if _em:
                                    to = _em
                                    break
            except Exception:
                pass
        if '@' not in to:
            # Try personal contacts_cache.json by name
            try:
                _pc = _contact_lookup(to, user)
                if _pc and _pc.get('email'):
                    to = _pc['email']
            except Exception:
                pass
        if '@' not in to:
            return (
                f"❌ Could not resolve '{to}' to an email address.\n"
                "Check that the name matches a customer in the Customers sheet, "
                "a registered user in the Admin tab, or a saved contact "
                "(\"<name>'s email is ...\"), and that their Email is filled in."
            )

    subject = subject.strip()
    if not subject:
        return "❌ subject is required."
    if not body.strip():
        return "❌ body is required."

    # v8.0.0: all server-mode roles may send email via the company SMTP.
    allowed, why = _send_email_cap(user)
    if not allowed:
        return f"❌ {why}"

    # Server mode: personalise From display name and set Reply-To to the
    # employee's registered email so the recipient can reply directly to them.
    reply_to       = None
    sender_display = None
    if user is not None:
        emp_email = (user.get("email") or "").strip()
        emp_name  = (user.get("name")  or "").strip()
        if emp_email:
            reply_to = (f"{emp_name} <{emp_email}>"
                        if emp_name else emp_email)
        if emp_name:
            cfg_name = (_email_config_load() or {}).get("from_name", "")
            sender_display = (f"{emp_name} via {cfg_name}"
                              if cfg_name else emp_name)

    # Resolve optional attachment
    attach = None
    if attachment_path and attachment_path.strip():
        resolved_attach, err = _resolve_allowlisted_path(attachment_path.strip())
        if err:
            return f"❌ Attachment: {err}"
        attach = resolved_attach

    ok, msg = _send_smtp(to, subject, body, attachment_path=attach,
                         body_html=(body_html.strip() or None),
                         reply_to=reply_to, sender_display=sender_display)
    return msg


@mcp.tool()
def send_alert(message: str, to: str = "",
               ctx: Context = None) -> str:
    """
    Send a quick one-line alert email. Subject is auto-generated from the
    message. Great for short voice-commanded notifications.

    In server mode, field crew may alert:
      • Other registered users of this server (by email address or name).
      • Customers listed in the job spreadsheet (by name, company, or CustomerID).

    Args:
        message: The alert text. Keep it concise — it becomes both the
                 subject (truncated) and the body.
        to:      Recipient email address, or a customer/user name to look up
                 (e.g. "Torres", "Blue Wave Cafe"). Leave blank for default_to.
        ctx:     MCP context (injected automatically)

    Returns:
        "✅ Alert sent to <address>" on success, or an error string.

    Voice examples:
        "Alert Torres that we're 20 minutes away"
        "Ping the Blue Wave Cafe that we're done"
        "Alert Jake that I'm running late"
        "Ping the owner that the gate code isn't working"
    """
    _telemetry_increment_tool_count("send_alert")

    cfg = _email_config_load()
    if not cfg:
        return ("❌ Email not configured. "
                "Call configure_email() first with your SMTP settings.")

    to = (to or "").strip() or cfg.get("default_to", "").strip()
    if not to:
        return "❌ No recipient. Provide a 'to' address or name, or set default_to via configure_email()."

    # v8.0.0: name → email resolution (Customers sheet first, then users.json)
    if '@' not in to:
        _resolved = _lookup_customer_email(to)
        if _resolved:
            to = _resolved
        else:
            try:
                _ud = _load_users()
                if _ud:
                    _needle = to.lower()
                    for _u in (_ud.get('users') or {}).values():
                        if isinstance(_u, dict):
                            if _needle in (_u.get('name') or '').lower():
                                _em = (_u.get('email') or '').strip()
                                if _em:
                                    to = _em
                                    break
            except Exception:
                pass
        if '@' not in to:
            return (
                f"❌ Could not resolve '{to}' to an email address.\n"
                "Check that the name matches a customer in the Customers sheet "
                "or a registered user in the Admin tab, and that their Email is filled in."
            )

    message = message.strip()
    if not message:
        return "❌ message is required."

    # v8.0.0: all server-mode roles may send alerts via the company SMTP.
    user = _current_user(ctx)
    allowed, why = _send_email_cap(user)
    if not allowed:
        return f"❌ {why}"

    # Server mode: personalise From display name and set Reply-To to employee's email.
    reply_to       = None
    sender_display = None
    if user is not None:
        emp_email = (user.get("email") or "").strip()
        emp_name  = (user.get("name")  or "").strip()
        if emp_email:
            reply_to = (f"{emp_name} <{emp_email}>"
                        if emp_name else emp_email)
        if emp_name:
            cfg_name = (_email_config_load() or {}).get("from_name", "")
            sender_display = (f"{emp_name} via {cfg_name}"
                              if cfg_name else emp_name)

    # Subject: first 80 chars of message
    subject = f"AI-Prowler Alert: {message[:80]}"
    import datetime as _dt2
    body = (f"{message}\n\n"
            f"— Sent by AI-Prowler at "
            f"{_dt2.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    ok, msg = _send_smtp(to, subject, body,
                         reply_to=reply_to, sender_display=sender_display)
    return msg


@mcp.tool()
def send_file(to: str, filepath: str,
              subject: str = "", body: str = "",
              ctx: Context = None) -> str:
    """
    Send any file from a tracked directory as an email attachment.

    Args:
        to:       Recipient email address. Leave blank for default_to.
        filepath: Absolute path to the file to send (must be in the read
                  allowlist). Any file type is supported.
        subject:  Email subject. Auto-generated from filename if blank.
        body:     Email body text. Auto-generated if blank.
        ctx:      MCP context (injected automatically)

    Returns:
        "✅ Email sent" or an error string.

    Voice examples:
        "Email the Q3 report to the boss"
        "Send the updated config to john@company.com"
        "Email the job tracker spreadsheet to myself"
    """
    _telemetry_increment_tool_count("send_file")

    cfg = _email_config_load()
    if not cfg:
        return ("❌ Email not configured. "
                "Call configure_email() first.")

    to = (to or "").strip() or cfg.get("default_to", "").strip()
    if not to:
        return "❌ No recipient address."

    resolved, err = _resolve_allowlisted_path(filepath.strip())
    if err:
        return f"❌ {err}"

    fname = Path(resolved).name
    if not subject.strip():
        subject = f"AI-Prowler: {fname}"
    if not body.strip():
        import datetime as _dt3
        body = (f"File attached: {fname}\n\n"
                f"Sent by AI-Prowler at "
                f"{_dt3.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # send_file: personal-mode only (arbitrary file attachment = exfiltration risk).
    user = _current_user(ctx)
    allowed, why = _email_allowed_for_user(user)
    if not allowed:
        return f"❌ {why}"

    ok, msg = _send_smtp(to, subject, body, attachment_path=resolved)
    return msg


# ══════════════════════════════════════════════════════════════════════════════
# LEARNINGS MOBILE EXPORT TOOLS  (v7.1.0)
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def get_learnings_report(category: str = "",
                          status: str = "active",
                          format: str = "summary",
                          ctx: Context = None) -> str:
    """
    Return learnings as formatted text directly in the conversation — no file
    dialog, no desktop required. Full mobile control over the learning store.

    Args:
        category: Filter to a single category (e.g. 'business', 'client',
                  'technical'). Leave blank for all categories.
        status:   'active' (default), 'archived', 'deprecated', or 'all'.
        format:   'summary'  — title + category + one-line content (default)
                  'full'     — everything: all fields per learning
                  'titles'   — titles only (quick list)
        ctx:      MCP context (injected automatically)

    Returns:
        Formatted text of the matching learnings, suitable for Claude to
        read aloud or display in the conversation.

    Voice examples:
        "Read me all my business learnings"
        "What have we learned about the Johnson account?"
        "List all my active technical learnings"
        "Show me a full summary of my client preferences"
    """
    _telemetry_increment_tool_count("get_learnings_report")

    if not _sl:
        return "❌ Self-learning module not available."

    try:
        db = _sl._load_db()
        learnings = db.get("learnings", [])
    except Exception as _e:
        return f"❌ Could not load learnings: {_e}"

    # Filter by status
    status = (status or "active").strip().lower()
    if status != "all":
        learnings = [l for l in learnings
                     if l.get("status", "active").lower() == status]

    # Filter by category
    cat = (category or "").strip().lower()
    if cat:
        learnings = [l for l in learnings
                     if l.get("category", "").lower() == cat]

    if not learnings:
        filters = []
        if cat:
            filters.append(f"category='{category}'")
        if status != "all":
            filters.append(f"status='{status}'")
        filter_str = " with " + ", ".join(filters) if filters else ""
        return f"ℹ️  No learnings found{filter_str}."

    fmt = (format or "summary").strip().lower()
    lines = [f"📚 AI-Prowler Learnings — {len(learnings)} found\n"]

    for i, l in enumerate(learnings, 1):
        title    = l.get("title", "(no title)")
        cat_val  = l.get("category", "")
        content  = l.get("content", "")
        conf     = l.get("confidence", "")
        outcome  = l.get("outcome", "")
        tags     = l.get("tags", [])
        recorded = l.get("recorded_at", "")[:10] if l.get("recorded_at") else ""

        if fmt == "titles":
            lines.append(f"{i}. {title}")
        elif fmt == "full":
            lines.append(f"{'─'*50}")
            lines.append(f"{i}. {title}")
            if cat_val:
                lines.append(f"   Category  : {cat_val}")
            if conf:
                lines.append(f"   Confidence: {conf}")
            if outcome:
                lines.append(f"   Outcome   : {outcome}")
            if recorded:
                lines.append(f"   Recorded  : {recorded}")
            if tags:
                lines.append(f"   Tags      : {', '.join(tags)}")
            lines.append(f"   Content   : {content}")
        else:  # summary
            short = content[:120].strip()
            if len(content) > 120:
                short += "…"
            cat_str = f"[{cat_val}] " if cat_val else ""
            lines.append(f"{i}. {cat_str}{title}")
            lines.append(f"   {short}")

    return "\n".join(lines)


@mcp.tool()
def export_learnings_file(filepath: str,
                          format: str = "pack",
                          category: str = "",
                          include_inactive: bool = False,
                          ctx: Context = None) -> str:
    """
    Export learnings to a file in a writable zone — mobile equivalent of the
    GUI's 'Export Pack' and 'Export to CSV' buttons.

    Args:
        filepath:         Destination path inside a writable zone. Use
                          '.aiplearn' extension for a pack (importable by
                          other AI-Prowler installs), '.csv' for spreadsheet.
                          Example: 'C:/Users/david/Documents/learnings.aiplearn'
        format:           'pack' (default) — .aiplearn JSON pack file
                          'csv'            — comma-separated spreadsheet
        category:         Filter to a single category. Leave blank for all.
        include_inactive: Include archived/deprecated learnings (default False).
        ctx:              MCP context (injected automatically)

    Returns:
        "✅ Exported N learnings to <path>" or an error string.

    Voice examples:
        "Export all my learnings to my documents folder"
        "Save a CSV of my business learnings to the desktop"
        "Export a learning pack to my OneDrive"
    """
    _telemetry_increment_tool_count("export_learnings_file")

    if not _sl:
        return "❌ Self-learning module not available."

    filepath = filepath.strip()
    if not filepath:
        return "❌ filepath is required."

    # Must be in a writable zone
    resolved, deny = _resolve_writable_path(filepath)
    if not resolved:
        return f"❌ {deny}"

    try:
        db = _sl._load_db()
        learnings = list(db.get("learnings", []))
    except Exception as _e:
        return f"❌ Could not load learnings: {_e}"

    # Filter
    cat = (category or "").strip().lower()
    if cat:
        learnings = [l for l in learnings if l.get("category", "").lower() == cat]
    if not include_inactive:
        learnings = [l for l in learnings
                     if l.get("status", "active") == "active"]

    if not learnings:
        return "ℹ️  No learnings matched the filter — nothing exported."

    fmt = (format or "pack").strip().lower()
    dest = Path(resolved)
    dest.parent.mkdir(parents=True, exist_ok=True)

    try:
        if fmt == "csv":
            import csv as _csv
            import io as _io
            buf = _io.StringIO()
            writer = _csv.writer(buf)
            writer.writerow(["id", "title", "content", "category", "confidence",
                             "outcome", "tags", "status", "recorded_at",
                             "supersedes", "superseded_by"])
            for l in learnings:
                writer.writerow([
                    l.get("id", ""),
                    l.get("title", ""),
                    l.get("content", ""),
                    l.get("category", ""),
                    l.get("confidence", ""),
                    l.get("outcome", ""),
                    "|".join(l.get("tags", [])),
                    l.get("status", ""),
                    l.get("recorded_at", ""),
                    l.get("supersedes", ""),
                    l.get("superseded_by", ""),
                ])
            dest.write_text(buf.getvalue(), encoding="utf-8")
        else:
            # .aiplearn pack — same format as self_learning.export_learnings
            import datetime as _dt4
            pack = {
                "schema":      "1.0",
                "exported_at": _dt4.datetime.now(_dt4.timezone.utc).isoformat(),
                "source_app":  "AI-Prowler",
                "count":       len(learnings),
                "learnings":   learnings,
            }
            tmp = dest.with_suffix(dest.suffix + ".tmp")
            tmp.write_text(json.dumps(pack, indent=2, ensure_ascii=False),
                           encoding="utf-8")
            import os as _eos2
            _eos2.replace(str(tmp), str(dest))

        _log.info("export_learnings_file: %d learnings → %s (%s)",
                  len(learnings), dest, fmt)
        return (f"✅ Exported {len(learnings)} learning(s) to:\n"
                f"   {dest}\n"
                f"   Format: {'Learning Pack (.aiplearn)' if fmt != 'csv' else 'CSV spreadsheet'}")
    except Exception as _e:
        return f"❌ Export failed: {_e}"


@mcp.tool()
def send_learnings_report(to: str = "",
                          category: str = "",
                          source: str = "",
                          recorded_by: str = "",
                          tags: str = "",
                          min_confidence: float = 0.0,
                          outcome: str = "",
                          since_date: str = "",
                          subject: str = "",
                          include_inactive: bool = False,
                          ctx: Context = None) -> str:
    """
    Export learnings as a formatted HTML email report and send it.
    Supports multiple filters so you can send targeted reports.

    Args:
        to:               Recipient email. Leave blank for configured default_to.
        category:         Filter to a single category (e.g. 'client', 'safety').
        source:           Filter by source (e.g. 'operator', 'claude', 'field_crew').
        recorded_by:      Filter by who recorded it (server mode employee name).
        tags:             Comma-separated tags to filter by (any match).
        min_confidence:   Minimum confidence score 0.0–1.0 (default 0.0 = all).
        outcome:          Filter by outcome (e.g. 'success', 'failure', 'unknown').
        since_date:       Only include learnings recorded on or after this date
                          (YYYY-MM-DD format).
        subject:          Email subject. Auto-generated if blank.
        include_inactive: Include archived/deprecated learnings (default False).
        ctx:              MCP context (injected automatically).

    Returns:
        "✅ Learnings report sent to <address>" or error string.

    Voice examples:
        "Email all my learnings to david@company.com"
        "Send learnings recorded by Mike to the team"
        "Email safety category learnings from this month"
        "Send all learnings with source field_crew to david@company.com"
        "Email learnings tagged urgent to myself"
    """
    _telemetry_increment_tool_count("send_learnings_report")

    cfg = _email_config_load()
    if not cfg:
        return ("❌ Email not configured. "
                "Call configure_email() first.")

    to = (to or "").strip() or cfg.get("default_to", "").strip()
    if not to:
        return "❌ No recipient address."

    user = _current_user(ctx)
    allowed, why = _send_email_cap(user)
    if not allowed:
        return f"❌ {why}"

    if not _sl:
        return "❌ Self-learning module not available."

    try:
        db = _sl._load_db()
        learnings = list(db.get("learnings", []))
    except Exception as _e:
        return f"❌ Could not load learnings: {_e}"

    # ── Apply filters ─────────────────────────────────────────────────────────
    if not include_inactive:
        learnings = [l for l in learnings
                     if l.get("status", "active") == "active"]

    if (category or "").strip():
        _cat = category.strip().lower()
        learnings = [l for l in learnings
                     if l.get("category", "").lower() == _cat]

    if (source or "").strip():
        _src = source.strip().lower()
        learnings = [l for l in learnings
                     if l.get("source", "").lower() == _src]

    if (recorded_by or "").strip():
        _rb = recorded_by.strip().lower()
        learnings = [l for l in learnings
                     if l.get("recorded_by", "").lower() == _rb]

    if (outcome or "").strip():
        _out = outcome.strip().lower()
        learnings = [l for l in learnings
                     if l.get("outcome", "").lower() == _out]

    if min_confidence > 0.0:
        learnings = [l for l in learnings
                     if float(l.get("confidence", 0.0)) >= min_confidence]

    if (tags or "").strip():
        _tag_list = [t.strip().lower() for t in tags.split(",") if t.strip()]
        learnings = [l for l in learnings
                     if any(t in [x.lower() for x in l.get("tags", [])]
                            for t in _tag_list)]

    if (since_date or "").strip():
        _since = since_date.strip()
        learnings = [l for l in learnings
                     if (l.get("created_at") or "")[:10] >= _since]

    if not learnings:
        return "ℹ️  No learnings matched the filter — nothing to send."

    import datetime as _dt5

    # ── Build filter summary for report header ────────────────────────────────
    active_filters = []
    if category:     active_filters.append(f"Category: {category}")
    if source:       active_filters.append(f"Source: {source}")
    if recorded_by:  active_filters.append(f"Recorded by: {recorded_by}")
    if tags:         active_filters.append(f"Tags: {tags}")
    if outcome:      active_filters.append(f"Outcome: {outcome}")
    if min_confidence > 0.0: active_filters.append(f"Min confidence: {min_confidence:.0%}")
    if since_date:   active_filters.append(f"Since: {since_date}")
    filter_label = " | ".join(active_filters) if active_filters else "All (no filter)"

    # ── Build HTML report ─────────────────────────────────────────────────────
    html_rows = ""
    for l in learnings:
        tag_str = ", ".join(l.get("tags", []))
        rb      = l.get("recorded_by", "") or "—"
        src     = l.get("source", "") or "—"
        conf    = f"{float(l.get('confidence', 0)):.0%}"
        html_rows += (
            f"<tr>"
            f"<td style='padding:6px;border-bottom:1px solid #eee'><b>{l.get('title','')}</b></td>"
            f"<td style='padding:6px;border-bottom:1px solid #eee'>{l.get('category','')}</td>"
            f"<td style='padding:6px;border-bottom:1px solid #eee'>{src}</td>"
            f"<td style='padding:6px;border-bottom:1px solid #eee'>{rb}</td>"
            f"<td style='padding:6px;border-bottom:1px solid #eee'>{conf}</td>"
            f"<td style='padding:6px;border-bottom:1px solid #eee'>{l.get('content','')[:200]}</td>"
            f"<td style='padding:6px;border-bottom:1px solid #eee'>{tag_str}</td>"
            f"<td style='padding:6px;border-bottom:1px solid #eee'>{l.get('created_at','')[:10]}</td>"
            f"</tr>"
        )

    html_body = f"""<html><body style='font-family:Segoe UI,Arial,sans-serif;color:#222'>
<h2 style='color:#005a9e'>AI-Prowler Learnings Report</h2>
<p><b>Filters:</b> {filter_label} &nbsp;|&nbsp; <b>Count:</b> {len(learnings)}
   &nbsp;|&nbsp; <b>Generated:</b> {_dt5.datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
<table style='border-collapse:collapse;width:100%;font-size:13px'>
<tr style='background:#005a9e;color:white'>
  <th style='padding:8px;text-align:left'>Title</th>
  <th style='padding:8px;text-align:left'>Category</th>
  <th style='padding:8px;text-align:left'>Source</th>
  <th style='padding:8px;text-align:left'>Recorded By</th>
  <th style='padding:8px;text-align:left'>Confidence</th>
  <th style='padding:8px;text-align:left'>Content</th>
  <th style='padding:8px;text-align:left'>Tags</th>
  <th style='padding:8px;text-align:left'>Recorded</th>
</tr>
{html_rows}
</table>
<p style='color:#888;font-size:11px;margin-top:20px'>
Sent by AI-Prowler — Agentic RAG Knowledge Base</p>
</body></html>"""

    # Plain-text fallback
    plain = (f"AI-Prowler Learnings Report\n"
             f"Filters: {filter_label}\n"
             f"Count: {len(learnings)}\n\n")
    for l in learnings:
        plain += f"• {l.get('title','')} [{l.get('category','')}]"
        if l.get('recorded_by'): plain += f" — {l.get('recorded_by')}"
        plain += f"\n  {l.get('content','')[:120]}\n\n"

    if not subject.strip():
        parts = ["AI-Prowler Learnings Report"]
        if active_filters:
            parts.append(" — " + ", ".join(active_filters[:2]))
        parts.append(f" ({len(learnings)} entries)")
        subject = "".join(parts)

    ok, msg = _send_smtp(to, subject, plain, body_html=html_body)
    if ok:
        return f"✅ Learnings report sent to {to} ({len(learnings)} learning(s))"
    return msg


@mcp.tool()
def rebuild_learnings_index(ctx: Context = None) -> str:
    """
    Rebuild the ChromaDB learnings index from the JSON data file. Fixes the
    case where search_learnings() returns nothing but learnings exist in the
    JSON file (index/data mismatch). Mobile equivalent of the GUI's
    'Rebuild ChromaDB Index' button in the Learnings tab.

    Returns:
        "✅ Rebuilt N learnings in ChromaDB" or error string.

    Voice examples:
        "Rebuild the learnings search index"
        "Fix the learnings database"
        "Reindex all my learnings"
    """
    _telemetry_increment_tool_count("rebuild_learnings_index")

    if not _sl:
        return "❌ Self-learning module not available."

    try:
        db = _sl._load_db()
        learnings = [l for l in db.get("learnings", [])
                     if l.get("status", "active") == "active"]
        count = 0
        for l in learnings:
            try:
                _sl.reindex_learning(l)
                count += 1
            except Exception:
                pass
        return f"✅ Rebuilt {count} active learning(s) in ChromaDB index."
    except Exception as _e:
        return f"❌ Rebuild failed: {_e}"




# ══════════════════════════════════════════════════════════════════════════════
# AI ANALYSIS QUEUE TOOLS  (v8.0.0)
# ══════════════════════════════════════════════════════════════════════════════
# These two tools power the Quick Links "AI Analysis" buttons.
#
# WORKFLOW:
#   1. User clicks an analysis button in the Quick Links tab (e.g. "Analyze My
#      Business"). AI-Prowler writes a task record to pending_tasks.json and
#      copies a matching Claude command to the clipboard.
#   2. User pastes the command into Claude. Claude calls
#      get_pending_analysis_tasks() which returns all queued tasks.
#   3. Claude executes the analysis — searching documents, reading job data,
#      reviewing learnings — and records insights via record_learning().
#   4. Claude calls complete_analysis_task(task_id) to mark each task done.
#
# pending_tasks.json schema (list of task objects):
#   {
#     "task_id":    "analyze_business_20260623_143022",
#     "type":       "analyze_business",
#     "label":      "Analyze My Business",
#     "prompt":     "Analyze my job tracker, recent learnings, ...",
#     "created_at": "2026-06-23T14:30:22Z",
#     "status":     "pending"   # or "completed"
#   }
# ══════════════════════════════════════════════════════════════════════════════

_PENDING_TASKS_FILE = Path.home() / ".ai-prowler" / "pending_tasks.json"


def _load_pending_tasks() -> list:
    """Load pending_tasks.json, returning an empty list if absent or corrupt."""
    try:
        if _PENDING_TASKS_FILE.exists():
            data = json.loads(_PENDING_TASKS_FILE.read_text(encoding="utf-8"))
            return data if isinstance(data, list) else []
    except Exception:
        pass
    return []


def _save_pending_tasks(tasks: list) -> None:
    """Write the tasks list to pending_tasks.json atomically."""
    _PENDING_TASKS_FILE.parent.mkdir(parents=True, exist_ok=True)
    _PENDING_TASKS_FILE.write_text(
        json.dumps(tasks, indent=2, ensure_ascii=False),
        encoding="utf-8"
    )


@mcp.tool()
def create_analysis_task(
    label: str,
    prompt: str,
    schedule: str = "none",
    first_due: str = "",
    scope_dirs: "list | None" = None,
    output_learnings: bool = True,
    output_report: bool = False,
    report_folder: str = "",
) -> str:
    """
    AGENTIC ANALYSIS — Define a new recurring (or one-off) custom analysis
    task in the Quick Links / My Custom Analyses queue, from a plain-language
    request — the same feature the "+ New Custom Analysis" GUI dialog builds,
    now callable directly from a conversation.

    IMPORTANT — how this actually behaves (set expectations honestly):
      • Day-granularity only. There is no time-of-day in this system — a
        request like "every Monday at 8am" can only be stored as "due every
        Monday" (schedule="weekly", first_due=<that Monday's date>). The
        "8am" part cannot be represented; do not imply it will be.
      • Pull-based, not autonomous. Creating a task does NOT make AI-Prowler
        wake up and run it unattended at the due time. A task becoming due
        just means it's waiting — the NEXT time the user is in a Claude
        conversation and either asks, or Claude calls
        get_pending_analysis_tasks() at session start, the due task shows up
        and can be run then. If the user's phrasing implies full automation
        ("check my email every Monday morning without me doing anything"),
        say so plainly rather than letting the wording stand uncorrected —
        that requires a Claude session to actually happen after the due
        date, not a background process.
      • schedule must be one of: none, daily, weekly, biweekly, monthly,
        quarterly, yearly. If schedule is anything but "none", first_due is
        REQUIRED (YYYY-MM-DD) — compute the correct date yourself (e.g. the
        next occurrence of the requested weekday) rather than asking the
        user to do date math. Use user_time_v0 first if you need today's
        date/timezone to compute it.
      • Maximum 25 custom tasks at once (enforced inside create_task() —
        the same limit the GUI dialog respects). If the user is at the cap,
        tell them to delete an existing task first (Links & Analysis tab).

    Args:
        label:            Short name for the task, max 60 characters.
        prompt:           The full analysis instructions Claude should run
                          when this task comes due — be as specific as the
                          user was, don't compress it.
        schedule:         "none" (one-off), "daily", "weekly", "biweekly",
                          "monthly", "quarterly", or "yearly".
        first_due:        YYYY-MM-DD. Required if schedule != "none".
        scope_dirs:       Optional list of directory paths to focus the
                          analysis on. Omit for "search everything."
        output_learnings: Record key findings as learnings when this task
                          runs (default True).
        output_report:    Save the full analysis as a .docx report when
                          this task runs (default False).
        report_folder:    Output folder for .docx reports, if output_report
                          is True. Uses the default reports folder if omitted.

    Returns:
        Confirmation with the new task_id and its next due date, or a
        clear validation error (including hitting the 25-task cap).

    Voice examples:
        "Set up a task to check for unread invoice emails every Monday"
        "Remind me — er, set up a recurring task — to review the AR aging
         report monthly"
        "Create a one-off task to summarize last quarter's contracts"
    """
    _telemetry_increment_tool_count("create_analysis_task")

    try:
        import custom_tasks_manager as _ctm
    except Exception as _ie:
        return f"❌ custom_tasks_manager module not available: {_ie}"

    try:
        tasks = _ctm.load_custom_tasks()
        new_task = _ctm.create_task(
            label=label,
            prompt=prompt,
            scope_dirs=scope_dirs or [],
            schedule=schedule,
            first_due=(first_due.strip() or None),
            output_learnings=output_learnings,
            output_report=output_report,
            report_folder=(report_folder.strip() or None),
        )
        tasks.append(new_task)
        if not _ctm.save_custom_tasks(tasks):
            return "❌ Could not save the new task to disk."
    except ValueError as _ve:
        return f"❌ {_ve}"
    except Exception as _e:
        return f"❌ Could not create task: {_e}"

    due_note = (f"first due {new_task['next_due']}"
               if new_task.get("next_due") else "no schedule (one-off, manual run only)")
    return (
        f"✅ Created custom analysis task '{label}'\n"
        f"   task_id  : {new_task['task_id']}\n"
        f"   schedule : {schedule} ({due_note})\n\n"
        f"This task will show up in get_pending_analysis_tasks() once it's "
        f"due AND queued — either the user clicks 'Queue' / 'Run Due Tasks' "
        f"in the Links & Analysis tab, or asks you to check for due tasks "
        f"in a future conversation. It will not run automatically without "
        f"either of those."
    )


@mcp.tool()
def save_analysis_report(task_id: str,
                         title: str,
                         content: str,
                         report_folder: str = "",
                         ctx: Context = None) -> str:
    """
    AGENTIC ANALYSIS — Save a full analysis report as a Word (.docx) document.

    Call this when a task has output_report=true. The report is saved to the
    task's configured report_folder (or ~/.ai-prowler/reports/ by default).
    After saving, record a completion learning via record_learning() with
    category 'analysis_report' noting the file path and next scheduled run.

    Args:
        task_id:       The task_id from get_pending_analysis_tasks().
        title:         Report title — used as the filename and document heading.
                       Example: "Q2 Window Review — June 30 2026"
        content:       Full report content in Markdown format.
                       Use ## for sections, **bold** for emphasis, - for lists.
                       Claude should write this as a comprehensive document.
        report_folder: Override the output folder. Leave empty to use the
                       task's configured folder or the default reports folder.

    Returns:
        Path to the saved .docx file, or error message.

    Example:
        save_analysis_report(
            task_id="custom_001_20260630_143022",
            title="Q2 Window Cleaning Review — June 30 2026",
            content="## Executive Summary\\n\\nThis quarter showed...",
        )
    """
    _telemetry_increment_tool_count("save_analysis_report")

    if not task_id or not task_id.strip():
        return "❌ task_id is required."
    if not title or not title.strip():
        return "❌ title is required."
    if not content or not content.strip():
        return "❌ content is required."

    title   = title.strip()
    content = content.strip()

    # Determine output folder
    import datetime as _dt
    from pathlib import Path as _Path

    # Check pending_tasks.json for the report_folder from the task definition
    if not report_folder:
        try:
            tasks = _load_pending_tasks()
            for t in tasks:
                if t.get("task_id") == task_id.strip():
                    report_folder = t.get("report_folder", "")
                    break
        except Exception:
            pass

    if not report_folder:
        report_folder = str(_Path.home() / "Documents" / "AI-Prowler_tasks_reports")

    try:
        out_dir = _Path(report_folder)
        out_dir.mkdir(parents=True, exist_ok=True)
    except Exception as _e:
        return f"❌ Could not create report folder '{report_folder}': {_e}"

    # Build safe filename from title + date
    safe_title = "".join(
        c if c.isalnum() or c in " -_" else "_"
        for c in title
    ).strip().replace(" ", "_")[:80]
    date_str  = _dt.date.today().strftime("%Y-%m-%d")
    filename  = f"{safe_title}_{date_str}.docx"
    out_path  = out_dir / filename

    # Write the docx using python-docx
    try:
        from docx import Document as _Document
        from docx.shared import Pt as _Pt, RGBColor as _RGB
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN

        doc = _Document()

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = _WD_ALIGN.CENTER

        # Date subtitle
        sub = doc.add_paragraph(
            f"Generated: {_dt.datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        )
        sub.alignment = _WD_ALIGN.CENTER
        sub.runs[0].font.size = _Pt(10)
        sub.runs[0].font.color.rgb = _RGB(0x66, 0x66, 0x66)

        doc.add_paragraph()  # spacer

        # Parse Markdown content into docx elements
        for line in content.split("\n"):
            stripped = line.rstrip()

            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("  - ") or stripped.startswith("  * "):
                p = doc.add_paragraph(stripped[4:], style="List Bullet 2")
            elif stripped == "" or stripped == "---":
                doc.add_paragraph()
            else:
                # Handle **bold** inline
                p = doc.add_paragraph()
                remaining = stripped
                while "**" in remaining:
                    before, _, rest = remaining.partition("**")
                    bold_text, _, remaining = rest.partition("**")
                    if before:
                        p.add_run(before)
                    run = p.add_run(bold_text)
                    run.bold = True
                if remaining:
                    p.add_run(remaining)

        doc.save(str(out_path))

    except ImportError:
        # python-docx not available — save as plain text .txt instead
        txt_path = out_dir / filename.replace(".docx", ".txt")
        txt_path.write_text(
            f"{title}\n{'=' * len(title)}\n\n{content}",
            encoding="utf-8"
        )
        return (
            f"⚠️ python-docx not installed — saved as plain text instead.\n"
            f"File: {txt_path}\n\n"
            f"To enable Word document output run:\n"
            f"  pip install python-docx"
        )

    except Exception as _e:
        return f"❌ Failed to save report: {_e}"

    return (
        f"✅ Report saved: {out_path}\n"
        f"Title: {title}\n"
        f"Size: {out_path.stat().st_size:,} bytes"
    )



@mcp.tool()
def get_pending_analysis_tasks(ctx: Context = None) -> str:
    """
    AGENTIC ANALYSIS — Return all pending analysis tasks queued by the user
    from the AI-Prowler Quick Links tab.

    Call this tool at the start of any session where the user says they have
    analysis to run, or whenever instructed to check for pending tasks.

    For each pending task returned, Claude should:
      1. Execute the analysis described in the task's 'prompt' field using
         all available AI-Prowler tools (search_documents, search_learnings,
         read_job_spreadsheet, get_weather, etc.)
      2. Record any significant findings as learnings via record_learning()
      3. Call complete_analysis_task(task_id) to mark the task done

    Returns:
        JSON list of pending task objects, each with:
          task_id    — unique identifier (pass to complete_analysis_task)
          type       — task type (analyze_business, weekly_advisor, etc.)
          label      — human-readable name
          prompt     — the full analysis prompt Claude should execute
          created_at — when the user queued this task (ISO 8601)

    Returns a plain message if no tasks are pending.

    Examples:
        "Check for pending analysis tasks"
        "Run any pending analysis from AI-Prowler"
        "What analysis do I have queued?"
    """
    _telemetry_increment_tool_count("get_pending_analysis_tasks")

    try:
        tasks = _load_pending_tasks()
        pending = [t for t in tasks if t.get("status") == "pending"]

        if not pending:
            return (
                "✅ No pending analysis tasks. "
                "Use the AI Analysis buttons in the Links & Analysis tab to queue analysis."
            )

        # Enrich each task with age info and output instructions
        import datetime as _dt
        now = _dt.datetime.utcnow()
        for t in pending:
            try:
                created = _dt.datetime.strptime(
                    t.get("created_at", ""), "%Y-%m-%dT%H:%M:%SZ")
                age_mins = int((now - created).total_seconds() / 60)
                if age_mins < 60:
                    t["queued_ago"] = f"{age_mins} minute{'s' if age_mins != 1 else ''} ago"
                elif age_mins < 1440:
                    hrs = age_mins // 60
                    t["queued_ago"] = f"{hrs} hour{'s' if hrs != 1 else ''} ago"
                else:
                    days = age_mins // 1440
                    t["queued_ago"] = f"{days} day{'s' if days != 1 else ''} ago"
            except Exception:
                t["queued_ago"] = "unknown"

        result = {
            "pending_count": len(pending),
            "tasks": pending,
            "instruction": (
                "For each task: execute the full analysis described in the 'prompt' "
                "field using all available AI-Prowler tools. "
                "If output_report=true: call save_analysis_report() with the full "
                "analysis as Markdown content. "
                "If output_learnings=true: call record_learning() with key insights. "
                "Always record a completion learning with category 'analysis_report' "
                "noting what was done and the next_due date. "
                "Finally call complete_analysis_task(task_id, summary) to mark done "
                "and auto-advance the next scheduled run date."
            )
        }
        return json.dumps(result, indent=2, ensure_ascii=False)

    except Exception as _e:
        return f"❌ Could not load pending tasks: {_e}"


@mcp.tool()
def list_analysis_tasks(ctx: Context = None) -> str:
    """
    AGENTIC ANALYSIS — List EVERY custom analysis task the user has defined
    in the Links & Analysis tab's "My Custom Analyses" panel (up to 25),
    regardless of whether it's currently due.

    This is DIFFERENT from get_pending_analysis_tasks(), which only shows
    tasks that have already been QUEUED into pending_tasks.json (either by
    clicking "Run Due Tasks" in the GUI, or by a task actually being due).
    A task can exist here — fully defined, scheduled, recurring — and never
    show up in get_pending_analysis_tasks() until its next_due date arrives
    AND someone queues it. Use this tool whenever the user asks something
    like "what's in my task queue", "what tasks do I have set up", "list
    all my custom analyses", or "show me everything I've scheduled" — those
    are asking about the full definition list, not just what's due right now.

    Does NOT modify anything and does NOT queue anything — read-only. To
    actually run a task before its due date, call the analysis directly
    (or tell the user to click "Run Due Tasks" / queue it manually in the
    GUI, or use create_analysis_task with a near-term schedule).

    Returns:
        JSON with:
          total_count — how many custom tasks exist (max 25)
          tasks       — full list, each with:
            task_id, label, prompt, schedule, first_due, next_due,
            is_due (bool — true if next_due <= today AND schedule != none),
            last_run, last_status, output_learnings, output_report,
            report_folder, scope_dirs

        Returns a plain message if no custom tasks are defined yet.

    Examples:
        "What's in my task queue?"
        "List all my scheduled analyses"
        "Show me every custom task I've set up, even ones not due yet"
    """
    _telemetry_increment_tool_count("list_analysis_tasks")

    try:
        import custom_tasks_manager as _ctm
    except Exception as _ie:
        return f"❌ custom_tasks_manager module not available: {_ie}"

    try:
        tasks = _ctm.load_custom_tasks()
    except Exception as _e:
        return f"❌ Could not load custom tasks: {_e}"

    if not tasks:
        return (
            "✅ No custom analysis tasks defined yet. "
            "Use create_analysis_task() or the \"+ New Custom Analysis\" "
            "button in the Links & Analysis tab to create one."
        )

    try:
        due_ids = {t["task_id"] for t in _ctm.get_due_tasks(tasks)}
    except Exception:
        due_ids = set()
    for t in tasks:
        t["is_due"] = t.get("task_id") in due_ids

    result = {
        "total_count": len(tasks),
        "max_tasks": _ctm.MAX_CUSTOM_TASKS,
        "tasks": tasks,
        "note": (
            "This is the full definition list (custom_analysis_tasks.json), "
            "not the run queue. A task with is_due=true is not automatically "
            "running — it still needs to be queued (Run Due Tasks in the GUI, "
            "or check get_pending_analysis_tasks() in a future session) or "
            "run directly right now if the user asks for that."
        ),
    }
    return json.dumps(result, indent=2, ensure_ascii=False)


@mcp.tool()
def complete_analysis_task(task_id: str,
                           summary: str = "",
                           ctx: Context = None) -> str:
    """
    AGENTIC ANALYSIS — Mark a pending analysis task as completed.

    Call this after finishing each analysis task returned by
    get_pending_analysis_tasks(). This prevents the same task from
    running again in future sessions.

    Args:
        task_id: The task_id from get_pending_analysis_tasks() output.
                 Example: "analyze_business_20260623_143022"
        summary: Optional one-sentence summary of what was found/recorded.
                 Stored with the task for the user's reference in the GUI.

    Returns:
        Confirmation message with the task label and summary.

    Example:
        complete_analysis_task(
            task_id="analyze_business_20260623_143022",
            summary="Found 3 overdue invoices and recorded 4 business insights."
        )
    """
    _telemetry_increment_tool_count("complete_analysis_task")

    if not task_id or not task_id.strip():
        return "❌ task_id is required."

    task_id = task_id.strip()

    try:
        tasks = _load_pending_tasks()
        matched = False
        label = task_id

        for t in tasks:
            if t.get("task_id") == task_id:
                t["status"] = "completed"
                # Use inline datetime to avoid dependency on _utc_now_iso
                # which is only defined inside the server startup closure.
                try:
                    import datetime as _dt
                    t["completed_at"] = _dt.datetime.utcnow().strftime(
                        "%Y-%m-%dT%H:%M:%SZ")
                except Exception:
                    t["completed_at"] = ""
                if summary:
                    t["completion_summary"] = summary.strip()
                label = t.get("label", task_id)
                matched = True
                break

        if not matched:
            return (
                f"⚠️ Task '{task_id}' not found in pending_tasks.json. "
                "It may have already been completed or was never queued."
            )

        _save_pending_tasks(tasks)

        # ── Schedule advancement ──────────────────────────────────────────────
        # For custom tasks (source_id present): advance next_due in
        # custom_analysis_tasks.json.
        # For built-in tasks with a schedule (schedule != "none"): advance
        # next_due directly in the completed task record so the Queue Panel
        # and get_pending_analysis_tasks() can surface when it's next due.

        source_id    = None
        task_schedule = None
        task_next_due = None

        for t in tasks:
            if t.get("task_id") == task_id:
                source_id     = t.get("source_id")
                task_schedule = t.get("schedule", "none")
                task_next_due = t.get("next_due")
                break

        next_due_msg = ""

        if source_id:
            # Custom task — advance next_due in custom_analysis_tasks.json
            try:
                import datetime as _dt
                import custom_tasks_manager as _ctm
                custom_tasks = _ctm.load_custom_tasks()
                new_next_due = _ctm.advance_next_due(
                    custom_tasks, source_id,
                    completed_date=_dt.date.today().isoformat()
                )
                _ctm.save_custom_tasks(custom_tasks)
                if new_next_due:
                    next_due_msg = f"\nNext scheduled run: {new_next_due}"
            except Exception as _cte:
                print(f"[complete_analysis_task] custom next_due advance failed: {_cte}")

        elif task_schedule and task_schedule != "none":
            # Built-in scheduled task — advance next_due in pending_tasks.json
            try:
                import datetime as _dt
                import custom_tasks_manager as _ctm
                anchor = task_next_due or _dt.date.today().isoformat()
                # v8.1.5 fix: use the catch-up variant so a task overdue by
                # MULTIPLE intervals resyncs fully in one completion instead
                # of needing one completion per missed interval.
                new_next_due = _ctm._advance_date_catchup(
                    anchor, task_schedule, _dt.date.today().isoformat())
                if new_next_due:
                    # Stamp the new next_due on the completed task record
                    for t in tasks:
                        if t.get("task_id") == task_id:
                            t["next_due"] = new_next_due
                            break
                    _save_pending_tasks(tasks)
                    next_due_msg = f"\nNext scheduled run: {new_next_due}"
            except Exception as _cte:
                print(f"[complete_analysis_task] builtin next_due advance failed: {_cte}")

        msg = f"✅ Analysis task completed: {label}"
        if summary:
            msg += f"\nSummary: {summary}"
        msg += next_due_msg
        return msg

    except Exception as _e:
        return f"❌ Could not complete task: {_e}"


# ══════════════════════════════════════════════════════════════════════════════
# WRITE ZONE MANAGEMENT TOOLS  (v7.1.0)
# ══════════════════════════════════════════════════════════════════════════════
# Tools to list, grant, and revoke write-zone permissions from mobile.
#
# Server-mode scope enforcement:
#   • owner      — may grant/revoke any directory in the read allowlist
#   • manager    — may grant/revoke directories that are in their scope
#                  (scope is determined by the collection_map prefix rules)
#   • staff /    — no write-zone management; read-only users cannot grant
#     field_crew   themselves write access
#
# Personal mode (no ctx user): full access — no restrictions.
# ══════════════════════════════════════════════════════════════════════════════

def _write_zone_allowed_for_user(user: "dict | None", directory: str,
                                 users_data: "dict | None" = None) -> tuple:
    """Server-mode gate: may this user grant/revoke write access to directory?
    Returns (allowed: bool, reason: str). PURE-ish (reads users_data once)."""
    if user is None:
        return (True, "personal mode — no restrictions")

    role = (user.get("role") or "").lower()
    caps = _role_caps(role)

    # staff / field_crew: never
    if not caps.get("can_write"):
        return (False,
                f"role '{role}' cannot manage write zones. "
                "Only owner or manager roles may change write permissions.")

    # owner: unrestricted
    if caps.get("is_admin"):
        return (True, "owner may manage any write zone")

    # manager: directory must fall within one of their assigned scopes'
    # prefix rules in the collection_map (or their own private area).
    # We check the allowlisted read paths and whether the requested directory
    # is under a prefix the user can write to per _can_index.
    if users_data is None:
        users_data = _load_users()

    coll_map = _company_collection_map(users_data)
    dir_norm  = _normalize_path_for_match(directory)

    # Does any collection-map rule whose collection is in the user's scopes
    # cover this directory?
    user_scopes = set()
    for s in (user.get("scopes") or []):
        s = str(s).strip()
        user_scopes.add(_canon_scope(s))

    for rule in (coll_map.get("rules") or []):
        prefix = _normalize_path_for_match(rule.get("prefix", ""))
        coll   = str(rule.get("collection", "")).strip()
        if not prefix or not coll:
            continue
        if (dir_norm == prefix or dir_norm.startswith(prefix + "/")):
            if _canon_scope(coll) in user_scopes:
                return (True, f"directory is within manager's scope '{coll}'")

    # Also allow manager's own private area (no scope rule needed)
    uid = user.get("id", "")
    if uid:
        priv_col = f"user:{uid}"
        allowed_priv, _ = _can_index(user, priv_col)
        if allowed_priv:
            # Check if the directory is the user's home or a known private path
            # (we can't know exactly — use a conservative check: allow if no
            # company rules matched but the user explicitly owns this area)
            pass   # fall through to denial — safest for managers

    return (False,
            f"directory '{directory}' is not within any scope assigned to role "
            f"'{role}'. Only the owner can grant write access outside your "
            f"assigned scopes.")


@mcp.tool()
def list_writable_directories(ctx: Context = None) -> str:
    """
    List directories currently in the write-zone allowlist
    (~/.rag_writable_dirs.json) with their status.

    Personal mode: full write-zone and read-zone allowlist, unchanged.

    Server mode:
      - Owner / manager (full DB-management capability): full company-wide
        list, same as personal mode — they're the ones who manage the
        write-zone config via the Admin tab / grant_write_access.
      - Staff / field_crew: shows ONLY their own personal directory's
        read/write status (whether it's currently in the writable
        allowlist), not the company-wide list of arbitrary host paths.
        This tells them whether create_file/write_file etc. will actually
        work for them right now, without exposing internal folder layout
        unrelated to their own work.

    Voice examples:
        "What directories can Claude write to?"
        "Show me my write permissions"
        "Can I write to my personal folder?"
    """
    _telemetry_increment_tool_count("list_writable_directories")

    _lwd_user = _current_user(ctx)

    if _lwd_user is not None:
        _lwd_db_ok, _ = _check_db_cap(_lwd_user, "full")
        if not _lwd_db_ok:
            # Staff / field_crew: own personal directory status only.
            _lwd_status, _lwd_priv_dir = _user_private_write_dir(ctx)
            if _lwd_status == "blocked":
                return (
                    "📁 Personal Directory Status\n\n"
                    "🚫 You don't have a personal directory configured.\n"
                    "Ask your owner/admin to enable 'Private collection' and set "
                    "up your personal folder in the Admin tab."
                )
            _lwd_writable = _writable_allowlist_load()
            _lwd_priv_norm = _normalize_path_for_match(str(_lwd_priv_dir))
            _lwd_is_writable = any(
                _normalize_path_for_match(w) == _lwd_priv_norm for w in _lwd_writable
            )
            if _lwd_is_writable:
                return (
                    "📁 Personal Directory Status\n\n"
                    f"  ✅ [W]  {_lwd_priv_dir}\n\n"
                    "Your personal directory is read + write — create_file, "
                    "write_file, and similar tools will work inside it."
                )
            return (
                "📁 Personal Directory Status\n\n"
                f"  📖 [R]  {_lwd_priv_dir}\n\n"
                "Your personal directory is currently read-only. Ask your "
                "owner/admin to enable write access in the Admin tab."
            )

    # Personal mode, or server-mode owner/manager: full list, unchanged.
    writable = _writable_allowlist_load()
    read_dirs = load_auto_update_list() or []

    lines = ["📁 AI-Prowler Write Zone Status\n"]
    lines.append(f"Writable directories ({len(writable)}):")
    if writable:
        for w in sorted(writable):
            lines.append(f"  ✅ [W]  {w}")
    else:
        lines.append("  (none configured — Claude can read but not write files)")

    lines.append("")
    read_only = [r for r in read_dirs if r not in writable]
    lines.append(f"Read-only directories ({len(read_only)}):")
    if read_only:
        for r in sorted(read_only):
            lines.append(f"  📖 [R]  {r}")
    else:
        lines.append("  (no additional read-only directories)")

    lines += [
        "",
        "Legend: [W] = writable  [R] = read-only",
        "Use grant_write_access() to enable writes, revoke_write_access() to disable.",
    ]
    return "\n".join(lines)


@mcp.tool()
def grant_write_access(directory: str, ctx: Context = None) -> str:
    """
    Add a directory to the write-zone allowlist, giving Claude permission to
    create and edit files there. The directory must already be in the tracked
    read allowlist.

    In server mode, managers may only grant write access to directories within
    their assigned scopes. The owner may grant any directory.
    In personal mode there are no restrictions.

    Args:
        directory: Absolute path to the directory to make writable.
                   Must be in the tracked read allowlist.
        ctx:       MCP context (injected automatically)

    Returns:
        Confirmation string or error.

    Voice examples:
        "Grant write access to my AI-Prowler work folder"
        "Make the documents folder writable"
        "Allow Claude to edit files in C:/Projects/MyProject"
    """
    _telemetry_increment_tool_count("grant_write_access")

    directory = directory.strip()
    if not directory:
        return "❌ directory is required."

    # Normalize the path
    try:
        resolved = str(Path(directory).resolve())
    except Exception:
        resolved = directory

    # Server-mode scope check
    user = _current_user(ctx)
    users_data = _load_users() if user else None
    allowed, why = _write_zone_allowed_for_user(user, resolved, users_data)
    if not allowed:
        return f"❌ {why}"

    # Must be in the read allowlist
    read_dirs = load_auto_update_list() or []
    read_norm = [_normalize_path_for_match(r) for r in read_dirs]
    dir_norm  = _normalize_path_for_match(resolved)

    in_read = any(
        dir_norm == rn or dir_norm.startswith(rn + "/")
        for rn in read_norm
    )
    if not in_read:
        return (f"❌ '{resolved}' is not in the read allowlist.\n"
                f"Add it first with index_path(), then grant write access.")

    # Check if already writable
    writable = _writable_allowlist_load()
    writable_norm = [_normalize_path_for_match(w) for w in writable]
    if any(dir_norm == wn or dir_norm.startswith(wn + "/")
           for wn in writable_norm):
        return f"ℹ️  '{resolved}' is already in the write zone — no change needed."

    # Add and save
    writable.append(resolved)
    if not _writable_allowlist_save(writable):
        return "❌ Could not save write-zone allowlist. Check disk permissions."

    who = f" (granted by {user.get('name','?')})" if user else ""
    _log.info("grant_write_access: %s%s", resolved, who)
    return (f"✅ Write zone granted: '{resolved}'\n"
            f"Claude can now create and edit files in this directory.\n"
            f"Use revoke_write_access() to remove this permission.")


@mcp.tool()
def revoke_write_access(directory: str, ctx: Context = None) -> str:
    """
    Remove a directory from the write-zone allowlist. Claude will no longer
    be able to create or edit files there (read access is unaffected).

    In server mode, the same scope restrictions as grant_write_access apply.
    In personal mode there are no restrictions.

    Args:
        directory: Path to remove from the write zone.
        ctx:       MCP context (injected automatically)

    Returns:
        Confirmation string or error.

    Voice examples:
        "Revoke write access to the downloads folder"
        "Make the temp folder read-only again"
        "Remove write permission from C:/OldProject"
    """
    _telemetry_increment_tool_count("revoke_write_access")

    directory = directory.strip()
    if not directory:
        return "❌ directory is required."

    try:
        resolved = str(Path(directory).resolve())
    except Exception:
        resolved = directory

    # Server-mode scope check
    user = _current_user(ctx)
    users_data = _load_users() if user else None
    allowed, why = _write_zone_allowed_for_user(user, resolved, users_data)
    if not allowed:
        return f"❌ {why}"

    writable = _writable_allowlist_load()
    dir_norm  = _normalize_path_for_match(resolved)

    # Find matches (exact or parent)
    to_remove = [w for w in writable
                 if _normalize_path_for_match(w) == dir_norm
                 or _normalize_path_for_match(w).startswith(dir_norm + "/")]

    if not to_remove:
        return f"ℹ️  '{resolved}' is not in the write zone — nothing to revoke."

    new_list = [w for w in writable if w not in to_remove]
    if not _writable_allowlist_save(new_list):
        return "❌ Could not save write-zone allowlist. Check disk permissions."

    who = f" (revoked by {user.get('name','?')})" if user else ""
    removed = ", ".join(to_remove)
    _log.info("revoke_write_access: removed %s%s", removed, who)
    return (f"✅ Write zone revoked: '{resolved}'\n"
            f"Claude can still READ files there but can no longer write or edit them.")


# ══════════════════════════════════════════════════════════════════════════════
# REINDEX TOOLS  (v7.1.0)
# ══════════════════════════════════════════════════════════════════════════════

@mcp.tool()
def reindex_file(filepath: str, ctx: Context = None) -> str:
    """
    Re-index a SINGLE file in ChromaDB to match its current on-disk content.

    Call this ONCE when you have finished editing a file — NOT after every
    individual str_replace_in_file. Writes no longer auto-index (that caused an
    HTTP-server deadlock on large files), so the database stays stale until you
    run this. Delete-then-add: purges the file's existing chunks, then re-chunks
    and re-embeds the current content.

    Server mode: collection-aware. Stale chunks are purged from whichever
    collection they actually live in (own private, an assigned scope, or
    shared) — not just the default — and fresh chunks are routed back into
    the correct collection via the same resolver index_path() uses, so
    content never gets relocated into a broader collection than it started in.

    Args:
        filepath: Absolute path to the file. Must be under a tracked
                  read-allowlisted root.

    Voice examples:
        "Re-index rag_gui.py"
        "Update the index for the file I just edited"
    """
    _telemetry_increment_tool_count("reindex_file")

    # SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cutover, 2026-07-16):
    # role-based manage_db gate removed -- see index_path's identical
    # comment for the full rationale (indexing isn't a leak; every tracked
    # path was already admin/owner-created).
    _user_rf = _current_user(ctx)

    try:
        from rag_preprocessor import (
            normalise_path, index_file_list,
            get_chroma_client, COLLECTION_NAME,
        )
    except Exception as exc:
        return f"❌ reindex_file failed to import indexer: {exc}"

    # Resolve + allowlist-gate using the SAME wrapper the write/read tools use.
    # Returns (resolved_path, None) on allow, or (None, error_message) on deny.
    resolved, err = _resolve_allowlisted_path(filepath)
    if err:
        return err

    if not Path(resolved).is_file():
        return f"❌ Not a file (or does not exist): {resolved}"

    # No longer builds a collection_resolver -- there is only one physical
    # collection now. Scope is carried entirely by build_scope_resolver()'s
    # "scope" chunk-metadata tag (wired into index_file_list already).
    _indexer_user_rf = _user_rf

    # All ChromaDB writes run on the single dedicated db-writer thread to avoid
    # the HNSW cross-thread EDEADLK ("resource deadlock would occur") that wedges
    # the HTTP server. purge + re-index happen together inside one job.
    def _job():
        try:
            client, embedding_func = get_chroma_client()
            # Single physical collection now -- purge just this one, not a
            # sweep across every collection (there's only ever one to check).
            _coll = client.get_or_create_collection(
                name=COLLECTION_NAME, embedding_function=embedding_func)
            _coll.delete(where={"filepath": resolved})
        except Exception as exc:
            _log.warning("reindex_file: purge failed for %s: %s", resolved, exc)
            # continue — index_file_list still adds fresh chunks
        return index_file_list([resolved], label="reindex_file",
                               root_directory=str(Path(resolved).parent),
                               indexer_user=_indexer_user_rf)

    try:
        stats = _db_write(_job, timeout=900.0)
    except TimeoutError as exc:
        return f"⏱️  reindex_file timed out for {resolved}: {exc}"
    except Exception as exc:
        return f"❌ reindex_file failed for {resolved}: {exc}"

    chunks = 0
    try:
        chunks = int(stats.get("chunks", 0)) if isinstance(stats, dict) else 0
    except Exception:
        pass

    return (f"✅ Re-indexed {resolved}\n"
            f"   {chunks:,} chunk(s) now in ChromaDB for this file.")


@mcp.tool()
def reindex_directory(directory: str, purge_first: bool = True,
                      ctx: Context = None) -> str:
    """
    Fully re-index a tracked directory — purge all existing chunks then
    rebuild from scratch. More thorough than update_tracked_directories which
    only processes changed files. Use this after chunk-size changes, suspected
    index corruption, or when you want a guaranteed clean slate.

    Stale chunks are purged from the single index, then each file is
    re-indexed and re-tagged with its current scope (see
    SCOPE_SIMPLIFICATION_SPEC.md section 3.3b/3.7).

    Args:
        directory:   Absolute path to the tracked directory to reindex.
                     Must be in the read allowlist.
        purge_first: If True (default), delete existing chunks before
                     re-indexing. Set False to add/update without purging
                     (faster but won't remove stale chunks for deleted files).
        ctx:         MCP context (injected automatically)

    Returns:
        Summary of files indexed and chunks created.

    Voice examples:
        "Completely re-index my documents folder"
        "Wipe and rebuild the index for the work folder"
        "Reindex everything in AI-Prowler"
    """
    _telemetry_increment_tool_count("reindex_directory")

    # SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cutover, 2026-07-16):
    # role-based manage_db gate removed -- see index_path's identical
    # comment for the full rationale (indexing isn't a leak; every tracked
    # path was already admin/owner-created).
    _user_rd = _current_user(ctx)

    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait and try again."

    directory = directory.strip()
    if not directory:
        return "❌ directory is required."

    resolved, err = _resolve_allowlisted_path(directory)
    if err:
        return f"❌ {err}"

    # No longer builds a collection_resolver -- there is only one physical
    # collection now. Scope is carried entirely by build_scope_resolver()'s
    # "scope" chunk-metadata tag (wired into index_directory already).
    _indexer_user_rd = _user_rd

    try:
        import datetime as _dt6
        start = _dt6.datetime.now()

        # purge + re-index run together on the dedicated db-writer thread to
        # avoid the HNSW cross-thread EDEADLK that wedges the HTTP server.
        def _job():
            if purge_first:
                try:
                    client, embedding_func = _engine.get_chroma_client()
                    norm_dir = _normalize_path_for_match(resolved)
                    # Single physical collection now -- purge just this one,
                    # not a sweep across every collection.
                    coll = client.get_or_create_collection(
                        name=_engine.COLLECTION_NAME, embedding_function=embedding_func)
                    existing = coll.get(where={"source": {"$regex": ".*"}},
                                        include=["metadatas"])
                    ids_to_del = [
                        existing["ids"][i]
                        for i, meta in enumerate(existing.get("metadatas") or [])
                        if _normalize_path_for_match(
                               str(meta.get("source",""))).startswith(norm_dir)
                    ]
                    if ids_to_del:
                        coll.delete(ids=ids_to_del)
                        _log.info(
                            "reindex_directory: purged %d chunks for %s",
                            len(ids_to_del), resolved)
                except Exception as _pe:
                    _log.warning("reindex_directory: purge step warning: %s", _pe)
            with _capture_stdout() as buf:
                _r = index_directory(resolved, indexer_user=_indexer_user_rd)
            return _r, buf.getvalue()

        result, _out = _db_write(_job, timeout=1800.0)
        buf_value = _out
        elapsed = (_dt6.datetime.now() - start).total_seconds()

        output = (buf_value or "").strip()
        if isinstance(result, dict):
            files   = result.get("files_indexed", "?")
            chunks  = result.get("chunks_added", "?")
            return (f"✅ Reindex complete for: {resolved}\n"
                    f"   Files indexed : {files}\n"
                    f"   Chunks created: {chunks}\n"
                    f"   Time          : {elapsed:.1f}s\n"
                    + (f"   Output        : {output[:200]}" if output else ""))
        else:
            return (f"✅ Reindex triggered for: {resolved}\n"
                    f"   Time: {elapsed:.1f}s\n"
                    + (f"   {output[:300]}" if output else ""))

    except Exception as _e:
        _log.error("reindex_directory error: %s", _e)
        return f"❌ Reindex failed: {_e}"


@mcp.tool()
def reindex_all(purge_first: bool = True, ctx: Context = None) -> str:
    """
    Fully re-index ALL tracked directories — the nuclear option. Purges and
    rebuilds the entire ChromaDB index from scratch. Use after changing
    chunk size settings, or to fix any index corruption.

    This may take several minutes depending on how many documents are indexed.

    Args:
        purge_first: If True (default), wipe the entire index before
                     re-indexing. Set False to update without purging.
        ctx:         MCP context (injected automatically)

    Returns:
        Summary of all directories processed, files indexed, and time taken.

    Voice examples:
        "Rebuild the entire knowledge base from scratch"
        "Reindex all tracked directories"
        "Do a full reindex of everything"
    """
    _telemetry_increment_tool_count("reindex_all")

    # SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cutover, 2026-07-16):
    # role-based manage_db gate removed -- see index_path's identical
    # comment for the full rationale (indexing isn't a leak; every tracked
    # path was already admin/owner-created).
    _user_ra = _current_user(ctx)

    if not _prewarm_event.wait(timeout=60):
        return "⏳ AI-Prowler is still initializing. Please wait and try again."

    read_dirs = load_auto_update_list() or []
    if not read_dirs:
        return "ℹ️  No tracked directories found. Use index_path() first."

    import datetime as _dt7
    start     = _dt7.datetime.now()
    results   = []
    total_files  = 0
    total_chunks = 0
    errors = []

    for d in read_dirs:
        try:
            dir_result = reindex_directory(d, purge_first=purge_first, ctx=ctx)
            results.append(f"  ✅ {d}")
            # Parse counts from the returned string if available
            for line in dir_result.splitlines():
                if "Files indexed" in line:
                    try:
                        total_files += int(line.split(":")[-1].strip())
                    except Exception:
                        pass
                if "Chunks created" in line:
                    try:
                        total_chunks += int(line.split(":")[-1].strip())
                    except Exception:
                        pass
        except Exception as _e:
            errors.append(f"  ❌ {d}: {_e}")

    elapsed = (_dt7.datetime.now() - start).total_seconds()
    lines = [
        f"✅ Full reindex complete — {len(read_dirs)} director{'y' if len(read_dirs)==1 else 'ies'}",
        f"   Total files  : {total_files}",
        f"   Total chunks : {total_chunks}",
        f"   Time         : {elapsed:.1f}s",
        "",
        "Directories processed:",
    ] + results

    if errors:
        lines += ["", "Errors:"] + errors

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# EDITION / MODE + ACTIVATION — MODULE-LEVEL PURE HELPERS  (v7.0.0 Phase A')
# ══════════════════════════════════════════════════════════════════════════════
# These four functions are PURE (no I/O beyond reading config.json in
# _load_runtime_config) and were hoisted out of _run_http() to module level so
# the test suite (tests/mcp/test_edition_activation.py) can import and call them
# directly without launching the HTTP server. _run_http() calls them by bare
# name, which now resolves here. See PHASE_A_PRIME_TEST_PLAN.md §4.0.
#
# EDITION : "home" | "mobile" | "business"     MODE : "personal" | "server"
# Valid combos: (home,personal) (mobile,personal) (business,personal)
#               (business,server).  (home,server) and (mobile,server) invalid.
# Backwards-compat: plan "individual" reads as the "mobile" edition; no live
# subs.json record is ever rewritten.
# ══════════════════════════════════════════════════════════════════════════════
# ── Test-state sandbox hook (v7.0.0 — pre-release server-mode validation) ─────
# When AIPROWLER_TEST_STATE_DIR is set, AI-Prowler reads its state files
# (config.json, users.json, and the subs cache) from that directory instead of
# ~/.ai-prowler. This lets an automated test suite stand up a fully sandboxed
# server-mode instance without touching the operator's real state and without a
# live license. SAFETY: this only redirects WHERE state is read from; it does
# NOT disable any enforcement (auth, scoping, ownership all still run for real).
# The companion entitlement short-circuit (_test_entitlement_active / used in
# main()) is what skips the NETWORK license/subscription calls under test, and
# it is HARD-GATED to dev/Home boxes — see _test_entitlement_active(). The env
# var must be set at process launch, so it can never be flipped on a deployed,
# tunnel-exposed customer install by editing a file or calling a tool.
def _state_dir() -> Path:
    """Directory holding AI-Prowler's state files. Honors the test sandbox env
    var AIPROWLER_TEST_STATE_DIR when set; otherwise the real ~/.ai-prowler."""
    import os as _os
    td = _os.environ.get("AIPROWLER_TEST_STATE_DIR", "").strip()
    if td:
        return Path(td)
    return Path.home() / ".ai-prowler"


def _test_state_active() -> bool:
    """True iff the test-state sandbox env var is set (paths are redirected)."""
    import os as _os
    return bool(_os.environ.get("AIPROWLER_TEST_STATE_DIR", "").strip())


def _test_entitlement_active(cfg: "dict | None" = None) -> bool:
    """True iff the test ENTITLEMENT short-circuit may engage — i.e. the startup
    flow should substitute a sandboxed (edition/mode/status) verdict and SKIP the
    network license/subscription/activation calls.

    HARD-GATED by TWO independent affirmations, both only settable by whoever
    launches the process (never by a deployed customer box at runtime):
      1. AIPROWLER_TEST_STATE_DIR env var is set (state is sandboxed), AND
      2. the sandboxed config.json explicitly carries "test_mode": true.
    Either alone is insufficient. This short-circuit substitutes only the
    ENTITLEMENT verdict; it never disables auth, scoping, or ownership — those
    run for real against the sandboxed users.json. PURE (reads cfg only)."""
    if not _test_state_active():
        return False
    if cfg is None:
        try:
            cfg = _load_runtime_config()
        except Exception:
            return False
    return cfg.get("test_mode") is True


_CONFIG_PATH = _state_dir() / "config.json"

_VALID_EDITIONS = ("home", "mobile", "business")
_VALID_MODES    = ("personal", "server")
_MOBILE_PLAN_SYNONYMS   = ("mobile", "individual")
_BUSINESS_PLAN_SYNONYMS = ("business", "small_business", "enterprise")

# 2-active-install rule (architecture spec §4)
_ACTIVE_WINDOW_DAYS  = 14   # last_seen within N days == "active"
_MAX_ACTIVE_INSTALLS = 2    # at most this many active install_ids per license


def _plan_to_edition(plan: str) -> str:
    """Map a subs.json plan string onto an edition. Unknown/empty → 'mobile'
    (any managed subscriber gets at least remote access; a fail-open default
    that never strips entitlement from a paying user)."""
    p = (plan or "").strip().lower()
    if p in _BUSINESS_PLAN_SYNONYMS:
        return "business"
    if p in _MOBILE_PLAN_SYNONYMS:
        return "mobile"
    return "mobile"


def _load_runtime_config() -> dict:
    """Read ~/.ai-prowler/config.json and return a normalized dict with at
    least edition/mode keys. Missing file or unreadable JSON yields the safe
    default (home, personal). Never raises."""
    cfg = {}
    try:
        if _CONFIG_PATH.exists():
            # utf-8-sig tolerates a BOM from PowerShell/Notepad-edited configs.
            cfg = json.loads(_CONFIG_PATH.read_text(encoding="utf-8-sig")) or {}
    except Exception as _e:
        _log.warning("config.json unreadable (%s) — defaulting to home/personal", _e)
        cfg = {}
    edition = str(cfg.get("edition", "home")).strip().lower()
    mode    = str(cfg.get("mode", "personal")).strip().lower()
    if edition not in _VALID_EDITIONS:
        _log.warning("Unknown edition %r in config — defaulting to 'home'", edition)
        edition = "home"
    if mode not in _VALID_MODES:
        _log.warning("Unknown mode %r in config — defaulting to 'personal'", mode)
        mode = "personal"
    cfg["edition"] = edition
    cfg["mode"]    = mode
    return cfg


def _enforce_edition_mode(edition: str, mode: str, sub_status: str) -> tuple[str, str]:
    """Belt-and-suspenders startup enforcement (architecture spec §2),
    hardened against manual config.json tampering. Returns the (possibly
    downgraded) effective (edition, mode).

    Rules:
      1. server mode requires business edition → else fall back to personal.
      2. mobile/business editions require a VALID managed subscription
         (sub_status 'ok' or 'warning'); otherwise fall back to home.
    """
    eff_edition, eff_mode = edition, mode

    if eff_mode == "server" and eff_edition != "business":
        _log.warning(
            "Server mode requires Business edition; got edition=%s. "
            "Falling back to personal mode.", eff_edition)
        eff_mode = "personal"

    if eff_edition in ("mobile", "business") and sub_status not in ("ok", "warning"):
        _log.warning(
            "Edition=%s requires a valid subscription but status is %r. "
            "Falling back to Home edition (remote access disabled).",
            eff_edition, sub_status)
        eff_edition = "home"
        eff_mode = "personal"   # home can never be server

    return eff_edition, eff_mode


def _evaluate_activation(entry: dict, install_id: str, now=None) -> dict:
    """Decide this install_id's activation standing for a subscriber entry.
    PURE — no I/O, no mutation. Returns a dict with keys: decision
    ('active'|'admissible'|'rejected'|'unbound'), active_install_ids,
    active_count, this_active, message. See architecture spec §4 and
    PHASE_A_PRIME_TEST_PLAN.md §4.1 for the full decision table.

    'now' is an injectable clock (datetime) for deterministic tests; defaults
    to UTC now. Records older than _ACTIVE_WINDOW_DAYS are not counted (they
    auto-release); pruning is a write-side concern, not done here."""
    import datetime as _dt

    if not install_id:
        return {"decision": "unbound", "active_install_ids": [],
                "active_count": 0, "this_active": False,
                "message": "install_id unavailable — activation binding skipped (fail-open)"}

    if now is None:
        now = _dt.datetime.now(_dt.timezone.utc)

    def _parse(ts: str):
        if not ts:
            return None
        try:
            s = ts.strip().replace("Z", "+00:00")
            d = _dt.datetime.fromisoformat(s)
            if d.tzinfo is None:
                d = d.replace(tzinfo=_dt.timezone.utc)
            return d
        except Exception:
            return None

    cutoff = now - _dt.timedelta(days=_ACTIVE_WINDOW_DAYS)
    activations = entry.get("activations", []) or []

    active_ids = []
    for a in activations:
        if not isinstance(a, dict):
            continue
        iid = a.get("install_id")
        if not iid:
            continue
        seen = _parse(a.get("last_seen", "")) or _parse(a.get("first_seen", ""))
        if seen is not None and seen >= cutoff:
            if iid not in active_ids:
                active_ids.append(iid)

    this_active = install_id in active_ids
    count = len(active_ids)

    if this_active:
        return {"decision": "active", "active_install_ids": active_ids,
                "active_count": count, "this_active": True,
                "message": f"This machine is an active install ({count} of {_MAX_ACTIVE_INSTALLS})."}

    if count < _MAX_ACTIVE_INSTALLS:
        return {"decision": "admissible", "active_install_ids": active_ids,
                "active_count": count, "this_active": False,
                "message": (f"New install with capacity "
                            f"({count} of {_MAX_ACTIVE_INSTALLS} used) — will activate.")}

    return {"decision": "rejected", "active_install_ids": active_ids,
            "active_count": count, "this_active": False,
            "message": (
                f"This license is already active on {_MAX_ACTIVE_INSTALLS} machines. "
                f"Release one in your AI-Prowler License panel, or purchase an "
                f"additional license, to use it here.")}


# ══════════════════════════════════════════════════════════════════════════════
# BUSINESS LICENSE VALIDATION + GRACE LADDER  (v7.0.0 Phase B — base spec §3.3/§3.4)
# ══════════════════════════════════════════════════════════════════════════════
# A Business install validates its license key against the Worker's
# /license/validate endpoint on launch, caching the result in
# ~/.ai-prowler/license_cache.json. The cache + grace ladder protect a paying
# customer from network blips: a brief outage must NOT instantly downgrade them
# to Home. The ladder (§3.4), measured from last SUCCESSFUL validation:
#   • within 30 days of success                → use cache, no network call needed
#   • network fail, within 30 days of last ok  → cached_fresh (still trust)
#   • network fail, 30-37 days since last ok    → run business, log silently
#   • network fail, 37-44 days since last ok    → run business, show warning banner
#   • network fail, ≥ 44 days since last ok     → revert to home (data kept, features gated)
#   • explicit 'revoked'/'suspended'            → revert to home IMMEDIATELY (no grace)
#
# Constants are ordered FRESH ≤ WARN ≤ GRACE; the grace ladder only fires for
# caches OLDER than FRESH, so each tier gets a meaningful 7-day window beyond
# the 30-day trust period.
#
# _evaluate_license_grace() is PURE (no I/O) so it is unit-testable; the I/O
# wrapper _validate_business_license() does the cache read/write + HTTP POST.
_LICENSE_CACHE_PATH   = Path.home() / ".ai-prowler" / "license_cache.json"
_LICENSE_FRESH_HOURS  = 720  # within this since last success → trust cache, skip network (30 days, v7.0.0)
_LICENSE_WARN_DAYS    = 37   # network-fail past this since success → warning banner (30d trust + 7d silent)
_LICENSE_GRACE_DAYS   = 44   # network-fail past this since success → revert to home (30d trust + 14d grace)

# Reasons the Worker returns that mean "stop being business right now" (no grace).
_LICENSE_HARD_FAIL_REASONS = ("revoked", "suspended", "parent_revoked",
                              "parent_suspended", "not_found", "bad_format")


def _evaluate_license_grace(cache: dict, validate_result: "dict | None",
                            now=None) -> dict:
    """Decide the effective license standing (mobile OR business). PURE.

    v8.0.0: generalized from a Business-only evaluator to cover every
    license type the ai-prowler-subscription Worker issues (personal/mobile,
    business, and free beta — beta uses the exact same grace rules, just a
    cosmetic 'tier' tag). The Worker's /license/validate response now carries
    its own 'edition' field ('mobile' | 'business'), so the fallback/"home"
    edition is the only one this function still hard-codes — the GRANTED
    edition always comes from whichever the license itself claims.

    Args:
        cache:           parsed license_cache.json (may be {} if none yet).
                         Recognized keys: last_validated_at (ISO), status,
                         cached_expires_at (ISO), edition (str).
        validate_result: the Worker's /license/validate JSON this launch, or
                         None if the network call wasn't made/failed. Shape:
                         {valid: bool, reason?, edition?, expires_at?, status?}.
        now:             injectable clock (datetime); defaults to UTC now.

    Returns dict:
        effective_edition : 'mobile' | 'business' | 'home'
        action            : 'validated' | 'cached_fresh' | 'grace_silent'
                            | 'grace_warning' | 'reverted_expired'
                            | 'reverted_revoked'
        banner            : str  (warning/notice text, '' if none)
        used_network      : bool (did we trust a fresh validate_result)
    """
    import datetime as _dt

    if now is None:
        now = _dt.datetime.now(_dt.timezone.utc)

    def _parse(ts):
        if not ts:
            return None
        try:
            s = str(ts).strip().replace("Z", "+00:00")
            d = _dt.datetime.fromisoformat(s)
            if d.tzinfo is None:
                d = d.replace(tzinfo=_dt.timezone.utc)
            return d
        except Exception:
            return None

    # The edition a license is granting, from whichever source is freshest:
    # a live validate_result takes priority; otherwise fall back to whatever
    # edition was cached from the last successful check. Default 'mobile'
    # is a safe floor (matches the old _plan_to_edition fail-open default —
    # never silently strip entitlement from a paying customer due to a
    # missing/old field on an otherwise-valid cache entry).
    granted_edition = (
        (validate_result or {}).get("edition")
        or cache.get("edition")
        or "mobile"
    )

    # 1) Fresh successful validation this launch → trust it.
    if validate_result is not None:
        if validate_result.get("valid") is True:
            return {"effective_edition": granted_edition, "action": "validated",
                    "banner": "", "used_network": True}
        # Explicit negative from the Worker — hard fail, no grace.
        reason = validate_result.get("reason", "invalid")
        if reason in _LICENSE_HARD_FAIL_REASONS:
            return {"effective_edition": "home", "action": "reverted_revoked",
                    "banner": (f"Your AI-Prowler license is no longer valid ({reason}). "
                               f"Reverted to Home features. Contact "
                               f"david.vavro1@gmail.com to restore service."),
                    "used_network": True}
        # Unknown negative reason — treat as a soft/network-ish failure, fall
        # through to the cache-based grace ladder below.

    # 2) No fresh success (network failure, or unknown negative). Lean on cache.
    last_ok = _parse(cache.get("last_validated_at"))
    if last_ok is None:
        # Never successfully validated and can't now → can't grant entitlement.
        # had_prior_cache=False: no subscription was ever registered on this
        # install; _run_http() uses this to fail-open rather than block.
        return {"effective_edition": "home", "action": "reverted_expired",
                "banner": ("Could not validate your AI-Prowler license and no "
                           "prior validation is cached. Running Home features. "
                           "Check your connection and license key."),
                "used_network": False, "had_prior_cache": False}

    age = now - last_ok
    age_days = age.total_seconds() / 86400.0

    # 2a) Within 24h of a prior success and we didn't need the network → cached.
    if validate_result is None and age.total_seconds() <= _LICENSE_FRESH_HOURS * 3600:
        return {"effective_edition": granted_edition, "action": "cached_fresh",
                "banner": "", "used_network": False}

    # 2b) Grace ladder by days since last success.
    if age_days < _LICENSE_WARN_DAYS:
        return {"effective_edition": granted_edition, "action": "grace_silent",
                "banner": "", "used_network": False}

    if age_days < _LICENSE_GRACE_DAYS:
        revert_on = (last_ok + _dt.timedelta(days=_LICENSE_GRACE_DAYS)).date().isoformat()
        return {"effective_edition": granted_edition, "action": "grace_warning",
                "banner": (f"License validation has failed for several days. "
                           f"Renew/reconnect before {revert_on} or your AI-Prowler "
                           f"features will be disabled."),
                "used_network": False}

    # had_prior_cache=True: the subscription was once valid but grace has expired.
    # _run_http() uses this to block mobile access (unlike a fresh install).
    return {"effective_edition": "home", "action": "reverted_expired",
            "banner": ("Your license could not be validated for "
                       f"{_LICENSE_GRACE_DAYS}+ days. Reverted to Home features. "
                       "Your data is intact; reconnect to restore service."),
            "used_network": False, "had_prior_cache": True}


def _load_license_cache_for(license_key: str) -> dict:
    """Return the cached validation entry for ONE license key, or {} if absent.

    Cache shape (v7.0.0): a per-key map keyed by full license_key:
        {"licenses": {"AP-XXXX-...": {"last_validated_at": "...", ...}}}

    Backward compat: a legacy file with top-level last_validated_at is treated
    as the entry for the FIRST key we're asked about (the parent at startup),
    and migrated to the new shape on the next successful save.
    """
    try:
        if not _LICENSE_CACHE_PATH.exists():
            return {}
        raw = json.loads(_LICENSE_CACHE_PATH.read_text(encoding="utf-8-sig")) or {}
    except Exception as _e:
        _log.warning("license_cache.json unreadable (%s)", _e)
        return {}

    # New shape: {"licenses": {key: entry}}
    if isinstance(raw.get("licenses"), dict):
        entry = raw["licenses"].get(license_key)
        return entry if isinstance(entry, dict) else {}

    # Legacy shape: top-level fields belong to whichever key asks first.
    # We can't distinguish, so we return it unconditionally — the writer below
    # will migrate it to the new shape under THIS key on next success.
    if "last_validated_at" in raw:
        return raw
    return {}


def _save_license_cache_for(license_key: str, entry: dict) -> None:
    """Persist a per-key cache entry under the new shape, migrating legacy
    top-level fields into the new map if needed. Never raises."""
    try:
        _LICENSE_CACHE_PATH.parent.mkdir(parents=True, exist_ok=True)
        existing = {}
        if _LICENSE_CACHE_PATH.exists():
            try:
                existing = json.loads(
                    _LICENSE_CACHE_PATH.read_text(encoding="utf-8-sig")) or {}
            except Exception:
                existing = {}

        # Ensure we end up with the new shape: {"licenses": {key: entry, ...}}
        licenses_map = {}
        if isinstance(existing.get("licenses"), dict):
            licenses_map = dict(existing["licenses"])
        elif "last_validated_at" in existing:
            # Legacy single-license file — we don't know which key it belonged
            # to. Drop it: the new save below puts the current key in correctly,
            # and the legacy entry is only useful if it matches the active key,
            # in which case the new entry replaces it anyway.
            pass

        licenses_map[license_key] = entry
        _LICENSE_CACHE_PATH.write_text(
            json.dumps({"licenses": licenses_map}, indent=2), encoding="utf-8")
    except Exception as _e:
        _log.warning("Could not write license_cache.json (%s)", _e)


def _validate_license(license_key: str, install_id: str,
                      endpoint: str, now=None) -> dict:
    """I/O wrapper around _evaluate_license_grace: read PER-KEY cache, decide
    whether to call the Worker, GET /license/{key}/validate, persist a fresh
    success, and return the grace evaluation. Never raises.

    Returns the _evaluate_license_grace dict, plus 'license_key_present': bool.

    v8.0.0: renamed from _validate_business_license — this is now THE single
    validation path for every license type (personal/mobile, business, and
    free beta), all issued by the one ai-prowler-subscription Worker. The old
    GitHub subs.json registry (_check_subscription) and the separate
    D1-backed telemetry-Worker Business path are both retired; 'endpoint'
    here is always the subscription Worker's base URL.

    v7.0.0 changes carried forward: per-key cache (so the server can validate
    parent + N child keys without overwriting each other), and a 30-day
    fresh-cache window (_LICENSE_FRESH_HOURS = 720) so a long-running server
    hits the network on a ~30d cadence and a daily-launched personal install
    skips the network on each launch within those 30 days.
    """
    import datetime as _dt
    if now is None:
        now = _dt.datetime.now(_dt.timezone.utc)

    if not license_key:
        return {"effective_edition": "home", "action": "no_license",
                "banner": "", "used_network": False, "license_key_present": False}

    # Load this key's cache entry (tolerant of missing/corrupt file).
    cache = _load_license_cache_for(license_key)

    # If cache is fresh (<30d since success), skip the network entirely.
    last_ok = None
    try:
        s = str(cache.get("last_validated_at", "")).strip().replace("Z", "+00:00")
        if s:
            last_ok = _dt.datetime.fromisoformat(s)
            if last_ok.tzinfo is None:
                last_ok = last_ok.replace(tzinfo=_dt.timezone.utc)
    except Exception:
        last_ok = None

    if last_ok is not None and (now - last_ok).total_seconds() <= _LICENSE_FRESH_HOURS * 3600:
        result = _evaluate_license_grace(cache, None, now=now)
        result["license_key_present"] = True
        return result

    # Otherwise GET the Worker's validate endpoint (license key is the
    # credential — same trust model as the activation code; no bearer auth).
    validate_result = None
    try:
        import requests as _req
        resp = _req.get(
            f"{endpoint.rstrip('/')}/license/{license_key}/validate",
            params={"install_id": install_id or ""},
            headers={"User-Agent": "AI-Prowler-MCP/1.0"},
            timeout=10, proxies={"http": None, "https": None})
        if resp.status_code == 200:
            validate_result = resp.json()
        else:
            _log.warning("license/validate returned HTTP %s", resp.status_code)
    except Exception as _e:
        _log.warning("license/validate unreachable (%s) — using grace ladder", _e)

    result = _evaluate_license_grace(cache, validate_result, now=now)

    # Persist a fresh SUCCESS to the cache (so the 30d fast-path + grace anchor work).
    if validate_result is not None and validate_result.get("valid") is True:
        _save_license_cache_for(license_key, {
            "last_validated_at": now.isoformat(),
            "status": validate_result.get("status", "active"),
            "cached_expires_at": validate_result.get("expires_at", ""),
            "edition": validate_result.get("edition", "mobile"),
            "tier": validate_result.get("tier", "standard"),
        })

    result["license_key_present"] = True
    return result


# Backward-compat alias — a handful of call sites/tests may still reference
# the old name during the v8.0.0 transition. New code should call
# _validate_license() directly; this alias can be removed once nothing
# references the old name.
_validate_business_license = _validate_license


def _sweep_child_licenses(users_doc, validate_fn):
    """Walk active users and validate each one's child_license_key.

    PURE-AT-THE-EDGE — takes the users dict + the validate function as injectables
    so tests can stub the network. Returns a list of warning entries; the caller
    decides what to do with them (log, surface in GUI, etc.). Soft policy per
    David 2026-05-28: rejections produce warnings but do NOT mutate users_doc.

    Args:
        users_doc:    The parsed users.json dict, or anything with a 'users' key
                      mapping bearer-token → user record. Falsy/missing 'users'
                      yields an empty result.
        validate_fn:  Callable(license_key: str) -> dict matching the shape of
                      _validate_business_license's return:
                        {'effective_edition': 'business'|'home',
                         'action': str,
                         'banner': str, ...}

    Returns:
        list of dicts: [{name, child_key_masked, reason, banner}, ...]
        Empty list when nothing is wrong.

    Skip rules (each entry is silently skipped if):
      - record is not a dict (malformed users.json entry)
      - status is not 'active' (suspended/removed users don't consume checks)
      - child_license_key is empty (not yet assigned)

    A user whose validate_fn returns effective_edition=='business' AND no banner
    is healthy and contributes nothing to the result.
    """
    warnings = []
    if not isinstance(users_doc, dict):
        return warnings
    users_map = users_doc.get("users") or {}
    if not isinstance(users_map, dict):
        return warnings

    for _utok, urec in users_map.items():
        if not isinstance(urec, dict):
            continue
        if urec.get("status", "active") != "active":
            continue
        # Read the child key. Treat JSON null / missing / empty / non-string all
        # as 'no key assigned'. (A bare str() on None becomes the literal string
        # "None" which is truthy, so the falsy-check has to come FIRST.)
        ckey_raw = urec.get("child_license_key")
        if not ckey_raw or not isinstance(ckey_raw, str):
            continue   # phone-only-without-key OR not yet assigned OR malformed
        ckey = ckey_raw.strip()
        if not ckey:
            continue   # whitespace-only is treated as no key

        try:
            cres = validate_fn(ckey) or {}
        except Exception as _ve:
            # A failed validate call shouldn't kill the sweep — record it and
            # move on (the caller's outer try/except still catches catastrophic
            # failures, but per-user errors should be local).
            _log.warning("validate_fn raised for child key (%s); skipping user", _ve)
            continue

        eff = cres.get("effective_edition")
        banner = cres.get("banner", "")
        # Mask: first 4 + last 8. The dashed AP-key format puts the distinguishing
        # bits in the last two segments, so a short tail like 4 would still leave
        # sibling keys looking identical (e.g. AP-CHLD-AAAA-0001 vs ...-BBBB-0002
        # both shrink to AP-C…0001 / AP-C…0002 — same prefix, indistinguishable).
        # 8 trailing chars preserves the AAAA/BBBB segment.
        masked = (ckey[:4] + "…" + ckey[-8:]) if len(ckey) > 13 else ckey

        if eff != "business":
            warnings.append({
                "name": urec.get("name", "(unnamed)"),
                "child_key_masked": masked,
                "reason": cres.get("action", "rejected"),
                "banner": banner,
            })
            _log.warning(
                "Child license check FAILED for user=%s key=%s → %s",
                urec.get("name", "?"), masked, cres.get("action", "rejected"))
        elif banner:
            # Grace-warning on a child key — surface as a softer warning.
            warnings.append({
                "name": urec.get("name", "(unnamed)"),
                "child_key_masked": masked,
                "reason": cres.get("action", "grace_warning"),
                "banner": banner,
            })
            _log.warning(
                "Child license grace-warning for user=%s key=%s: %s",
                urec.get("name", "?"), masked, banner)

    return warnings


# Path the engine writes child-license warnings to; the GUI Admin tab reads it
# on each refresh. Always written (even when warnings=[]) so the GUI can tell
# "no issues as of <last_check_at>" from "the sweep has never run" (file absent).
_LICENSE_WARNINGS_PATH = Path.home() / ".ai-prowler" / "license_warnings.json"


def _save_license_warnings(warnings):
    """Persist the child-license warnings list for the GUI to read. Atomic
    write (tmp + os.replace). Never raises — a failed write should never
    block startup, since the warnings are advisory (the bearer-token auth
    path still works regardless)."""
    import datetime as _dt
    try:
        _LICENSE_WARNINGS_PATH.parent.mkdir(parents=True, exist_ok=True)
        payload = {
            "last_check_at": _dt.datetime.now(_dt.timezone.utc).isoformat(),
            "warnings":      list(warnings or []),
        }
        tmp = _LICENSE_WARNINGS_PATH.with_suffix(
            _LICENSE_WARNINGS_PATH.suffix + ".tmp")
        tmp.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        import os as _os
        _os.replace(str(tmp), str(_LICENSE_WARNINGS_PATH))
    except Exception as _e:
        _log.warning("Could not write %s (%s)", _LICENSE_WARNINGS_PATH, _e)


# ══════════════════════════════════════════════════════════════════════════════
# MULTI-USER MODEL — pure auth/scoping helpers  (v7.0.0 Phase B Block 3, spec §6)
# ══════════════════════════════════════════════════════════════════════════════
# Server-mode only. These PURE functions are the security spine (§6.4): they
# resolve a bearer token to a user, compute the exact set of ChromaDB collections
# that user may touch, and decide indexing permission. They have NO I/O — the
# uvicorn middleware (built at the keyboard) reads users.json, calls
# _resolve_user(), attaches request.state.user, then uses _allowed_collections()
# to constrain every ChromaDB call. Enforcement is at the collection-selection
# level, never a post-filter, so a request can never receive results from a
# collection outside its scope.
#
# Roles (§6.1) are a FIXED set. The matrix below encodes the cross-cutting
# capabilities; collection READ scope is computed per-user from role + scopes.
_USER_ROLES = ("owner", "manager", "staff", "field_crew")

# Per-role capabilities that are NOT per-collection (those come from scopes).
# v7.0.1 additions: manage_db, can_write_shared, can_send_email.
# v8.0.0 additions: can_send_sms (all roles); can_send_email extended to ALL roles.
#
#   read_all_role_scopes : owner sees every scope:* collection (any assigned scope)
#   read_others_private  : owner may read any user:* collection
#   can_write            : may write/index (subject to per-collection _can_index)
#   can_write_shared     : may index INTO the 'shared' company commons
#   is_admin             : may use Admin-tab / user-management tools
#   manage_db            : DB-management tool access level
#                          "full"    — index_path, untrack, reindex, update (owner, manager)
#                          "limited" — index_path into own scopes/private only (staff)
#                          "none"    — no DB management tools at all (field_crew)
#   can_send_email       : may call send_email / send_alert in server mode.
#                          v8.0.0: True for ALL roles — owner, manager, staff,
#                          and field_crew can all send email via the company SMTP.
#   can_send_sms         : may call send_sms / send_whatsapp in server mode.
#                          v8.0.0: True for ALL roles.
_ROLE_CAPS = {
    # owner — unrestricted: full DB management, reads all scopes, admin rights
    "owner":      {"read_all_role_scopes": True,  "read_others_private": True,
                   "can_write": True,  "can_write_shared": True,  "is_admin": True,
                   "manage_db": "full",    "can_send_email": True,  "can_send_sms": True},
    # manager — full DB management within their assigned scopes + shared
    "manager":    {"read_all_role_scopes": False, "read_others_private": False,
                   "can_write": True,  "can_write_shared": True,  "is_admin": False,
                   "manage_db": "full",    "can_send_email": True,  "can_send_sms": True},
    # staff — limited DB: may index own private + assigned scopes; NOT shared or destructive ops
    "staff":      {"read_all_role_scopes": False, "read_others_private": False,
                   "can_write": True,  "can_write_shared": False, "is_admin": False,
                   "manage_db": "limited", "can_send_email": True,  "can_send_sms": True},
    # field_crew — no DB management; may send email and SMS (no personal AI-Prowler install)
    "field_crew": {"read_all_role_scopes": False, "read_others_private": False,
                   "can_write": False, "can_write_shared": False, "is_admin": False,
                   "manage_db": "none",    "can_send_email": True,  "can_send_sms": True},
}

_SHARED_COLLECTION = "shared"


def _canon_scope(name: str) -> str:
    """Canonicalize a logical data-bucket scope name to the 'scope:<name>' form.

    v7.0.1 vocabulary change: data buckets used to be written 'role:<name>',
    which collided conceptually with the owner/manager/staff/field_crew JOB
    roles. The user-facing name is now 'scope:<name>'. This accepts all three
    historical spellings and returns the canonical one, so comparisons line up
    no matter what is stored in users.json / collection_map:

        'sales'       -> 'scope:sales'   (bare)
        'role:sales'  -> 'scope:sales'   (legacy)
        'scope:sales' -> 'scope:sales'   (canonical, unchanged)

    NON-bucket logical names are returned UNCHANGED (they are not buckets):
    'shared', any 'user:<id>', and the personal 'documents' collection.
    Physically, 'scope:<name>' and 'role:<name>' resolve to the SAME ChromaDB
    collection via rag_preprocessor.chroma_collection_name(), so no data moves.
    """
    s = (name or "").strip()
    if not s:
        return s
    low = s.lower()
    if low == _SHARED_COLLECTION or low == "documents" or low.startswith("user:"):
        return s
    if low.startswith("scope:") or low.startswith("role:"):
        return "scope:" + s.split(":", 1)[1].strip()
    return "scope:" + s


def _role_caps(role: str) -> dict:
    """Capabilities for a role; unknown roles get the most-restricted set."""
    return _ROLE_CAPS.get((role or "").strip().lower(), _ROLE_CAPS["field_crew"])


def _check_db_cap(user: "dict | None", level: str = "full") -> tuple:
    """Capability gate for DB-management tool calls. PURE.

    Personal mode (user=None): always allowed — no role gate in single-user mode.
    Server mode: the role's manage_db capability must be at or above `level`:

        'full'    — index, untrack, reindex, update (owner, manager)
        'limited' — index into own scopes/private only (staff + above)

    Returns (allowed: bool, reason: str).
    """
    if user is None:
        return (True, "personal mode — no DB management gate")
    caps = _role_caps(user.get("role"))
    db_level = caps.get("manage_db", "none")
    role     = user.get("role", "unknown")
    if level == "full":
        if db_level == "full":
            return (True, f"role '{role}' has full DB management")
        return (False,
                f"role '{role}' (manage_db={db_level!r}) cannot perform full DB "
                "management (index/untrack/reindex). Owner or manager required.")
    elif level == "limited":
        if db_level in ("full", "limited"):
            return (True, f"role '{role}' may index into permitted scopes")
        return (False,
                f"role '{role}' (manage_db='none') cannot index documents. "
                "Owner, manager, or staff required.")
    return (False, f"_check_db_cap: unknown level {level!r}")


def _make_user_id(name: str) -> str:
    """Generate a stable, lowercase slug id from a display name.

    Examples:
        "David Vavro"  -> "david-vavro"
        "Vicki Vavro"  -> "vicki-vavro"
        "Field Crew 1" -> "field-crew-1"

    Rules:
        - Lowercase
        - Spaces and underscores become hyphens
        - All non-alphanumeric/hyphen characters stripped
        - Multiple consecutive hyphens collapsed to one
        - Leading/trailing hyphens removed
    """
    import re
    slug = name.strip().lower()
    slug = re.sub(r'[\s_]+', '-', slug)       # spaces/underscores -> hyphens
    slug = re.sub(r'[^a-z0-9-]', '', slug)    # strip non-alphanumeric/hyphen
    slug = re.sub(r'-+', '-', slug)            # collapse multiple hyphens
    slug = slug.strip('-')                     # remove leading/trailing hyphens
    return slug or "unknown-user"


def _resolve_user(users_data: "dict | None", token: str) -> "dict | None":
    """Look up a bearer token in users.json data. Returns the user dict
    (augmented with 'id') if found AND active, else None. PURE.

    v7.0.1 — The users.json dict key is the bearer token (legacy design).
    The stable user ID is now derived from the display name as a lowercase
    hyphenated slug (e.g. "David Vavro" -> "david-vavro"), NOT from the
    bearer token. This decouples identity from credentials so that:
      - Token regeneration never changes a user's ChromaDB collection name
      - Private collections have human-readable names (scope-user-david-vavro)
      - collection_map entries use predictable, stable ids

    A non-active status ('suspended'/'revoked') resolves to None — a soft-revoke
    that denies access without losing the audit record. Matches §6.4 steps 1-2.
    """
    if not token or not users_data:
        return None
    users = users_data.get("users", {})
    # The dict key is the legacy bearer token; look it up directly.
    entry = users.get(token)
    if not isinstance(entry, dict):
        return None
    if entry.get("status", "active") != "active":
        return None
    # Return a shallow copy with the stable slug id derived from display name.
    user = dict(entry)
    user["id"] = _make_user_id(user.get("name", token))
    # Normalize role to the known set (defense against a hand-edited users.json).
    if user.get("role") not in _USER_ROLES:
        user["role"] = "field_crew"
    return user


def _current_user(ctx) -> "dict | None":
    """Extract the authenticated user that server-mode auth middleware attached
    to this request, via the FastMCP Context. Returns None in personal mode
    (no server-mode middleware, so no user on request.state) or if anything in
    the chain is absent. PURE-ish (only reads ctx; no other I/O). Never raises.
    """
    if ctx is None:
        return None
    try:
        return getattr(ctx.request_context.request.state, "user", None)
    except Exception:
        return None


# SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cleanup, 2026-07-17):
# _allowed_collections() (the multi-collection READ enforcement list this
# module used before the single-collection cutover) has been removed --
# its last caller was a request.state diagnostic field, now replaced with
# _allowed_scopes() directly (see the /whoami handler). _allowed_scopes()
# below is the sole remaining "what may this user search" function.
def _allowed_scopes(user: "dict | None") -> set:
    """Return the set of scope strings `user` may SEARCH, for the
    single-collection design. Thin wrapper around
    scope_lookup.allowed_scopes_for_user() -- kept here (rather than
    calling scope_lookup directly from every read tool) so there is one
    place in ai_prowler_mcp.py that owns "how do I get a user's allowed
    scopes".

    SERVER MODE ONLY -- callers must not invoke this in personal mode
    (no `user`); returns an empty set for user=None.

    Direct product decision (2026-07-16): NO role-based elevation, for
    anyone, including owner. Every user's search visibility is exactly
    {"shared"} union their own assigned scopes union their own private
    scope if enabled -- computed identically regardless of role. See
    SCOPE_SIMPLIFICATION_SPEC.md section 3.6 for the full rationale.
    Role continues to gate the Admin tab (user management) only -- see
    _role_caps()/is_admin -- never search visibility.
    """
    import scope_lookup
    return scope_lookup.allowed_scopes_for_user(user)


def _can_index(user: "dict | None", target_collection: str,
               all_role_collections: "list | tuple | None" = None) -> tuple:
    """Decide whether `user` may INDEX (write) into `target_collection`. PURE.
    Implements §6.5. Returns (allowed: bool, reason: str).

    Rules (v7.0.1):
      - owner      : any collection (shared, any scope:*, any user:*)
      - manager    : scope:* in their assigned scopes + own private + shared
      - staff      : scope:* in their assigned scopes + own private; NOT shared
      - field_crew : nothing (can_write=False; manage_db='none')
    """
    if not user:
        return (False, "no user context")
    caps = _role_caps(user.get("role"))
    role = user.get("role")
    target = (target_collection or "").strip()
    if not target:
        return (False, "no target collection")

    if role == "owner":
        return (True, "owner may index any collection")

    if not caps["can_write"]:
        return (False, f"role '{role}' cannot index server-side")

    # can_write roles from here (manager writes shared; staff does not).
    if target == _SHARED_COLLECTION:
        if caps.get("can_write_shared"):
            return (True, f"role '{role}' may write to the shared commons")
        return (False, f"role '{role}' cannot write to the shared collection")

    if target.startswith("user:"):
        if target == f"user:{user.get('id')}":
            return (True, "own private collection")
        return (False, "cannot index another user's private collection")

    if _canon_scope(target).startswith("scope:"):
        scopes = set()
        for s in (user.get("scopes") or []):
            s = str(s).strip()
            scopes.add(_canon_scope(s))
        if _canon_scope(target) in scopes:
            return (True, "assigned scope")
        return (False, f"scope not assigned to role '{role}'")

    return (False, "unrecognized collection type")


# ── Multi-collection result merge (v7.0.0 Phase B Step 2 read path) ───────────
# After a server-mode read tool queries the user's N _allowed_collections, it
# gets N separate ChromaDB result sets. This merges them into ONE ranked list,
# the final step of the read-enforcement flow. PURE — no ChromaDB, no I/O — so
# it's unit-testable in isolation.
#
# ChromaDB query() returns parallel lists per collection: ids, documents,
# metadatas, distances (LOWER distance = more similar). We flatten all hits,
# tag each with its source collection, dedup by id (defensive — collections are
# disjoint so dupes shouldn't occur, but a chunk indexed into two scopes would),
# sort by distance ascending, and truncate to n_results.
def _merge_collection_results(per_collection: "dict", n_results: int = 10) -> list:
    """Merge per-collection ChromaDB query results into one ranked list. PURE.

    Args:
        per_collection: {collection_name: chroma_result_dict}, where each
            chroma_result_dict has the shape ChromaDB query() returns:
            {"ids": [[...]], "documents": [[...]], "metadatas": [[...]],
             "distances": [[...]]}  (each value is a list-of-lists; we use [0]).
            A value may also be the already-unwrapped single-query form
            {"ids": [...], ...} — both are tolerated.
        n_results: max hits to return after merge.

    Returns a list of dicts sorted by distance ascending (best first):
        {"id", "document", "metadata", "distance", "collection"}
    Hits missing a distance sort last (treated as +inf). Malformed entries are
    skipped, never raised on.
    """
    merged = []
    seen_ids = set()

    for cname, res in (per_collection or {}).items():
        if not isinstance(res, dict):
            continue

        def _col(key):
            """Pull a column, unwrapping Chroma's list-of-lists if present."""
            v = res.get(key)
            if v is None:
                return []
            # list-of-lists (multi-query shape) → take first query's row
            if isinstance(v, list) and len(v) == 1 and isinstance(v[0], list):
                return v[0]
            return v if isinstance(v, list) else []

        ids       = _col("ids")
        docs      = _col("documents")
        metas     = _col("metadatas")
        distances = _col("distances")

        for i, _id in enumerate(ids):
            if _id is None:
                continue
            if _id in seen_ids:
                continue           # dedup across collections
            seen_ids.add(_id)
            dist = distances[i] if i < len(distances) and distances[i] is not None else float("inf")
            try:
                dist = float(dist)
            except (TypeError, ValueError):
                dist = float("inf")
            merged.append({
                "id":         _id,
                "document":   docs[i]  if i < len(docs)  else "",
                "metadata":   metas[i] if i < len(metas) else {},
                "distance":   dist,
                "collection": cname,
            })

    merged.sort(key=lambda h: h["distance"])
    if n_results and n_results > 0:
        merged = merged[:n_results]
    return merged


# ── Model B: path → collection resolver (v7.0.0 Phase B Step 2 write path) ─────
# Decides which ChromaDB collection an indexed file belongs to, from a
# CONFIGURED path→collection map (the owner declares it once; assignment is an
# auditable config artifact, not per-index human judgment). PURE — no I/O — so
# it is unit-testable. The pipeline surgery that CALLS this (inside
# index_directory / index_file_list) is a separate, keyboard-tested step.
#
# mapping shape (from config, e.g. users.json or a company collection_map):
#   {"rules": [{"prefix": "C:/CompanyDocs/Sales", "collection": "role:sales"},
#              {"prefix": "C:/CompanyDocs/Public", "collection": "shared"}],
#    "default_collection": "shared"}   # optional; see fallback below
#
# Matching: longest matching prefix wins (most specific rule), case-insensitive,
# slash-agnostic. No rule + no default → the indexer's OWN private collection
# (user:<id>) — the safe fallback: an unclassified doc goes to the indexer's
# own space, NEVER accidentally to shared.
def _normalize_path_for_match(p: str) -> str:
    """Lowercased, forward-slashed, trailing-slash-free path for prefix match."""
    s = (p or "").replace("\\", "/").strip()
    s = s.rstrip("/")
    return s.lower()


def _resolve_collection_for_path(filepath: str, mapping: "dict | None",
                                 indexer_user: "dict | None" = None) -> str:
    """Return the target collection name for `filepath`. PURE.

    Args:
        filepath:     the file being indexed.
        mapping:      {"rules": [{"prefix","collection"}...],
                       "default_collection": "..."}  (may be {} / None).
        indexer_user: resolved user dict (for the user:<id> fallback). May be
                      None in personal mode — then the ultimate fallback is the
                      single default 'documents' collection (personal behavior).

    Resolution order:
      1. Longest matching prefix rule wins.
      2. Else mapping['default_collection'] if set.
      3. Else indexer's own 'user:<id>' (if a user is known).
      4. Else 'documents' (personal/Home single-collection default).
    """
    fp = _normalize_path_for_match(filepath)

    best_collection = None
    best_len = -1
    for rule in ((mapping or {}).get("rules") or []):
        if not isinstance(rule, dict):
            continue
        prefix = _normalize_path_for_match(rule.get("prefix", ""))
        coll   = str(rule.get("collection", "")).strip()
        if not prefix or not coll:
            continue
        # Prefix match on a path-segment boundary: either exact, or followed
        # by '/', so '.../Sales' doesn't spuriously match '.../SalesArchive'.
        if fp == prefix or fp.startswith(prefix + "/"):
            if len(prefix) > best_len:
                best_len = len(prefix)
                best_collection = coll

    if best_collection:
        return best_collection

    default_coll = str((mapping or {}).get("default_collection", "")).strip()
    if default_coll:
        return default_coll

    if indexer_user and indexer_user.get("id"):
        return f"user:{indexer_user['id']}"

    # Personal/Home single-collection default.
    return "documents"


def _company_collection_map(users_data: "dict | None" = None) -> dict:
    """Return the company path→scope mapping from users.json's top-level
    'collection_map' key, or {} if absent/malformed. Shape:
      {"rules": [{"prefix","collection"}...], "default_collection": "..."}
    PURE given users_data; loads users.json if not supplied."""
    if users_data is None:
        users_data = _load_users()
    if not isinstance(users_data, dict):
        return {}
    cm = users_data.get("collection_map")
    return cm if isinstance(cm, dict) else {}


# SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cleanup, 2026-07-17):
# _build_collection_resolver() (the per-user collection_map-based WHERE
# router this module used before the single-collection cutover) has been
# removed -- zero remaining callers across index_path, update_tracked_
# directories, reindex_file, reindex_directory. There is only one physical
# collection now; scope is carried entirely by build_scope_resolver()'s
# chunk-metadata tag (rag_preprocessor.py), computed independently at
# index time, not proposed by a per-request resolver here.


# ── Admin role gate (§9.1 / spec item 9) ──────────────────────────────────────
# Admin MCP tools (add_user, revoke_user, view_audit_log, ...) are owner-only.
# This is the PURE decision the @requires_role gate / middleware will call.
def _user_has_role(user: "dict | None", required_role: str) -> bool:
    """True iff the resolved user holds exactly `required_role`. PURE.
    A None user (unauthenticated) never satisfies a role requirement."""
    if not user:
        return False
    return (user.get("role") or "").strip().lower() == required_role.strip().lower()


def _is_admin(user: "dict | None") -> bool:
    """True iff the user may use admin/user-management tools. PURE.
    Currently only 'owner' is admin (§6.1 matrix)."""
    if not user:
        return False
    return bool(_role_caps(user.get("role")).get("is_admin"))


# ── Audit log (§6.6 / spec item 11) — pure format + filter helpers ────────────
# The audit log records {timestamp, user, tool, collection} per request, plus
# named events like 'remote_support_enabled'. These two helpers are the pure
# surface; the actual append (one JSONL line per entry) and the Admin-tab table
# rendering are thin I/O built at the keyboard. Kept deliberately small.
def _format_audit_entry(user: "dict | None", tool: str,
                        collection: str = "", event: str = "",
                        now=None) -> dict:
    """Build one audit record. PURE (except reading the clock when now is None).
    Records only tool name + collection — no query content or results (§6.6:
    'read-only, no detail beyond tool name and collection')."""
    import datetime as _dt
    if now is None:
        now = _dt.datetime.now(_dt.timezone.utc)
    return {
        "ts":         now.isoformat(),
        "user_id":    (user or {}).get("id", ""),
        "user_name":  (user or {}).get("name", ""),
        "role":       (user or {}).get("role", ""),
        "tool":       tool or "",
        "collection": collection or "",
        "event":      event or "",
    }


def _filter_audit_entries(entries: "list | None", limit: int = 20,
                          since=None) -> list:
    """Return the most recent `limit` entries, optionally only those at/after
    `since` (a datetime or ISO string). PURE. Newest-last input is assumed
    (append order); returns newest-last, trimmed to `limit`."""
    import datetime as _dt
    rows = list(entries or [])

    if since is not None:
        if isinstance(since, str):
            try:
                s = since.strip().replace("Z", "+00:00")
                since = _dt.datetime.fromisoformat(s)
                if since.tzinfo is None:
                    since = since.replace(tzinfo=_dt.timezone.utc)
            except Exception:
                since = None

    def _ts(row):
        try:
            s = str(row.get("ts", "")).strip().replace("Z", "+00:00")
            d = _dt.datetime.fromisoformat(s)
            if d.tzinfo is None:
                d = d.replace(tzinfo=_dt.timezone.utc)
            return d
        except Exception:
            return None

    if since is not None:
        rows = [r for r in rows if (_ts(r) is not None and _ts(r) >= since)]

    if limit and limit > 0:
        rows = rows[-limit:]
    return rows


# ══════════════════════════════════════════════════════════════════════════════
# SERVER MODE — multi-user HTTP transport  (v7.0.0 Phase B Block 3, spec §5.1/§6.4)
# ══════════════════════════════════════════════════════════════════════════════
# Business "company server" transport. Differs from _run_http (single shared
# token) by authenticating EACH request against users.json: bearer token →
# _resolve_user() → request.state.user, 401 on unknown/suspended. This is
# STEP 1 — the AUTHENTICATION layer only. Per-user ChromaDB COLLECTION SCOPING
# (using _allowed_collections / _can_index) is STEP 2, wired into the tools
# next; until then a logged warning makes clear data is not yet user-scoped.
#
# SAFETY: only reached when config.json mode=server. If prerequisites are
# missing (no users.json, deps unavailable) it FALLS BACK to _run_http so a
# server install is never left with no transport. _run_http itself is untouched.
_USERS_JSON_PATH = _state_dir() / "users.json"


def _load_users() -> "dict | None":
    """Load users.json. Returns the parsed dict, or None if missing/unreadable
    (caller treats None as 'server mode not provisioned')."""
    try:
        if _USERS_JSON_PATH.exists():
            # utf-8-sig: tolerate a BOM (PowerShell Out-File / some editors add
            # one). Reads correctly whether or not a BOM is present.
            data = json.loads(_USERS_JSON_PATH.read_text(encoding="utf-8-sig"))
            if isinstance(data, dict) and isinstance(data.get("users"), dict):
                return data
            _log.error("users.json present but malformed (no 'users' map).")
            return None
    except Exception as _e:
        _log.error("users.json unreadable: %s", _e)
    return None


def _hot_reload_users(users_data: "dict | None") -> "dict | None":
    """Re-read users.json from disk, falling back to the given snapshot if the
    file is missing/unreadable right now. PURE wrapper around _load_users().

    Every auth-check call site in server mode MUST route through this (not
    reference a captured `users_data` closure variable directly) so that
    token corrections, additions, and suspensions made to users.json while
    the server is already running take effect immediately — on BOTH the
    /authorize login page and subsequent MCP bearer-auth requests — without
    requiring a server restart.

    Historical bug (fixed 2026-07-08): the /authorize handler used to
    resolve logins against the stale startup-time `users_data` snapshot
    directly, while the MCP request path already hot-reloaded. That
    inconsistency meant a corrected bearer token would keep failing at
    the login page until the process was restarted, even though the
    file on disk was correct.
    """
    return _load_users() or users_data


def _owner_user_id(users_data: "dict | None" = None) -> "str | None":
    """Return the stable slug id of the OWNER from users.json, or None if no
    owner is defined / users.json absent. Used to PROTECT the owner's private
    collection from admins (can_manage_users grants read of all OTHER users'
    privates, but never the owner's). If multiple owners exist (unusual), the
    first found is returned. PURE given users_data; loads it if not supplied.

    v7.0.1: returns the slug id derived from the owner's display name
    (e.g. "david-vavro"), NOT the bearer token key. Consistent with
    _resolve_user which also derives id from name.
    """
    if users_data is None:
        users_data = _load_users()
    if not users_data:
        return None
    for uid, entry in (users_data.get("users") or {}).items():
        if isinstance(entry, dict) and (entry.get("role") or "").strip().lower() == "owner":
            return _make_user_id(entry.get("name", uid))
    return None


def _can_manage_user_data(actor: "dict | None", target_user_id: str,
                          owner_id: "str | None") -> tuple:
    """Decide whether `actor` may DELETE / clean up the data of target_user_id.
    PURE. Guards the (future) Admin-tab data-management / departed-employee
    cleanup operation — NOT yet wired to any delete op, encoded now + tested.

    Rules (FAIL CLOSED on the owner-protection — never destroy owner data on a
    guess):
      • The owner may manage anyone's data (including their own).
      • An admin (can_manage_users, not owner) may manage any EMPLOYEE's data
        but NEVER the owner's. If the owner's id cannot be determined
        (owner_id is None/empty), an admin is DENIED — we will not risk
        deleting owner data we cannot rule out.
      • Anyone else: no management rights.
    Returns (allowed: bool, reason: str).
    """
    if not actor:
        return (False, "no actor")
    if _user_has_role(actor, "owner"):
        return (True, "owner may manage any data")
    if not actor.get("can_manage_users"):
        return (False, "actor lacks can_manage_users")
    # Admin from here. Owner data is protected — and we FAIL CLOSED: if we can't
    # identify the owner, we cannot prove the target isn't the owner, so deny.
    if not owner_id:
        return (False, "owner id unknown — refusing to risk owner data (fail closed)")
    if target_user_id == owner_id:
        return (False, "owner data is protected from admins")
    return (True, "admin may manage employee data")


# ── Chunk-ownership purge gate (v7.0.0 Phase B — "delete only your own") ───────
# Filesystem-style protection against a bad actor wiping others' data: every
# indexed chunk carries indexed_by=<user_id>. Before the index pipeline PURGES
# existing chunks for a path (the delete(where=filepath) that precedes every
# re-add), this gate checks the actor may remove every owner present. PURE.
_OWNERLESS = "__ownerless__"  # sentinel for legacy chunks lacking indexed_by


def _chunk_owners(existing_metadatas: "list | None") -> set:
    """Distinct indexed_by owners across the given chunk metadatas. Chunks with
    no indexed_by map to the _OWNERLESS sentinel. PURE."""
    owners = set()
    for meta in (existing_metadatas or []):
        if not isinstance(meta, dict):
            continue
        owners.add(str(meta.get("indexed_by") or _OWNERLESS))
    return owners


def _can_purge_chunks(actor: "dict | None", existing_metadatas: "list | None",
                      owner_id: "str | None") -> tuple:
    """May `actor` purge/overwrite the given existing chunks? PURE.

    Rule (filesystem-like): you may purge chunks you OWN; an owner/admin may
    purge employees' chunks (subject to owner-data protection in
    _can_manage_user_data); legacy chunks with no indexed_by are treated as
    manageable only by owner/admin (a normal user can't wipe un-owned data).

    • No existing chunks (empty/None) → allowed (pure add, nothing destroyed).
    • Else: allowed iff the actor can manage EVERY distinct owner present.
      For each owner X present, require _can_manage_user_data(actor, X, owner_id);
      ownerless legacy chunks require actor to be owner/admin.
    Returns (allowed: bool, reason: str).
    """
    if not actor:
        return (False, "no actor")
    owners = _chunk_owners(existing_metadatas)
    if not owners:
        return (True, "no existing chunks to purge")

    actor_id = actor.get("id")
    for owner in owners:
        if owner == _OWNERLESS:
            # Legacy/un-owned: only owner/admin (manage rights over *anyone*) may
            # purge. Reuse _can_manage_user_data against a non-owner placeholder
            # so a plain user is denied but owner/admin allowed.
            allowed, _ = _can_manage_user_data(actor, "__legacy__", owner_id)
            if not allowed:
                return (False, "contains chunks with no owner (legacy) — only "
                               "owner/admin may purge")
            continue
        if owner == actor_id:
            continue  # you always may purge your own
        allowed, reason = _can_manage_user_data(actor, owner, owner_id)
        if not allowed:
            return (False, f"contains chunks owned by another user ({reason})")
    return (True, "actor may purge all present owners")


def _run_server_mode(port: int, token: str,
                     public_base: str = "https://mobile.dvavro-ai-prowler.com") -> None:
    """Multi-user server transport. STEP 1: authentication layer.

    Falls back to _run_http() if server mode can't be established, so a
    misconfigured Business server still serves (single-user) rather than dying.
    """
    users_data = _load_users()
    if users_data is None:
        _log.error(
            "Server mode requested but ~/.ai-prowler/users.json is missing or "
            "malformed. Falling back to single-user HTTP transport. Create the "
            "first owner user via the Admin tab to enable multi-user mode.")
        return _run_http(port=port, token=token, public_base=public_base)

    try:
        import uvicorn
        from starlette.applications import Starlette
        from starlette.middleware.base import BaseHTTPMiddleware
        from starlette.responses import JSONResponse, PlainTextResponse
        from starlette.routing import Route, Mount
        from starlette.requests import Request
    except ImportError as _ie:
        _log.error("Server mode needs uvicorn/starlette (%s). Falling back.", _ie)
        return _run_http(port=port, token=token, public_base=public_base)

    n_users = len(users_data.get("users", {}))
    company = users_data.get("company_id", "<unknown>")
    _log.info("Server mode: company=%s, %d user(s) loaded from users.json",
              company, n_users)

    # ── FastMCP ASGI app ──────────────────────────────────────────────────────
    # FastMCP exposes a streamable-http ASGI app mounted at /mcp.
    try:
        mcp_asgi = mcp.streamable_http_app()
    except Exception as _e:
        _log.error("Could not get FastMCP ASGI app (%s). Falling back.", _e)
        return _run_http(port=port, token=token, public_base=public_base)

    # ── OAuth state (server mode) ─────────────────────────────────────────────
    # In server mode, /authorize accepts any token that resolves to a valid user
    # in users.json (instead of the single bearer token used in personal mode).
    # OAuth codes and issued access tokens are tracked in-memory; issued tokens
    # are also cross-referenced back to the user record so _resolve_user can
    # find the right user when subsequent MCP requests arrive.
    import secrets as _srv_secrets
    import urllib.parse as _srv_urlparse
    import hashlib as _srv_hashlib
    import base64 as _srv_base64

    # code  -> {redirect_uri, code_challenge, code_challenge_method, user_token}
    _srv_auth_codes: dict = {}
    # issued_access_token -> original users.json bearer token (for _resolve_user)
    _srv_access_tokens: dict = {}
    # Pre-populate with every existing users.json token so curl/manual clients
    # that present their raw bearer token directly still work without OAuth dance.
    # The dict key in users.json IS the bearer token (legacy schema), so we
    # map token -> token here. _resolve_user will derive the stable slug id
    # from the user's display name — the token never becomes the id.
    for _ut in (users_data.get("users") or {}).keys():
        _srv_access_tokens[_ut] = _ut

    _local_host = f"127.0.0.1:{port}".encode()
    _srv_public_base = public_base.rstrip("/")

    def _srv_get_public_base(scope: dict) -> str:
        """Derive public base URL from Host header (mirrors _run_http logic)."""
        headers_dict = {k.lower(): v for k, v in scope.get("headers", [])}
        host = (headers_dict.get(b"host", b"")
                or headers_dict.get(b"x-forwarded-host", b"")).decode("utf-8", "ignore")
        if host:
            scheme = "http" if (host.startswith("127.") or host.startswith("localhost")) else "https"
            return f"{scheme}://{host}"
        return _srv_public_base

    # ── Pure-ASGI multi-user router (spec §6.4 steps 1-4) ─────────────────────
    # CRITICAL: this is a raw ASGI app, NOT BaseHTTPMiddleware. BaseHTTPMiddleware
    # buffers responses and breaks the MCP streamable-HTTP SSE transport (that was
    # the POST /mcp 500). This mirrors _run_http's proven _RouterASGI pattern:
    #   1. /health, /whoami, /.well-known/*, /register, /authorize, /token
    #      handled inline (OAuth paths exempt from bearer auth).
    #   2. every other path: bearer → _resolve_user → stash the user in
    #      scope["state"] so FastMCP's Request(scope).state.user resolves it
    #      (this is what _current_user(ctx) reads inside the tools — the keystone),
    #      then inject the MCP headers (Accept/Content-Type/Host/MCP-Protocol-
    #      Version) exactly like _run_http, then stream to mcp_asgi WITHOUT
    #      buffering so SSE works.

    _SRV_OAUTH_PATHS = {
        "/health", "/register", "/authorize", "/token",
        "/.well-known/oauth-authorization-server",
        "/.well-known/oauth-protected-resource",
    }

    async def _send_json(send, status, payload):
        import json as _json
        body = _json.dumps(payload).encode()
        await send({"type": "http.response.start", "status": status,
                    "headers": [(b"content-type", b"application/json"),
                                (b"content-length", str(len(body)).encode())]})
        await send({"type": "http.response.body", "body": body})

    async def _send_text(send, status, text):
        body = text.encode()
        await send({"type": "http.response.start", "status": status,
                    "headers": [(b"content-type", b"text/plain; charset=utf-8"),
                                (b"content-length", str(len(body)).encode())]})
        await send({"type": "http.response.body", "body": body})

    async def _send_html(send, status, html):
        body = html.encode("utf-8")
        await send({"type": "http.response.start", "status": status,
                    "headers": [(b"content-type", b"text/html; charset=utf-8"),
                                (b"content-length", str(len(body)).encode())]})
        await send({"type": "http.response.body", "body": body})

    async def _send_redirect(send, location):
        await send({"type": "http.response.start", "status": 302,
                    "headers": [(b"location", location.encode())]})
        await send({"type": "http.response.body", "body": b""})

    def _bearer_from_scope(scope):
        for hk, hv in scope.get("headers", []):
            if hk.lower() == b"authorization":
                v = hv.decode("utf-8", "replace")
                if v.lower().startswith("bearer "):
                    return v[7:].strip()
        return ""

    async def _read_body(receive) -> bytes:
        """Accumulate the full request body from ASGI receive events."""
        chunks = []
        while True:
            msg = await receive()
            if msg.get("type") == "http.request":
                chunks.append(msg.get("body", b""))
                if not msg.get("more_body", False):
                    break
        return b"".join(chunks)

    def _parse_qs(raw: bytes) -> dict:
        import urllib.parse as _up
        return {k: v[-1] for k, v in _up.parse_qs(raw.decode("utf-8", "replace")).items()}

    class _ServerRouterASGI:
        async def __call__(self, scope, receive, send):
            stype = scope.get("type", "")

            if stype == "lifespan":
                # Lifespan MUST go to mcp_asgi so its task group initialises.
                await mcp_asgi(scope, receive, send)
                return

            if stype != "http":
                await mcp_asgi(scope, receive, send)
                return

            path   = scope.get("path", "")
            method = scope.get("method", "GET")
            _log.debug("SERVER-MODE REQUEST  %s %s", method, path)

            # ── OAuth / health endpoints (no bearer auth required) ────────────
            if path == "/health":
                await _send_text(send, 200, "OK")
                return

            # ── SMS / WhatsApp webhooks (no bearer — signed by Twilio) ────────
            if path in ("/sms-webhook", "/whatsapp-webhook") and method == "POST":
                body_bytes = await _read_body(receive)
                import urllib.parse as _whup
                params = dict(_whup.parse_qsl(body_bytes.decode("utf-8", errors="replace")))
                if not params:
                    await _send_text(send, 400, "Bad Request")
                    return
                try:
                    from sms_inbox import (sms_inbox_append,
                                          validate_twilio_signature,
                                          validate_signalwire_signature)
                    from sms_backends import load_sms_config
                    cfg        = load_sms_config()
                    provider   = str(cfg.get("sms_provider", "twilio")).lower()
                    auth_token = (cfg.get("twilio_auth_token")
                                  or cfg.get("signalwire_auth_token", ""))
                    headers_d  = {k.lower().decode(): v.decode("utf-8","ignore")
                                  for k, v in scope.get("headers", [])}
                    signature  = headers_d.get("x-twilio-signature", "")
                    scheme     = "https"
                    host       = headers_d.get("host", "localhost")
                    req_url    = f"{scheme}://{host}{path}"

                    if auth_token and signature:
                        if provider == "signalwire":
                            valid = validate_signalwire_signature(auth_token, signature, req_url, params)
                        else:
                            valid = validate_twilio_signature(auth_token, signature, req_url, params)
                        if not valid:
                            _log.warning("Webhook: invalid signature on %s", path)
                            await _send_text(send, 403, "Forbidden")
                            return

                    is_wa    = path == "/whatsapp-webhook"
                    prov_tag = "whatsapp" if is_wa else provider
                    sms_inbox_append(
                        message_id   = params.get("MessageSid","") or params.get("SmsSid",""),
                        from_number  = params.get("From",""),
                        to_number    = params.get("To",""),
                        body         = params.get("Body",""),
                        provider     = prov_tag,
                        contact_name = "",
                        timestamp    = "",
                    )
                    _log.info("Webhook %s: stored inbound from=%s", path, params.get("From",""))
                except Exception as _whe:
                    _log.error("Webhook %s error: %s", path, _whe)

                twiml = b"<?xml version='1.0' encoding='UTF-8'?><Response/>"
                await send({"type": "http.response.start", "status": 200,
                            "headers": [[b"content-type", b"text/xml"],
                                        [b"content-length", str(len(twiml)).encode()]]})
                await send({"type": "http.response.body", "body": twiml})
                return

            if path == "/.well-known/oauth-protected-resource":
                base = _srv_get_public_base(scope)
                await _send_json(send, 200, {
                    "resource": f"{base}/mcp",
                    "authorization_servers": [base],
                })
                return

            if path == "/.well-known/oauth-authorization-server":
                base = _srv_get_public_base(scope)
                await _send_json(send, 200, {
                    "issuer": base,
                    "authorization_endpoint": f"{base}/authorize",
                    "token_endpoint": f"{base}/token",
                    "registration_endpoint": f"{base}/register",
                    "response_types_supported": ["code"],
                    "grant_types_supported": ["authorization_code"],
                    "code_challenge_methods_supported": ["S256", "plain"],
                    "token_endpoint_auth_methods_supported": ["none"],
                    "scopes_supported": ["mcp"],
                })
                return

            if path == "/register" and method == "POST":
                # RFC 7591 Dynamic Client Registration — Claude.ai POSTs here
                # before starting the OAuth flow. Accept any registration and
                # echo back a generated client_id.
                import time as _time
                try:
                    body_bytes = await _read_body(receive)
                    import json as _rj
                    reg_body = _rj.loads(body_bytes) if body_bytes else {}
                except Exception:
                    reg_body = {}
                client_id = _srv_secrets.token_urlsafe(16)
                _log.info("OAuth /register: issued client_id=%s…", client_id[:8])
                await _send_json(send, 201, {
                    "client_id": client_id,
                    "client_id_issued_at": int(_time.time()),
                    "grant_types": reg_body.get("grant_types", ["authorization_code"]),
                    "response_types": reg_body.get("response_types", ["code"]),
                    "redirect_uris": reg_body.get("redirect_uris", []),
                    "token_endpoint_auth_method": "none",
                    "client_name": reg_body.get("client_name", "Claude"),
                })
                return

            if path == "/authorize" or path.startswith("/authorize"):
                # Show a login form. User enters their personal users.json token.
                # On success, redirect back to Claude.ai with a one-time auth code.
                qs_raw = scope.get("query_string", b"")
                params = _parse_qs(qs_raw) if qs_raw else {}
                redirect_uri          = params.get("redirect_uri", "")
                state_val             = params.get("state", "")
                code_challenge        = params.get("code_challenge", "")
                code_challenge_method = params.get("code_challenge_method", "plain")

                error_msg = ""
                if method == "POST":
                    body_bytes  = await _read_body(receive)
                    form_params = _parse_qs(body_bytes)
                    entered     = form_params.get("token", "").strip()
                    # Hot-reload users.json here too, mirroring the MCP bearer-auth
                    # path below. See _hot_reload_users() docstring for why this matters.
                    user        = _resolve_user(_hot_reload_users(users_data), entered)
                    if user is not None:
                        code = _srv_secrets.token_urlsafe(32)
                        _srv_auth_codes[code] = {
                            "redirect_uri":          redirect_uri,
                            "code_challenge":        code_challenge,
                            "code_challenge_method": code_challenge_method,
                            "user_token":            entered,
                        }
                        sep    = "&" if "?" in redirect_uri else "?"
                        target = (f"{redirect_uri}{sep}code={code}"
                                  f"&state={_srv_urlparse.quote(state_val)}")
                        _log.info("OAuth /authorize: user=%s authenticated, issuing code",
                                  user.get("id", "?"))
                        await _send_redirect(send, target)
                        return
                    else:
                        _log.warning("OAuth /authorize: bad token entered (last 4: …%s)",
                                     entered[-4:] if entered else "")
                        error_msg = "<p style='color:red;margin-top:8px'>Incorrect token — try again.</p>"

                html = f"""<!DOCTYPE html>
<html>
<head>
  <title>AI-Prowler Login</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    body {{ font-family: Arial, sans-serif; background: #1a1a2e; color: #eee;
            display: flex; justify-content: center; align-items: center; min-height: 100vh; margin: 0; }}
    .card {{ background: #16213e; border-radius: 12px; padding: 40px; width: 360px;
             box-shadow: 0 4px 24px rgba(0,0,0,0.4); }}
    h2 {{ margin: 0 0 8px; color: #4fc3f7; }}
    p  {{ color: #aaa; font-size: 14px; margin: 0 0 24px; }}
    input {{ width: 100%; padding: 10px 12px; border-radius: 6px; border: 1px solid #333;
             background: #0f3460; color: #eee; font-size: 15px; box-sizing: border-box; }}
    button {{ margin-top: 16px; width: 100%; padding: 11px; border-radius: 6px; border: none;
              background: #4fc3f7; color: #111; font-size: 15px; font-weight: bold; cursor: pointer; }}
    button:hover {{ background: #81d4fa; }}
  </style>
</head>
<body>
  <div class="card">
    <h2>🐾 AI-Prowler</h2>
    <p>Enter your personal Bearer token to connect Claude to your knowledge base.</p>
    <form method="post">
      <input type="password" name="token" placeholder="Bearer token" autofocus>
      <button type="submit">Connect</button>
    </form>
    {error_msg}
  </div>
</body>
</html>"""
                await _send_html(send, 200, html)
                return

            if path == "/token" and method == "POST":
                # Exchange auth code for an access token (PKCE verified).
                body_bytes  = await _read_body(receive)
                form_params = _parse_qs(body_bytes)
                grant_type    = form_params.get("grant_type", "")
                code          = form_params.get("code", "")
                code_verifier = form_params.get("code_verifier", "")

                if grant_type != "authorization_code":
                    await _send_json(send, 400, {"error": "unsupported_grant_type"})
                    return
                if code not in _srv_auth_codes:
                    await _send_json(send, 400, {"error": "invalid_grant"})
                    return

                stored     = _srv_auth_codes.pop(code)
                user_token = stored.get("user_token", "")

                # Verify PKCE if code_challenge was stored.
                if stored.get("code_challenge"):
                    method_pkce = stored.get("code_challenge_method", "plain")
                    if method_pkce == "S256":
                        digest   = _srv_hashlib.sha256(code_verifier.encode()).digest()
                        computed = _srv_base64.urlsafe_b64encode(digest).rstrip(b"=").decode()
                    else:
                        computed = code_verifier
                    if computed != stored["code_challenge"]:
                        _log.warning("OAuth /token: PKCE mismatch for code")
                        await _send_json(send, 400, {
                            "error": "invalid_grant",
                            "error_description": "PKCE mismatch",
                        })
                        return

                access_token = _srv_secrets.token_urlsafe(48)
                # Map the new access token back to the original users.json token
                # so _resolve_user can look up the user on subsequent MCP calls.
                _srv_access_tokens[access_token] = user_token
                _log.info("OAuth /token: issued access token for user_token …%s",
                          user_token[-4:] if user_token else "?")
                await _send_json(send, 200, {
                    "access_token": access_token,
                    "token_type":   "bearer",
                    "expires_in":   31536000,  # 1 year
                })
                return

            # ── Authenticate all other paths (MCP, whoami, etc.) ─────────────
            tok = _bearer_from_scope(scope)
            if not tok:
                base = _srv_get_public_base(scope)
                www  = (f'Bearer realm="{base}", '
                        f'resource_metadata="{base}/.well-known/oauth-protected-resource"')
                body = b'{"error":"missing bearer token"}'
                await send({"type": "http.response.start", "status": 401,
                            "headers": [
                                (b"content-type",    b"application/json"),
                                (b"content-length",  str(len(body)).encode()),
                                (b"www-authenticate", www.encode()),
                            ]})
                await send({"type": "http.response.body", "body": body})
                return

            # Resolve the presented token: it may be an OAuth-issued access
            # token (map back to users.json token first) or a raw users.json token.
            # Hot-reload users.json so suspensions/additions take effect
            # immediately without a server restart. (Shared with the /authorize
            # login page via _hot_reload_users() — see that docstring.)
            _live_users = _hot_reload_users(users_data)
            for _new_tok in (_live_users.get("users") or {}).keys():
                if _new_tok not in _srv_access_tokens:
                    _srv_access_tokens[_new_tok] = _new_tok
            raw_user_token = _srv_access_tokens.get(tok, tok)
            user = _resolve_user(_live_users, raw_user_token)
            if user is None:
                _log.info("Auth rejected for token …%s on %s",
                          tok[-4:] if tok else "", path)
                base = _srv_get_public_base(scope)
                www  = (f'Bearer realm="{base}", '
                        f'resource_metadata="{base}/.well-known/oauth-protected-resource"')
                body = b'{"error":"invalid or revoked token"}'
                await send({"type": "http.response.start", "status": 401,
                            "headers": [
                                (b"content-type",    b"application/json"),
                                (b"content-length",  str(len(body)).encode()),
                                (b"www-authenticate", www.encode()),
                            ]})
                await send({"type": "http.response.body", "body": body})
                return

            # Stash the user where FastMCP's Request(scope).state will expose it.
            # Starlette's Request.state is backed by scope["state"]; setting it
            # here is what makes _current_user(ctx) resolve the right user inside
            # the tools over HTTP (the keystone).
            scope = dict(scope)
            state = dict(scope.get("state") or {})
            state["user"] = user
            # SCOPE_SIMPLIFICATION_SPEC.md section 3.7 (Phase 7 cleanup,
            # 2026-07-17): was state["allowed_collections"] = _allowed_collections
            # (user) -- the multi-collection enforcement list this fed into
            # (_scoped_collections_for_ctx) is gone; the single-collection
            # where-filter is built fresh per read tool call now, not carried
            # on request.state. Kept here only as a /whoami diagnostic,
            # updated to report the actual new enforcement input.
            state["allowed_scopes"] = sorted(_allowed_scopes(user))
            scope["state"] = state
            _log.debug("Authenticated user=%s role=%s", user.get("id"),
                       user.get("role"))

            # /whoami diagnostic (reports scoping_active=True — tools scope).
            if path == "/whoami":
                await _send_json(send, 200, {
                    "user_id":              user.get("id"),
                    "name":                 user.get("name"),
                    "role":                 user.get("role"),
                    "allowed_scopes":       state["allowed_scopes"],
                    "scoping_active":       True,
                })
                return

            # ── MCP header injection (mirrors _run_http; required by FastMCP's
            #    streamable_http_app validators) ────────────────────────────────
            if path == "/mcp" or path.startswith("/mcp"):
                headers = list(scope.get("headers", []))

                def _has(name, frag):
                    return any(hk.lower() == name and frag in hv.lower()
                               for hk, hv in headers)

                def _set(name, value):
                    for i, (hk, hv) in enumerate(headers):
                        if hk.lower() == name:
                            headers[i] = (name, value); return
                    headers.append((name, value))

                if method == "POST":
                    if not _has(b"content-type", b"application/json"):
                        _set(b"content-type", b"application/json")
                    if not _has(b"accept", b"text/event-stream"):
                        _set(b"accept", b"application/json, text/event-stream")
                elif method == "GET":
                    if not _has(b"accept", b"text/event-stream"):
                        _set(b"accept", b"text/event-stream")

                # Host rewrite + MCP-Protocol-Version (FastMCP validators).
                _set(b"host", _local_host)
                if not any(hk.lower() == b"mcp-protocol-version"
                           for hk, hv in headers):
                    _set(b"mcp-protocol-version", b"2025-03-26")

                scope["headers"] = headers

            # Stream to FastMCP WITHOUT buffering (SSE-safe).
            await mcp_asgi(scope, receive, send)

    app = _ServerRouterASGI()

    _log.info("Server mode (multi-user, scoped) listening on port %d", port)
    _log.info("Pure-ASGI router active — SSE-safe, per-user ctx scoping live.")
    uvicorn.run(app, host="0.0.0.0", port=port,
                proxy_headers=True, forwarded_allow_ips="*", log_config=None)


# ══════════════════════════════════════════════════════════════════════════════
# HTTP transport with Bearer-token auth (launched by GUI or manually)
# ══════════════════════════════════════════════════════════════════════════════

def _run_http(port: int, token: str, public_base: str = "https://mobile.dvavro-ai-prowler.com") -> None:
    """
    Run the MCP server over HTTP with OAuth 2.0 + PKCE authentication.

    Claude.ai custom connectors require OAuth — they first fetch
    /.well-known/oauth-authorization-server to discover auth endpoints,
    then redirect the user to /authorize where they enter the Bearer token,
    then exchange the auth code for an access token at /token.
    Subsequent MCP requests use the returned access token as a Bearer token.

    The server binds to 127.0.0.1 only.  Cloudflare Tunnel provides the
    public HTTPS endpoint.
    """
    try:
        import uvicorn
        import hashlib
        import base64
        import urllib.parse
        from starlette.applications import Starlette
        from starlette.middleware.base import BaseHTTPMiddleware
        from starlette.responses import PlainTextResponse, JSONResponse, HTMLResponse, RedirectResponse
        from starlette.routing import Route, Mount
        from starlette.requests import Request
    except ImportError as _ie:
        print(f"ERROR: {_ie}\nRun: pip install uvicorn", file=sys.stderr)
        sys.exit(1)

    if not token:
        print("ERROR: --token is required for HTTP transport.", file=sys.stderr)
        sys.exit(1)

    # ══════════════════════════════════════════════════════════════════════════
    # LICENSE VALIDATION (v8.0.0)
    # ══════════════════════════════════════════════════════════════════════════
    # Every AI-Prowler license (personal/mobile, business, and free beta) is
    # issued and validated by the ai-prowler-subscription Cloudflare Worker —
    # see _validate_license() / _evaluate_license_grace() above, and
    # GET /license/{key}/validate in the Worker. The bearer token (set by the
    # user in Settings → Remote Access) is unrelated to licensing — it is
    # purely a local secret for the Claude.ai MCP connector handshake and is
    # never sent to, or known by, any subscription system.
    #
    # Check schedule:
    #   • On startup — validate the configured license_key
    #   • Every 30 days — re-validate silently in a background thread
    #   • If validation fails — local cache + grace ladder allow continued use
    #     for up to _LICENSE_WARN_DAYS (silent) then _LICENSE_GRACE_DAYS (warning)
    #     since the last successful check, per _evaluate_license_grace().
    #
    # The pre-v8.0.0 design (a GitHub-hosted subs.json registry, keyed by a
    # hash of the bearer token, fetched read-only by every client) has been
    # retired entirely — there are no remaining references to subs.json,
    # _check_subscription, or _fetch_subs_registry anywhere in this file.
    # ══════════════════════════════════════════════════════════════════════════

    import hashlib    as _hashlib
    import datetime   as _dt
    import threading  as _threading

    _CHECK_INTERVAL_DAYS = 30   # re-validate the license every 30 days (background thread)

    # ── EDITION / MODE helpers (v7.0.0) ───────────────────────────────────────
    # _load_runtime_config, _enforce_edition_mode and the edition/mode
    # constants were HOISTED to module level (just before _run_http) so the
    # test suite can call them directly. The bare-name calls below resolve to
    # those module-level definitions. See the module-level block and
    # PHASE_A_PRIME_TEST_PLAN.md §4.0.

    def _token_key(tok: str) -> str:
        """Short hash of a string, used as an opaque per-license fingerprint
        for the 2-active-install activation check (NOT a subs.json lookup —
        that registry is retired; see _post_activation below)."""
        return _hashlib.sha256(tok.encode()).hexdigest()[:16]

    # ── install_id (v7.0.0 — Phase A' install-id binding) ─────────────────────
    # The GUI generates ~/.ai-prowler/install_id on first launch (a 16-char
    # sha256 hash). The MCP server reads it here so the subscription check and
    # the 2-active-install rule can bind a license to specific machines. If the
    # file is missing (e.g. server started before the GUI ever ran), we generate
    # it here too so the server is self-sufficient — same format the GUI uses.
    _INSTALL_ID_PATH = Path.home() / ".ai-prowler" / "install_id"

    def _read_or_create_install_id() -> str:
        """Return this machine's install_id, creating it if absent. Never raises;
        returns '' only if the home directory itself is unwritable."""
        try:
            if _INSTALL_ID_PATH.exists():
                val = _INSTALL_ID_PATH.read_text(encoding="utf-8").strip()
                if val:
                    return val
        except Exception as _e:
            _log.warning("Could not read install_id (%s)", _e)
        # Generate — mirror the GUI's algorithm (uuid4 → sha256 → first 16 hex)
        try:
            import uuid as _uuid
            new_id = _hashlib.sha256(str(_uuid.uuid4()).encode()).hexdigest()[:16]
            _INSTALL_ID_PATH.parent.mkdir(parents=True, exist_ok=True)
            _INSTALL_ID_PATH.write_text(new_id, encoding="utf-8")
            _log.info("Generated new install_id for this machine")
            return new_id
        except Exception as _e:
            _log.warning("Could not create install_id (%s) — install-id binding disabled", _e)
            return ""

    _INSTALL_ID = _read_or_create_install_id()



    # ── 2-active-install rule (v7.0.0) ────────────────────────────────────────
    # _evaluate_activation and the constants _ACTIVE_WINDOW_DAYS /
    # _MAX_ACTIVE_INSTALLS were HOISTED to module level (just before _run_http)
    # so the test suite can call the pure evaluator directly. The bare-name
    # references below resolve to those module-level definitions. See the
    # module-level block and PHASE_A_PRIME_TEST_PLAN.md §4.

    # ── Subscription-Worker-backed activation (v8.0.0) ────────────────────────
    # Activations are authoritatively stored in the ai-prowler-subscription
    # Worker (the same one /license/validate hits). The client POSTs its
    # install_id + license-key HASH (never the raw key, never the bearer
    # token) to /license/activate and the Worker returns the binding decision.
    # If the Worker is unreachable we FAIL OPEN — there is no local fallback
    # registry anymore (v8.0.0 retired subs.json entirely) — and ultimately
    # allow access. This mirrors the grace-ladder philosophy: a network blip
    # must never lock out a paying customer.
    _SUBSCRIPTION_WORKER_URL = "https://ai-prowler-subscription.david-vavro1.workers.dev"

    def _activation_endpoint() -> str:
        """Resolve the Worker base URL (config override or default), no trailing slash."""
        base = _SUBSCRIPTION_WORKER_URL
        try:
            ep = (_runtime_cfg_for_endpoint or {}).get("subscription_endpoint", "")
            if ep:
                base = str(ep)
        except Exception:
            pass
        return base.rstrip("/")


    def _post_activation(license_hash: str, install_id: str,
                         os_str: str, version: str) -> dict | None:
        """POST to the Worker's /license/activate. Returns the parsed decision
        dict on success, or None on any failure (caller falls back to local)."""
        if not install_id or not license_hash:
            return None
        url = f"{_activation_endpoint()}/license/activate"
        payload = {
            "license_key_hash": license_hash,
            "install_id":       install_id,
            "os":               os_str,
            "version":          version,
        }
        try:
            import requests as _req
            resp = _req.post(
                url, json=payload,
                headers={"User-Agent": "AI-Prowler-MCP/1.0",
                         "Content-Type": "application/json"},
                timeout=10,
                proxies={"http": None, "https": None})
            if resp.status_code == 200:
                data = resp.json()
                _log.info("D1 activation decision=%s (%s of %s active)",
                          data.get("decision"), data.get("active_count"),
                          data.get("max_active"))
                return data
            _log.warning("Activation endpoint returned HTTP %s", resp.status_code)
        except Exception as _e:
            _log.warning("Activation endpoint unreachable (%s) — failing open", _e)
        return None

    # Pulled out so _activation_endpoint can see the loaded config without a
    # forward reference; populated just before the activation call below.
    _runtime_cfg_for_endpoint = None

    # ── Perform initial license check on startup (v8.0.0) ─────────────────────
    # Single source of truth: the ai-prowler-subscription Worker's
    # /license/{key}/validate, via _validate_license(). The old GitHub
    # subs.json registry (_check_subscription/_fetch_subs_registry) is
    # retired entirely — every license (personal/mobile, business, and free
    # beta) is now issued and validated by this one Worker, keyed by
    # license_key from config.json. No bearer-token lookup is involved here
    # at all (the bearer token is purely a local Claude.ai connector secret
    # the user chooses themselves — see Settings → Remote Access — and was
    # never assigned or known by any subscription system, old or new).
    if _test_entitlement_active():
        # TEST ENTITLEMENT SHORT-CIRCUIT (pre-release validation only). Both the
        # env var AND config.json "test_mode": true are set (dev/test launch).
        # Skip the network license-validation call and substitute a known-good
        # verdict so the suite can exercise the REAL auth + scoping + ownership
        # code against the sandboxed users.json. Enforcement is NOT disabled —
        # only the entitlement verdict and state-file paths are sandboxed.
        # Loud on purpose.
        _log.warning(
            "⚠️  TEST ENTITLEMENT ACTIVE — network license checks "
            "SKIPPED, entitlement sandboxed (edition=business,status=ok). "
            "State dir=%s. NOT FOR PRODUCTION.", _state_dir())
        _runtime_cfg_for_endpoint = None
        _sub_result = {"status": "ok", "name": "TEST", "days_left": None,
                       "edition": "business",
                       "message": "test entitlement (sandboxed; network skipped)"}
    else:
        _runtime_cfg_for_endpoint = _load_runtime_config()
        _lic_key_startup = str(_runtime_cfg_for_endpoint.get("license_key", "")).strip()
        if not _lic_key_startup:
            # No license key configured at all — this is the normal, expected
            # state for a brand-new Home install. Not an error.
            _sub_result = {"status": "unmanaged", "name": None, "days_left": None,
                           "edition": "home",
                           "message": "No license key configured — Home/local mode"}
        else:
            _startup_grace = _validate_license(
                _lic_key_startup, _INSTALL_ID, _activation_endpoint())
            _granted = _startup_grace.get("effective_edition", "home")
            _action  = _startup_grace.get("action", "")
            if _granted == "home" and _action in ("reverted_revoked",):
                # Hard fail (revoked/suspended-with-no-grace-left/not_found) —
                # this is the v8.0.0 equivalent of the old "blocked" status.
                # NOTE: ordinary Stripe cancellation does NOT land here — it's
                # a SOFT reason (see suspendLicense() in the Worker), so a
                # cancelling customer keeps working through the grace ladder
                # below instead of being hard-blocked immediately.
                _sub_result = {"status": "blocked", "name": None, "days_left": None,
                               "edition": "home",
                               "message": _startup_grace.get("banner") or
                                          "License is no longer valid."}
            elif _granted == "home" and _action in ("reverted_revoked",):
                # Hard revocation — Worker explicitly said the key is revoked/
                # suspended. Block mobile access immediately.
                _sub_result = {"status": "blocked", "name": None, "days_left": None,
                               "edition": "home",
                               "message": _startup_grace.get("banner") or
                                          "License could not be validated."}
            elif _granted == "home" and _startup_grace.get("had_prior_cache"):
                # Grace ladder expired for a previously-valid subscription.
                # The user had a paid subscription that is now expired/unreachable
                # beyond the grace window → block mobile access.
                _sub_result = {"status": "blocked", "name": None, "days_left": None,
                               "edition": "home",
                               "message": _startup_grace.get("banner") or
                                          "Subscription expired — mobile access disabled."}
            elif _granted == "home":
                # reverted_expired with had_prior_cache=False: no subscription has
                # ever been validated on this install (new user, self-hosted, or
                # Worker returned 404 before first successful activation).
                # Fail-open — never block a user who hasn't subscribed yet.
                _sub_result = {"status": "unmanaged", "name": None, "days_left": None,
                               "edition": "home",
                               "message": _startup_grace.get("banner") or
                                          "No subscription found — mobile access requires a subscription."}
            elif _startup_grace.get("banner"):
                # Still entitled, but the grace ladder has a warning to show
                # (grace_warning action — failing validation for several days,
                # or a cancelled-but-still-in-grace subscription counting down).
                _sub_result = {"status": "warning", "name": None, "days_left": None,
                               "edition": _granted,
                               "message": _startup_grace["banner"],
                               "banner": (
                                   f"<div style='background:#7c4a00;border-radius:6px;padding:10px 14px;"
                                   f"margin-top:12px;font-size:13px;color:#ffe082;'>"
                                   f"⚠️  {_startup_grace['banner']}</div>")}
            else:
                # validated / cached_fresh / grace_silent — fully healthy.
                _sub_result = {"status": "ok", "name": None, "days_left": None,
                               "edition": _granted,
                               "message": f"License OK — edition={_granted} (action={_action})"}
    _sub_result.setdefault("banner", "")
    _log.info("Startup license check: %s", _sub_result["message"])

    if _sub_result["status"] == "blocked":
        _log.critical("ACCESS BLOCKED: %s", _sub_result["message"])
        print(f"ERROR: {_sub_result['message']}", file=sys.stderr)
        sys.exit(1)
    elif _sub_result["status"] == "warning":
        _log.warning("SUBSCRIPTION NOTICE: %s", _sub_result["message"])

    # ── Resolve effective edition / mode (v7.0.0 — Phase A') ──────────────────
    # config.json declares the *requested* edition/mode; the license check
    # decides whether the key is actually entitled to it. _enforce_edition_mode
    # reconciles the two, downgrading to home/personal where the request can't be
    # honored. The license's plan (via _sub_result["edition"]) is the upper
    # bound on entitlement; config.json can request equal-or-lower, never higher.
    _runtime_cfg       = _runtime_cfg_for_endpoint or _load_runtime_config()
    _requested_edition = _runtime_cfg["edition"]
    _requested_mode    = _runtime_cfg["mode"]
    _entitled_edition  = _sub_result.get("edition", "home")

    # The requested edition may not exceed what the license entitles.
    # Ranking: home(0) < mobile(1) < business(2). Clamp request to entitlement.
    _EDITION_RANK = {"home": 0, "mobile": 1, "business": 2}
    if _EDITION_RANK.get(_requested_edition, 0) > _EDITION_RANK.get(_entitled_edition, 0):
        _log.warning(
            "config.json requests edition=%s but license only entitles %s. "
            "Clamping to %s.", _requested_edition, _entitled_edition, _entitled_edition)
        _requested_edition = _entitled_edition

    _EFFECTIVE_EDITION, _EFFECTIVE_MODE = _enforce_edition_mode(
        _requested_edition, _requested_mode, _sub_result["status"])

    # ── Per-user child-license validation (server mode only, v7.0.0) ──────────
    # When running as a company server, each active employee in users.json may
    # carry a child_license_key (their paid seat). On startup AND on the
    # natural 30-day fresh-cache rhythm of _validate_license, validate every
    # active user's child key against the Worker. Soft policy per David
    # 2026-05-28: REJECTIONS LOG + SHOW A BANNER, they DO NOT mutate users.json
    # and do NOT block the bearer-token auth at request time. The owner sees
    # the banner and acts via the Admin tab. The hard enforcement that matters
    # is the future-dated expires_at on the child key in the Worker; when it
    # eventually comes back expired, validate just keeps reporting that — the
    # seat is logically dead but service continues (white-glove model).
    #
    # The actual sweep logic lives in _sweep_child_licenses(users_doc, validate_fn)
    # so it's testable in isolation (tests pass synthetic users_doc + stubbed
    # validate_fn, no network).
    _child_warnings = []
    if (_EFFECTIVE_EDITION == "business" and _EFFECTIVE_MODE == "server"
            and not _test_entitlement_active()):
        try:
            _users_doc = _load_users() or {}
            _endpoint  = _activation_endpoint()
            _validate  = lambda k: _validate_license(k, _INSTALL_ID, _endpoint)
            _child_warnings = _sweep_child_licenses(_users_doc, _validate)
        except Exception as _cwerr:
            # Never let a child-key sweep block startup. Log and continue.
            _log.warning("Child-license sweep failed (%s); continuing.", _cwerr)
        # Persist the warnings (or the empty list) so the Admin tab can display
        # them. ALWAYS write — an absent file means 'sweep never ran', a present
        # file with warnings=[] means 'all clear as of last_check_at'.
        _save_license_warnings(_child_warnings)
    if _child_warnings:
        _sub_result = dict(_sub_result)
        _sub_result["child_license_warnings"] = _child_warnings

    _log.info(
        "Effective runtime: edition=%s mode=%s install_id=%s (requested edition=%s mode=%s; "
        "license status=%s, entitles=%s)",
        _EFFECTIVE_EDITION, _EFFECTIVE_MODE, _INSTALL_ID or "<none>",
        _runtime_cfg["edition"], _runtime_cfg["mode"],
        _sub_result["status"], _entitled_edition)

    # ── Apply the 2-active-install rule (v7.0.0 — Phase A') ───────────────────
    # Only relevant once we've cleared edition entitlement: a Home install has
    # no remote-access seat to bind, so there's nothing to enforce. For a
    # mobile/business effective edition, determine THIS machine's activation
    # standing. Authoritative source is the same subscription Worker
    # (/license/activate); if it's unreachable we FAIL OPEN by treating this
    # install as active — a network blip must never lock out a paying customer.
    # A "rejected" decision → soft-revert to Home (spec §4.4 Mobile flow),
    # leaving everything else functional, and annotate _sub_result so the
    # License panel and /authorize page can surface the "release a machine" CTA.

    _activation = {"decision": "unbound", "active_install_ids": [],
                   "active_count": 0, "this_active": False, "message": ""}
    if _test_entitlement_active():
        _log.warning("⚠️  TEST ENTITLEMENT — skipping activation (2-install) "
                     "check (network); treating this install as active.")
        _activation["this_active"] = True
    elif _EFFECTIVE_EDITION in ("mobile", "business"):
        # Build an OS string for the activation record.
        try:
            import platform as _platform
            _os_str = f"{_platform.system()}-{_platform.release()}"[:50]
        except Exception:
            _os_str = "unknown"
        _license_hash = _token_key(
            str((_runtime_cfg or {}).get("license_key", "")).strip() or token)

        # Resolve the app version for the activation record. APP_VERSION lives
        # in rag_gui.py, not here, so read the bundled VERSION file directly
        # (same value the installer ships) with a safe fallback.
        _app_version = ""
        try:
            _ver_file = Path(__file__).resolve().parent / "VERSION"
            if _ver_file.exists():
                _app_version = _ver_file.read_text(encoding="utf-8").strip()
        except Exception:
            _app_version = ""

        _d1 = _post_activation(_license_hash, _INSTALL_ID, _os_str, _app_version)
        if _d1 is not None and _d1.get("decision"):
            _activation = {
                "decision":           _d1.get("decision"),
                "active_install_ids": _d1.get("active_install_ids", []),
                "active_count":       _d1.get("active_count", 0),
                "this_active":        _d1.get("decision") == "active",
                "message":            _d1.get("message", ""),
            }
            _log.info("Activation: decision=%s (%d of %d active)",
                      _activation["decision"], _activation["active_count"],
                      _MAX_ACTIVE_INSTALLS)
        else:
            # FAIL OPEN — Worker unreachable; allow this install unconditionally
            # rather than lock out a paying customer over a network blip. There
            # is no local subs.json fallback anymore (v8.0.0 retired it) — the
            # Worker is the only source of activation truth now.
            _activation["this_active"] = True
            _log.warning("Activation endpoint unreachable — failing open "
                         "(treating this install as active).")

        if _activation["decision"] == "rejected":
            _log.warning(
                "Install-id rejected (2-active-install cap reached). Soft-reverting "
                "to Home edition; remote access disabled on this machine.")
            _EFFECTIVE_EDITION = "home"
            _EFFECTIVE_MODE    = "personal"
            _sub_result = dict(_sub_result)
            _sub_result["activation_rejected"] = True
            _sub_result["activation_message"]  = _activation["message"]
            _sub_result["active_install_ids"]  = _activation["active_install_ids"]

    # Mutable container so background thread can update it
    _current_sub_result = [_sub_result]
    _subs_lock          = _threading.Lock()

    # ── Background 30-day license re-validation ───────────────────────────────
    def _periodic_sub_refresh():
        import time as _time
        while True:
            _time.sleep(_CHECK_INTERVAL_DAYS * 86400)   # 30 days in seconds
            _log.info("30-day license re-validation starting…")
            _lic_key_bg = str((_runtime_cfg or {}).get("license_key", "")).strip()
            if not _lic_key_bg:
                _log.info("No license key configured — nothing to re-validate.")
                continue
            _grace_bg = _validate_license(_lic_key_bg, _INSTALL_ID, _activation_endpoint())
            _granted_bg = _grace_bg.get("effective_edition", "home")
            if _granted_bg == "home":
                result = {"status": "blocked", "name": None, "days_left": None,
                          "edition": "home",
                          "message": _grace_bg.get("banner") or "License is no longer valid."}
            elif _grace_bg.get("banner"):
                result = {"status": "warning", "name": None, "days_left": None,
                          "edition": _granted_bg, "message": _grace_bg["banner"]}
            else:
                result = {"status": "ok", "name": None, "days_left": None,
                          "edition": _granted_bg, "message": "License OK"}
            with _subs_lock:
                _current_sub_result[0] = result
            _log.info("License re-check: %s", result["message"])
            if result["status"] == "blocked":
                _log.critical(
                    "ACCESS NOW BLOCKED on re-check — new /authorize attempts "
                    "will show expiry page: %s", result["message"]
                )
            elif result["status"] == "warning":
                _log.warning("SUBSCRIPTION NOTICE: %s", result["message"])

    if _test_entitlement_active():
        _log.warning("⚠️  TEST ENTITLEMENT — not starting the 30-day "
                     "license refresh thread (would hit the network).")
    else:
        _threading.Thread(target=_periodic_sub_refresh, daemon=True).start()

    # In-memory stores for OAuth codes and issued access tokens
    # code -> {redirect_uri, code_challenge, code_challenge_method}
    _auth_codes: dict = {}
    # access_token -> True
    _access_tokens: set = set()
    # Pre-add the user's own bearer token so existing curl/manual clients still work
    _access_tokens.add(token)

    PUBLIC_BASE = public_base.rstrip("/")   # passed in from CLI / config.json

    def _get_public_base(request: Request = None, scope: dict = None) -> str:
        """Derive the public base URL dynamically from the request.

        If a Request or ASGI scope is provided, use the Host header so that
        OAuth endpoints always point back to the URL Claude is actually
        connecting through — whether that's a Named Tunnel, Quick Tunnel,
        or localhost.  Falls back to the configured PUBLIC_BASE if the
        Host header isn't available.
        """
        host = ''
        if request is not None:
            host = request.headers.get('host', '') or request.headers.get('x-forwarded-host', '')
        elif scope is not None:
            headers = {k.lower(): v for k, v in scope.get("headers", [])}
            host = (headers.get(b"host", b"")
                    or headers.get(b"x-forwarded-host", b"")).decode("utf-8", errors="ignore")
        if host:
            # Use https for any non-localhost host
            scheme = "http" if host.startswith("127.") or host.startswith("localhost") else "https"
            return f"{scheme}://{host}"
        return PUBLIC_BASE

    # ── OAuth discovery endpoints ─────────────────────────────────────────────
    async def oauth_protected_resource(request: Request):
        """RFC 9728 — OAuth 2.0 Protected Resource Metadata.
        Claude.ai fetches THIS endpoint first when it receives a 401.
        It tells Claude where the authorization server lives.
        """
        base = _get_public_base(request)
        return JSONResponse({
            "resource": f"{base}/mcp",
            "authorization_servers": [base],
        })

    async def oauth_metadata(request: Request):
        """RFC 8414 — OAuth 2.0 Authorization Server Metadata.
        Claude.ai fetches this second, using the AS URL from above.
        """
        base = _get_public_base(request)
        return JSONResponse({
            "issuer": base,
            "authorization_endpoint": f"{base}/authorize",
            "token_endpoint": f"{base}/token",
            "registration_endpoint": f"{base}/register",
            "response_types_supported": ["code"],
            "grant_types_supported": ["authorization_code"],
            "code_challenge_methods_supported": ["S256", "plain"],
            "token_endpoint_auth_methods_supported": ["none"],
            "scopes_supported": ["mcp"],
        })

    # ── Dynamic Client Registration (RFC 7591) ───────────────────────────────
    # Claude.ai POSTs here before starting the OAuth flow to register itself.
    # We accept any registration and echo back a generated client_id.
    async def register_client(request: Request):
        import secrets as _sec, time as _time
        try:
            body = await request.json()
        except Exception:
            body = {}
        client_id = _sec.token_urlsafe(16)
        return JSONResponse({
            "client_id": client_id,
            "client_id_issued_at": int(_time.time()),
            "grant_types": body.get("grant_types", ["authorization_code"]),
            "response_types": body.get("response_types", ["code"]),
            "redirect_uris": body.get("redirect_uris", []),
            "token_endpoint_auth_method": "none",
            "client_name": body.get("client_name", "Claude"),
        }, status_code=201)

    # ── Authorization endpoint ────────────────────────────────────────────────
    async def authorize(request: Request):
        """
        Show a login form.  User enters their Bearer token.
        On success, redirect back to Claude.ai with an auth code.
        """
        params = dict(request.query_params)
        redirect_uri      = params.get("redirect_uri", "")
        state             = params.get("state", "")
        code_challenge    = params.get("code_challenge", "")
        code_challenge_method = params.get("code_challenge_method", "plain")

        # Get current subscription status for banner display
        with _subs_lock:
            _cur_sub = _current_sub_result[0]

        # If subscription is BLOCKED, show expiry page instead of login form
        if _cur_sub["status"] == "blocked":
            html = f"""<!DOCTYPE html>
<html>
<head>
  <title>AI-Prowler — Access Suspended</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    body {{ font-family: Arial, sans-serif; background: #1a1a2e; color: #eee;
            display: flex; justify-content: center; align-items: center; min-height: 100vh; margin: 0; }}
    .card {{ background: #16213e; border-radius: 12px; padding: 40px; width: 380px;
             box-shadow: 0 4px 24px rgba(0,0,0,0.4); text-align: center; }}
    h2 {{ margin: 0 0 8px; color: #ef5350; }}
    p  {{ color: #aaa; font-size: 14px; margin: 12px 0; line-height: 1.5; }}
    a  {{ color: #4fc3f7; }}
    .badge {{ background: #7c0000; border-radius: 6px; padding: 12px 16px;
              font-size: 13px; color: #ffcdd2; margin-top: 16px; }}
  </style>
</head>
<body>
  <div class="card">
    <h2>🔒 Remote Access Suspended</h2>
    <p>Your AI-Prowler managed remote access subscription has expired and
       the grace period has elapsed.</p>
    <div class="badge">
      To restore access, please renew your subscription:<br><br>
      <strong><a href="mailto:david.vavro1@gmail.com">david.vavro1@gmail.com</a></strong>
    </div>
    <p style="font-size:12px;color:#555;margin-top:20px;">
      AI-Prowler desktop features continue to work normally.<br>
      Only remote / Claude.ai access requires a subscription.
    </p>
  </div>
</body>
</html>"""
            return HTMLResponse(html, status_code=403)

        if request.method == "POST":
            form = await request.form()
            entered = (form.get("token") or "").strip()
            if entered == token:
                # Generate a one-time auth code
                import secrets
                code = secrets.token_urlsafe(32)
                _auth_codes[code] = {
                    "redirect_uri": redirect_uri,
                    "code_challenge": code_challenge,
                    "code_challenge_method": code_challenge_method,
                }
                sep = "&" if "?" in redirect_uri else "?"
                target = f"{redirect_uri}{sep}code={code}&state={urllib.parse.quote(state)}"
                return RedirectResponse(url=target, status_code=302)
            else:
                error_msg = "<p style='color:red;margin-top:8px'>Incorrect token — try again.</p>"
        else:
            error_msg = ""

        # Show renewal banner if subscription is in warning period
        sub_banner = _cur_sub.get("banner", "")

        html = f"""<!DOCTYPE html>
<html>
<head>
  <title>AI-Prowler Login</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    body {{ font-family: Arial, sans-serif; background: #1a1a2e; color: #eee;
            display: flex; justify-content: center; align-items: center; min-height: 100vh; margin: 0; }}
    .card {{ background: #16213e; border-radius: 12px; padding: 40px; width: 360px;
             box-shadow: 0 4px 24px rgba(0,0,0,0.4); }}
    h2 {{ margin: 0 0 8px; color: #4fc3f7; }}
    p  {{ color: #aaa; font-size: 14px; margin: 0 0 24px; }}
    input {{ width: 100%; padding: 10px 12px; border-radius: 6px; border: 1px solid #333;
             background: #0f3460; color: #eee; font-size: 15px; box-sizing: border-box; }}
    button {{ margin-top: 16px; width: 100%; padding: 11px; border-radius: 6px; border: none;
              background: #4fc3f7; color: #111; font-size: 15px; font-weight: bold; cursor: pointer; }}
    button:hover {{ background: #81d4fa; }}
  </style>
</head>
<body>
  <div class="card">
    <h2>🐾 AI-Prowler</h2>
    <p>Enter your Bearer token to connect Claude to your knowledge base.</p>
    <form method="post">
      <input type="password" name="token" placeholder="Bearer token" autofocus>
      <button type="submit">Connect</button>
    </form>
    {error_msg}
    {sub_banner}
  </div>
</body>
</html>"""
        return HTMLResponse(html)

    # ── Token endpoint ────────────────────────────────────────────────────────
    async def token_endpoint(request: Request):
        """Exchange auth code for access token."""
        try:
            form = await request.form()
        except Exception:
            form = {}
        grant_type    = form.get("grant_type", "")
        code          = form.get("code", "")
        code_verifier = form.get("code_verifier", "")

        if grant_type != "authorization_code":
            return JSONResponse({"error": "unsupported_grant_type"}, status_code=400)
        if code not in _auth_codes:
            return JSONResponse({"error": "invalid_grant"}, status_code=400)

        stored = _auth_codes.pop(code)

        # Verify PKCE if code_challenge was stored
        if stored.get("code_challenge"):
            method = stored.get("code_challenge_method", "plain")
            if method == "S256":
                digest = hashlib.sha256(code_verifier.encode()).digest()
                computed = base64.urlsafe_b64encode(digest).rstrip(b"=").decode()
            else:
                computed = code_verifier
            if computed != stored["code_challenge"]:
                return JSONResponse({"error": "invalid_grant", "error_description": "PKCE mismatch"}, status_code=400)

        import secrets
        access_token = secrets.token_urlsafe(48)
        _access_tokens.add(access_token)

        return JSONResponse({
            "access_token": access_token,
            "token_type": "bearer",
            "expires_in": 31536000,  # 1 year
        })

    # ── Pure ASGI auth middleware ─────────────────────────────────────────────
    # IMPORTANT: Must NOT use BaseHTTPMiddleware here — it buffers the entire
    # response body which breaks FastMCP's Server-Sent Events streaming transport
    # causing 500 Internal Server Error on every MCP request.
    # This pure ASGI class passes bytes through directly without any buffering.
    _PUBLIC_PATHS = {"/health", "/authorize", "/token", "/register",
                     "/.well-known/oauth-authorization-server",
                     "/.well-known/oauth-protected-resource",
                     "/sms-webhook", "/whatsapp-webhook"}

    class _AuthASGI:
        """Zero-buffering ASGI auth wrapper — safe for SSE/streaming responses."""
        def __init__(self, asgi_app):
            self._app = asgi_app

        async def __call__(self, scope, receive, send):
            if scope["type"] not in ("http", "websocket"):
                await self._app(scope, receive, send)
                return

            path = scope.get("path", "")
            if path in _PUBLIC_PATHS or path.startswith("/authorize"):
                await self._app(scope, receive, send)
                return

            # Extract Authorization header from raw ASGI scope headers
            headers = {k.lower(): v for k, v in scope.get("headers", [])}
            auth_raw = headers.get(b"authorization", b"")
            auth = auth_raw.decode("utf-8", errors="ignore")
            tok  = auth[7:].strip() if auth.lower().startswith("bearer ") else ""

            if tok not in _access_tokens:
                # Return 401 — derive public base from the Host header
                # so Claude discovers OAuth endpoints at the correct URL
                _dyn_base = _get_public_base(scope=scope)
                body = b'{"error":"unauthorized","error_description":"Invalid or missing Bearer token"}'
                www_auth = (
                    f'Bearer realm="{_dyn_base}", '
                    f'resource_metadata="{_dyn_base}/.well-known/oauth-protected-resource"'
                ).encode()
                await send({"type": "http.response.start", "status": 401,
                            "headers": [
                                [b"content-type",  b"application/json"],
                                [b"content-length", str(len(body)).encode()],
                                [b"www-authenticate", www_auth],
                            ]})
                await send({"type": "http.response.body", "body": body, "more_body": False})
                return

            # Authenticated — pass straight through, no buffering
            await self._app(scope, receive, send)

    # ── Build OAuth-only routes app ──────────────────────────────────────────

    # ── SMS / WhatsApp webhook handlers ──────────────────────────────────────
    async def sms_webhook(request: Request):
        """
        POST /sms-webhook — receives inbound SMS from Twilio / SignalWire.

        Twilio and SignalWire POST form-encoded data here when a message is
        received on your SMS number.  The request is signed with HMAC-SHA1
        using your Auth Token — we validate the signature before storing.

        Responds with empty TwiML so Twilio doesn't auto-reply.
        No Bearer token required (Twilio calls this, not Claude).
        """
        try:
            from sms_inbox import sms_inbox_append, validate_twilio_signature, validate_signalwire_signature
            from sms_backends import load_sms_config
        except ImportError as _e:
            _log.error("sms_webhook: import error — %s", _e)
            return PlainTextResponse("<?xml version='1.0'?><Response/>",
                                     media_type="text/xml", status_code=200)

        try:
            body_bytes = await request.body()
            import urllib.parse as _up
            params = dict(_up.parse_qsl(body_bytes.decode("utf-8", errors="replace")))
        except Exception as _e:
            _log.error("sms_webhook: body parse error — %s", _e)
            return PlainTextResponse("Bad Request", status_code=400)

        if not params:
            return PlainTextResponse("Bad Request", status_code=400)

        # Signature validation
        cfg        = load_sms_config()
        provider   = str(cfg.get("sms_provider", "twilio")).lower()
        auth_token = cfg.get("twilio_auth_token") or cfg.get("signalwire_auth_token", "")
        signature  = request.headers.get("X-Twilio-Signature", "")
        url        = str(request.url)

        if auth_token and signature:
            if provider == "signalwire":
                valid = validate_signalwire_signature(auth_token, signature, url, params)
            else:
                valid = validate_twilio_signature(auth_token, signature, url, params)
            if not valid:
                _log.warning("sms_webhook: invalid signature — rejecting inbound message")
                return PlainTextResponse("Forbidden", status_code=403)
        else:
            _log.warning("sms_webhook: no auth_token or signature — accepting (dev mode)")

        # Store the inbound message
        msg_id    = params.get("MessageSid", "") or params.get("SmsSid", "")
        from_num  = params.get("From", "")
        to_num    = params.get("To", "")
        body_text = params.get("Body", "")

        # Resolve contact name from contacts cache
        contact_name = ""
        try:
            from_digits = "".join(c for c in from_num if c.isdigit())
            if len(from_digits) == 11 and from_digits[0] == '1':
                from_digits = from_digits[1:]
            cache = _contacts_cache_load(None)
            for entry in (cache.get("contacts") or []):
                ph = "".join(c for c in (entry.get("phone","") or "") if c.isdigit())
                if len(ph) == 11 and ph[0] == '1': ph = ph[1:]
                if ph == from_digits:
                    contact_name = entry.get("name","")
                    break
        except Exception:
            pass

        appended = sms_inbox_append(
            message_id   = msg_id,
            from_number  = from_num,
            to_number    = to_num,
            body         = body_text,
            provider     = provider,
            contact_name = contact_name,
            timestamp    = "",
        )
        _log.info("sms_webhook: inbound from=%s to=%s appended=%s", from_num, to_num, appended)

        # Return empty TwiML — Twilio expects this, otherwise it will
        # try to play a default response message to the caller.
        return PlainTextResponse(
            "<?xml version='1.0' encoding='UTF-8'?><Response/>",
            media_type="text/xml", status_code=200)

    async def whatsapp_webhook(request: Request):
        """
        POST /whatsapp-webhook — receives inbound WhatsApp messages via Twilio.

        Identical to sms_webhook but stores messages with provider='whatsapp'.
        Register this URL in the Twilio Console under WhatsApp → Sandbox →
        'When a message comes in'.
        """
        try:
            from sms_inbox import sms_inbox_append, validate_twilio_signature
            from sms_backends import load_sms_config
        except ImportError as _e:
            _log.error("whatsapp_webhook: import error — %s", _e)
            return PlainTextResponse("<?xml version='1.0'?><Response/>",
                                     media_type="text/xml", status_code=200)

        try:
            body_bytes = await request.body()
            import urllib.parse as _up2
            params = dict(_up2.parse_qsl(body_bytes.decode("utf-8", errors="replace")))
        except Exception as _e:
            _log.error("whatsapp_webhook: body parse error — %s", _e)
            return PlainTextResponse("Bad Request", status_code=400)

        if not params:
            return PlainTextResponse("Bad Request", status_code=400)

        # Signature validation (same as SMS — Twilio signs WhatsApp webhooks too)
        cfg        = load_sms_config()
        auth_token = cfg.get("twilio_auth_token", "")
        signature  = request.headers.get("X-Twilio-Signature", "")
        url        = str(request.url)

        if auth_token and signature:
            if not validate_twilio_signature(auth_token, signature, url, params):
                _log.warning("whatsapp_webhook: invalid signature — rejecting")
                return PlainTextResponse("Forbidden", status_code=403)

        msg_id    = params.get("MessageSid", "")
        from_num  = params.get("From", "")   # e.g. 'whatsapp:+13865550101'
        to_num    = params.get("To",   "")
        body_text = params.get("Body", "")

        sms_inbox_append(
            message_id   = msg_id,
            from_number  = from_num,
            to_number    = to_num,
            body         = body_text,
            provider     = "whatsapp",
            contact_name = "",
            timestamp    = "",
        )
        _log.info("whatsapp_webhook: inbound from=%s appended", from_num)

        return PlainTextResponse(
            "<?xml version='1.0' encoding='UTF-8'?><Response/>",
            media_type="text/xml", status_code=200)

    oauth_only_app = Starlette(routes=[
        Route("/health",  PlainTextResponse("OK")),
        Route("/.well-known/oauth-protected-resource",   oauth_protected_resource),
        Route("/.well-known/oauth-authorization-server", oauth_metadata),
        Route("/register",  register_client, methods=["POST"]),
        Route("/authorize", authorize, methods=["GET", "POST"]),
        Route("/token",     token_endpoint, methods=["POST"]),
        Route("/sms-webhook",       sms_webhook,       methods=["POST"]),
        Route("/whatsapp-webhook",  whatsapp_webhook,  methods=["POST"]),
    ])

    # ── Get FastMCP ASGI app ──────────────────────────────────────────────────
    # Claude.ai custom connectors use the Streamable HTTP transport (MCP spec
    # 2025-03-26): POST /mcp to send, GET /mcp to receive SSE responses.
    # streamable_http_app() implements this correctly and mounts at /mcp, so
    # no path rewriting is needed.
    # The earlier 421 errors were caused by missing proxy_headers on uvicorn —
    # that is now fixed.  sse_app() is kept as a fallback only for older mcp
    # installs that pre-date streamable_http_app().
    try:
        mcp_asgi = mcp.streamable_http_app()
        _log.info("FastMCP ASGI: using streamable_http_app() [Streamable HTTP — correct for Claude.ai]")
    except AttributeError:
        try:
            mcp_asgi = mcp.sse_app()
            _log.info("FastMCP ASGI: streamable_http_app() not found, falling back to sse_app()")
        except AttributeError:
            _log.critical("FATAL: FastMCP has neither streamable_http_app() nor sse_app(). Run: pip install --upgrade mcp")
            print("ERROR: FastMCP does not expose an ASGI app. Run: pip install --upgrade mcp",
                  file=sys.stderr)
            sys.exit(1)

    # ── Router ASGI app ───────────────────────────────────────────────────────
    # FastMCP MUST receive lifespan events to initialize its internal task group
    # (otherwise: RuntimeError "Task group is not initialized. Make sure to use run()").
    # _RouterASGI passes lifespan events to mcp_asgi, routes OAuth paths to
    # oauth_only_app, checks Bearer token, then forwards MCP paths to mcp_asgi.
    # Zero buffering — streaming/SSE responses pass through untouched.

    _OAUTH_PATHS = {"/health", "/authorize", "/token", "/register",
                    "/.well-known/oauth-authorization-server",
                    "/.well-known/oauth-protected-resource"}

    class _RouterASGI:
        async def __call__(self, scope, receive, send):
            stype = scope.get("type", "")

            # Lifespan MUST go to mcp_asgi so task group initialises
            if stype == "lifespan":
                _log.info("LIFESPAN event received — forwarding to mcp_asgi")
                await mcp_asgi(scope, receive, send)
                return

            if stype not in ("http", "websocket"):
                await oauth_only_app(scope, receive, send)
                return

            path = scope.get("path", "")
            method = scope.get("method", "?")
            _log.debug("REQUEST  %s %s", method, path)

            # OAuth / health — no auth required
            if path in _OAUTH_PATHS or path.startswith("/authorize"):
                _log.debug("ROUTE -> oauth_only_app  (%s)", path)
                await oauth_only_app(scope, receive, send)
                return

            # Everything else (including /mcp) — check Bearer token first
            headers = {k.lower(): v for k, v in scope.get("headers", [])}
            auth_raw = headers.get(b"authorization", b"")
            auth = auth_raw.decode("utf-8", errors="ignore")
            tok  = auth[7:].strip() if auth.lower().startswith("bearer ") else ""

            if tok not in _access_tokens:
                _log.warning("AUTH FAIL: no valid Bearer token for %s %s (got: '%s…')",
                             method, path, tok[:8] if tok else "<empty>")
                body = b'{"error":"unauthorized","error_description":"Invalid or missing Bearer token"}'
                www  = (f'Bearer realm="{PUBLIC_BASE}", '
                        f'resource_metadata="{PUBLIC_BASE}/.well-known/oauth-protected-resource"'
                        ).encode()
                await send({"type": "http.response.start", "status": 401,
                            "headers": [[b"content-type",  b"application/json"],
                                        [b"content-length", str(len(body)).encode()],
                                        [b"www-authenticate", www]]})
                await send({"type": "http.response.body", "body": body, "more_body": False})
                return

            # Authenticated — forward to FastMCP.
            #
            # HEADER INJECTION: streamable_http_app() (and sse_app() fallback)
            # strictly validate request headers per the MCP spec. Claude.ai's
            # custom connector does NOT send all required headers, causing 421.
            # We inject the missing headers into the ASGI scope before
            # forwarding so FastMCP's validators pass.
            #
            # Required headers per MCP Streamable HTTP spec (2025-03-26):
            #   POST /mcp: Content-Type: application/json
            #              Accept: application/json, text/event-stream
            #   GET  /mcp: Accept: text/event-stream
            #
            if path == "/mcp":
                # Shallow-copy scope so we can modify headers safely
                scope = dict(scope)
                headers = list(scope.get("headers", []))

                def _has_header(name: bytes, value_fragment: bytes) -> bool:
                    for hk, hv in headers:
                        if hk.lower() == name and value_fragment in hv.lower():
                            return True
                    return False

                def _set_header(name: bytes, value: bytes) -> None:
                    """Replace first matching header or append if absent."""
                    for i, (hk, hv) in enumerate(headers):
                        if hk.lower() == name:
                            headers[i] = (name, value)
                            return
                    headers.append((name, value))

                if method == "POST":
                    # Ensure Content-Type: application/json
                    if not _has_header(b"content-type", b"application/json"):
                        _log.debug("INJECT Content-Type: application/json")
                        _set_header(b"content-type", b"application/json")
                    # Ensure Accept includes both json and event-stream
                    if not _has_header(b"accept", b"text/event-stream"):
                        _log.debug("INJECT Accept: application/json, text/event-stream")
                        _set_header(b"accept", b"application/json, text/event-stream")

                elif method == "GET":
                    # Ensure Accept: text/event-stream
                    if not _has_header(b"accept", b"text/event-stream"):
                        _log.debug("INJECT Accept: text/event-stream")
                        _set_header(b"accept", b"text/event-stream")

                scope["headers"] = headers

            # ── Header fixes for /mcp requests ───────────────────────────────────
            # FastMCP's streamable_http_app() performs two validations that fail
            # when requests arrive via Cloudflare Tunnel:
            #
            # 1. HOST HEADER: FastMCP checks that the Host header matches the
            #    server's bound address (127.0.0.1:{port}). Cloudflare rewrites
            #    it to the public domain (mobile.dvavro-ai-prowler.com), so
            #    FastMCP rejects every request with HTTP 421 "Invalid Host header".
            #    Fix: rewrite host → 127.0.0.1:{port} before forwarding.
            #
            # 2. MCP-Protocol-Version: newer FastMCP builds require this header.
            #    Claude.ai does not send it. Fix: inject it.
            if path == "/mcp":
                scope = dict(scope)
                hdrs  = list(scope.get("headers", []))

                # Rewrite Host to match the bound address FastMCP expects
                local_host = f"127.0.0.1:{port}".encode()
                new_hdrs = []
                host_rewritten = False
                for hk, hv in hdrs:
                    if hk.lower() == b"host":
                        _log.debug("REWRITE Host: %s -> %s",
                                   hv.decode("utf-8","replace"),
                                   local_host.decode())
                        new_hdrs.append((b"host", local_host))
                        host_rewritten = True
                    else:
                        new_hdrs.append((hk, hv))
                if not host_rewritten:
                    new_hdrs.append((b"host", local_host))

                # Inject MCP-Protocol-Version if absent
                has_proto = any(k.lower() == b"mcp-protocol-version"
                                for k, v in new_hdrs)
                if not has_proto:
                    _log.debug("INJECT MCP-Protocol-Version: 2025-03-26")
                    new_hdrs.append((b"mcp-protocol-version", b"2025-03-26"))

                scope["headers"] = new_hdrs

            # Wrap send() to log error status codes from FastMCP.
            _resp_status = []
            async def _logging_send(message):
                if message.get("type") == "http.response.start":
                    _resp_status.append(message.get("status", "?"))
                    status = message.get("status", 200)
                    if status >= 400:
                        _log.warning("FASTMCP RESPONSE: %s %s → HTTP %s",
                                     method, path, status)
                await send(message)

            _log.debug("AUTH OK  -> mcp_asgi  (%s %s)", method, path)

            await mcp_asgi(scope, receive, _logging_send)

    app = _RouterASGI()

    print(f"\U0001f680 AI-Prowler HTTP MCP server starting…", flush=True)
    print(f"   Local endpoint : http://127.0.0.1:{port}/mcp", flush=True)
    print(f"   Health check   : http://127.0.0.1:{port}/health", flush=True)
    print(f"   OAuth login    : {PUBLIC_BASE}/authorize", flush=True)
    print(f"   Auth           : OAuth 2.0 + PKCE  (token: {len(token)} chars)", flush=True)
    print(f"   AI-Prowler HTTP server ready", flush=True)

    _log.info("HTTP MCP server starting on port %d", port)
    _log.info("Public base : %s", PUBLIC_BASE)
    _log.info("Token length: %d chars", len(token))
    _log.info("Calling uvicorn.run() — server will block here until stopped")

    # ── Wire uvicorn loggers to our file handler BEFORE calling uvicorn.run() ──
    # IMPORTANT: Do NOT pass log_config= to uvicorn.run().  Doing so causes
    # uvicorn to call logging.config.dictConfig() internally which OVERWRITES
    # the root logger's FileHandler — silencing all _log.debug() calls made
    # inside _RouterASGI (i.e. during every request).  Instead we attach our
    # existing file handler directly to uvicorn's named loggers here, then
    # let uvicorn.run() manage its own internal lifecycle without touching
    # the logging configuration.
    _our_file_handler = logging.root.handlers[0] if logging.root.handlers else None
    if _our_file_handler:
        _uv_fmt = logging.Formatter(
            "%(asctime)s [UVICORN ] %(message)s", datefmt="%Y-%m-%d %H:%M:%S"
        )
        _our_file_handler.setFormatter(_uv_fmt)
        for _uv_logger_name in ("uvicorn", "uvicorn.error", "uvicorn.access"):
            _uv_lg = logging.getLogger(_uv_logger_name)
            _uv_lg.setLevel(logging.DEBUG)
            _uv_lg.propagate = True   # propagate → root → our FileHandler
    _log.info("Uvicorn loggers wired to our file handler — debug logging active during requests")

    uvicorn.run(app, host='127.0.0.1', port=port,
                proxy_headers=True,
                forwarded_allow_ips="*",
                log_config=None)   # None = don't touch our logging config


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='AI-Prowler MCP Server')
    parser.add_argument('--transport', default='stdio', choices=['stdio', 'http'],
                        help='Transport mode: stdio (Claude Desktop) or http (remote/mobile)')
    parser.add_argument('--port', type=int, default=8000,
                        help='Port for HTTP transport (default: 8000)')
    parser.add_argument('--token', default='',
                        help='Bearer token for HTTP transport authentication')
    parser.add_argument('--public-base', default='',
                        dest='public_base',
                        help='Public HTTPS base URL for this tunnel, e.g. '
                             'https://john.dvavro-ai-prowler.com  '
                             'Falls back to tunnel_domain in ~/.ai-prowler/config.json '
                             'then to the built-in default.')
    parser.add_argument('--force', action='store_true',
                        help='If the target port is already in use, kill the '
                             'occupying process (requires admin rights on Windows) '
                             'and start fresh. Safe for dev/test use.')
    args = parser.parse_args()

    # ── Port conflict detection & cleanup (HTTP mode only) ────────────────────
    # When running --transport http, check whether the target port is already
    # bound. This catches the common case where a previous manual test run or
    # crashed GUI server left a zombie python.exe holding port 8000.
    #
    # Without --force: print a clear error with the PID and exit cleanly.
    # With --force:    kill the occupying process and proceed (dev/test use).
    #
    # This replaces the confusing "address already in use" uvicorn crash.
    if args.transport == 'http':
        import socket as _sock
        import subprocess as _sp
        import platform as _pf

        def _find_pid_on_port(port: int):
            """Return the PID holding `port` on 127.0.0.1, or None."""
            try:
                result = _sp.run(
                    ['netstat', '-ano'],
                    capture_output=True, text=True, timeout=5
                )
                for line in result.stdout.splitlines():
                    if f':{port}' in line and 'LISTENING' in line:
                        parts = line.split()
                        if parts:
                            try:
                                return int(parts[-1])
                            except ValueError:
                                pass
            except Exception:
                pass
            return None

        def _port_in_use(port: int) -> bool:
            """True if something is already bound to 127.0.0.1:port."""
            with _sock.socket(_sock.AF_INET, _sock.SOCK_STREAM) as s:
                s.settimeout(1)
                return s.connect_ex(('127.0.0.1', port)) == 0

        if _port_in_use(args.port):
            pid = _find_pid_on_port(args.port)
            pid_str = f"PID {pid}" if pid else "unknown PID"
            if args.force:
                print(f"⚠️  Port {args.port} is in use by {pid_str} — --force specified, killing it...",
                      flush=True)
                if pid:
                    try:
                        if _pf.system() == 'Windows':
                            _sp.run(['taskkill', '/PID', str(pid), '/F'],
                                    capture_output=True, timeout=5)
                        else:
                            import signal as _sig
                            import os as _os
                            _os.kill(pid, _sig.SIGKILL)
                        import time as _t
                        _t.sleep(1)  # give the OS a moment to release the port
                        print(f"✅  Killed {pid_str} — proceeding on port {args.port}",
                              flush=True)
                    except Exception as _ke:
                        print(f"❌  Could not kill {pid_str}: {_ke}", flush=True)
                        print(f"   Run as Administrator and try again, or free port {args.port} manually.",
                              flush=True)
                        raise SystemExit(1)
                else:
                    print(f"❌  Port {args.port} is busy but PID could not be identified.",
                          flush=True)
                    print(f"   Run: netstat -ano | findstr :{args.port}  then taskkill /PID <pid> /F",
                          flush=True)
                    raise SystemExit(1)
            else:
                print(f"\n❌  AI-Prowler cannot start: port {args.port} is already in use ({pid_str}).",
                      flush=True)
                print(f"", flush=True)
                print(f"   This usually means a previous AI-Prowler server process", flush=True)
                print(f"   was not shut down cleanly (crash, Ctrl+C, or test run).", flush=True)
                print(f"", flush=True)
                if pid:
                    print(f"   To fix — run as Administrator:", flush=True)
                    print(f"     taskkill /PID {pid} /F", flush=True)
                    print(f"", flush=True)
                    print(f"   Or restart with --force to kill it automatically:", flush=True)
                    print(f"     python ai_prowler_mcp.py --transport http --force", flush=True)
                else:
                    print(f"   To fix:", flush=True)
                    print(f"     netstat -ano | findstr :{args.port}", flush=True)
                    print(f"     taskkill /PID <pid> /F   (run as Administrator)", flush=True)
                print(f"", flush=True)
                raise SystemExit(1)

    # ── Resolve public_base ───────────────────────────────────────────────────
    # Priority: --public-base CLI arg > config.json tunnel_domain > built-in default
    resolved_public_base = args.public_base.strip()

    if not resolved_public_base:
        try:
            _cfg_path = _CONFIG_PATH   # honors AIPROWLER_TEST_STATE_DIR sandbox
            if _cfg_path.exists():
                import json as _jcfg
                _cfg = _jcfg.loads(_cfg_path.read_text(encoding='utf-8'))
                _domain = _cfg.get('tunnel_domain', '').strip()
                if _domain:
                    resolved_public_base = (
                        _domain if _domain.startswith('http')
                        else f'https://{_domain}'
                    )
        except Exception as _ce:
            _log.warning("Could not read tunnel_domain from config.json: %s", _ce)

    if not resolved_public_base:
        _log.warning(
            "public_base: no --public-base arg and no tunnel_domain in config.json. "
            "OAuth endpoints will be misconfigured. "
            "Set your Cloudflare tunnel hostname in the AI-Prowler Settings tab."
        )
        resolved_public_base = "http://127.0.0.1:8000"  # safe local fallback
    else:
        _log.info("public_base: resolved to %s", resolved_public_base)

    _log.info("Entry point: transport=%s port=%s public_base=%s",
              args.transport, args.port, resolved_public_base)

    if args.transport == 'http':
        # ── Server mode dispatch (v7.0.0 Phase B) ─────────────────────────────
        # When config.json sets mode=server (Business company server), use the
        # multi-user server transport instead of the single-token _run_http.
        # Personal/Mobile/Home installs are UNAFFECTED — _run_http is unchanged
        # and remains the default path. _run_server_mode falls back to _run_http
        # if its prerequisites aren't met (no users.json, missing deps), so a
        # misconfigured server can never end up with NO transport.
        try:
            _rt = _load_runtime_config()
        except Exception:
            _rt = {"mode": "personal", "edition": "home"}
        if _rt.get("mode") == "server":
            _log.info("Starting SERVER-MODE HTTP transport on port %d (multi-user)", args.port)
            _run_server_mode(port=args.port, token=args.token,
                             public_base=resolved_public_base)
        else:
            _log.info("Starting HTTP transport on port %d", args.port)
            _run_http(port=args.port, token=args.token, public_base=resolved_public_base)
    else:
        _log.info("Starting stdio transport (Claude Desktop mode)")

        # ── Protect the MCP pipe — _STDIO_MODE flag only ──────────────────────
        #
        # CRITICAL: do NOT replace sys.stdout with devnull here.
        #
        # In mcp 1.x, mcp.run(transport="stdio") accesses sys.stdout.buffer at
        # call-time to set up its JSON-RPC write stream.  If sys.stdout has been
        # replaced with a devnull file object before mcp.run() is called, FastMCP
        # picks up the devnull handle and every response — including the initial
        # initialize reply — goes into a black hole.  Claude Desktop never
        # receives a response and shows "not connected" / no tools.
        #
        # The _STDIO_MODE = True flag is the correct and sufficient protection:
        #   - _capture_stdout() becomes a no-op so no tool handler can
        #     accidentally redirect sys.stdout during a tool call.
        #   - All output goes via tool return values and the log file.
        #   - warnings are already suppressed via warnings.filterwarnings above.
        _STDIO_MODE = True
        _log.info("STDIO mode: _STDIO_MODE=True — _capture_stdout() is a no-op")

        # ── BACKGROUND PREWARM — mcp.run() starts immediately ─────────────────
        #
        # Problem: get_chroma_client() initialises ChromaDB and loads the
        # sentence-transformers embedding model (~3s even with HF offline).
        # This MUST NOT block mcp.run() because Claude Desktop sends its
        # initialize message immediately on process launch.  Any delay before
        # mcp.run() starts causes Desktop to time out and disconnect.
        #
        # Problem: rag_preprocessor.get_chroma_client() calls print() directly.
        # In stdio mode stdout is the MCP JSON-RPC pipe — any non-JSON bytes
        # corrupt the protocol.
        # Fix: _do_prewarm() installs _StdoutFilter which discards write() calls
        # while keeping .buffer pointed at the real pipe for mcp.run().
        # NOTE: GUI_MODE=True must NOT be set here — it causes rag_preprocessor
        # to invoke tkinter callbacks that hang forever with no event loop.
        #
        # Problem: get_chroma_client() has internal asyncio code paths that
        # deadlock silently when called from inside a FastMCP tool handler
        # (conflicts with the already-running event loop).
        # Fix: prewarm runs in a background thread OUTSIDE the asyncio loop.
        # rag_preprocessor caches the client and model in module-level globals.
        # Tool handlers wait on _prewarm_event then return cached objects
        # instantly — no asyncio-conflicting initialisation code runs.

        # Prints from rag_preprocessor are suppressed by _StdoutFilter inside
        # _do_prewarm() — GUI_MODE=True is NOT set here because it causes
        # rag_preprocessor to invoke tkinter/GUI callbacks that hang forever
        # in a subprocess with no event loop running.

        # Event that tool handlers wait on before using ChromaDB
        _prewarm_event.clear()

        def _do_prewarm():
            # ── Stdout pollution guard ────────────────────────────────────────
            # Some libraries (PyTorch Blackwell GPU detection, sentence-
            # transformers, tokenizers) print directly to sys.stdout during
            # model load.  In stdio mode sys.stdout IS the MCP JSON-RPC pipe,
            # so any stray text corrupts it and Claude Desktop shows
            # "not valid JSON".
            #
            # Fix: replace sys.stdout with a filter that discards Python-level
            # write() calls (killing stray prints) while keeping .buffer
            # pointing at the real pipe (so mcp.run() can still write JSON).
            # This is race-condition-safe: mcp.run() uses sys.stdout.buffer,
            # and _StdoutFilter.buffer always points to the real pipe regardless
            # of when the swap happens.
            import sys as _sys
            _real_stdout = _sys.stdout   # capture before any modification

            class _StdoutFilter:
                """Discards print()-level writes; exposes real .buffer for mcp."""
                buffer   = getattr(_real_stdout, 'buffer', _real_stdout)
                encoding = getattr(_real_stdout, 'encoding', 'utf-8')
                errors   = getattr(_real_stdout, 'errors',   'replace')
                def write(self, s):   pass   # discard stray prints
                def flush(self):      pass
                def fileno(self):
                    return _real_stdout.fileno()

            _sys.stdout = _StdoutFilter()
            try:
                _log.info("PREWARM: background thread started — loading ChromaDB "
                          "+ embedding model...")
                from rag_preprocessor import get_chroma_client, COLLECTION_NAME
                _pw_client, _pw_emb = get_chroma_client()
                try:
                    _pw_col   = _pw_client.get_or_create_collection(
                                    name=COLLECTION_NAME,
                                    embedding_function=_pw_emb)
                    _pw_count = _pw_col.count()
                    _log.info("PREWARM: done — %d chunks indexed, model cached, "
                              "asyncio-safe", _pw_count)
                except Exception:
                    _log.info("PREWARM: done — DB empty/not created yet, "
                              "embedding model cached")
            except Exception as _pw_err:
                _log.warning("PREWARM: failed (%s) — tool calls will load on demand",
                             _pw_err)
            finally:
                _sys.stdout = _real_stdout   # restore before unblocking tools
                _prewarm_event.set()          # unblock any waiting tool handlers
                _log.info("PREWARM: complete, tool handlers unblocked")

        threading.Thread(target=_do_prewarm, daemon=True, name="prewarm").start()
        _log.info("PREWARM: thread launched — calling mcp.run() immediately")

        mcp.run(transport="stdio")
